"""
Вставка изображений и вложений в заполняемые Word-формы ТО.
"""
from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff", ".webp"}


def build_attachments_map(
    document_files: Optional[List[Dict[str, Any]]],
) -> Dict[str, str]:
    """document_number → file_path."""
    result: Dict[str, str] = {}
    if not document_files:
        return result
    for f in document_files:
        if not isinstance(f, dict):
            continue
        dn = str(f.get("document_number") or "").strip()
        fp = f.get("file_path")
        if dn and isinstance(fp, str) and fp.strip():
            result[dn] = fp.strip()
    return result


def resolve_image_path(
    path: Optional[str],
    find_image: Optional[Callable[[str], Optional[str]]] = None,
) -> Optional[str]:
    if not path or not isinstance(path, str):
        return None
    path = path.strip().replace("\\", "/")
    if find_image:
        resolved = find_image(path)
        if resolved and os.path.isfile(resolved):
            return resolved
    if os.path.isfile(path):
        return path
    # Относительные пути uploads
    for base in ("/app/uploads", "uploads", "backend/uploads"):
        candidate = Path(base) / path.lstrip("/")
        if candidate.is_file():
            return str(candidate)
    return None


def is_image_file(path: str) -> bool:
    return Path(path).suffix.lower() in _IMAGE_EXTS


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    """Вставить новый абзац сразу после указанного."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def add_picture_after_paragraph(
    paragraph: Paragraph,
    image_path: str,
    width_inches: float = 5.2,
    caption: Optional[str] = None,
) -> Optional[Paragraph]:
    """Вставить картинку (и опционально подпись) после абзаца."""
    try:
        if caption:
            cap = insert_paragraph_after(paragraph, caption)
            try:
                cap.runs[0].bold = True
                cap.runs[0].font.size = Pt(10)
            except Exception:
                pass
            anchor = cap
        else:
            anchor = paragraph
        pic_p = insert_paragraph_after(anchor, "")
        run = pic_p.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return pic_p
    except Exception as exc:
        logger.warning("Не удалось вставить изображение %s: %s", image_path, exc)
        return None


def find_paragraph_containing(doc: Document, needle: str) -> Optional[Paragraph]:
    needle_l = needle.lower()
    for p in doc.paragraphs:
        if needle_l in (p.text or "").lower():
            return p
    return None


def find_all_paragraphs_containing(doc: Document, needle: str) -> List[Paragraph]:
    needle_l = needle.lower()
    return [p for p in doc.paragraphs if needle_l in (p.text or "").lower()]


def add_heading_block(
    doc: Document,
    title: str,
    after_paragraph: Optional[Paragraph] = None,
) -> Paragraph:
    """Добавить заголовок в конец документа или после абзаца."""
    if after_paragraph is not None:
        p = insert_paragraph_after(after_paragraph, title)
    else:
        p = doc.add_paragraph(title)
    try:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)
    except Exception:
        pass
    return p


def collect_scheme_paths(
    data: Dict[str, Any],
    attachments: Dict[str, str],
) -> List[Dict[str, str]]:
    """Схемы контроля: ВИК + УЗТ."""
    items: List[Dict[str, str]] = []
    for key in ("control_scheme_image", "control_scheme", "base_vessel_scheme_image"):
        path = attachments.get(key) or data.get(key)
        if isinstance(path, str) and path.strip():
            items.append({"label": "Базовая схема сосуда (ВИК)", "path": path})
            break
    # Базовая схема из структурированного объекта
    base = data.get("base_vessel_scheme")
    if isinstance(base, dict):
        bp = base.get("image_path") or base.get("scheme_image_path") or base.get("path")
        if isinstance(bp, str) and bp.strip():
            if not any(it["path"] == bp for it in items):
                items.append({"label": "Базовая схема сосуда (ВИК)", "path": bp})
    # Схема подключения (файл)
    for key, label in (
        ("connection_scheme_file", "Схема подключения сосуда"),
        ("connection_scheme_image", "Схема подключения сосуда"),
    ):
        path = attachments.get(key) or data.get(key)
        if isinstance(path, str) and path.strip():
            if not any(it["path"] == path for it in items):
                items.append({"label": label, "path": path})
    schemes = data.get("uzt_schemes") or []
    if isinstance(schemes, list):
        for i, s in enumerate(schemes):
            if not isinstance(s, dict):
                continue
            path = (
                s.get("scheme_image_path")
                or s.get("image_path")
                or s.get("path")
                or attachments.get(f"uzt_scheme_{i}")
            )
            if isinstance(path, str) and path.strip():
                label = str(s.get("label") or s.get("name") or f"Схема УЗТ №{i + 1}")
                items.append({"label": label, "path": path})
    # Техкарта / схема с задания (веб)
    for key, label in (
        ("tech_card_file_path", "Технологическая карта / схема контроля"),
        ("tech_card_scheme_path", "Схема контроля (техкарта)"),
        ("scheme_file_path", "Схема контроля"),
    ):
        path = attachments.get(key) or data.get(key)
        if isinstance(path, str) and path.strip():
            if not any(it["path"] == path for it in items):
                items.append({"label": label, "path": path})
    # Вложения uzt_scheme_* без записи в uzt_schemes
    for key, path in attachments.items():
        if key.startswith("uzt_scheme_") and "_point_" not in key:
            if not any(it["path"] == path for it in items):
                items.append({"label": f"Схема УЗТ ({key})", "path": path})
    return items


def collect_photo_paths(
    data: Dict[str, Any],
    attachments: Dict[str, str],
) -> List[Dict[str, str]]:
    """Фото точек УЗТ, дефектов ВИК, объекта."""
    items: List[Dict[str, str]] = []
    thickness = data.get("thickness_measurements") or data.get("thicknessMeasurements") or []
    if isinstance(thickness, list):
        for i, point in enumerate(thickness):
            if not isinstance(point, dict):
                continue
            photos = point.get("photos") or []
            if isinstance(photos, list):
                for j, ph in enumerate(photos):
                    if isinstance(ph, str) and ph.strip():
                        num = point.get("point_number") or point.get("number") or (i + 1)
                        items.append({"label": f"Фото точки УЗТ №{num}", "path": ph})
            att = attachments.get(f"uzt_point_{i}_0") or attachments.get(f"uzt_point_{i}_1")
            if att and not any(it["path"] == att for it in items):
                items.append({"label": f"Фото точки УЗТ #{i + 1}", "path": att})
    for key, path in attachments.items():
        if key.startswith("uzt_point_") and not any(it["path"] == path for it in items):
            items.append({"label": f"Фото УЗТ ({key})", "path": path})
        if key.startswith("vd_") and not any(it["path"] == path for it in items):
            items.append({"label": f"Фото дефекта ВИК ({key})", "path": path})
        if key.startswith("object_photo_") and not any(it["path"] == path for it in items):
            items.append({"label": f"Фото объекта ({key})", "path": path})
    defects = data.get("visual_defects") or data.get("defects") or []
    if isinstance(defects, list):
        for i, d in enumerate(defects):
            if not isinstance(d, dict):
                continue
            photos = d.get("photos") or []
            if isinstance(photos, list):
                for ph in photos:
                    if isinstance(ph, str) and ph.strip():
                        items.append(
                            {
                                "label": f"Фото дефекта: {d.get('location') or d.get('description') or i + 1}",
                                "path": ph,
                            }
                        )
    return items


def collect_hydraulic_act_paths(
    data: Dict[str, Any],
    attachments: Dict[str, str],
) -> List[str]:
    """Сканы акта гидравлического испытания."""
    paths: List[str] = []
    for key in (
        "hydraulic_test_act",
        "hydro_act",
        "hydraulic_act",
        "act_hydraulic",
        "hydrostatic_act",
    ):
        p = attachments.get(key) or data.get(key)
        if isinstance(p, str) and p.strip():
            paths.append(p)
    # Часто акт кладут как документ №15 (ремонтная) или отдельный ключ
    for key, path in attachments.items():
        kl = key.lower()
        if "hydraulic" in kl or "гидро" in kl or "hydro" in kl:
            if path not in paths:
                paths.append(path)
    # Явный список в data
    extra = data.get("hydraulic_act_files") or data.get("hydro_act_scans") or []
    if isinstance(extra, list):
        for p in extra:
            if isinstance(p, str) and p.strip() and p not in paths:
                paths.append(p)
            elif isinstance(p, dict):
                fp = p.get("file_path") or p.get("path")
                if isinstance(fp, str) and fp.strip() and fp not in paths:
                    paths.append(fp)
    return paths


def insert_media_block(
    doc: Document,
    anchor_text: str,
    items: Sequence[Dict[str, str]],
    find_image: Optional[Callable[[str], Optional[str]]] = None,
    width_inches: float = 5.2,
    max_items: int = 12,
) -> int:
    """Вставить изображения после абзаца, содержащего anchor_text. Возвращает число вставленных."""
    if not items:
        return 0
    anchors = find_all_paragraphs_containing(doc, anchor_text)
    if not anchors:
        # В конец документа
        title = doc.add_paragraph(anchor_text)
        try:
            title.runs[0].bold = True
        except Exception:
            pass
        anchors = [title]
    inserted = 0
    anchor = anchors[0]
    last = anchor
    for item in list(items)[:max_items]:
        path = resolve_image_path(item.get("path"), find_image)
        if not path:
            continue
        if is_image_file(path):
            pic = add_picture_after_paragraph(
                last, path, width_inches=width_inches, caption=item.get("label")
            )
            if pic is not None:
                last = pic
                inserted += 1
        else:
            note = insert_paragraph_after(
                last,
                f"{item.get('label') or 'Документ'}: {Path(path).name}",
            )
            last = note
            inserted += 1
    return inserted
