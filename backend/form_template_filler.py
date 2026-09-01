"""
Заполнение Word-формы ТО «Приложение № 1. Обследование сосудов и аппаратов»
данными обследования из мобильного приложения.

Структура шаблона (таблицы/приложения) сохраняется; подставляются
паспортные данные сосуда и результаты измерений (ВИК, УЗТ, твердость, УЗК, МПК).
"""
from __future__ import annotations

import logging
import re
import shutil
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from report_forms_registry import resolve_form_path
from report_org_settings import load_report_org_settings
from technical_report_builder import TO_DOCUMENT_NAMES
from form_media_helpers import (
    build_attachments_map,
    collect_hydraulic_act_paths,
    collect_photo_paths,
    collect_scheme_paths,
    find_all_paragraphs_containing,
    find_paragraph_containing,
    insert_media_block,
    insert_paragraph_after,
    is_image_file,
    resolve_image_path,
    add_picture_after_paragraph,
    clear_pictures_after_paragraph,
)
from scheme_ndt_overlays import (
    LAYER_ORDER,
    layer_title,
    png_to_tempfile,
    render_all_layer_pngs,
)

logger = logging.getLogger(__name__)

MISSING = "—"
NOT_PROVIDED = "Не предоставлено"
_BLANK_RE = re.compile(r"_+")


def _merge_report_instruments(
    verification_equipment: Optional[List[Dict[str, Any]]],
    data: Dict[str, Any],
    ndt_methods: Optional[List[Dict[str, Any]]] = None,
) -> List[Dict[str, Any]]:
    """Собрать приборы для табл. «Перечень приборов»: реестр + ручной ввод + НК."""
    out: List[Dict[str, Any]] = []
    seen: set = set()

    def _key(eq: Dict[str, Any]) -> Tuple[str, str]:
        name = str(eq.get("name") or eq.get("equipment") or "").strip().lower()
        serial = str(
            eq.get("serial_number")
            or eq.get("factory_number")
            or eq.get("equipment_serial")
            or ""
        ).strip().lower()
        return name, serial

    def _add(eq: Any) -> None:
        if not isinstance(eq, dict):
            return
        name = str(eq.get("name") or eq.get("equipment") or "").strip()
        if not name:
            return
        k = _key(eq)
        for s in seen:
            if s[0] == k[0] and (not k[1] or not s[1] or s[1] == k[1]):
                return
        seen.add(k)
        out.append(eq)

    for e in verification_equipment or []:
        _add(e)

    ad = data.get("additional_data") if isinstance(data.get("additional_data"), dict) else {}
    for e in ad.get("manual_verification_equipment") or ad.get("manualVerificationEquipment") or []:
        _add(e)

    for e in data.get("_ndt_instruments") or []:
        _add(e)

    for key in ("instruments", "selected_instruments", "verification_equipment"):
        raw = data.get(key)
        if isinstance(raw, list):
            for e in raw:
                _add(e)

    for m in ndt_methods or []:
        if not isinstance(m, dict):
            continue
        eq_name = m.get("equipment") or (m.get("additional_data") or {}).get("equipment")
        if not eq_name:
            continue
        ad_m = m.get("additional_data") if isinstance(m.get("additional_data"), dict) else {}
        _add(
            {
                "name": eq_name,
                "serial_number": m.get("equipment_serial")
                or m.get("serial_number")
                or ad_m.get("serial_number")
                or ad_m.get("device_serial")
                or "",
                "equipment_type": m.get("method_code") or m.get("method_name") or "",
                "model": ad_m.get("device_type") or ad_m.get("model") or "",
                "verification_certificate_number": ad_m.get("verification_certificate_number")
                or "",
                "next_verification_date": ad_m.get("next_verification_date") or "",
            }
        )
    return out


def _renumber_table_column(table: Table, start_row: int = 1, col: int = 0) -> None:
    """Перенумеровать № п/п после удаления пустых строк."""
    n = 1
    for r in range(start_row, len(table.rows)):
        cells = table.rows[r].cells
        if col >= len(cells):
            continue
        # пропускаем полностью пустые (на случай)
        texts = [(c.text or "").strip() for c in cells]
        if not any(t and t not in ("—", "-", "–") for i, t in enumerate(texts) if i != col):
            continue
        _set(table, r, col, f"{n}.")
        n += 1


_TABLE_CAPTION_RE = re.compile(r"^\s*Таблица\s*(?:№|No|N)?\s*\d", re.IGNORECASE)


def _keep_para(p: Paragraph) -> None:
    try:
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.widow_control = True
    except Exception:
        pass


def _norm_ws(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\xa0", " ")).strip()


def _force_horizontal_text(cell: _Cell) -> None:
    """Горизонтальное направление текста в ячейке (шаблон to-1 часто задаёт вертикальное)."""
    try:
        tcPr = cell._tc.get_or_add_tcPr()
        for td in list(tcPr.findall(qn("w:textDirection"))):
            tcPr.remove(td)
    except Exception:
        pass


def _strip_parens(text: str) -> str:
    return re.sub(r"\s+", " ", re.sub(r"[()]", " ", text or "")).strip()


def _doc_is_incomplete(text: str) -> bool:
    t = (text or "").lower()
    return any(
        k in t
        for k in (
            "частичн",
            "отсутств",
            "не предостав",
            "не в полном",
            "неполн",
        )
    )


def _doc_analysis_result_cell(ctx: Dict[str, Any]) -> str:
    raw = _strip_parens(str(ctx.get("conclusion_doc") or ""))
    if not raw:
        return "Соответствует требованиям"
    if _doc_is_incomplete(raw):
        return raw[0].upper() + raw[1:] if raw else raw
    if "полном" in raw.lower():
        return "Соответствует требованиям"
    return raw[0].upper() + raw[1:] if raw else raw


def _doc_verdict_word(ctx: Dict[str, Any]) -> str:
    raw = str(ctx.get("conclusion_doc") or "")
    if _doc_is_incomplete(raw):
        if "отсутств" in raw.lower() or "не предостав" in raw.lower():
            return "не соответствует"
        return "соответствует частично"
    return "соответствует"


def _ndt_items_have_defects(items: Any) -> bool:
    if not isinstance(items, list):
        return False
    for it in items:
        if not isinstance(it, dict):
            continue
        blob = " ".join(
            str(it.get(k) or "")
            for k in (
                "conclusion",
                "assessment",
                "quality",
                "defects",
                "description",
                "defect_description",
                "uzk_defect",
                "pvk_defect",
                "indication",
            )
        ).lower()
        if not blob.strip():
            continue
        if any(k in blob for k in ("ремонт", "негоден", "не годен", "недопустим", "брак")):
            if "не обнаружен" not in blob:
                return True
        if "дефект" in blob and "не обнаружен" not in blob and "дефектов не" not in blob:
            return True
    return False


def _ndt_result_summary(
    ctx: Dict[str, Any],
    *,
    data_keys: Sequence[str],
    custom_key: str,
    ok_text: str,
    defect_text: str,
) -> str:
    g = ctx["g"]
    custom = str(g(custom_key, default="") or "").strip()
    items = None
    for k in data_keys:
        items = g(k, default=None)
        if isinstance(items, list) and items:
            break
    has_defects = _ndt_items_have_defects(items)
    if has_defects:
        if custom and any(
            k in custom.lower() for k in ("дефект", "ремонт", "негоден", "не годен", "брак")
        ):
            return custom
        return defect_text
    if custom:
        return custom
    return ok_text


def _keep_table_head(table: Table, max_rows: int = 2) -> None:
    """Не отрывать шапку/первые строки таблицы от заголовка раздела."""
    try:
        for r in range(min(max_rows, len(table.rows))):
            for cell in table.rows[r].cells:
                for p in cell.paragraphs:
                    _keep_para(p)
    except Exception:
        pass


def _is_table_caption(text: str) -> bool:
    """«Таблица № 3», «Таблица No 3», «Таблица 3», «Таблица№3» — всё это подписи."""
    return bool(_TABLE_CAPTION_RE.match(text or ""))


def _apply_heading_keep_with_next(doc: Document) -> None:
    """Не отрывать заголовок раздела / «Таблица №» / СОДЕРЖАНИЕ от следующего контента."""
    section_re = re.compile(r"^\d+\.\s+\S")
    # Основной отчёт: заголовки вида «6. Сведения…», не пункты приложений с длинным текстом
    main_section_re = re.compile(
        r"^\d{1,2}\.\s+(Основания|Сроки|Перечень|Сведения|Краткая|Результаты|"
        r"Заключение|Выводы|Рекомендации|Оценка|Расчёт|Расчет)"
    )
    paras = list(_iter_all_paragraphs(doc))
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if not t:
            continue
        keep = False
        if t == "СОДЕРЖАНИЕ":
            keep = True
            for j in range(i + 1, min(i + 20, len(paras))):
                tj = (paras[j].text or "").strip()
                if not tj:
                    continue
                if "…" in tj or "..." in tj or re.match(r"^\d+\.", tj):
                    _keep_para(paras[j])
                elif tj.startswith("1.") or "Основания" in tj:
                    _keep_para(paras[j])
                    break
                else:
                    break
        elif main_section_re.match(t) and len(t) < 160 and "…" not in t and "..." not in t:
            keep = True
        elif section_re.match(t) and len(t) < 120 and "…" not in t and "..." not in t:
            # короткие нумерованные заголовки разделов основного отчёта
            keep = True
        elif _is_table_caption(t):
            keep = True
        elif t.startswith("Схема ") or t.startswith("4. Схема") or "Параметры контроля" in t:
            keep = True
        elif t.startswith("Результаты контроля") or t.startswith("5. Результаты") or t.startswith("3. Результаты") or "Материалы элементов" in t:
            keep = True
        elif t.startswith("Результаты визуального") or t.startswith("4. Результаты"):
            keep = True
        if keep:
            _keep_para(p)
            # «Таблица № N» сразу после заголовка раздела
            if i + 1 < len(paras):
                nxt = (paras[i + 1].text or "").strip()
                if nxt.startswith("Таблица"):
                    _keep_para(paras[i + 1])
            # первая таблица после заголовка в том же SDT/родителе
            try:
                cur = p._p.getnext()
                hops = 0
                while cur is not None and hops < 8:
                    hops += 1
                    tag = cur.tag
                    if tag == qn("w:p"):
                        txt = "".join(x.text or "" for x in cur.iter(qn("w:t"))).strip()
                        if txt.startswith("Таблица"):
                            cur = cur.getnext()
                            continue
                        if txt:
                            break
                        cur = cur.getnext()
                        continue
                    if tag == qn("w:tbl"):
                        _keep_table_head(Table(cur, doc), 2)
                        break
                    if tag == qn("w:sdt"):
                        content = cur.find(qn("w:sdtContent"))
                        if content is not None:
                            tbl = content.find(qn("w:tbl"))
                            if tbl is not None:
                                _keep_table_head(Table(tbl, doc), 2)
                                break
                    cur = cur.getnext()
            except Exception:
                pass

_DATE_HEADER_RE = re.compile(r"^\s*дата\b")


def _widen_date_columns_everywhere(doc: Document) -> None:
    """Во всех таблицах расширить колонки с заголовком «Дата» (обычный шрифт)."""
    seen: set = set()
    tables: List[Table] = []
    try:
        for t in _main_sdt_tables(doc):
            if id(t._tbl) not in seen:
                seen.add(id(t._tbl))
                tables.append(t)
    except Exception:
        pass
    for t in doc.tables:
        if id(t._tbl) not in seen:
            seen.add(id(t._tbl))
            tables.append(t)

    for table in tables:
        if not table.rows:
            continue
        try:
            header_cells = table.rows[0].cells
        except Exception:
            continue
        for col, cell in enumerate(header_cells):
            if _DATE_HEADER_RE.match((cell.text or "").strip().lower()):
                _widen_date_column(table, col, 3.4)


def finalize_official_form(doc: Document, form_id: str = "") -> None:
    """
    Общая «полировка» любой официальной формы ТО — то же, что получает to-1:
    ширина колонок «Дата», единая типографика таблиц с landscape для широких
    и keep-with-next, чтобы заголовок «Таблица № N» не отрывался от таблицы.
    """
    tag = form_id or "form"
    try:
        _widen_date_columns_everywhere(doc)
    except Exception:
        logger.exception("%s: не удалось расширить колонки «Дата»", tag)
    try:
        _finalize_table_typography(doc)
    except Exception:
        logger.exception("%s: не удалось унифицировать шрифт/ориентацию таблиц", tag)
    try:
        _apply_heading_keep_with_next(doc)
    except Exception:
        logger.exception("%s: не удалось применить keep-with-next к заголовкам", tag)


def _table_header_blob(table: Table, rows: int = 2) -> str:
    parts: List[str] = []
    try:
        for r in table.rows[:rows]:
            for c in r.cells:
                parts.append((c.text or "").replace("\n", " ").strip().lower())
    except Exception:
        return ""
    return " | ".join(parts)


def apply_ndt_protocol_tables(doc: Document, ctx: Dict[str, Any]) -> set:
    """
    Заполнить протокольные таблицы НК (ВИК/УЗТ/твёрдость/УЗК/МПК) в любой форме.

    В to-1 таблицы адресуются фиксированными индексами; здесь они находятся по
    заголовкам, поэтому те же данные попадают в формы to-2…to-44, у которых
    порядок и количество таблиц другие.

    Возвращает id(table._tbl) заполненных таблиц, чтобы более грубые эвристики
    формы их потом не перезаписали.
    """
    filled: set = set()
    seen: set = set()
    tables: List[Table] = []
    try:
        for t in _main_sdt_tables(doc):
            if id(t._tbl) not in seen:
                seen.add(id(t._tbl))
                tables.append(t)
    except Exception:
        pass
    for t in doc.tables:
        if id(t._tbl) not in seen:
            seen.add(id(t._tbl))
            tables.append(t)

    for table in tables:
        if not table.rows:
            continue
        blob = _table_header_blob(table)
        if not blob:
            continue
        try:
            if "шероховатость поверхности" in blob and "освещ" in blob:
                _fill_vik_parameters(table, ctx)
            elif "способ контроля" in blob and "уровень чувствительности" in blob:
                _fill_mpk_parameters(table, ctx)
            elif "№ стыка по карте контроля" in blob or "номер дефекта" in blob:
                _fill_uzk_results(table, ctx)
                _strip_empty_rows(table, 1, ignore_cols=(0,))
            elif "тип сварного соединения" in blob and "пэп" in blob:
                _fill_uzk_parameters(table, ctx)
                _strip_empty_rows(table, 1)
            elif "участок контроля согласно схемы измерения" in blob:
                _fill_hardness_list(table, ctx)
                _strip_empty_rows(table, 1, ignore_cols=(0,))
            elif "допустимая твердость металла в точке" in blob:
                _fill_hardness_matrix(table, ctx)
                _strip_empty_rows(table, 3, ignore_cols=(0,))
            elif "наименование элемента" in blob and "№ точки" in blob and "толщин" in blob:
                _fill_uzt_results(table, ctx)
            elif "оценка качества" in blob and "объем контроля" in blob or (
                "оценка качества" in blob and "объём контроля" in blob
            ):
                # МПК отличается от ВИК наличием колонки «Зона контроля»
                if "зона" in blob:
                    _fill_mpk_results(table, ctx)
                    _strip_empty_rows(table, 1, ignore_cols=(0,))
                else:
                    _fill_vik_results(table, ctx)
            else:
                continue
        except Exception:
            logger.debug("НК-таблица не заполнена: %s", blob[:120], exc_info=True)
            continue
        filled.add(id(table._tbl))
    return filled


def fill_vessel_form_to1(
    inspection_data: Dict[str, Any],
    equipment_data: Dict[str, Any],
    output_path: str,
    verification_equipment: Optional[List[Dict[str, Any]]] = None,
    org_settings: Optional[Dict[str, Any]] = None,
    specialist_docs: Optional[List[Dict[str, Any]]] = None,
    document_files: Optional[List[Dict[str, Any]]] = None,
    find_image: Optional[Any] = None,
    ndt_methods: Optional[List[Dict[str, Any]]] = None,
) -> str:
    """
    Скопировать шаблон to-1 и заполнить данными обследования.
    Возвращает путь к готовому файлу.
    """
    template = resolve_form_path("to-1")
    if template is None or not template.exists():
        raise FileNotFoundError(
            "Шаблон формы to-1 не найден в backend/report_forms/"
        )

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, out)

    doc = Document(str(out))
    if org_settings is None:
        org_settings = load_report_org_settings()

    # Нормализуем/обогащаем data (мобильные ключи + методы НК + схемы УЗТ)
    inspection_data = dict(inspection_data or {})
    raw_data = inspection_data.get("data")
    if not isinstance(raw_data, dict):
        raw_data = {}
    inspection_data["data"] = _enrich_inspection_data(raw_data, ndt_methods or [])

    # Приборы: реестр InspectionEquipment + ручной ввод + приборы из методов НК
    merged_instruments = _merge_report_instruments(
        verification_equipment,
        inspection_data["data"],
        ndt_methods or [],
    )

    attachments = build_attachments_map(document_files)
    ctx = _build_context(
        inspection_data,
        equipment_data,
        merged_instruments,
        org_settings,
        specialist_docs or [],
        attachments,
    )
    ctx["find_image"] = find_image
    ctx["ndt_methods"] = ndt_methods or []

    tables = doc.tables
    if len(tables) < 39:
        logger.warning(
            "Шаблон to-1: ожидалось ≥39 таблиц, найдено %s", len(tables)
        )

    # Основной отчёт (титул + разделы 1–15) лежит в content control (SDT)
    # и раньше не попадал в doc.tables / doc.paragraphs.
    _fill_main_report(doc, ctx)
    _fix_main_report_captions(doc)

    # Заголовки протоколов (исполнитель / заказчик / зав.№)
    for idx in (0, 11, 14, 19, 23, 28, 33):
        if idx < len(tables):
            _fill_protocol_header(tables[idx], ctx)

    if len(tables) > 1:
        _fill_documents_table(tables[1], ctx)
    if len(tables) > 2:
        _fill_general_data(tables[2], ctx)
    if len(tables) > 3:
        _fill_elements_table(tables[3], ctx)
        _strip_empty_rows(tables[3], 2)
    if len(tables) > 4:
        _fill_characteristics(tables[4], ctx)
    if len(tables) > 5:
        _fill_materials(tables[5], ctx)
        _strip_empty_rows(tables[5], 2)
    if len(tables) > 6:
        _fill_heat_treatment(tables[6], ctx)
        _strip_empty_rows(tables[6], 2)
    if len(tables) > 7:
        _fill_strength_tests(tables[7], ctx)
        _strip_empty_rows(tables[7], 1)
        _widen_date_column(tables[7], 0, 3.4)
    if len(tables) > 8:
        _fill_previous_inspections(tables[8], ctx)
        _strip_empty_rows(tables[8], 1)
        _widen_date_column(tables[8], 0, 3.4)
    if len(tables) > 9:
        _fill_additional_data(tables[9], ctx)

    # Подписи специалистов (анализ документации и др.) — с учётом того, какой
    # именно вид НК соответствует протоколу (см. SIGNATURE_METHOD_KEYS).
    for idx in (10, 13, 18, 22, 27, 32, 37):
        if idx < len(tables):
            _fill_signatures(tables[idx], ctx, SIGNATURE_METHOD_KEYS.get(idx))

    # Прил. 2 — оперативная диагностика
    if len(tables) > 12:
        _fill_operational_diagnostics(tables[12], ctx)

    # Прил. 3 — ВИК: оборудование + результаты
    if len(tables) > 15:
        _fill_instrument_table(
            tables[15],
            ctx,
            method_keys=("ВИК", "VIK", "ПВК", "ОСВЕЩ", "ШЕРОХ", "RZ"),
            defaults=[
                ("Комплект ВИК", ""),
                ("Образцы шероховатости", ""),
                ("Измеритель освещённости", ""),
            ],
        )
    if len(tables) > 16:
        _fill_vik_parameters(tables[16], ctx)
    if len(tables) > 17:
        _fill_vik_results(tables[17], ctx)

    # Прил. 4 — УЗТ
    if len(tables) > 20:
        _fill_instrument_table(
            tables[20],
            ctx,
            method_keys=("УЗТ", "UZT", "ТОЛЩИНОМЕР"),
            defaults=[("Толщиномер", ""), ("Настроечный образец", ""), ("Образцы шероховатости", "")],
        )
    if len(tables) > 21:
        _fill_uzt_results(tables[21], ctx)
        # strip выполняется внутри _fill_uzt_results после заполнения

    # Прил. 5 — твердость
    if len(tables) > 24:
        _fill_instrument_table(
            tables[24],
            ctx,
            method_keys=("ТВЕРД", "TVI", "HARD"),
            defaults=[("Твердомер", ""), ("Меры твердости", ""), ("Образцы шероховатости", "")],
        )
    if len(tables) > 25:
        _fill_hardness_matrix(tables[25], ctx)
        _strip_empty_rows(tables[25], 3, ignore_cols=(0,))
    if len(tables) > 26:
        _fill_hardness_list(tables[26], ctx)
        _strip_empty_rows(tables[26], 1, ignore_cols=(0,))

    # Прил. 6 — УЗК
    if len(tables) > 29:
        _fill_instrument_table(
            tables[29],
            ctx,
            method_keys=("УЗК", "UZK", "ДЕФЕКТОСКОП"),
            defaults=[("Дефектоскоп", ""), ("СОП", ""), ("Образцы шероховатости", "")],
        )
    if len(tables) > 30:
        _fill_uzk_parameters(tables[30], ctx)
        _strip_empty_rows(tables[30], 1)
    if len(tables) > 31:
        _fill_uzk_results(tables[31], ctx)
        _strip_empty_rows(tables[31], 1, ignore_cols=(0,))

    # Прил. 7 — МПК
    if len(tables) > 34:
        _fill_instrument_table(
            tables[34],
            ctx,
            method_keys=("МПК", "MPK", "МАГНИТ", "МПД", "MPI", "MK", "МК", "YOKEMAG"),
            defaults=None,
        )
    if len(tables) > 35:
        _fill_mpk_parameters(tables[35], ctx)
    if len(tables) > 36:
        _fill_mpk_results(tables[36], ctx)
        _strip_empty_rows(tables[36], 1, ignore_cols=(0,))

    _fill_paragraph_blanks(doc, ctx)
    _fill_hardness_steel_heading(doc, ctx)
    _fill_appendix_8_calculation(doc, ctx)
    _fill_appendix_9_hydraulic_act(doc, ctx)

    finalize_official_form(doc, "to-1")
    _insert_schemes_and_photos(doc, ctx)

    doc.save(str(out))
    logger.info("Форма to-1 заполнена: %s", out)
    return str(out)


# ---------------------------------------------------------------------------
# Контекст данных
# ---------------------------------------------------------------------------

def _enrich_inspection_data(
    data: Dict[str, Any],
    ndt_methods: Optional[List[Dict[str, Any]]] = None,
) -> Dict[str, Any]:
    """Свести мобильные ключи и данные методов НК к единому словарю для формы."""
    out = dict(data or {})

    # Материал корпуса из элементов сосуда
    elements = out.get("vessel_elements") or out.get("elements") or []
    if isinstance(elements, list) and elements:
        first = next((e for e in elements if isinstance(e, dict)), None)
        if first:
            mat = first.get("material") or first.get("steel_grade")
            if mat and not out.get("shell_material"):
                out["shell_material"] = mat
            if first.get("gost") and not out.get("material_gost"):
                out["material_gost"] = first.get("gost")

    # Вложенные характеристики среды (mobile: medium_characteristics)
    medium = out.get("medium_characteristics")
    if isinstance(medium, dict):
        for src, dst in (
            ("hazard_class", "hazard_class"),
            ("class_hazard", "hazard_class"),
            ("explosion_hazard", "explosion_hazard"),
            ("explosion_category", "explosion_hazard"),
            ("fire_hazard", "fire_hazard"),
            ("fire_category", "fire_hazard"),
            ("composition", "working_medium"),
            ("working_medium", "working_medium"),
            ("temperature", "working_medium_temperature"),
        ):
            if medium.get(src) and not out.get(dst):
                out[dst] = medium.get(src)

    # ОПО: плоские ключи из вложенного объекта (если API/mobile положили dict)
    opo = out.get("opo")
    if isinstance(opo, dict):
        for src, dst in (
            ("name", "opo_name"),
            ("hazard_class", "opo_hazard_class"),
            ("registration_number", "opo_reg_number"),
            ("code", "opo_code"),
            ("description", "opo_description"),
        ):
            if opo.get(src) and not out.get(dst):
                out[dst] = opo.get(src)

    # Пустые заготовки previous_inspections / ndt_control_history не считаем данными
    for key in ("previous_inspections", "ndt_control_history", "heat_treatment_records"):
        recs = out.get(key)
        if isinstance(recs, list):
            cleaned = []
            for rec in recs:
                if not isinstance(rec, dict):
                    continue
                if any(
                    str(v).strip()
                    for k, v in rec.items()
                    if v not in (None, "", [], {}) and k not in ("id",)
                ):
                    cleaned.append(rec)
            out[key] = cleaned

    # История НК (Б4) → таблица предыдущих контролей, если previous_inspections пуст
    if not out.get("previous_inspections") and out.get("ndt_control_history"):
        out["previous_inspections"] = list(out["ndt_control_history"])
    elif out.get("previous_inspections") and out.get("ndt_control_history"):
        # Дополнить уникальными записями из ndt_control_history
        existing = out["previous_inspections"]
        if isinstance(existing, list):
            merged = list(existing)
            for rec in out["ndt_control_history"]:
                if isinstance(rec, dict) and rec not in merged:
                    merged.append(rec)
            out["previous_inspections"] = merged

    # Гидравлика: унифицировать type ← test_type
    hydro = out.get("hydraulic_test_history")
    if isinstance(hydro, list):
        for rec in hydro:
            if isinstance(rec, dict) and not rec.get("type") and rec.get("test_type"):
                rec["type"] = rec.get("test_type")

    # УЗТ: точки из uzt_schemes[].measurements → thickness_measurements
    thickness = out.get("thickness_measurements") or out.get("thicknessMeasurements")
    if not isinstance(thickness, list):
        thickness = []
    else:
        thickness = list(thickness)
    # Нормализация ключей (mobile: point_number / point)
    for p in thickness:
        if not isinstance(p, dict):
            continue
        if not p.get("section_number"):
            alt = p.get("point_number") or p.get("number") or p.get("point")
            if alt not in (None, ""):
                p["section_number"] = alt
        if p.get("thickness") in (None, "") and p.get("value") not in (None, ""):
            p["thickness"] = p.get("value")
    schemes = out.get("uzt_schemes") or []
    if isinstance(schemes, list):
        for sch in schemes:
            if not isinstance(sch, dict):
                continue
            for m in sch.get("measurements") or []:
                if isinstance(m, dict):
                    thickness.append(dict(m))
            # путь схемы как control_scheme, если основной пуст
            sp = sch.get("scheme_image_path") or sch.get("scheme_path")
            if sp and not out.get("control_scheme_image"):
                out["control_scheme_image"] = sp

    # Методы НК (таблица ndt_methods): точки УЗТ, дефекты, приборы, даты
    for m in ndt_methods or []:
        if not isinstance(m, dict):
            continue
        code = str(m.get("method_code") or m.get("method_name") or "").upper()
        ad = m.get("additional_data") or {}
        if not isinstance(ad, dict):
            ad = {}

        # Прибор метода → verification-like список в data
        eq_name = m.get("equipment") or ad.get("equipment")
        if eq_name:
            instruments = out.setdefault("_ndt_instruments", [])
            if isinstance(instruments, list):
                instruments.append(
                    {
                        "name": eq_name,
                        "serial_number": ad.get("serial_number") or "",
                        "equipment_type": code,
                        "method_code": code,
                    }
                )

        # Дата выполнения → inspection_date / method dates
        if m.get("performed_date") and not out.get("inspection_date"):
            out["inspection_date"] = m.get("performed_date")

        # УЗТ measurement_points из экрана «добавить метод НК»
        if any(k in code for k in ("УЗТ", "UZT", "ТОЛЩ")):
            pts = ad.get("measurement_points") or ad.get("points") or []
            if isinstance(pts, list):
                for p in pts:
                    if not isinstance(p, dict):
                        continue
                    thickness.append(
                        {
                            "location": p.get("location") or p.get("element") or p.get("zone") or "",
                            "section_number": p.get("point") or p.get("section_number") or p.get("number") or "",
                            "thickness": p.get("thickness") or p.get("value"),
                            "nominal_thickness": p.get("nominal_thickness")
                            or ad.get("nominal_thickness"),
                            "min_allowed_thickness": p.get("min_allowed_thickness")
                            or ad.get("min_allowed_thickness"),
                        }
                    )
            if ad.get("nominal_thickness") and not out.get("wall_thickness"):
                out["wall_thickness"] = ad.get("nominal_thickness")

        # ВИК дефекты + параметры контроля (шероховатость/освещённость)
        if any(k in code for k in ("ВИК", "VIK", "ПВК")):
            defects = m.get("defects")
            if isinstance(defects, list) and defects:
                existing = out.get("visual_defects") or []
                if not isinstance(existing, list) or not existing:
                    out["visual_defects"] = defects
            if m.get("results") and not out.get("vik_results_text"):
                out["vik_results_text"] = m.get("results")
            if m.get("conclusion") and not out.get("vik_conclusion_text"):
                out["vik_conclusion_text"] = m.get("conclusion")
            if ad.get("illumination") and not out.get("vik_illumination"):
                out["vik_illumination"] = ad.get("illumination")
            if ad.get("additional_lighting") is not None and "vik_additional_lighting" not in out:
                out["vik_additional_lighting"] = ad.get("additional_lighting")
            if ad.get("roughness") and not out.get("vik_roughness"):
                out["vik_roughness"] = ad.get("roughness")

        # УЗК
        if any(k in code for k in ("УЗК", "UZK")):
            mapped: List[Dict[str, Any]] = []
            defects = m.get("defects")
            if isinstance(defects, list) and defects:
                for d in defects:
                    if not isinstance(d, dict):
                        continue
                    mapped.append(
                        {
                            "weld_number": d.get("weld_number") or d.get("joint") or d.get("seam"),
                            "defect_description": d.get("description") or d.get("defect"),
                            "conclusion": d.get("conclusion") or d.get("assessment"),
                            "uzk_defect": d.get("description") or d.get("defect"),
                            "control_method": "UZK",
                        }
                    )

            # Экран «Добавить метод НК»: результаты точечного сканирования
            # (additional_data.results_list — zone/coordinate/amplitude/equivalent_size)
            results_list = ad.get("results_list") or []
            if isinstance(results_list, list) and results_list:
                method_conclusion = m.get("conclusion") or ""
                for r in results_list:
                    if not isinstance(r, dict):
                        continue
                    parts = []
                    if r.get("coordinate"):
                        parts.append(f"коорд. {r.get('coordinate')}")
                    if r.get("amplitude"):
                        parts.append(f"амплитуда {r.get('amplitude')} дБ")
                    if r.get("equivalent_size"):
                        parts.append(f"экв. размер {r.get('equivalent_size')} мм")
                    desc = ", ".join(parts)
                    mapped.append(
                        {
                            "weld_number": r.get("joint")
                            or r.get("zone")
                            or ad.get("control_zone")
                            or "",
                            "joint": r.get("joint") or r.get("zone") or "",
                            "defect_number": r.get("defect_number") or "",
                            "defect_description": desc,
                            "uzk_defect": desc,
                            "area": r.get("equivalent_size")
                            or r.get("equivalent_area")
                            or r.get("area")
                            or "",
                            "equivalent_area": r.get("equivalent_size")
                            or r.get("equivalent_area")
                            or "",
                            "depth": r.get("depth") or "",
                            "length": r.get("length") or r.get("extent") or "",
                            "form": r.get("form") or r.get("character") or "",
                            "character": r.get("character") or r.get("form") or "",
                            "location": r.get("coordinate") or r.get("location") or "",
                            "location_on_control_map": r.get("coordinate")
                            or r.get("location_on_control_map")
                            or r.get("location")
                            or "",
                            "conclusion": method_conclusion,
                            "control_method": "UZK",
                        }
                    )

            # Свободный текст «Дефекты» (formData['defects']) — если структурированных
            # данных нет, заносим как единственную запись
            if not mapped and isinstance(defects, str) and defects.strip():
                mapped.append(
                    {
                        "weld_number": ad.get("control_zone") or "",
                        "defect_description": defects.strip(),
                        "uzk_defect": defects.strip(),
                        "conclusion": m.get("conclusion") or "",
                        "control_method": "UZK",
                    }
                )

            if mapped:
                existing = out.get("weld_inspections") or []
                if not isinstance(existing, list) or not existing:
                    out["weld_inspections"] = mapped
                else:
                    def _structured(w: Dict[str, Any]) -> bool:
                        return bool(
                            w.get("equivalent_area")
                            or w.get("area")
                            or w.get("depth")
                            or w.get("defect_number")
                        )

                    if not any(_structured(w) for w in existing if isinstance(w, dict)):
                        # Чек-лист дал упрощённые записи — дополняем структурированными из метода НК
                        out["weld_inspections"] = mapped + list(existing)

            # Параметры контроля УЗК → таблица № параметров в to-1
            param_row = {
                "joint_type": ad.get("joint_type") or ad.get("connection_type") or "",
                "element_thickness": ad.get("element_thickness")
                or ad.get("thickness")
                or "",
                "transducer_type": ad.get("transducer_type") or ad.get("pep_type") or "",
                "frequency_mhz": ad.get("frequency_mhz") or ad.get("frequency") or "",
                "angle_deg": ad.get("angle_deg") or ad.get("angle") or "",
                "max_equivalent_area": ad.get("max_equivalent_area")
                or ad.get("s_reject")
                or "",
                "notch_params": ad.get("notch_params")
                or ad.get("notch")
                or ad.get("reference_sample")
                or "",
                "device_type": ad.get("device_type") or m.get("equipment") or "",
            }
            if any(str(v).strip() for v in param_row.values()):
                params = out.setdefault("uzk_control_params", [])
                if isinstance(params, list):
                    params.append(param_row)

        # Твердость
        if any(k in code for k in ("ТВЕРД", "HARD", "TVI")):
            ht = ad.get("hardness_tests") or ad.get("points") or []
            if isinstance(ht, list) and ht:
                existing = out.get("hardness_tests") or []
                if not isinstance(existing, list) or not existing:
                    out["hardness_tests"] = ht

        # МПК / магнитный контроль
        if any(k in code for k in ("МПК", "MPK", "МАГНИТ", "MPI", "МПД", "МК")):
            mapped_mpk: List[Dict[str, Any]] = []
            defects = m.get("defects")
            if isinstance(defects, list):
                for d in defects:
                    if not isinstance(d, dict):
                        continue
                    mapped_mpk.append(
                        {
                            "object": d.get("element") or d.get("object") or d.get("zone") or "",
                            "zone": d.get("zone") or d.get("location") or "",
                            "scope": d.get("scope") or ad.get("control_volume") or "100%",
                            "defects": d.get("description") or d.get("defect") or "",
                            "assessment": d.get("assessment")
                            or d.get("conclusion")
                            or m.get("conclusion")
                            or "",
                        }
                    )
            indications = ad.get("indications_list") or ad.get("results_list") or []
            if isinstance(indications, list):
                for ind in indications:
                    if not isinstance(ind, dict):
                        continue
                    mapped_mpk.append(
                        {
                            "object": ind.get("element") or ind.get("object") or "",
                            "zone": ind.get("zone") or ind.get("location") or "",
                            "scope": ind.get("scope") or "100%",
                            "defects": ind.get("description") or ind.get("indication") or "",
                            "assessment": ind.get("assessment") or m.get("conclusion") or "",
                        }
                    )
            if mapped_mpk:
                existing = out.get("mpk_results") or []
                if not isinstance(existing, list) or not existing:
                    out["mpk_results"] = mapped_mpk
            param_row = {
                "control_method": ad.get("magnetization_type")
                or ad.get("control_method")
                or ad.get("method")
                or "",
                "sensitivity": ad.get("sensitivity")
                or ad.get("sensitivity_level")
                or ad.get("field_strength")
                or "",
                "field_strength": ad.get("field_strength") or "",
                "indicator": ad.get("indicator_suspension") or "",
            }
            if any(str(v).strip() for v in param_row.values()):
                params = out.setdefault("mpk_control_params", [])
                if isinstance(params, list):
                    params.append(param_row)
                if param_row["control_method"] and not out.get("mpk_control_method"):
                    out["mpk_control_method"] = param_row["control_method"]
                if param_row["sensitivity"] and not out.get("mpk_sensitivity"):
                    out["mpk_sensitivity"] = param_row["sensitivity"]

    # МПК из weld_inspections чек-листа
    welds = out.get("weld_inspections") or []
    if isinstance(welds, list):
        extra_mpk: List[Dict[str, Any]] = []
        for w in welds:
            if not isinstance(w, dict):
                continue
            method = str(w.get("control_method") or w.get("method") or "").upper()
            if method not in ("MPK", "МПК", "МПД", "MK", "МК", "MPI"):
                continue
            extra_mpk.append(
                {
                    "object": w.get("weld_number") or w.get("joint") or "",
                    "zone": w.get("location_on_control_map") or w.get("location") or "",
                    "scope": w.get("scope") or "100%",
                    "defects": w.get("pvk_defect")
                    or w.get("defect_description")
                    or w.get("uzk_defect")
                    or "",
                    "assessment": w.get("conclusion") or w.get("assessment") or "",
                }
            )
        if extra_mpk:
            existing = out.get("mpk_results") or []
            if not isinstance(existing, list) or not existing:
                out["mpk_results"] = extra_mpk
            else:
                out["mpk_results"] = list(existing) + extra_mpk

    if thickness:
        out["thickness_measurements"] = thickness

    return out


def _build_context(
    inspection_data: Dict[str, Any],
    equipment_data: Dict[str, Any],
    verification_equipment: List[Dict[str, Any]],
    org_settings: Dict[str, Any],
    specialist_docs: List[Dict[str, Any]],
    attachments: Optional[Dict[str, str]] = None,
) -> Dict[str, Any]:
    data = inspection_data.get("data") or {}
    if not isinstance(data, dict):
        data = {}
    attrs = equipment_data.get("attributes") or {}
    if not isinstance(attrs, dict):
        attrs = {}

    def g(*keys: str, default: Any = MISSING) -> Any:
        for k in keys:
            if k in data and data.get(k) not in (None, ""):
                return data.get(k)
        for k in keys:
            if k in attrs and attrs.get(k) not in (None, ""):
                return attrs.get(k)
        for k in keys:
            camel = "".join(
                w.capitalize() if i else w for i, w in enumerate(k.split("_"))
            )
            camel_l = camel[0].lower() + camel[1:] if camel else k
            if camel_l in data and data.get(camel_l) not in (None, ""):
                return data.get(camel_l)
            if camel_l in attrs and attrs.get(camel_l) not in (None, ""):
                return attrs.get(camel_l)
        return default

    date_perf = (
        inspection_data.get("date_performed")
        or data.get("inspection_date")
    )
    date_ru = _fmt_date_ru(date_perf) or datetime.now().strftime("%d.%m.%Y")

    contractor = (org_settings.get("contractor") or {}) if isinstance(org_settings, dict) else {}
    customer = (org_settings.get("customer") or {}) if isinstance(org_settings, dict) else {}
    lab = (org_settings.get("ndt_lab") or {}) if isinstance(org_settings, dict) else {}

    # Override из обследования (mobile customer_info / contractor_info)
    cust_ov = data.get("customer_info") if isinstance(data.get("customer_info"), dict) else {}
    contr_ov = data.get("contractor_info") if isinstance(data.get("contractor_info"), dict) else {}
    if cust_ov:
        customer = {**customer, **{k: v for k, v in cust_ov.items() if v not in (None, "")}}
    if contr_ov:
        contractor = {**contractor, **{k: v for k, v in contr_ov.items() if v not in (None, "")}}

    # Единый параметр ориентации сосуда
    orientation = str(g("orientation", default="") or "").strip().lower()
    if orientation not in ("horizontal", "vertical"):
        ct = str(g("construction_type", default="") or "").lower()
        if "горизонт" in ct or "horizontal" in ct:
            orientation = "horizontal"
        elif "вертикал" in ct or "vertical" in ct:
            orientation = "vertical"
        else:
            orientation = ""
    if orientation:
        data["orientation"] = orientation

    device_name = str(
        g(
            "equipment_device_name",
            "vessel_name",
            "device_name",
            default=equipment_data.get("name") or MISSING,
        )
    )
    serial = str(
        g(
            "serial_number",
            "factory_number",
            "equipment_serial",
            default=equipment_data.get("serial_number")
            or equipment_data.get("factory_number")
            or MISSING,
        )
    )
    # Карточка оборудования часто хранит рег.№ как registration_number
    reg_no = str(
        g(
            "reg_number",
            "regNumber",
            "registration_number",
            default=equipment_data.get("registration_number") or MISSING,
        )
    )
    inv_no = str(
        g(
            "inventory_number",
            "inv_number",
            "equipment_inventory_number",
            default=equipment_data.get("inventory_number") or MISSING,
        )
    )
    location = str(g("location", "equipment_location", default=equipment_data.get("location") or ""))
    org_name = str(
        g(
            "organization",
            "customer_name",
            "enterprise_name",
            default=customer.get("legal_name") or MISSING,
        )
    )
    org_name, location = _split_customer_location(org_name, location)

    docs_dict = g("documents", default={}) or {}
    docs_info = g("documents_info", default={}) or {}
    if not isinstance(docs_dict, dict):
        docs_dict = {}
    if not isinstance(docs_info, dict):
        docs_info = {}
    docs_info = _merge_document_sets_into_info(data, docs_info)

    specialists = _extract_specialists(data, specialist_docs)
    opo_name = str(g("opo_name", default=MISSING))
    opo_class = str(g("opo_hazard_class", "opo_class", default=MISSING))
    # Регистрационный № ОПО — официальный рег.номер; внутренний код (OPO-004)
    # не подставляем как «рег. №», чтобы не путать с реестровым номером Ростехнадзора.
    opo_reg = str(g("opo_reg_number", "opo_registration_number", default=MISSING))
    if opo_name in ("", MISSING):
        nested = data.get("opo") if isinstance(data.get("opo"), dict) else {}
        if nested.get("name"):
            opo_name = str(nested.get("name"))
    if opo_class in ("", MISSING):
        nested = data.get("opo") if isinstance(data.get("opo"), dict) else {}
        if nested.get("hazard_class"):
            opo_class = str(nested.get("hazard_class"))
    if opo_reg in ("", MISSING):
        nested = data.get("opo") if isinstance(data.get("opo"), dict) else {}
        if nested.get("registration_number"):
            opo_reg = str(nested.get("registration_number"))

    return {
        "g": g,
        "data": data,
        "attrs": attrs,
        "date_ru": date_ru,
        "device_name": device_name,
        "serial": serial,
        "reg_no": reg_no,
        "inv_no": inv_no,
        "location": location,
        "org_name": org_name,
        "opo_name": opo_name,
        "opo_class": opo_class,
        "opo_reg": opo_reg,
        "contractor_name": contractor.get("legal_name") or contractor.get("name") or "",
        "contractor_address": contractor.get("postal_address") or contractor.get("address") or "",
        "contractor_director": contractor.get("director_name") or contractor.get("director") or "",
        "contractor_phone": contractor.get("phone") or "",
        "contractor_email": contractor.get("email") or "",
        "customer_director": customer.get("director") or customer.get("director_name") or "",
        "customer_address": customer.get("address") or customer.get("legal_address") or "",
        "customer_phone": customer.get("phone") or "",
        "customer_email": customer.get("email") or "",
        "customer_legal_name": customer.get("legal_name") or customer.get("name") or "",
        "orientation": orientation,
        "lab_name": lab.get("name") or contractor.get("legal_name") or "",
        "lab_cert": lab.get("certificate")
        or lab.get("attestation_number")
        or contractor.get("certificate")
        or "",
        "docs_dict": docs_dict,
        "docs_info": docs_info,
        "verification_equipment": verification_equipment,
        "specialists": specialists,
        "specialist_docs": specialist_docs,
        "attachments": attachments or {},
        "org_settings": org_settings,
        "conclusion_doc": str(g("documentation_conclusion", "doc_analysis_conclusion", default="")),
        "conclusion_suitable": str(
            g("suitability_conclusion", "conclusion", default="соответствует")
        ),
        "operational_ok": str(
            g("operational_conclusion", "operational_ok", default="соответствует")
        ),
        "calculation_result": str(
            g(
                "calculation_result",
                "calc_assessment",
                default="сосуда при рабочих параметрах",
            )
        ),
        "tech_state": str(
            g(
                "technical_state",
                "suitability_status",
                default="работоспособное, пригодно к дальнейшей эксплуатации",
            )
        ),
    }



# Соответствие латинских/сокращённых кодов методов НК русскоязычным
# наименованиям области аттестации (кириллицей), как того требует НТД.
_METHOD_CODE_TO_RU = {
    "VIK": "ВИК",
    "VT": "ВИК",
    "VISUAL": "ВИК",
    "ПВК": "ВИК",
    "PVK": "ВИК",
    "UZK": "УЗК",
    "UT": "УЗК",
    "UZT": "УЗТ",
    "UTT": "УЗТ",
    "ТОЛЩИНОМЕТРИЯ": "УЗТ",
    "MPK": "МПК",
    "MT": "МПК",
    "МАГНИТ": "МПК",
    "TVI": "Твёрдометрия",
    "HARD": "Твёрдометрия",
    "HARDNESS": "Твёрдометрия",
    "ТВЕРД": "Твёрдометрия",
    "PT": "ПВК (капиллярный)",
    "ПВК-К": "ПВК (капиллярный)",
}


# Развёрнутые наименования приборов/оборудования по типу (коду метода НК),
# как того требует нормативная документация (не аббревиатура, а полное
# наименование + марка/модель прибора).
_INSTRUMENT_TYPE_FULL_NAME = {
    "ВИК": "Комплект для визуально-измерительного контроля",
    "VIK": "Комплект для визуально-измерительного контроля",
    "УЗК": "Ультразвуковой дефектоскоп",
    "UZK": "Ультразвуковой дефектоскоп",
    "УЗТ": "Ультразвуковой толщиномер",
    "UZT": "Ультразвуковой толщиномер",
    "МПК": "Дефектоскоп магнитопорошковый",
    "MPK": "Дефектоскоп магнитопорошковый",
    "ТВЕРД": "Твердомер",
    "HARD": "Твердомер",
    "TVI": "Твердомер",
    "ТВЕРДОМЕР": "Твердомер",
    "ПВК": "Комплект капиллярного контроля",
    "PVK": "Комплект капиллярного контроля",
    "ПВК (КАПИЛЛЯРНЫЙ)": "Комплект капиллярного контроля",
    "КОМПЛЕКТ КАПИЛЛЯРНОГО КОНТРОЛЯ (С РЕАГЕНТАМИ)": "Комплект капиллярного контроля",
    "ШЕРОХ": "Образец шероховатости поверхности",
    "RZ": "Образец шероховатости поверхности",
    "ОБРАЗЕЦ ШЕРОХОВАТОСТИ": "Образец шероховатости поверхности",
    "ОСВЕЩ": "Люксметр",
    "LUX": "Люксметр",
    "ЛЮКСМЕТР": "Люксметр",
}


def _instrument_full_name(eq: Dict[str, Any]) -> str:
    """Полное наименование прибора: <тип прибора> <марка/модель>."""
    et = str(eq.get("equipment_type") or "").strip()
    name = str(eq.get("name") or "").strip()
    model = str(eq.get("model") or "").strip()
    manufacturer = str(eq.get("manufacturer") or "").strip()
    prefix = _INSTRUMENT_TYPE_FULL_NAME.get(et.upper())
    # Марка — модель прибора, либо (если модели нет) собственное имя записи,
    # но только если оно не совпадает с кодом типа (иначе получим "ВИК ВИК").
    brand_parts = [manufacturer] if manufacturer else []
    if model:
        brand_parts.append(model)
    elif name and name.upper() != et.upper():
        brand_parts.append(name)
    brand = " ".join(p for p in brand_parts if p)
    if prefix:
        return f"{prefix} {brand}".strip() if brand else f"{prefix} __________ (марка прибора)"
    return name or et or MISSING


def method_label_ru(code: str) -> str:
    """Область аттестации кириллицей вместо латинского кода метода (напр. VIK → ВИК)."""
    raw = (code or "").strip()
    if not raw:
        return raw
    # Несколько кодов через запятую/слэш — переводим каждый
    parts = re.split(r"[,/;]+", raw)
    if len(parts) > 1:
        return ", ".join(method_label_ru(p.strip()) for p in parts if p.strip())
    return _METHOD_CODE_TO_RU.get(raw.upper(), raw)


def _extract_specialists(
    data: Dict[str, Any], specialist_docs: List[Dict[str, Any]]
) -> List[Dict[str, str]]:
    """Собрать специалистов, реально задействованных в ДАННОМ обследовании.

    Приоритет отдаётся специалистам, выбранным в мобильном приложении для
    этого обследования (``inspection_engineers``) — иначе в отчёт могли
    попасть посторонние ФИО (например, общий список исполнителей ОПО),
    не имеющие отношения к фактически выполненному контролю.
    """
    result: List[Dict[str, str]] = []

    def _upsert(
        name: str,
        cert: str = "",
        role: str = "",
        scan: str = "",
        valid_until: str = "",
    ) -> None:
        name = (name or "").strip()
        if not name:
            return
        for s in result:
            if s["name"].lower() == name.lower():
                if cert and not s.get("cert"):
                    s["cert"] = cert
                if role:
                    existing_roles = {r.strip().upper() for r in s.get("role", "").split(",") if r.strip()}
                    if role.strip().upper() not in existing_roles:
                        s["role"] = ", ".join(filter(None, [s.get("role", ""), role])).strip(", ")
                if scan and not s.get("scan"):
                    s["scan"] = scan
                if valid_until and not s.get("valid_until"):
                    s["valid_until"] = valid_until
                return
        result.append(
            {
                "name": name,
                "cert": cert or "",
                "role": role or "",
                "scan": scan or "",
                "valid_until": valid_until or "",
            }
        )

    def _cert_lookup(name: str) -> Dict[str, str]:
        """Найти № удостоверения / срок действия по ФИО в specialist_docs (справочник)."""
        for doc in specialist_docs or []:
            if not isinstance(doc, dict):
                continue
            doc_name = str(
                doc.get("inspector_name") or doc.get("specialist_name") or doc.get("name") or ""
            )
            if doc_name.strip().lower() != name.strip().lower():
                continue
            certs = doc.get("certifications") or []
            if isinstance(certs, list):
                for c in certs:
                    if not isinstance(c, dict):
                        continue
                    cert_no = str(c.get("certificate_number") or c.get("number") or "")
                    if cert_no:
                        return {
                            "cert": cert_no,
                            "scan": str(c.get("scan_file_path") or ""),
                            "valid_until": str(c.get("expiry_date") or ""),
                            "role": method_label_ru(str(c.get("method_code") or c.get("certification_type") or "")),
                        }
        return {}

    # 1) Специалисты, реально выбранные в мобильном приложении для этого обследования
    #    (по каждому виду НК — свой инженер).
    engineers = data.get("inspection_engineers") or []
    if isinstance(engineers, list):
        for eng in engineers:
            if not isinstance(eng, dict):
                continue
            name = str(eng.get("full_name") or eng.get("name") or "").strip()
            if not name:
                continue
            role_raw = str(eng.get("method") or "")
            _upsert(
                name,
                str(eng.get("certificate_number") or eng.get("cert") or ""),
                method_label_ru(role_raw),
                valid_until=str(eng.get("valid_until") or eng.get("expiry") or ""),
            )
            # Дозаполнить № удостоверения / срок действия из справочника (Certification),
            # если в чек-листе они не сохранились.
            found = _cert_lookup(name)
            if found:
                for s in result:
                    if s["name"].lower() == name.lower():
                        if not s.get("cert"):
                            s["cert"] = found.get("cert", "")
                        if not s.get("valid_until"):
                            s["valid_until"] = found.get("valid_until", "")
                        if not s.get("scan"):
                            s["scan"] = found.get("scan", "")
                        break

    if result:
        return result

    # 2) Фолбэк — specialist_docs из API (сертификаты по методам НК, выполненным в обследовании)
    for doc in specialist_docs or []:
        if not isinstance(doc, dict):
            continue
        name = str(
            doc.get("inspector_name")
            or doc.get("specialist_name")
            or doc.get("name")
            or ""
        )
        certs = doc.get("certifications") or []
        cert_no = ""
        scan = ""
        valid_until = ""
        role = str(doc.get("role") or "")
        if isinstance(certs, list) and certs:
            for c in certs:
                if not isinstance(c, dict):
                    continue
                cert_no = str(
                    c.get("certificate_number") or c.get("number") or c.get("cert") or ""
                )
                scan = str(c.get("scan_file_path") or c.get("scan") or "")
                valid_until = str(c.get("expiry_date") or c.get("valid_until") or "")
                method = str(c.get("method_code") or c.get("certification_type") or "")
                if method and not role:
                    role = method_label_ru(method)
                if cert_no:
                    break
        else:
            cert_no = str(doc.get("certificate_number") or doc.get("cert") or "")
            scan = str(doc.get("scan_file_path") or "")
        _upsert(name, cert_no, role, scan, valid_until)

    if result:
        return result

    # 3) Последний фолбэк — свободный текст исполнителей (напр. общий список по ОПО),
    #    используется, только если конкретных специалистов по обследованию нет вовсе.
    executors = data.get("executors") or data.get("specialists")
    if isinstance(executors, str) and executors.strip():
        for part in re.split(r"[,;/\n]+", executors):
            _upsert(part.strip())
    if isinstance(executors, list):
        for item in executors:
            if isinstance(item, dict):
                _upsert(
                    str(item.get("name") or item.get("full_name") or ""),
                    str(item.get("certificate") or item.get("cert_number") or ""),
                    method_label_ru(str(item.get("role") or "")),
                )
            elif isinstance(item, str) and item.strip():
                _upsert(item.strip())

    return result


# ---------------------------------------------------------------------------
# Низкоуровневые хелперы ячеек
# ---------------------------------------------------------------------------

def _set_cell(cell: _Cell, text: Any, *, nowrap: bool = False) -> None:
    value = "" if text is None else str(text)
    if value in ("None",):
        value = ""
    # Неразрывный дефис — чтобы «09Г2С» / даты не ломались посередине
    if nowrap:
        value = value.replace("-", "\u2011")
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.text = value
        return
    p0 = paragraphs[0]
    if p0.runs:
        p0.runs[0].text = value
        _set_run_font(p0.runs[0])
        for run in p0.runs[1:]:
            run.text = ""
    else:
        run = p0.add_run(value)
        _set_run_font(run)
    for p in paragraphs[1:]:
        p.clear()
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in list(tcPr.findall(qn("w:noWrap"))):
            tcPr.remove(old)
        if nowrap:
            tcPr.append(OxmlElement("w:noWrap"))
        # Иначе явно НЕ добавляем noWrap — длинный текст (напр. полное
        # наименование организации-исполнителя) должен переноситься по
        # словам, а не обрезаться/вылезать за пределы ячейки шаблона.
    except Exception:
        pass


def _cell(table: Table, row: int, col: int) -> Optional[_Cell]:
    try:
        return table.rows[row].cells[col]
    except (IndexError, KeyError):
        return None


def _set(table: Table, row: int, col: int, text: Any, *, nowrap: bool = False) -> None:
    c = _cell(table, row, col)
    if c is not None:
        _set_cell(c, text, nowrap=nowrap)


def _split_customer_location(org_name: str, location: str) -> Tuple[str, str]:
    """Отделить местонахождение от иерархии заказчика (… / Пункт подготовки…)."""
    org = (org_name or "").strip()
    loc = (location or "").strip()
    if loc in ("", "-", "—", MISSING):
        loc = ""
    if not loc and " / " in org:
        parts = [p.strip() for p in org.split(" / ") if p.strip()]
        if len(parts) >= 2:
            loc = parts[-1]
            org = " / ".join(parts[:-1])
    return org or MISSING, loc or MISSING


def _iter_all_paragraphs(doc: Document):
    """Все абзацы документа, включая content controls (SDT)."""
    body = doc.element.body
    for p_el in body.iter(qn("w:p")):
        yield Paragraph(p_el, doc)


def _main_sdt_tables(doc: Document) -> List[Table]:
    """Таблицы основного отчёта (титул + разд. 1–15) из первого SDT."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag != qn("w:sdt"):
            continue
        content = child.find(qn("w:sdtContent"))
        if content is None:
            continue
        texts = "".join(t.text or "" for t in content.iter(qn("w:t")))
        if "ТЕХНИЧЕСКИЙ ОТЧЕТ" not in texts and "УТВЕРЖДАЮ" not in texts:
            continue
        return [Table(tbl, doc) for tbl in content.iter(qn("w:tbl"))]
    return []


def _insert_page_break_before_paragraph(paragraph: Paragraph) -> None:
    """Вставить разрыв страницы перед абзацем (СОДЕРЖАНИЕ на новой странице)."""
    p = paragraph._p
    parent = p.getparent()
    if parent is None:
        return
    prev = p.getprevious()
    if prev is not None and prev.tag == qn("w:p"):
        for br in prev.iter(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return
    new_p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    new_p.append(r)
    parent.insert(list(parent).index(p), new_p)


# Минимальный нормальный размер шрифта таблиц (встреча 03.08.2026).
_TABLE_FONT_PT = 12.0
# Умеренное уменьшение — только после landscape / переноса (не ниже этого).
_TABLE_FONT_MIN_PT = 10.0
# Таблицы с таким числом колонок считаются «широкими» → landscape.
_WIDE_TABLE_MIN_COLS = 7
# Единый шрифт всего отчёта ТО.
_REPORT_FONT_NAME = "Times New Roman"


def _set_run_font(
    run,
    *,
    pt: Optional[float] = None,
    bold: Optional[bool] = None,
    name: str = _REPORT_FONT_NAME,
) -> None:
    """Times New Roman (+ ascii/hAnsi/eastAsia), опционально размер и bold."""
    try:
        run.font.name = name
    except Exception:
        pass
    try:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.rFonts
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(attr), name)
    except Exception:
        pass
    if pt is not None:
        try:
            run.font.size = Pt(pt)
        except Exception:
            pass
    if bold is not None:
        try:
            run.font.bold = bold
        except Exception:
            pass
    try:
        run.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        pass


def _shrink_table_font(table: Table, pt: float = _TABLE_FONT_PT) -> None:
    """Унифицировать размер шрифта ячеек (читаемый минимум ~12 pt)."""
    _apply_table_font(table, pt=pt)


def _apply_table_font(table: Table, pt: float = _TABLE_FONT_PT) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                pf = p.paragraph_format
                try:
                    pf.space_before = Pt(1)
                    pf.space_after = Pt(1)
                except Exception:
                    pass
                for r in p.runs:
                    _set_run_font(r, pt=pt)


def _style_table_header_row(table: Table, header_rows: int = 1, pt: float = _TABLE_FONT_PT) -> None:
    """Единый стиль заголовков таблиц приложений."""
    for ri in range(min(header_rows, len(table.rows))):
        for cell in table.rows[ri].cells:
            for p in cell.paragraphs:
                try:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                except Exception:
                    pass
                for r in p.runs:
                    _set_run_font(r, pt=pt, bold=True)


def _table_col_count(table: Table) -> int:
    if not table.rows:
        return 0
    return len(table.rows[0].cells)


def _set_section_orientation(section, landscape: bool) -> None:
    """Portrait/landscape с корректным обменом page_width/page_height."""
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    w, h = section.page_width, section.page_height
    if landscape and w < h:
        section.page_width, section.page_height = h, w
    elif not landscape and w > h:
        section.page_width, section.page_height = h, w


def _make_section_break_paragraph(landscape: bool):
    """Абзац с sectPr (nextPage). sectPr описывает секцию, которая им заканчивается."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    sectPr = OxmlElement("w:sectPr")
    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "nextPage")
    sectPr.append(type_el)
    pgSz = OxmlElement("w:pgSz")
    if landscape:
        pgSz.set(qn("w:w"), "16838")
        pgSz.set(qn("w:h"), "11906")
        pgSz.set(qn("w:orient"), "landscape")
    else:
        pgSz.set(qn("w:w"), "11906")
        pgSz.set(qn("w:h"), "16838")
        pgSz.set(qn("w:orient"), "portrait")
    sectPr.append(pgSz)
    pgMar = OxmlElement("w:pgMar")
    for attr, val in (
        ("w:top", "720"),
        ("w:right", "720"),
        ("w:bottom", "720"),
        ("w:left", "720"),
        ("w:header", "360"),
        ("w:footer", "360"),
        ("w:gutter", "0"),
    ):
        pgMar.set(qn(attr), val)
    sectPr.append(pgMar)
    pPr.append(sectPr)
    p.append(pPr)
    return p


def _insert_section_break_before(element, landscape: bool) -> None:
    """Вставить разрыв раздела (новая страница) перед элементом."""
    element.addprevious(_make_section_break_paragraph(landscape))


def _insert_section_break_after(element, landscape: bool) -> None:
    """Вставить разрыв раздела (новая страница) сразу после элемента."""
    element.addnext(_make_section_break_paragraph(landscape))


# A4 landscape: 297×210 мм. Поля 0.5" → поле ~10.2"×7.1".
# Оставляем запас сверху под заголовок (название + номер схемы).
_SCHEME_PAGE_MAX_W_IN = 9.0
_SCHEME_PAGE_MAX_H_IN = 5.6
_SCHEME_EXTRA_MAX_H_IN = 4.8


def _append_section_properties_to_paragraph(paragraph: Paragraph, *, landscape: bool) -> None:
    """
    sectPr в pPr задаёт свойства секции, которой принадлежит абзац (ECMA-376):
    секция ЗАКАНЧИВАЕТСЯ этим абзацем.
    """
    pPr = paragraph._p.get_or_add_pPr()
    for old in list(pPr.findall(qn("w:sectPr"))):
        pPr.remove(old)
    sectPr = OxmlElement("w:sectPr")
    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "nextPage")
    sectPr.append(type_el)
    pgSz = OxmlElement("w:pgSz")
    if landscape:
        pgSz.set(qn("w:w"), "16838")
        pgSz.set(qn("w:h"), "11906")
        pgSz.set(qn("w:orient"), "landscape")
    else:
        pgSz.set(qn("w:w"), "11906")
        pgSz.set(qn("w:h"), "16838")
        pgSz.set(qn("w:orient"), "portrait")
    sectPr.append(pgSz)
    pgMar = OxmlElement("w:pgMar")
    for attr, val in (
        ("w:top", "720"),
        ("w:right", "720"),
        ("w:bottom", "720"),
        ("w:left", "720"),
        ("w:header", "360"),
        ("w:footer", "360"),
        ("w:gutter", "0"),
    ):
        pgMar.set(qn(attr), val)
    sectPr.append(pgMar)
    pPr.append(sectPr)


def _fit_image_width_inches(image_path: str, max_width: float, max_height: float) -> float:
    """Подогнать ширину PNG под landscape-лист с запасом под заголовок."""
    try:
        from PIL import Image

        with Image.open(image_path) as im:
            w_px, h_px = im.size
        if w_px <= 0 or h_px <= 0:
            return max_width
        aspect = w_px / float(h_px)
        w = max_width
        if w / aspect > max_height:
            w = max_height * aspect
        return round(max(4.0, min(w, max_width)), 2)
    except Exception:
        return round(min(max_width, 7.0), 2)


def _style_scheme_title_paragraph(paragraph: Paragraph, title: str) -> None:
    """Заголовок схемы сверху: 12 pt, жирный, не отрывать от картинки."""
    _set_paragraph_text(paragraph, title, pt=12.0)
    try:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.keep_with_next = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.runs:
            _set_run_font(paragraph.runs[0], pt=12.0, bold=True)
    except Exception:
        pass


def _insert_scheme_landscape_block(
    anchor: Paragraph,
    title: str,
    image_path: str,
    extra_paths: List[Tuple[str, str]],
    *,
    close_prev_as_portrait: bool = True,
) -> Optional[Paragraph]:
    """
    Одна схема на landscape-листе: заголовок сверху + масштабированная картинка.

    OOXML: sectPr на абзаце описывает секцию, которая им заканчивается.
    1) перед блоком — разрыв секции, закрывающий предыдущий контент как portrait
       (иначе заголовок+схема остаются на portrait и PNG вылезает за край);
    2) в конце блока — landscape на последнем абзаце (картинка).
    """
    if close_prev_as_portrait:
        try:
            _insert_section_break_before(anchor._p, landscape=False)
        except Exception:
            _insert_page_break_before_paragraph(anchor)
    else:
        # Предыдущая схема уже закрыла секцию landscape — нужна только новая страница.
        _insert_page_break_before_paragraph(anchor)

    _style_scheme_title_paragraph(anchor, title)
    clear_pictures_after_paragraph(anchor)
    last = anchor
    w_main = _fit_image_width_inches(image_path, _SCHEME_PAGE_MAX_W_IN, _SCHEME_PAGE_MAX_H_IN)
    pic = add_picture_after_paragraph(last, image_path, width_inches=w_main, caption=None)
    if pic is None:
        return None
    last = pic
    for cap, ep in extra_paths:
        w_extra = _fit_image_width_inches(ep, _SCHEME_PAGE_MAX_W_IN - 0.3, _SCHEME_EXTRA_MAX_H_IN)
        cap_p = insert_paragraph_after(last, cap)
        _set_paragraph_font_size(cap_p, 12.0)
        try:
            if cap_p.runs:
                _set_run_font(cap_p.runs[0], pt=12.0, bold=True)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.keep_with_next = True
        except Exception:
            pass
        pic2 = add_picture_after_paragraph(cap_p, ep, width_inches=w_extra, caption=None)
        if pic2 is not None:
            last = pic2
    _append_section_properties_to_paragraph(last, landscape=True)
    return last


def _element_plain_text(el) -> str:
    """Текст XML-элемента абзаца без лишних пробелов."""
    try:
        return _norm_ws("".join(t.text or "" for t in el.iter(qn("w:t"))))
    except Exception:
        return ""


def _landscape_block_start(tbl) -> Any:
    """
    Начало блока «заголовок раздела + Таблица № N + таблица».

    Раньше section break ставился прямо перед таблицей — подпись «Таблица №»
    оставалась на portrait, а сама таблица уезжала на landscape.
    """
    start = tbl
    el = tbl.getprevious()
    hops = 0
    while el is not None and hops < 12:
        hops += 1
        if el.tag == qn("w:tbl"):
            break
        # Уже есть разрыв раздела — дальше не заходим
        if el.tag == qn("w:p"):
            pPr = el.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                break
            text = _element_plain_text(el).strip()
            if not text:
                el = el.getprevious()
                continue
            if _is_table_caption(text):
                start = el
                el = el.getprevious()
                continue
            # Короткий заголовок раздела/блока над подписью таблицы
            if len(text) < 160 and (
                text.startswith("Результаты")
                or text.startswith("Параметры")
                or "Материалы элементов" in text
                or bool(re.match(r"^\d{1,2}\.\s+\S", text))
            ):
                start = el
                el = el.getprevious()
                continue
            break
        el = el.getprevious()
    return start


def _finalize_table_typography(doc: Document) -> None:
    """12 pt; landscape только для широких таблиц приложений; титул — portrait."""
    main_ids: set = set()
    try:
        main_list = _main_sdt_tables(doc)
        for t in main_list:
            main_ids.add(id(t._tbl))
    except Exception:
        main_list = []

    seen: set = set()
    tables: List[Table] = []
    for t in main_list:
        tid = id(t._tbl)
        if tid not in seen:
            seen.add(tid)
            tables.append(t)
    for t in doc.tables:
        tid = id(t._tbl)
        if tid not in seen:
            seen.add(tid)
            tables.append(t)

    for t in tables:
        cols = _table_col_count(t)
        header_rows = 2 if cols >= 6 else 1
        in_main = id(t._tbl) in main_ids
        # Не трогаем основной SDT (титул/разд.1–15): там бывает «широкая»
        # однострочная сетка, из-за которой титул уезжал в landscape и
        # три колонки титула выглядели как несколько одинаковых листов.
        if not in_main and cols >= _WIDE_TABLE_MIN_COLS and len(t.rows) >= 2:
            try:
                # Заголовок «Таблица № N» (и при необходимости название раздела)
                # должны оказаться в той же landscape-секции, что и таблица.
                block_start = _landscape_block_start(t._tbl)
                _insert_section_break_before(block_start, landscape=False)
                _insert_section_break_after(t._tbl, landscape=True)
            except Exception:
                logger.exception("to-1: landscape для широкой таблицы")

        pt = _TABLE_FONT_PT
        if cols >= 10:
            pt = _TABLE_FONT_MIN_PT
        _apply_table_font(t, pt=pt)
        _style_table_header_row(t, header_rows=header_rows, pt=pt)

    try:
        for section in doc.sections:
            if section.orientation == WD_ORIENT.LANDSCAPE:
                continue
            _set_section_orientation(section, landscape=False)
    except Exception:
        pass

    _ensure_title_section_portrait(doc)
    _apply_document_font(doc)


def _ensure_title_section_portrait(doc: Document) -> None:
    """Если первый sectPr ошибочно landscape — сбросить в portrait (титул)."""
    try:
        body = doc.element.body
        first_sect = None
        for child in body.iter():
            if child.tag == qn("w:sectPr"):
                first_sect = child
                break
        if first_sect is None:
            return
        pgSz = first_sect.find(qn("w:pgSz"))
        if pgSz is None:
            return
        if pgSz.get(qn("w:orient")) == "landscape":
            pgSz.set(qn("w:orient"), "portrait")
            pgSz.set(qn("w:w"), "11906")
            pgSz.set(qn("w:h"), "16838")
            logger.info("to-1: первый sectPr был landscape — сброшен в portrait (титул)")
    except Exception:
        logger.exception("to-1: не удалось зафиксировать portrait титула")


def _iter_all_paragraphs(doc: Document):
    """Все абзацы: тело, SDT основного отчёта, таблицы, колонтитулы."""
    seen: set = set()

    def _yield_p(p: Paragraph):
        pid = id(p._p)
        if pid in seen:
            return
        seen.add(pid)
        yield p

    try:
        for sdt in doc.element.body.iter(qn("w:sdt")):
            content = sdt.find(qn("w:sdtContent"))
            if content is None:
                continue
            for p_el in content.iter(qn("w:p")):
                yield from _yield_p(Paragraph(p_el, doc))
    except Exception:
        pass

    for p in doc.paragraphs:
        yield from _yield_p(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield from _yield_p(p)

    try:
        for section in doc.sections:
            for part in (section.header, section.footer):
                if part is None:
                    continue
                for p in part.paragraphs:
                    yield from _yield_p(p)
                for table in part.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                yield from _yield_p(p)
    except Exception:
        pass


def _apply_document_font(doc: Document) -> None:
    """Весь отчёт — Times New Roman (стили + все runs)."""
    try:
        for style_name in (
            "Normal",
            "Heading 1",
            "Heading 2",
            "Heading 3",
            "Title",
            "Subtitle",
            "List Paragraph",
            "Caption",
            "Body Text",
            "No Spacing",
        ):
            try:
                st = doc.styles[style_name]
            except KeyError:
                continue
            try:
                st.font.name = _REPORT_FONT_NAME
                # eastAsia через rPr стиля
                rPr = st.element.get_or_add_rPr()
                rFonts = rPr.rFonts
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.insert(0, rFonts)
                for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                    rFonts.set(qn(attr), _REPORT_FONT_NAME)
            except Exception:
                pass
    except Exception:
        logger.exception("to-1: не удалось задать шрифт стилей")

    for p in _iter_all_paragraphs(doc):
        for r in p.runs:
            _set_run_font(r)


def _add_table_row(table: Table) -> None:
    """Добавить строку, копируя структуру последней и очищая текст."""
    tbl = table._tbl
    last_tr = tbl.tr_lst[-1]
    new_tr = deepcopy(last_tr)
    for tc in new_tr.tc_lst:
        for node in tc.iterchildren():
            if node.tag == qn("w:p"):
                for r in node.findall(qn("w:r")):
                    for t in r.findall(qn("w:t")):
                        t.text = ""
    tbl.append(new_tr)


def _ensure_rows(table: Table, needed: int) -> None:
    while len(table.rows) < needed:
        _add_table_row(table)


def _insert_row_after(table: Table, after_idx: int, values: Sequence[Any]) -> int:
    """Вставить пустую строку-копию сразу после ``after_idx`` и заполнить
    значениями, начиная с колонки 1. Возвращает индекс новой строки.

    Нужно для динамического добавления объектов контроля (напр. ВИК), когда
    их больше, чем предусмотрено строк-заготовок в шаблоне.
    """
    ref_tr = table.rows[after_idx]._tr
    new_tr = deepcopy(ref_tr)
    for tc in new_tr.tc_lst:
        for node in tc.iterchildren():
            if node.tag == qn("w:p"):
                for r in node.findall(qn("w:r")):
                    for t in r.findall(qn("w:t")):
                        t.text = ""
    ref_tr.addnext(new_tr)
    new_idx = after_idx + 1
    for c, v in enumerate(values, start=1):
        _set(table, new_idx, c, v)
    return new_idx


def _row_is_blank(row, ignore_cols: Optional[Sequence[int]] = None) -> bool:
    ignore = set(ignore_cols or ())
    for c, cell in enumerate(row.cells):
        if c in ignore:
            continue
        text = (cell.text or "").strip().strip(".")
        if text and text not in ("—", "-"):
            return False
    return True


def _strip_empty_rows(
    table: Table, start_row: int, ignore_cols: Optional[Sequence[int]] = None
) -> int:
    """Удалить полностью незаполненные строки таблицы (кроме заголовков).

    Нужно, чтобы в готовом отчёте не оставались пустые строки таблиц —
    например, если специалистов/точек контроля меньше, чем строк-заготовок
    в исходном шаблоне.
    """
    tbl = table._tbl
    removed = 0
    for idx in range(len(table.rows) - 1, start_row - 1, -1):
        if idx >= len(table.rows):
            continue
        row = table.rows[idx]
        if _row_is_blank(row, ignore_cols=ignore_cols):
            tbl.remove(row._tr)
            removed += 1
    return removed


def _replace_underscores(text: str, replacements: Sequence[str]) -> str:
    """Последовательно заменить группы подчёркиваний значениями."""
    result = text
    for val in replacements:
        m = _BLANK_RE.search(result)
        if not m:
            break
        result = result[: m.start()] + (val or "____") + result[m.end() :]
    return result


def _fmt_date_ru(value: Any) -> Optional[str]:
    if value is None or value == "":
        return None
    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y")
    s = str(value).strip()
    if not s:
        return None
    if re.match(r"^\d{2}\.\d{2}\.\d{4}$", s):
        return s
    try:
        return datetime.fromisoformat(s.replace("Z", "+00:00")).strftime("%d.%m.%Y")
    except Exception:
        pass
    for fmt, length in (("%Y-%m-%d", 10), ("%d.%m.%Y", 10)):
        try:
            return datetime.strptime(s[:length], fmt).strftime("%d.%m.%Y")
        except ValueError:
            continue
    return s


# ---------------------------------------------------------------------------
# Основной отчёт (SDT: титул + разделы 1–15)
# ---------------------------------------------------------------------------

def _fix_main_report_captions(doc: Document) -> None:
    """Исправить задвоение «Таблица № 6» у раздела 10 → «Таблица № 7»."""
    seen_works = False
    for p in _iter_all_paragraphs(doc):
        tx = (p.text or "").strip()
        # Не брать строку оглавления («…диагностирования5»)
        if tx.startswith("10. Перечень работ, выполненных") and not tx[-1:].isdigit():
            seen_works = True
            continue
        if seen_works and tx in ("Таблица № 6", "Таблица №6"):
            _set_paragraph_text(p, "Таблица № 7")
            break


def _uzt_smin_summary(ctx: Dict[str, Any]) -> str:
    g = ctx["g"]
    measurements = g("thickness_measurements", "thicknessMeasurements", default=[])
    if not isinstance(measurements, list) or not measurements:
        return "Обечайка Smin= — мм\nДнище верхнее Smin= — мм\nДнище нижнее Smin= — мм"
    by_el: Dict[str, List[float]] = {}
    for m in measurements:
        if not isinstance(m, dict):
            continue
        el = str(m.get("element") or m.get("element_name") or m.get("location") or "Элемент")
        raw = m.get("thickness") or m.get("value") or m.get("thickness_mm")
        try:
            val = float(str(raw).replace(",", "."))
        except (TypeError, ValueError):
            continue
        by_el.setdefault(el, []).append(val)
    if not by_el:
        return "Обечайка Smin= — мм\nДнище верхнее Smin= — мм\nДнище нижнее Smin= — мм"
    lines = []
    for el, vals in by_el.items():
        lines.append(f"{el} Smin= {min(vals):g} мм")
    return "\n".join(lines)


def _residual_life_text(ctx: Dict[str, Any]) -> str:
    g = ctx["g"]
    calc = g("calculation_data", "calculationData", default=None)
    if not isinstance(calc, dict):
        calc = {}
    residual = (
        calc.get("residual_life_years")
        or calc.get("residual_life")
        or g("residual_life_text", "residual_life_years", default="")
    )
    if residual in (None, "", MISSING):
        return ""
    s = str(residual).strip()
    if not s:
        return ""
    if "лет" in s.lower() or "год" in s.lower():
        return s
    return f"{s} лет"


def _conclusion_needs_repair(ctx: Dict[str, Any]) -> bool:
    blob = " ".join(
        str(ctx.get(k) or "")
        for k in ("conclusion_suitable", "tech_state", "calculation_result")
    ).lower()
    if any(k in blob for k in ("ремонт", "негоден", "не годен", "ограничен")):
        return True
    g = ctx["g"]
    for key in ("weld_inspections", "uzk_results", "mpk_results", "visual_defects"):
        if _ndt_items_have_defects(g(key, default=[])):
            return True
    return False


def _hardness_summary(ctx: Dict[str, Any]) -> str:
    g = ctx["g"]
    tests = g("hardness_tests", "hardnessTests", default=[])
    if not isinstance(tests, list) or not tests:
        return "—"
    blob = " ".join(
        str(t.get(k) or "")
        for t in tests
        if isinstance(t, dict)
        for k in ("conclusion", "assessment", "note", "remark")
    ).lower()
    if any(k in blob for k in ("превыш", "не соотв", "ремонт")):
        return "Значения твёрдости выходят за допускаемые пределы (см. приложение № 5)"
    return "Твёрдость металла в пределах допускаемых значений (см. приложение № 5)"


def _widen_date_column(table: Table, col: int = 0, width_cm: float = 3.4) -> None:
    """Колонка «Дата»: обычный шрифт, шире, дата в одну строку."""
    try:
        for row in table.rows:
            if col >= len(row.cells):
                continue
            cell = row.cells[col]
            cell.width = Cm(width_cm)
            _force_horizontal_text(cell)
            for p in cell.paragraphs:
                for run in p.runs:
                    _set_run_font(run, pt=12, bold=False)
    except Exception:
        logger.debug("widen date column failed", exc_info=True)


def _fill_main_report(doc: Document, ctx: Dict[str, Any]) -> None:
    """Заполнить титул и таблицы/абзацы разделов 1–15 внутри SDT."""
    main_tables = _main_sdt_tables(doc)
    if not main_tables:
        logger.warning("to-1: основной отчёт (SDT) не найден — титул/разд.1–15 не заполнены")
        return

    g = ctx["g"]
    device = ctx["device_name"]
    serial = ctx["serial"]
    reg_no = ctx["reg_no"]
    inv_no = ctx["inv_no"]
    location = ctx["location"]
    org_name = ctx["org_name"]

    # --- Таблица 0: титул ---
    if len(main_tables) > 0:
        title = main_tables[0]
        if len(title.rows) > 6:
            _set(title, 6, 1, device)
            if len(title.rows[6].cells) > 2:
                _set(title, 6, 2, device)
        if len(title.rows) > 7:
            _set(title, 7, 1, serial, nowrap=True)
            if len(title.rows[7].cells) > 2:
                _set(title, 7, 2, serial, nowrap=True)
        if len(title.rows) > 8:
            _set(title, 8, 1, reg_no, nowrap=True)
            if len(title.rows[8].cells) > 2:
                _set(title, 8, 2, reg_no, nowrap=True)
        if len(title.rows) > 9:
            # В одной ячейке шаблона склеены инв.№ / ОПО
            inv_block = (
                f"{inv_no}\n{ctx.get('opo_name') or MISSING}\n"
                f"{ctx.get('opo_class') or MISSING}\n{ctx.get('opo_reg') or MISSING}"
            )
            _set(title, 9, 1, inv_block)
            if len(title.rows[9].cells) > 2:
                _set(title, 9, 2, inv_block)
        if len(title.rows) > 11:
            # Местонахождение объекта на титуле — полный адрес: наименование
            # заказчика/структурного подразделения + место нахождения.
            addr_parts = [p for p in (org_name, location) if p and p != MISSING]
            loc_text = ", ".join(addr_parts) if addr_parts else MISSING
            _set(title, 11, 1, loc_text)
            if len(title.rows[11].cells) > 2:
                _set(title, 11, 2, loc_text)

    # --- Таблица 1: номер отчёта ---
    if len(main_tables) > 1:
        report_no = str(g("protocol_number", "report_number", default="") or "")
        if report_no:
            _set(main_tables[1], 0, 1, report_no, nowrap=True)

    # --- Таблица 2: заказчик ---
    if len(main_tables) > 2:
        cust = main_tables[2]
        vals = [
            org_name,
            ctx.get("customer_director") or MISSING,
            ctx.get("customer_address") or MISSING,
            location if location != MISSING else MISSING,
            ctx.get("customer_phone") or MISSING,
            ctx.get("customer_email") or MISSING,
        ]
        for i, v in enumerate(vals):
            if i < len(cust.rows):
                _set(cust, i, 1, v)

    # --- Таблица 3: исполнитель ---
    if len(main_tables) > 3:
        contr = main_tables[3]
        vals = [
            ctx.get("contractor_name") or MISSING,
            ctx.get("contractor_director") or MISSING,
            ctx.get("contractor_address") or MISSING,
            ctx.get("contractor_address") or MISSING,
            ctx.get("contractor_phone") or MISSING,
            ctx.get("contractor_email") or MISSING,
            ctx.get("lab_cert") or MISSING,
        ]
        for i, v in enumerate(vals):
            if i < len(contr.rows):
                _set(contr, i, 1, v)

    # --- Таблица 4: специалисты ---
    if len(main_tables) > 4:
        specs = [s for s in (ctx.get("specialists") or []) if isinstance(s, dict) and (s.get("name") or "").strip()]
        st = main_tables[4]
        _ensure_rows(st, 1 + max(len(specs), 1))
        # Очистить все строки данных
        for r in range(1, len(st.rows)):
            for c in range(len(st.rows[r].cells)):
                _set(st, r, c, "")
        for i, s in enumerate(specs):
            r = i + 1
            if r >= len(st.rows):
                _ensure_rows(st, r + 1)
            _set(st, r, 0, f"{i + 1}.")
            _set(st, r, 1, s.get("name") or "")
            _set(st, r, 2, s.get("cert") or MISSING, nowrap=True)
            _set(st, r, 3, s.get("role") or s.get("area") or MISSING)
            _set(
                st,
                r,
                4,
                _fmt_date_ru(s.get("valid_until") or s.get("expiry"))
                or (s.get("valid_until") or s.get("expiry"))
                or MISSING,
            )
        _strip_empty_rows(st, 1, ignore_cols=(0,))
        _renumber_table_column(st, 1, 0)

    # --- Таблица 5: приборы ---
    if len(main_tables) > 5:
        ve = [e for e in (ctx.get("verification_equipment") or []) if isinstance(e, dict)]
        it = main_tables[5]
        _ensure_rows(it, 1 + max(len(ve), 1))
        for r in range(1, len(it.rows)):
            for c in range(len(it.rows[r].cells)):
                _set(it, r, c, "")
        for i, eq in enumerate(ve):
            r = i + 1
            if r >= len(it.rows):
                _ensure_rows(it, r + 1)
            _set(it, r, 0, f"{i + 1}.")
            _set(it, r, 1, _instrument_full_name(eq))
            _set(
                it,
                r,
                2,
                eq.get("serial_number") or eq.get("factory_number") or MISSING,
                nowrap=True,
            )
            _set(
                it,
                r,
                3,
                eq.get("verification_certificate_number")
                or eq.get("certificate")
                or eq.get("verification_certificate")
                or MISSING,
                nowrap=True,
            )
            _set(
                it,
                r,
                4,
                _fmt_date_ru(
                    eq.get("next_verification_date")
                    or eq.get("valid_until")
                    or eq.get("verification_until")
                )
                or eq.get("next_verification_date")
                or MISSING,
                nowrap=True,
            )
        _strip_empty_rows(it, 1, ignore_cols=(0,))
        _renumber_table_column(it, 1, 0)

    # --- Таблица 6: перечень объектов ---
    if len(main_tables) > 6:
        ot = main_tables[6]
        obj_rows = [
            ("Наименование", device),
            ("Заводской №", serial),
            ("Регистрационный №", reg_no),
            ("Инвентарный №", inv_no),
            ("Местонахождение", location),
            ("Заказчик", org_name),
        ]
        _ensure_rows(ot, len(obj_rows))
        for i, (k, v) in enumerate(obj_rows):
            if i < len(ot.rows):
                if not (ot.rows[i].cells[0].text or "").strip():
                    _set(ot, i, 0, k)
                _set(ot, i, 1, v)
        # Удалить хвост пустых строк шаблона (лишние 3 строки и т.п.)
        _strip_empty_rows(ot, len(obj_rows), ignore_cols=None)
        # Если после 6-й есть ещё строки — вырезать
        while len(ot.rows) > len(obj_rows):
            ot._tbl.remove(ot.rows[-1]._tr)

    # --- Таблица 7: краткая теххарактеристика ---
    if len(main_tables) > 7:
        tech = main_tables[7]
        tech_vals = [
            device,
            g("purpose", "vessel_purpose", default=MISSING),
            g("designation", "conditional_designation", "scheme_index", default=MISSING),
            g("manufacturer", default=MISSING),
            g("manufacture_year", "year_of_manufacture", "manufacturing_year", default=MISSING),
            g("commissioning_year", default=MISSING),
            g("design_pressure", default=MISSING),
            g("working_pressure", default=MISSING),
            g("diameter", "inner_diameter", default=MISSING),
            g("working_medium", "medium", default=MISSING),
            g(
                "working_medium_temperature",
                "medium_temperature",
                "working_temperature",
                default=MISSING,
            ),
            g("shell_material", "material", default=MISSING),
            g("volume", "capacity", default=MISSING),
            g("connection_scheme", default=MISSING),
            g("climatic_version", default=MISSING),
            g("service_life", default=MISSING),
        ]
        for i, v in enumerate(tech_vals):
            if i < len(tech.rows):
                _set(tech, i, 1, v if v not in (None, "") else MISSING)

    # --- Таблица 8: перечень работ ---
    if len(main_tables) > 8:
        works = main_tables[8]
        ndt = ctx.get("ndt_methods") or []
        scope_by_key = {
            "ВИК": "100%",
            "УЗТ": "по схеме контроля",
            "УЗК": "по схеме контроля",
            "МПК": "100%",
            "ТВЕРД": "по схеме контроля",
        }
        for r in range(1, len(works.rows)):
            name = (works.rows[r].cells[1].text or "").upper()
            scope = "в полном объёме"
            ntd = "СТО Газпром 2-2.3-491-2010"
            if "ВИЗУАЛЬ" in name:
                scope = scope_by_key["ВИК"]
                ntd = "СТО 9701105632-003-2021"
            elif "ТОЛЩИНОМЕТР" in name:
                scope = scope_by_key["УЗТ"]
                ntd = "ГОСТ Р ИСО 16809-2015"
            elif "ТВЕРД" in name:
                scope = scope_by_key["ТВЕРД"]
                ntd = "ГОСТ 22761-77"
            elif "УЛЬТРАЗВУКОВОЙ КОНТРОЛЬ КАЧЕСТВА" in name or "УЗК" in name:
                scope = scope_by_key["УЗК"]
                ntd = "ГОСТ Р 55724-2013"
            elif "МАГНИТ" in name:
                scope = scope_by_key["МПК"]
                ntd = "ГОСТ Р 56512-2015"
            elif "ГИДРАВЛ" in name:
                scope = "в соответствии с программой"
                ntd = "ФНП №536"
            # Если метод не выполнялся — оставить пустым объём
            performed = True
            if isinstance(ndt, list) and ndt:
                # не блокируем стандартный перечень — всегда заполняем объём/НТД
                performed = True
            if performed:
                _set(works, r, 2, scope)
                _set(works, r, 3, ntd)

    # --- Таблица 9: рассмотренные документы (как прил.1 / табл.1) ---
    if len(main_tables) > 9:
        docs_tbl = main_tables[9]
        docs_dict = ctx["docs_dict"]
        docs_info = ctx["docs_info"]
        for r in range(1, len(docs_tbl.rows)):
            num_txt = (docs_tbl.rows[r].cells[0].text or "").strip().rstrip(".")
            if not num_txt.isdigit():
                continue
            info = docs_info.get(num_txt) or docs_info.get(int(num_txt)) or {}
            if not isinstance(info, dict):
                info = {}
            present = docs_dict.get(num_txt, docs_dict.get(int(num_txt)))
            ident, pages = _doc_ident_and_pages(present, info)
            _set(docs_tbl, r, 2, ident)
            _set(docs_tbl, r, 3, pages)

    # --- Таблица 10: предыдущие обследования ---
    if len(main_tables) > 10:
        prev_tbl = main_tables[10]
        records = g("previous_inspections", default=[])
        if not isinstance(records, list):
            records = []
        if not records:
            legacy = g("previous_inspection_result", default="")
            if legacy and legacy != MISSING:
                records = [{"kind": "Техническое диагностирование", "result": legacy}]
        for i, rec in enumerate(records[: max(0, len(prev_tbl.rows) - 1)]):
            if not isinstance(rec, dict):
                continue
            r = i + 1
            _set(prev_tbl, r, 0, f"{i + 1}.")
            _set(prev_tbl, r, 1, rec.get("kind") or rec.get("type") or "")
            _set(prev_tbl, r, 2, rec.get("result") or "")
            report = rec.get("report_number") or rec.get("report") or ""
            date = _fmt_date_ru(rec.get("date")) or rec.get("date") or ""
            doc_ref = f"{report} от {date}".strip(" от") if (report or date) else ""
            _set(prev_tbl, r, 3, doc_ref)
        _strip_empty_rows(prev_tbl, 1, ignore_cols=(0,))

    # --- Таблица 11: результаты ТД ---
    if len(main_tables) > 11:
        res = main_tables[11]
        smin = _uzt_smin_summary(ctx)
        uzk_cell = _ndt_result_summary(
            ctx,
            data_keys=("weld_inspections", "uzk_results", "weld_defects"),
            custom_key="uzk_conclusion_text",
            ok_text="Недопустимых дефектов не обнаружено",
            defect_text="Выявлены дефекты, требуется ремонт (см. приложение № 6)",
        )
        mpk_cell = _ndt_result_summary(
            ctx,
            data_keys=("mpk_results", "magnetic_results", "weld_inspections"),
            custom_key="mpk_conclusion_text",
            ok_text="Недопустимых дефектов не обнаружено",
            defect_text="Выявлены дефекты, требуется ремонт (см. приложение № 7)",
        )
        vik_cell = _ndt_result_summary(
            ctx,
            data_keys=("visual_defects", "vik_defects", "vik_control_objects"),
            custom_key="vik_conclusion_text",
            ok_text="Дефектов, препятствующих дальнейшей безопасной эксплуатации сосуда не выявлено",
            defect_text="Выявлены дефекты (см. приложение № 3)",
        )
        residual = _residual_life_text(ctx)
        calc_cell = ctx.get("calculation_result") or "Условия прочности выполняются"
        if residual:
            calc_cell = f"{calc_cell}. Прогнозный срок остаточного ресурса: {residual}"
        result_by_row = {
            1: _doc_analysis_result_cell(ctx),
            2: (ctx.get("operational_ok") or "Соответствует требованиям"),
            3: vik_cell,
            4: smin,
            5: _hardness_summary(ctx),
            6: uzk_cell,
            7: mpk_cell,
            8: calc_cell,
        }
        for r, val in result_by_row.items():
            if r < len(res.rows) and len(res.rows[r].cells) > 2:
                _set(res, r, 2, val)
        appendix_nums = {
            1: "1",
            2: "2",
            3: "3",
            4: "4",
            5: "5",
            6: "6",
            7: "7",
            8: "8",
        }
        for r, app in appendix_nums.items():
            if r < len(res.rows) and len(res.rows[r].cells) > 3:
                cell_txt = (res.rows[r].cells[3].text or "").strip()
                if "Приложение" in cell_txt and "_" in cell_txt:
                    _set(res, r, 3, f"Приложение № {app}")

    # --- Абзацы разделов 1, 2, 14, 15 + разрыв перед СОДЕРЖАНИЕ ---
    contract = str(g("contract_number", default="") or "")
    work_basis = str(g("work_basis", "basis", default="") or "")
    contract_date = _fmt_date_ru(g("contract_date", default="")) or str(g("contract_date", default="") or "")
    period_from = _fmt_date_ru(g("work_period_from", "date_from", default="")) or ""
    period_to = _fmt_date_ru(g("work_period_to", "date_to", "date_performed", default=ctx["date_ru"])) or ctx["date_ru"]
    calc_txt = ctx.get("calculation_result") or "сосуда при рабочих параметрах"
    tech_state = ctx.get("tech_state") or "работоспособное"
    conclusion = ctx.get("conclusion_suitable") or "соответствует"
    contractor_name = ctx.get("contractor_name") or ""
    residual = _residual_life_text(ctx)
    needs_repair = _conclusion_needs_repair(ctx)

    for p in _iter_all_paragraphs(doc):
        text = p.text or ""
        stripped = text.strip()
        norm = _norm_ws(text)
        if stripped == "СОДЕРЖАНИЕ":
            _insert_page_break_before_paragraph(p)
            try:
                p.paragraph_format.page_break_before = True
                p.paragraph_format.keep_with_next = True
            except Exception:
                pass
            continue
        if "проведены в соответствии с" in norm and "договор" in norm:
            parties = org_name if org_name != MISSING else "__________________________"
            if contractor_name:
                parties = f"{parties} и {contractor_name}"
            c_no = contract or "____________"
            c_dt = contract_date or "______"
            if work_basis and not contract:
                _set_paragraph_text(
                    p,
                    f"Работы по техническому диагностированию проведены в соответствии с {work_basis}.",
                )
            else:
                _set_paragraph_text(
                    p,
                    f"Работы по техническому диагностированию проведены в соответствии с договором между {parties} от {c_dt} № {c_no}.",
                )
        elif "проведены в период" in norm:
            pf = period_from or "__.__.____"
            pt = period_to or "__.__.____"
            loc = location if location != MISSING else "____________________"
            _set_paragraph_text(
                p,
                f"Работы по техническому диагностированию проведены в период с {pf} по {pt}, "
                f"на объекте {loc}.",
            )
        elif "По результатам работ произведена оценка работоспособности" in text or (
            "оценка работоспособности" in norm and "приложение" in norm
        ):
            extra = f" Прогнозный срок остаточного ресурса составляет {residual}." if residual else ""
            _set_paragraph_text(
                p,
                f"По результатам работ произведена оценка работоспособности {calc_txt}.{extra} (Приложение № 8)",
            )
        elif (
            "фактическое значение параметров" in norm.lower()
            or (
                "состояние эксплуатации" in norm.lower()
                and ("зав" in norm.lower() or "инв" in norm.lower())
            )
            or (
                "работающего под давлением" in norm.lower()
                and "____" in text
                and "удовлетвор" in norm.lower()
            )
        ):
            satisfy = (
                "не удовлетворяют требованиям нормативных документов без выполнения ремонта"
                if needs_repair
                else "удовлетворяют требованиям нормативных документов"
            )
            _set_paragraph_text(
                p,
                f"Фактическое значение параметров, определяющих состояние эксплуатации сосуда "
                f"работающего под давлением – {device}, зав.№ {serial}, рег.№ {reg_no}, "
                f"инв.№ {inv_no}, {satisfy}.",
            )
        elif "техническое состояние объекта диагностирования" in norm.lower():
            vik_s = _ndt_result_summary(
                ctx,
                data_keys=("visual_defects", "vik_defects", "vik_control_objects"),
                custom_key="vik_conclusion_text",
                ok_text="дефектов не выявлено",
                defect_text="выявлены дефекты (прил. № 3)",
            )
            uzk_s = _ndt_result_summary(
                ctx,
                data_keys=("weld_inspections", "uzk_results"),
                custom_key="uzk_conclusion_text",
                ok_text="дефектов не обнаружено",
                defect_text="выявлены дефекты, требуется ремонт (прил. № 6)",
            )
            mpk_s = _ndt_result_summary(
                ctx,
                data_keys=("mpk_results", "weld_inspections"),
                custom_key="mpk_conclusion_text",
                ok_text="дефектов не обнаружено",
                defect_text="выявлены дефекты, требуется ремонт (прил. № 7)",
            )
            summary = (
                f"Результаты по видам контроля: анализ документации — {_doc_analysis_result_cell(ctx)}; "
                f"ВИК — {vik_s}; УЗТ — {_uzt_smin_summary(ctx).replace(chr(10), '; ')}; "
                f"твёрдость — {_hardness_summary(ctx)}; УЗК — {uzk_s}; МПК — {mpk_s}; "
                f"расчёт на прочность — {calc_txt}"
                + (f", остаточный ресурс {residual}" if residual else "")
                + "."
            )
            _set_paragraph_text(
                p,
                f"Техническое состояние объекта диагностирования: {tech_state}. {conclusion}. {summary}".replace("..", "."),
            )


def _fill_protocol_header(table: Table, ctx: Dict[str, Any]) -> None:
    """Шапка протокола 8×3: исполнитель / заказчик / оборудование."""
    if len(table.rows) < 7:
        return
    _set(table, 0, 0, ctx.get("contractor_name") or "")
    _set(table, 0, 2, ctx.get("org_name") or "")
    _set(table, 2, 0, ctx.get("contractor_address") or "")
    # Место нахождения оборудования — отдельная графа (не путать с заказчиком)
    _set(table, 2, 2, ctx.get("location") or MISSING)
    _set(table, 4, 0, ctx.get("lab_name") or "")
    _set(table, 4, 2, ctx.get("device_name") or "")
    _set(table, 6, 0, ctx.get("lab_cert") or "")
    ids = f"Зав.№ {ctx['serial']}, рег.№ {ctx['reg_no']}, инв.№ {ctx['inv_no']}"
    _set(table, 6, 2, ids, nowrap=True)


def _merge_document_sets_into_info(
    data: Dict[str, Any], docs_info: Dict[str, Any]
) -> Dict[str, Any]:
    """Подтянуть номер/дату/листы из document_sets (ремонтная документация, акты УЗТ)."""
    out = dict(docs_info)
    extra = data.get("additional_data") if isinstance(data.get("additional_data"), dict) else {}
    sets = data.get("document_sets") or extra.get("document_sets") or {}
    if not isinstance(sets, dict):
        return out
    for num, lst in sets.items():
        if not isinstance(lst, list) or not lst:
            continue
        idents: List[str] = []
        pages_parts: List[str] = []
        for item in lst:
            if not isinstance(item, dict):
                continue
            n = str(item.get("number") or "").strip()
            d = _fmt_date_ru(item.get("date")) or str(item.get("date") or "").strip()
            p = str(item.get("pages") or item.get("volume") or "").strip()
            ident = n
            if d:
                ident = f"{n} от {d}".strip() if n else f"от {d}"
            if ident:
                idents.append(ident)
            if p:
                pages_parts.append(p)
        key = str(num)
        info = dict(out.get(key) or out.get(num) or {})
        if not isinstance(info, dict):
            info = {}
        if idents and not str(info.get("number") or "").strip():
            info["number"] = "; ".join(idents)
            info["date"] = ""
        if pages_parts and not str(info.get("pages") or info.get("volume") or "").strip():
            info["pages"] = "; ".join(pages_parts)
        out[key] = info
    return out


def _ndt_methods_joined(ctx: Dict[str, Any]) -> str:
    """Список методов НК из чек-листа / таблицы ndt_methods."""
    names: List[str] = []
    data = ctx.get("data") or {}
    raw = data.get("ndt_methods") or ctx.get("ndt_methods") or []
    if isinstance(raw, list):
        for m in raw:
            if isinstance(m, str):
                val = m.strip()
            elif isinstance(m, dict):
                val = str(
                    m.get("method_code") or m.get("code") or m.get("name") or ""
                ).strip()
            else:
                val = ""
            if val and val not in names:
                names.append(val)
    return ", ".join(names)


def _hide_empty_row_borders(table: Table, start: int, end: int) -> None:
    """Убрать горизонтальные линии у полностью пустых строк (лишняя черта в блоке температур)."""
    for r in range(start, min(end, len(table.rows))):
        texts = [(c.text or "").strip() for c in table.rows[r].cells]
        if any(t and t not in ("—", "-", "–") for t in texts):
            continue
        for cell in table.rows[r].cells:
            try:
                tcPr = cell._tc.get_or_add_tcPr()
                borders = tcPr.find(qn("w:tcBorders"))
                if borders is None:
                    borders = OxmlElement("w:tcBorders")
                    tcPr.append(borders)
                for edge in ("top", "bottom"):
                    el = borders.find(qn(f"w:{edge}"))
                    if el is None:
                        el = OxmlElement(f"w:{edge}")
                        borders.append(el)
                    el.set(qn("w:val"), "nil")
            except Exception:
                pass


def _doc_ident_and_pages(present: Any, info: Dict[str, Any]) -> Tuple[str, str]:
    """Идентификатор и объём документа; при отсутствии — «Не предоставлено»."""
    doc_number = str(info.get("number") or info.get("doc_number") or "").strip()
    doc_date = _fmt_date_ru(info.get("date") or info.get("doc_date")) or ""
    pages = str(info.get("pages") or info.get("volume") or "").strip()
    ident = doc_number
    if doc_date:
        ident = f"{ident} от {doc_date}".strip() if ident else f"от {doc_date}"
    provided = present is True or (present is not False and bool(ident))
    if present is False or not provided:
        # Если документ «Не предоставлено» — соседняя колонка должна быть
        # тире, а не дублировать тот же текст.
        return NOT_PROVIDED, MISSING
    return ident or NOT_PROVIDED, pages or MISSING


def _fill_documents_table(table: Table, ctx: Dict[str, Any]) -> None:
    docs_dict = ctx["docs_dict"]
    docs_info = ctx["docs_info"]
    keys = set(str(k) for k in docs_dict.keys()) | set(str(k) for k in docs_info.keys())
    keys |= set(TO_DOCUMENT_NAMES.keys())
    ordered = sorted(keys, key=lambda x: int(x) if x.isdigit() else 999)

    # Строка 0 — заголовок; данные с 1
    needed = len(ordered) + 1
    _ensure_rows(table, needed)

    for i, num in enumerate(ordered):
        if not str(num).isdigit():
            continue
        row = i + 1
        name = TO_DOCUMENT_NAMES.get(str(num), f"Документ {num}")
        info = docs_info.get(str(num)) or docs_info.get(num) or {}
        if not isinstance(info, dict):
            info = {}
        present = docs_dict.get(str(num), docs_dict.get(num))
        ident, pages = _doc_ident_and_pages(present, info)
        _set(table, row, 0, f"{num}.")
        _set(table, row, 1, name)
        _set(table, row, 2, ident)
        if len(table.rows[row].cells) > 3:
            _set(table, row, 3, pages)


def _fill_general_data(table: Table, ctx: Dict[str, Any]) -> None:
    g = ctx["g"]
    rows = [
        g("vessel_name", "equipment_device_name", default=ctx["device_name"]),
        g("designation", "conditional_designation", default=MISSING),
        g("manufacturer", default=MISSING),
        g("manufacture_year", "year_of_manufacture", default=MISSING),
        g("commissioning_year", default=MISSING),
        g("working_pressure", default=MISSING),
        g("diameter", "inner_diameter", default=MISSING),
        g("working_temperature", default=MISSING),
        g("working_medium", "medium", default=MISSING),
        g("shell_material", "material", default=MISSING),
        g("volume", "capacity", default=MISSING),
        g("connection_scheme", default=MISSING),
        g("climatic_version", default=MISSING),
    ]
    for i, val in enumerate(rows):
        if i < len(table.rows):
            _set(table, i, 1, val)


def _fill_elements_table(table: Table, ctx: Dict[str, Any]) -> None:
    g = ctx["g"]
    elements = g("vessel_elements", "elements", default=[])
    if not isinstance(elements, list):
        elements = []
    # Строки 0–1 заголовки; данные с 2
    data_start = 2
    if not elements:
        # Заполнить корпус из паспортных полей
        defaults = [
            ("Корпус", g("shell_qty", default="1"), g("diameter", default=""), g("shell_length", "height", default=""),
             g("wall_thickness", "thickness", default=""), g("calc_thickness", default=""),
             g("shell_material", "material", default=""), g("material_gost", default=""),
             g("weld_type", default=""), g("electrodes", default=""), g("ndt_method", default="") or _ndt_methods_joined(ctx)),
        ]
        for i, row_vals in enumerate(defaults):
            r = data_start + i
            if r >= len(table.rows):
                break
            for c, v in enumerate(row_vals):
                if c < len(table.rows[r].cells):
                    _set(table, r, c, v, nowrap=(c in (2, 3, 4, 5, 6, 7)))
        return

    needed = data_start + len(elements)
    _ensure_rows(table, needed)
    for i, el in enumerate(elements):
        if not isinstance(el, dict):
            continue
        r = data_start + i
        vals = [
            el.get("name") or el.get("element_name") or "",
            el.get("quantity") or el.get("qty") or "1",
            el.get("diameter_mm")
            or el.get("inner_diameter")
            or el.get("diameter")
            or "",
            el.get("length_mm") or el.get("length") or el.get("height") or "",
            el.get("wall_thickness_mm")
            or el.get("nominal_thickness")
            or el.get("thickness")
            or el.get("wall_thickness")
            or "",
            # Расчётная толщина до прибавки на коррозию
            el.get("calc_thickness")
            or el.get("calculated_thickness")
            or el.get("design_thickness")
            or "",
            el.get("steel_grade") or el.get("material") or "",
            el.get("gost") or el.get("material_gost") or "",
            el.get("weld_type") or el.get("weld_data") or "",
            el.get("electrodes") or "",
            el.get("ndt_method") or _ndt_methods_joined(ctx),
        ]
        for c, v in enumerate(vals):
            if c < len(table.rows[r].cells):
                # Размеры и марка стали — без «ломания» по символу в узких колонках
                _set(table, r, c, v, nowrap=(c in (2, 3, 4, 5, 6, 7)))
                _force_horizontal_text(table.rows[r].cells[c])
    # Очистить оставшиеся строки-заготовки шаблона (напр. «Нижнее днище»)
    for r in range(data_start + len(elements), len(table.rows)):
        for c in range(len(table.rows[r].cells)):
            _set(table, r, c, "")


def _fill_characteristics(table: Table, ctx: Dict[str, Any]) -> None:
    g = ctx["g"]
    # Колонки: наименование | подпоказатель | проектные | фактические | примечание
    # Строки: 1 раб.давл, 2 расч.давл, 3 пневмо, 4 гидро,
    # 5 расч.t стенки, 6 t рабочей среды, 7 min t стенки, 8 состав среды,
    # 9 класс опасности, 10 взрывоопасность, 11 пожароопасность, …
    design_temp = g("design_temperature", default="")
    # Температура рабочей среды — отдельное поле; не путать с t стенки
    medium_temp = g(
        "working_medium_temperature",
        "medium_temperature",
        "working_temperature",
        default="",
    )
    wall_work_temp = g("wall_working_temperature", "working_temperature_wall", default="")
    mapping = {
        1: (g("working_pressure_design", "working_pressure", default=""), g("working_pressure", default="")),
        2: (g("design_pressure", default=""), g("design_pressure_fact", "design_pressure", default="")),
        3: (g("test_pressure_pneumo", default=""), g("test_pressure_pneumo_fact", default="")),
        4: (g("test_pressure", "test_pressure_hydro", default=""), g("test_pressure_fact", "test_pressure", default="")),
        5: (design_temp, g("design_temperature_fact", "design_temperature", default=design_temp)),
        6: (medium_temp, g("working_medium_temperature_fact", "working_medium_temperature", "medium_temperature", default=medium_temp)),
        7: (
            g("min_wall_temp", default=wall_work_temp),
            g("min_wall_temp_fact", "min_wall_temp", default=wall_work_temp),
        ),
        8: (g("working_medium", "medium", default=""), g("working_medium_fact", "working_medium", "medium", default="")),
        9: (
            g("hazard_class", "medium_hazard_class", default=""),
            g("hazard_class_fact", "hazard_class", "medium_hazard_class", default=""),
        ),
        10: (
            g("explosion_hazard", "explosion_category", default=""),
            g("explosion_hazard_fact", "explosion_hazard", "explosion_category", default=""),
        ),
        11: (
            g("fire_hazard", "fire_category", default=""),
            g("fire_hazard_fact", "fire_hazard", "fire_category", default=""),
        ),
        12: (g("volume", "capacity", default=""), g("volume_fact", "volume", "capacity", default="")),
        13: (g("empty_mass", "mass", default=""), g("empty_mass_fact", "empty_mass", "mass", default="")),
        14: (g("corrosion_allowance", default=""), g("corrosion_allowance_fact", "corrosion_allowance", default="")),
        15: (g("load_cycles", default=""), g("load_cycles_fact", "load_cycles", default="")),
        16: (g("service_life", default=""), g("service_life_fact", "service_life", default="")),
    }
    for row, (proj, fact) in mapping.items():
        if row < len(table.rows):
            cols = len(table.rows[row].cells)
            if cols >= 4:
                _set(table, row, 2, proj if proj not in (None, MISSING) else "", nowrap=True)
                _set(table, row, 3, fact if fact not in (None, MISSING) else "", nowrap=True)
    # Лишняя горизонтальная черта в блоке температур (пустая строка шаблона)
    _hide_empty_row_borders(table, 5, 9)


def _fill_materials(table: Table, ctx: Dict[str, Any]) -> None:
    g = ctx["g"]
    materials = g("materials", "element_materials", default=[])
    if not isinstance(materials, list) or not materials:
        materials = []
        elements = g("vessel_elements", "elements", default=[])
        if isinstance(elements, list):
            for el in elements:
                if not isinstance(el, dict):
                    continue
                mat = el.get("material") or el.get("steel_grade")
                if not mat:
                    continue
                materials.append(
                    {
                        "element": el.get("name") or el.get("element_name") or "",
                        "grade": mat,
                        "gost": el.get("gost") or el.get("material_gost") or "",
                        "yield_strength": el.get("yield_strength") or "",
                        "tensile_strength": el.get("tensile_strength") or "",
                        "elongation": el.get("elongation") or "",
                        "reduction": el.get("reduction") or "",
                        "impact": el.get("impact") or "",
                        "temperature": el.get("temperature") or el.get("test_temperature") or "",
                        "specimen_type": el.get("specimen_type") or "",
                    }
                )
    if isinstance(materials, list) and materials:
        start = 2
        _ensure_rows(table, start + len(materials))
        for i, m in enumerate(materials):
            if not isinstance(m, dict):
                continue
            r = start + i
            vals = [
                m.get("element") or m.get("name") or "",
                m.get("grade") or m.get("material") or "",
                m.get("gost") or "",
                m.get("yield_strength") or "",
                m.get("tensile_strength") or "",
                m.get("elongation") or "",
                m.get("reduction") or "",
                m.get("impact") or "",
                m.get("temperature") or "",
                m.get("specimen_type") or "",
            ]
            for c, v in enumerate(vals):
                if c < len(table.rows[r].cells):
                    _set(table, r, c, v)
        # Если у элемента той же марки стали пустые мех. свойства — взять с любой строки той же марки
        mech_cols = (3, 4, 5, 6, 7, 8, 9)
        last = start + len(materials)
        for r in range(start, last):
            if r >= len(table.rows):
                break
            grade = (table.rows[r].cells[1].text or "").strip()
            if not grade:
                continue
            empty = all(
                not (table.rows[r].cells[c].text or "").strip()
                for c in mech_cols
                if c < len(table.rows[r].cells)
            )
            if not empty:
                continue
            for src in range(start, last):
                if src == r or src >= len(table.rows):
                    continue
                if (table.rows[src].cells[1].text or "").strip() != grade:
                    continue
                src_empty = all(
                    not (table.rows[src].cells[c].text or "").strip()
                    for c in mech_cols
                    if c < len(table.rows[src].cells)
                )
                if src_empty:
                    continue
                for c in mech_cols:
                    if c < len(table.rows[r].cells):
                        _set(table, r, c, table.rows[src].cells[c].text)
                break
        for r in range(start + len(materials), len(table.rows)):
            for c in range(len(table.rows[r].cells)):
                _set(table, r, c, "")
        return
    # Минимум — корпус
    if len(table.rows) > 2:
        _set(table, 2, 0, "Корпус")
        _set(table, 2, 1, g("shell_material", "material", default=""))
        _set(table, 2, 2, g("material_gost", default=""))
        for r in range(3, len(table.rows)):
            for c in range(len(table.rows[r].cells)):
                _set(table, r, c, "")


def _fill_heat_treatment(table: Table, ctx: Dict[str, Any]) -> None:
    g = ctx["g"]
    records = g("heat_treatment_records", "heat_treatment", default=[])
    if not isinstance(records, list) or not records:
        return
    start = 2
    _ensure_rows(table, start + len(records))
    for i, rec in enumerate(records):
        if not isinstance(rec, dict):
            continue
        r = start + i
        vals = [
            rec.get("element") or rec.get("name") or "",
            rec.get("type") or rec.get("kind") or "",
            rec.get("mode") or rec.get("regime") or rec.get("heat_mode") or "",
            rec.get("temperature") or "",
            rec.get("duration") or "",
            rec.get("cooling") or "",
        ]
        # Если в шаблоне 5 колонок (без отдельного «режима») — склеиваем вид+режим.
        ncols = len(table.rows[r].cells) if r < len(table.rows) else 5
        if ncols <= 5:
            kind = vals[1]
            mode = vals[2]
            merged = kind if not mode else (f"{kind}; режим: {mode}" if kind else mode)
            vals = [vals[0], merged, vals[3], vals[4], vals[5]]
        for c, v in enumerate(vals):
            if c < len(table.rows[r].cells):
                _set(table, r, c, v)


def _fill_strength_tests(table: Table, ctx: Dict[str, Any]) -> None:
    g = ctx["g"]
    records = g("hydraulic_test_history", "strength_tests", default=[])
    if not isinstance(records, list):
        return
    # Отбросить полностью пустые записи
    records = [
        rec
        for rec in records
        if isinstance(rec, dict)
        and any(
            str(rec.get(k) or "").strip()
            for k in ("date", "type", "test_type", "kind", "pressure", "medium", "temperature")
        )
    ]
    if not records:
        _strip_empty_rows(table, 1)
        return
    start = 1
    _ensure_rows(table, start + len(records))
    for i, rec in enumerate(records):
        r = start + i
        vals = [
            _fmt_date_ru(rec.get("date")) or rec.get("date") or "",
            rec.get("type") or rec.get("test_type") or rec.get("kind") or "гидравлическое",
            rec.get("pressure") or "",
            rec.get("medium") or rec.get("test_medium") or "",
            rec.get("temperature") or rec.get("medium_temperature") or "",
            rec.get("note") or rec.get("remark") or "",
        ]
        for c, v in enumerate(vals):
            if c < len(table.rows[r].cells):
                _set(table, r, c, v, nowrap=(c in (0, 2, 4)))
    # Очистить лишние строки-заготовки шаблона
    for r in range(start + len(records), len(table.rows)):
        for c in range(len(table.rows[r].cells)):
            _set(table, r, c, "")


def _fill_previous_inspections(table: Table, ctx: Dict[str, Any]) -> None:
    g = ctx["g"]
    records = g("previous_inspections", "ndt_control_history", default=[])
    if not isinstance(records, list):
        records = []
    records = [
        rec
        for rec in records
        if isinstance(rec, dict)
        and any(
            str(rec.get(k) or "").strip()
            for k in (
                "date",
                "type",
                "kind",
                "control_type",
                "scope",
                "volume",
                "result",
                "results",
                "report_number",
                "organization",
                "executor",
            )
        )
    ]
    # Если есть только legacy-строка — разложить в одну запись
    if not records:
        legacy = g("previous_inspection_result", default="")
        if legacy and legacy != MISSING:
            records = [{"type": "Техническое диагностирование", "result": legacy}]
    start = 1
    if not records:
        for r in range(start, len(table.rows)):
            for c in range(len(table.rows[r].cells)):
                _set(table, r, c, "")
        return
    _ensure_rows(table, start + len(records))
    for i, rec in enumerate(records):
        r = start + i
        vals = [
            _fmt_date_ru(rec.get("date")) or rec.get("date") or "",
            rec.get("type") or rec.get("kind") or rec.get("control_type") or "",
            rec.get("scope") or rec.get("volume") or "",
            rec.get("result") or rec.get("results") or rec.get("report_number") or "",
            rec.get("organization") or rec.get("executor") or "",
        ]
        for c, v in enumerate(vals):
            if c < len(table.rows[r].cells):
                _set(table, r, c, v, nowrap=(c == 0))
    for r in range(start + len(records), len(table.rows)):
        for c in range(len(table.rows[r].cells)):
            _set(table, r, c, "")

def _fill_additional_data(table: Table, ctx: Dict[str, Any]) -> None:
    g = ctx["g"]
    mapping = {
        1: g("vessel_installed", "installation_info", default=""),
        2: g("load_cycles", default=""),
        3: g("supervisory_remarks", default=""),
        4: g("accidents_info", "incidents_info", default=""),
        5: g("repair_info", default=""),
    }
    for row, val in mapping.items():
        if row < len(table.rows) and len(table.rows[row].cells) > 2:
            # Явно показываем «—», если поле не заполнено техником в мобильном
            # приложении, вместо пустой ячейки.
            _set(table, row, 2, val if val not in (None, "") else MISSING)


# Индексы таблиц-подписей протоколов (doc.tables) → коды методов НК, которым
# соответствует данный протокол/приложение. Нужно, чтобы под протоколом ВИК
# расписывался специалист по ВИК, а не случайный человек из общего списка.
SIGNATURE_METHOD_KEYS: Dict[int, Tuple[str, ...]] = {
    18: ("ВИК", "VIK", "ПВК", "PVK"),
    22: ("УЗТ", "UZT"),
    27: ("ТВЕРД", "TVI", "HARD"),
    32: ("УЗК", "UZK"),
    37: ("МПК", "MPK", "МК", "MK", "MT", "MPI", "ПВК", "PVK"),
}


def _specialists_for_methods(
    ctx: Dict[str, Any], method_keys: Optional[Tuple[str, ...]]
) -> List[Dict[str, str]]:
    specs = ctx.get("specialists") or []
    if not method_keys:
        return specs
    matched = [
        s
        for s in specs
        if any(k.upper() in str(s.get("role") or "").upper() for k in method_keys)
    ]
    # Если по конкретному методу никто не привязан явно — не подставляем
    # посторонних специалистов, но если ролей вообще не было указано (одна
    # запись без разбивки по методам) — используем общий список.
    if matched:
        return matched
    if all(not str(s.get("role") or "").strip() for s in specs):
        return specs
    # Нет специалиста именно по этому методу — берём исполнителей обследования,
    # иначе блок подписи в протоколе остаётся пустым.
    return specs


def _fill_signatures(
    table: Table, ctx: Dict[str, Any], method_keys: Optional[Tuple[str, ...]] = None
) -> None:
    specs = _specialists_for_methods(ctx, method_keys)
    # Строки 1,2 — контроль; 4 — заключение
    slots = [1, 2, 4]
    for i, row in enumerate(slots):
        if row >= len(table.rows):
            continue
        if specs:
            # Если специалистов меньше, чем строк подписи (частый случай —
            # один специалист выполнял весь контроль), используем их по кругу
            # вместо того, чтобы оставлять строку с посторонними ФИО из шаблона.
            s = specs[i % len(specs)]
            name = s.get("name") or ""
            cert = s.get("cert") or ""
            label = f"Специалист {name}"
            if cert:
                label += f"  квал. уд. № {cert}"
            else:
                label += "  квал. уд. № ________________"
            _set(table, row, 0, label)
            if len(table.rows[row].cells) > 1:
                _set(table, row, 1, name)
            if len(table.rows[row].cells) > 2 and name:
                _set(table, row, 2, "Ф.И.О.")
        else:
            # Данных нет — не оставляем в отчёте посторонние ФИО из шаблона.
            _set(table, row, 0, "Специалист  квал. уд. № ________________")
            if len(table.rows[row].cells) > 1:
                _set(table, row, 1, "")
            if len(table.rows[row].cells) > 2:
                _set(table, row, 2, "")


def _fill_operational_diagnostics(table: Table, ctx: Dict[str, Any]) -> None:
    g = ctx["g"]
    op = g("operational_diagnostics", "functional_diagnostics", default={})
    if not isinstance(op, dict):
        op = {}
    # Строки 1–5: оценка / примечания в кол. 2 и 3
    defaults = {
        1: (op.get("params_ok") or "Соответствуют", op.get("params_note") or ""),
        2: (op.get("vibration") or "Не выявлена", op.get("vibration_note") or ""),
        3: (op.get("foundation") or "Не выявлена", op.get("foundation_note") or ""),
        4: (op.get("supports") or "Работоспособное", op.get("supports_note") or g("support_state", default="")),
        5: (op.get("kip") or "Работоспособное", op.get("kip_note") or ""),
    }
    for row, (eval_, note) in defaults.items():
        if row < len(table.rows):
            if len(table.rows[row].cells) > 2:
                _set(table, row, 2, eval_)
            if len(table.rows[row].cells) > 3:
                _set(table, row, 3, note)


def _fill_instrument_table(
    table: Table,
    ctx: Dict[str, Any],
    method_keys: Tuple[str, ...],
    defaults: Optional[List[Tuple[str, str]]] = None,
) -> None:
    ve = list(ctx.get("verification_equipment") or [])
    for item in (ctx.get("data") or {}).get("_ndt_instruments") or []:
        if isinstance(item, dict):
            ve.append(item)
    matched: List[Dict[str, Any]] = []
    for eq in ve:
        if not isinstance(eq, dict):
            continue
        et = str(
            eq.get("equipment_type") or eq.get("type") or eq.get("method_code") or ""
        ).upper()
        name = str(eq.get("name") or "")
        blob = f"{et} {name}".upper()
        if any(k.upper() in blob for k in method_keys):
            matched.append(eq)
    if not matched and defaults:
        matched = [{"name": n, "serial_number": s} for n, s in defaults]
    elif not matched and ve:
        # Нет точного совпадения по методу — берём приборы, выбранные в обследовании
        matched = [eq for eq in ve if isinstance(eq, dict)]
    # Очистим строки шаблона и заполним фактическими приборами
    for r in range(1, len(table.rows)):
        for c in range(len(table.rows[r].cells)):
            _set(table, r, c, "")
    for i, eq in enumerate(matched[: max(0, len(table.rows) - 1)]):
        r = i + 1
        if r >= len(table.rows):
            break
        serial = str(eq.get("serial_number") or eq.get("factory_number") or "")
        _set(table, r, 0, f"{i + 1}.")
        if len(table.rows[r].cells) > 1:
            _set(table, r, 1, _instrument_full_name(eq) if eq.get("equipment_type") or eq.get("name") else (eq.get("name") or ""))
        if len(table.rows[r].cells) > 2:
            _set(table, r, 2, serial, nowrap=True)
        if len(table.rows[r].cells) > 3:
            _set(
                table,
                r,
                3,
                eq.get("verification_certificate_number")
                or eq.get("certificate")
                or eq.get("verification_certificate")
                or "",
                nowrap=True,
            )
        if len(table.rows[r].cells) > 4:
            _set(
                table,
                r,
                4,
                _fmt_date_ru(
                    eq.get("next_verification_date")
                    or eq.get("valid_until")
                    or eq.get("verification_until")
                )
                or eq.get("next_verification_date")
                or "",
                nowrap=True,
            )
    _strip_empty_rows(table, 1, ignore_cols=(0,))


def _fill_vik_parameters(table: Table, ctx: Dict[str, Any]) -> None:
    """Таблица «Параметры контроля» ВИК — шероховатость/освещённость по факту,
    вместо статичных примерных значений («Rz 80», «500 Лк») из шаблона."""
    g = ctx["g"]
    roughness = g("vik_roughness", "roughness_rz", default="")
    illumination = g("vik_illumination", "illumination", default="")
    extra_light = g("vik_additional_lighting", default=None)
    if len(table.rows) > 0 and len(table.rows[0].cells) > 1:
        _set(table, 0, 1, roughness if roughness else MISSING)
    if len(table.rows) > 1 and len(table.rows[1].cells) > 1:
        illum_text = f"{illumination} лк" if illumination else MISSING
        if extra_light is True:
            illum_text = f"{illum_text} (с доп. освещением)" if illum_text != MISSING else "с дополнительным освещением"
        elif extra_light is False:
            illum_text = f"{illum_text} (без доп. освещения)" if illum_text != MISSING else MISSING
        _set(table, 1, 1, illum_text)


def _fill_vik_results(table: Table, ctx: Dict[str, Any]) -> None:
    """Таблица результатов ВИК: строки 2/3 — базовые объекты наружного
    осмотра (фундаменты, сварные соединения), строка 4 — доп. объекты
    наружного осмотра, строка 6 — объекты внутреннего осмотра.

    Поддерживает произвольный список объектов контроля
    ``vik_control_objects`` (каждый со своей «зоной» — наружный/внутренний),
    который техник может дополнять в мобильном приложении, а не только
    2 предустановленные категории.
    """
    g = ctx["g"]

    def _row_from(o: Dict[str, Any]) -> Tuple[str, str, str, str]:
        return (
            str(o.get("object") or o.get("location") or o.get("element") or ""),
            str(o.get("scope") or o.get("volume") or "100%"),
            str(o.get("description") or o.get("defects") or o.get("defect") or "Дефектов не обнаружено"),
            str(o.get("assessment") or o.get("quality") or "Годен"),
        )

    # Базовые (всегда присутствующие) категории наружного осмотра
    _set(table, 2, 1, "фундаментов")
    _set(table, 2, 2, "100%")
    _set(table, 2, 3, "Дефектов не обнаружено")
    _set(table, 2, 4, "Годен")
    _set(table, 3, 1, "сварных соединений")
    _set(table, 3, 2, "100%")
    _set(table, 3, 3, "Дефектов не обнаружено")
    _set(table, 3, 4, "Годен")

    # Список объектов контроля ВИК: технику доступно добавление произвольных
    # объектов (не только 2 предустановленные категории) через раздел
    # «Дефекты ВИК» мобильного приложения — каждая запись может нести
    # зону (наружный/внутренний), объём контроля и оценку.
    objects = g("vik_control_objects", "inspection_objects", "visual_defects", "vik_defects", "defects", default=None)
    external_extra: List[Tuple[str, str, str, str]] = []
    internal_extra: List[Tuple[str, str, str, str]] = []

    if isinstance(objects, list) and objects:
        for o in objects:
            if not isinstance(o, dict):
                continue
            zone = str(o.get("zone") or o.get("area") or "external").strip().lower()
            row_vals = _row_from(o)
            if zone.startswith("intern") or "внутр" in zone:
                internal_extra.append(row_vals)
            else:
                external_extra.append(row_vals)

    # Строка 4 (по умолчанию placeholder «….») — доп. объекты наружного осмотра
    ext_row_idx = 4
    if external_extra:
        for c, v in enumerate(external_extra[0], start=1):
            _set(table, ext_row_idx, c, v)
        for extra in external_extra[1:]:
            ext_row_idx = _insert_row_after(table, ext_row_idx, extra)
    else:
        for c in range(len(table.rows[ext_row_idx].cells)):
            _set(table, ext_row_idx, c, "")

    # Строка после «Внутренний осмотр:» — объекты внутреннего осмотра
    int_label_idx = ext_row_idx + 1
    int_row_idx = int_label_idx + 1 if int_label_idx + 1 < len(table.rows) else None
    if int_row_idx is not None:
        if internal_extra:
            for c, v in enumerate(internal_extra[0], start=1):
                _set(table, int_row_idx, c, v)
            for extra in internal_extra[1:]:
                int_row_idx = _insert_row_after(table, int_row_idx, extra)
        else:
            for c in range(len(table.rows[int_row_idx].cells)):
                _set(table, int_row_idx, c, "")

    _strip_empty_rows(table, 2, ignore_cols=(0,))


def _set_row_tc(table: Table, row_idx: int, col_idx: int, text: Any) -> None:
    """Запись в конкретный tc строки (обходит coalescing merged cells в python-docx)."""
    try:
        tr = table.rows[row_idx]._tr
        tcs = tr.tc_lst
        if col_idx < 0 or col_idx >= len(tcs):
            return
        _set_cell(_Cell(tcs[col_idx], table), text)
    except Exception:
        _set(table, row_idx, col_idx, text)


def _fill_uzt_results(table: Table, ctx: Dict[str, Any]) -> None:
    """Таблица результатов УЗТ: элемент | № точки | толщина | … (по 3 точки в ряд)."""
    g = ctx["g"]
    points: List[Dict[str, Any]] = []
    raw = g("thickness_measurements", "thicknessMeasurements", default=[])
    if isinstance(raw, list):
        points.extend(p for p in raw if isinstance(p, dict))

    def _fallbacks() -> List[Dict[str, Any]]:
        extra: List[Dict[str, Any]] = []
        for sch in g("uzt_schemes", default=[]) or []:
            if isinstance(sch, dict):
                for m in sch.get("measurements") or []:
                    if isinstance(m, dict):
                        extra.append(m)
        for m in ctx.get("ndt_methods") or []:
            if not isinstance(m, dict):
                continue
            code = str(m.get("method_code") or m.get("method_name") or "").upper()
            if not any(k in code for k in ("УЗТ", "UZT", "ТОЛЩ")):
                continue
            ad = m.get("additional_data") or {}
            if not isinstance(ad, dict):
                continue
            for p in ad.get("measurement_points") or ad.get("points") or []:
                if isinstance(p, dict):
                    extra.append(
                        {
                            "location": p.get("location") or p.get("element") or p.get("zone") or "",
                            "section_number": p.get("point")
                            or p.get("section_number")
                            or p.get("point_number")
                            or p.get("number")
                            or "",
                            "thickness": p.get("thickness") or p.get("value"),
                        }
                    )
        return extra

    def _to_usable(src: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        out: List[Dict[str, Any]] = []
        seen_keys: set = set()
        for p in src:
            num = (
                p.get("section_number")
                or p.get("point_number")
                or p.get("number")
                or p.get("point")
                or ""
            )
            thick = p.get("thickness")
            if thick is None:
                thick = p.get("value") or p.get("measured_thickness")
            loc = str(
                p.get("element_name")
                or p.get("element")
                or p.get("location")
                or p.get("zone")
                or p.get("name")
                or ""
            ).strip()
            if not str(num).strip() and thick in (None, "") and not loc:
                continue
            sec = str(num).strip() or str(len(out) + 1)
            key = (loc, sec, str(thick if thick not in (None, "") else ""))
            if key in seen_keys:
                continue
            seen_keys.add(key)
            out.append(
                {
                    **p,
                    "location": loc or "Элемент",
                    "section_number": sec,
                    "thickness": thick if thick not in (None, "") else "",
                }
            )
        return out

    # enrich уже подмешивает схемы/НК — fallback только если список пуст
    usable = _to_usable(points)
    if not usable:
        usable = _to_usable(_fallbacks())
    if not usable:
        return

    by_element: Dict[str, List[Dict[str, Any]]] = {}
    for p in usable:
        by_element.setdefault(str(p.get("location") or "Элемент"), []).append(p)

    # Снять vMerge у колонки «элемент» на уровне каждого tc строки
    try:
        for r_i in range(1, len(table.rows)):
            tr = table.rows[r_i]._tr
            if not tr.tc_lst:
                continue
            tcPr = tr.tc_lst[0].get_or_add_tcPr()
            for vm in list(tcPr.findall(qn("w:vMerge"))):
                tcPr.remove(vm)
    except Exception:
        logger.exception("uzt: не удалось снять vMerge")

    rows_needed = 1 + sum(max(1, (len(pts) + 2) // 3) for pts in by_element.values())
    _ensure_rows(table, rows_needed)

    # Очистить строки данных перед заполнением (через tc_lst — без merge-coalesce)
    for r in range(1, len(table.rows)):
        n_cols = len(table.rows[r]._tr.tc_lst)
        for c in range(n_cols):
            _set_row_tc(table, r, c, "")

    row = 1
    for el_name, pts in by_element.items():
        for chunk_start in range(0, len(pts), 3):
            chunk = pts[chunk_start : chunk_start + 3]
            if row >= len(table.rows):
                _ensure_rows(table, row + 1)
            _set_row_tc(table, row, 0, el_name)
            for j, p in enumerate(chunk):
                num = p.get("section_number") or (chunk_start + j + 1)
                thick = p.get("thickness")
                if thick in (None, ""):
                    thick = ""
                base = 1 + j * 2
                n_cols = len(table.rows[row]._tr.tc_lst)
                if base < n_cols:
                    _set_row_tc(table, row, base, num)
                if base + 1 < n_cols:
                    _set_row_tc(table, row, base + 1, thick)
            row += 1

    # Удалить только хвост полностью пустых строк (не трогая заполненные)
    _strip_empty_rows(table, row if row > 1 else 1, ignore_cols=(0,))


def _fill_hardness_matrix(table: Table, ctx: Dict[str, Any]) -> None:
    g = ctx["g"]
    tests = g("hardness_tests", "hardnessTests", default=[])
    if not isinstance(tests, list) or not tests:
        return
    # Строки 3.. — Т1..Т6; колонки 2..6 — значения 1..5
    for i, t in enumerate(tests[:6]):
        if not isinstance(t, dict):
            continue
        r = 3 + i
        if r >= len(table.rows):
            break
        zone = (
            t.get("weld_number")
            or t.get("location")
            or t.get("zone")
            or t.get("section")
            or t.get("element")
            or f"Т{i + 1}"
        )
        _set(table, r, 1, zone)
        vals = [
            t.get("hardness_base_t1") or t.get("hardness_base") or "",
            t.get("hardness_haz_t2") or t.get("hardness_haz") or "",
            t.get("hardness_weld") or "",
            t.get("hardness_haz_t4") or t.get("hardness_haz") or "",
            t.get("hardness_base_t5") or t.get("hardness_base") or "",
        ]
        for c, v in enumerate(vals):
            if 2 + c < len(table.rows[r].cells):
                _set(table, r, 2 + c, v)


def _fill_hardness_list(table: Table, ctx: Dict[str, Any]) -> None:
    g = ctx["g"]
    tests = g("hardness_tests", "hardnessTests", default=[])
    if not isinstance(tests, list) or not tests:
        return
    # Автоподстановка марки стали из элементов корпуса / паспорта
    default_steel = ""
    elements = g("vessel_elements", "elements", default=[])
    if isinstance(elements, list):
        for el in elements:
            if isinstance(el, dict) and (el.get("material") or el.get("steel_grade")):
                default_steel = str(el.get("material") or el.get("steel_grade"))
                break
    if not default_steel:
        default_steel = str(g("shell_material", "material", default="") or "")
    start = 1
    _ensure_rows(table, start + len(tests))
    for i, t in enumerate(tests):
        if not isinstance(t, dict):
            continue
        if not (t.get("steel_grade") or t.get("material") or t.get("grade")):
            t = {**t, "steel_grade": default_steel}
        r = start + i
        _set(
            table,
            r,
            0,
            t.get("element")
            or t.get("element_name")
            or t.get("location")
            or t.get("zone")
            or t.get("section")
            or "",
        )
        _set(
            table,
            r,
            1,
            t.get("point_number")
            or t.get("area_number")
            or t.get("point")
            or t.get("weld_number")
            or (i + 1),
        )
        steel = (
            t.get("steel_grade")
            or t.get("material")
            or t.get("grade")
            or default_steel
            or ""
        )
        hardness_val = (
            t.get("hardness_base")
            or t.get("hardness_weld")
            or t.get("value")
            or ""
        )
        if steel and 4 < len(table.rows[r].cells):
            _set(table, r, 2, hardness_val)
            _set(table, r, 3, steel)
            _set(
                table,
                r,
                4 if len(table.rows[r].cells) > 4 else 3,
                t.get("allowed_hardness_base")
                or t.get("allowed_hardness_weld")
                or t.get("allowed")
                or "",
            )
        else:
            _set(table, r, 2, hardness_val)
            _set(
                table,
                r,
                3,
                t.get("allowed_hardness_base")
                or t.get("allowed_hardness_weld")
                or t.get("allowed")
                or "",
            )


def _fill_uzk_parameters(table: Table, ctx: Dict[str, Any]) -> None:
    """Таблица параметров контроля УЗК (прил. 6, табл. до результатов).

    Колонки: № | тип соединения | толщина | ПЭП (тип, частота, угол) | Sбрак | зарубка.
    """
    g = ctx["g"]
    rows = g("uzk_control_params", default=[])
    if not isinstance(rows, list):
        rows = []
    else:
        rows = [r for r in rows if isinstance(r, dict)]

    if not rows:
        # Fallback напрямую из методов НК в контексте
        for m in ctx.get("ndt_methods") or []:
            if not isinstance(m, dict):
                continue
            code = str(m.get("method_code") or m.get("method_name") or "").upper()
            if not any(k in code for k in ("УЗК", "UZK")):
                continue
            ad = m.get("additional_data") or {}
            if not isinstance(ad, dict):
                ad = {}
            rows.append(
                {
                    "joint_type": ad.get("joint_type") or "",
                    "element_thickness": ad.get("element_thickness") or "",
                    "transducer_type": ad.get("transducer_type") or "",
                    "frequency_mhz": ad.get("frequency_mhz") or ad.get("frequency") or "",
                    "angle_deg": ad.get("angle_deg") or ad.get("angle") or "",
                    "max_equivalent_area": ad.get("max_equivalent_area") or "",
                    "notch_params": ad.get("notch_params")
                    or ad.get("notch")
                    or ad.get("reference_sample")
                    or "",
                }
            )

    usable = [
        r
        for r in rows
        if any(
            str(r.get(k) or "").strip()
            for k in (
                "joint_type",
                "element_thickness",
                "transducer_type",
                "frequency_mhz",
                "angle_deg",
                "max_equivalent_area",
                "notch_params",
                "pep",
            )
        )
    ]
    if not usable:
        return

    _ensure_rows(table, 1 + len(usable))
    for i, p in enumerate(usable):
        r = 1 + i
        pep = str(p.get("pep") or "").strip()
        if not pep:
            parts = []
            if p.get("transducer_type"):
                parts.append(str(p.get("transducer_type")))
            if p.get("frequency_mhz"):
                parts.append(f"{p.get('frequency_mhz')} МГц")
            if p.get("angle_deg"):
                parts.append(f"{p.get('angle_deg')}°")
            pep = ", ".join(parts)
        vals = [
            i + 1,
            p.get("joint_type") or "",
            p.get("element_thickness") or "",
            pep,
            p.get("max_equivalent_area") or "",
            p.get("notch_params") or "",
        ]
        for c, v in enumerate(vals):
            if c < len(table.rows[r].cells):
                _set(table, r, c, v)


def _fill_uzk_results(table: Table, ctx: Dict[str, Any]) -> None:
    g = ctx["g"]
    welds = g("weld_inspections", "uzk_results", "weld_defects", default=[])
    if not isinstance(welds, list) or not welds:
        if len(table.rows) > 1:
            _set(table, 1, 0, "—")
            if len(table.rows[1].cells) > 7:
                _set(table, 1, 7, "Дефектов не обнаружено")
        return
    start = 1
    _ensure_rows(table, start + len(welds))
    for i, w in enumerate(welds):
        if not isinstance(w, dict):
            continue
        r = start + i
        # Мобильное приложение (карта обследования, раздел «УЗК/ПВК») собирает
        # свободный текст дефекта под ключами uzk_defect/defect_description и
        # место контроля под location_on_control_map — сопоставляем их со
        # структурными колонками шаблона (форма/характер и место).
        defect_text = (
            w.get("defect_description")
            or w.get("uzk_defect")
            or w.get("pvk_defect")
            or ""
        )
        location = w.get("location") or w.get("location_on_control_map") or ""
        form_char = (
            w.get("character")
            or w.get("form")
            or w.get("defect_character")
            or w.get("defect_form")
            or ""
        )
        # Не подставлять весь текст дефекта в колонку «характер» (объёмный/плоскостной).
        if form_char and form_char == defect_text and form_char not in (
            "объёмный",
            "объемный",
            "плоскостной",
        ):
            form_char = ""
        conclusion = (
            w.get("conclusion")
            or w.get("assessment")
            or (defect_text if defect_text else "Дефектов не обнаружено")
        )
        vals = [
            w.get("joint") or w.get("weld_number") or w.get("seam") or "",
            w.get("defect_number") or (i + 1 if defect_text else ""),
            w.get("area") or w.get("equivalent_area") or "",
            w.get("depth") or "",
            w.get("length") or "",
            form_char,
            location,
            conclusion,
        ]
        for c, v in enumerate(vals):
            if c < len(table.rows[r].cells):
                _set(table, r, c, v)


def _fill_mpk_results(table: Table, ctx: Dict[str, Any]) -> None:
    g = ctx["g"]
    items = g("mpk_results", "magnetic_results", default=[])
    if not isinstance(items, list) or not items:
        if len(table.rows) > 1:
            _set(table, 1, 0, "Сварные соединения")
            _set(table, 1, 2, "100%")
            _set(table, 1, 3, "Дефектов не обнаружено")
            _set(table, 1, 4, "Годен")
        return
    start = 1
    _ensure_rows(table, start + len(items))
    for i, item in enumerate(items):
        if not isinstance(item, dict):
            continue
        r = start + i
        vals = [
            item.get("object") or item.get("element") or "",
            item.get("zone") or "",
            item.get("scope") or item.get("volume") or "",
            item.get("defects") or item.get("description") or "",
            item.get("assessment") or item.get("quality") or "",
        ]
        for c, v in enumerate(vals):
            if c < len(table.rows[r].cells):
                _set(table, r, c, v)


def _fill_mpk_parameters(table: Table, ctx: Dict[str, Any]) -> None:
    """Таблица параметров МПК: способ контроля и уровень чувствительности."""
    g = ctx["g"]
    method = str(g("mpk_control_method", "magnetization_type", default="") or "")
    sensitivity = str(g("mpk_sensitivity", "sensitivity_level", default="") or "")
    params = g("mpk_control_params", default=[])
    if isinstance(params, list) and params and isinstance(params[0], dict):
        method = method or str(params[0].get("control_method") or "")
        sensitivity = sensitivity or str(
            params[0].get("sensitivity") or params[0].get("field_strength") or ""
        )
    if not method:
        for m in ctx.get("ndt_methods") or []:
            if not isinstance(m, dict):
                continue
            code = str(m.get("method_code") or m.get("method_name") or "").upper()
            if not any(k in code for k in ("МПК", "MPK", "МПД", "МК", "MPI")):
                continue
            ad = m.get("additional_data") or {}
            if isinstance(ad, dict):
                method = str(ad.get("magnetization_type") or ad.get("control_method") or method)
                sensitivity = str(ad.get("sensitivity") or ad.get("field_strength") or sensitivity)
            break
    for r, needles, val in (
        (0, ("способ", "намагнич", "метод"), method),
        (1, ("чувствитель", "уровень", "напряж"), sensitivity),
    ):
        if r >= len(table.rows) or not val:
            continue
        blob = " ".join((c.text or "").lower() for c in table.rows[r].cells)
        col = 1 if len(table.rows[r].cells) > 1 else 0
        if col == 0 and any(n in blob for n in needles):
            continue
        _set(table, r, col, val)


def _fill_hardness_steel_heading(doc: Document, ctx: Dict[str, Any]) -> None:
    """Подставить марку стали в заголовок «для стали …….»."""
    g = ctx["g"]
    steel = str(g("shell_material", "material", default="") or "")
    if not steel:
        elements = g("vessel_elements", "elements", default=[])
        if isinstance(elements, list):
            for el in elements:
                if isinstance(el, dict) and (el.get("material") or el.get("steel_grade")):
                    steel = str(el.get("material") or el.get("steel_grade"))
                    break
    if not steel:
        tests = g("hardness_tests", default=[])
        if isinstance(tests, list):
            for t in tests:
                if isinstance(t, dict) and (t.get("steel_grade") or t.get("material")):
                    steel = str(t.get("steel_grade") or t.get("material"))
                    break
    if not steel:
        return
    pat = re.compile(r"(для стали\s*)[.…_—\-–\s]*", re.IGNORECASE)

    def _sub(text: str) -> str:
        if "для стали" not in (text or "").lower():
            return text
        return pat.sub(rf"\g<1>{steel} ", text)

    for p in _iter_all_paragraphs(doc):
        t = p.text or ""
        nt = _sub(t)
        if nt != t:
            _set_paragraph_text(p, nt)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text or ""
                nt = _sub(t)
                if nt != t:
                    _set_cell(cell, nt)


def _fill_paragraph_blanks(doc: Document, ctx: Dict[str, Any]) -> None:
    """Подставить номера протоколов, даты и выводы в абзацы с подчёркиваниями."""
    date_ru = ctx["date_ru"]
    serial = ctx["serial"]
    reg_no = ctx["reg_no"]
    inv_no = ctx["inv_no"]
    device = ctx["device_name"]
    conclusion = ctx.get("conclusion_suitable") or "соответствует"
    doc_concl = _strip_parens(str(ctx.get("conclusion_doc") or "в полном объёме"))
    doc_verdict = _doc_verdict_word(ctx)
    op_ok = ctx.get("operational_ok") or "соответствует"
    g = ctx["g"]
    protocol_no = str(g("protocol_number", "report_number", default="") or "")
    tech_card = str(g("tech_card_number", "technological_card", default="") or "")

    uzk_concl = _ndt_result_summary(
        ctx,
        data_keys=("weld_inspections", "uzk_results", "weld_defects"),
        custom_key="uzk_conclusion_text",
        ok_text=(
            "По результатам обследования сварных соединений сосуда, "
            "недопустимых дефектов не обнаружено, объект контроля соответствует требованиям НТД."
        ),
        defect_text=(
            "По результатам ультразвукового контроля сварных соединений выявлены недопустимые дефекты. "
            "Требуется ремонт сосуда / снижение максимально допустимого рабочего давления "
            "(см. таблицу результатов)."
        ),
    )
    mpk_concl = _ndt_result_summary(
        ctx,
        data_keys=("mpk_results", "magnetic_results", "weld_inspections"),
        custom_key="mpk_conclusion_text",
        ok_text=(
            "По результатам магнитопорошкового контроля дефектов в сварных "
            "соединениях сосуда не обнаружено, объект контроля соответствует "
            "требованиям нормативно-технической документации."
        ),
        defect_text=(
            "По результатам магнитопорошкового контроля выявлены дефекты. "
            "Требуется ремонт сосуда / снижение максимально допустимого рабочего давления "
            "(см. таблицу результатов)."
        ),
    )
    vik_concl = str(g("vik_conclusion_text", default="") or "").strip()
    if not vik_concl:
        vik_concl = _ndt_result_summary(
            ctx,
            data_keys=("visual_defects", "vik_defects", "vik_control_objects"),
            custom_key="vik_conclusion_text",
            ok_text=(
                "По результатам визуального и измерительного контроля основного "
                "металла и сварных соединений сосуда, недопустимых дефектов не обнаружено, "
                "что удовлетворяет требованиям нормативно-технической документации "
                f"{g('vik_ntd', default='СТО 9701105632-003-2021')}."
            ),
            defect_text=(
                "По результатам визуального и измерительного контроля выявлены дефекты, "
                "препятствующие дальнейшей безопасной эксплуатации без ремонта."
            ),
        )

    # Важно: основной отчёт в SDT — doc.paragraphs его не видит
    for p in _iter_all_paragraphs(doc):
        text = p.text
        if not text:
            continue
        new_text = text
        stripped = text.strip()

        # «№ _____ от _____ г.» — дату подставляем всегда; номер — если есть
        if stripped.startswith("№") and "от" in stripped and "г." in stripped:
            no_part = protocol_no if protocol_no else "_________"
            new_text = f"№ {no_part} от {date_ru} г."
        elif "При анализе технической документации установлено" in text:
            concl = doc_concl.strip() or "в полном объёме"
            new_text = f"При анализе технической документации установлено, {concl}"
        elif "ВЫВОД:" in text or "Представленная техническая документация" in text:
            new_text = (
                f"ВЫВОД: Представленная техническая документация на сосуд, "
                f"работающий под давлением – {device} зав. № {serial}, рег. № {reg_no}, "
                f"инв. № {inv_no} {doc_verdict} требованиям действующей "
                f"нормативно-технической документации."
            )
        elif "функциональной (оперативной) диагностики установлено" in text:
            op = str(op_ok).strip() or "соответствует"
            if "соответств" in op.lower():
                new_text = (
                    f"В результате функциональной (оперативной) диагностики установлено, "
                    f"что сосуд {op} паспортным характеристикам и требованиям действующей НТД."
                )
            else:
                new_text = (
                    f"В результате функциональной (оперативной) диагностики установлено, "
                    f"что сосуд {op} паспортным характеристикам. "
                    f"Сосуд соответствует требованиям действующей НТД."
                )
        elif "Технологическая карта №" in text:
            card_val = tech_card if tech_card else "—"
            if "_" in text:
                new_text = _replace_underscores(text, [card_val])
            else:
                new_text = re.sub(
                    r"(Технологическая карта №)\s*$",
                    rf"\1 {card_val}",
                    text,
                )
        elif (
            "фактическое значение параметров" in text.lower()
            or (
                "состояние эксплуатации" in text.lower()
                and ("зав" in text.lower() or "____" in text)
            )
        ):
            needs_repair = _conclusion_needs_repair(ctx)
            satisfy = (
                "не удовлетворяют требованиям нормативных документов без выполнения ремонта"
                if needs_repair
                else "удовлетворяют требованиям нормативных документов"
            )
            new_text = (
                f"Фактическое значение параметров, определяющих состояние эксплуатации сосуда "
                f"работающего под давлением – {device}, зав.№ {serial}, рег.№ {reg_no}, "
                f"инв.№ {inv_no}, {satisfy}."
            )
        elif "техническое состояние объекта диагностирования" in text.lower() and (
            "____" in text or len(stripped) < 80
        ):
            tech_state = ctx.get("tech_state") or "работоспособное"
            new_text = (
                f"Техническое состояние объекта диагностирования: {tech_state}. {conclusion}."
            )
        elif "недопустимых дефектов не обнаружено" in text.lower() or (
            "объект контроля соответствует" in text.lower() and "сварн" in text.lower()
        ) or ("магнитопорошкового контроля" in text.lower() and "заключен" not in stripped.lower()[:20]):
            low = text.lower()
            if "магнитопорошков" in low:
                new_text = mpk_concl
            elif "ультразвуков" in low or "сварных соединений сосуда" in low:
                new_text = uzk_concl
            else:
                new_text = vik_concl
        else:
            continue

        if new_text != text:
            _set_paragraph_text(p, new_text)


def _set_paragraph_text(paragraph: Paragraph, text: str, *, pt: float = 12.0) -> None:
    """Заменить текст абзаца, сохранив единый шрифт отчёта (12 pt)."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        try:
            _set_run_font(paragraph.runs[0], pt=pt)
        except Exception:
            pass
    else:
        paragraph.text = text
        if paragraph.runs:
            try:
                _set_run_font(paragraph.runs[0], pt=pt)
            except Exception:
                pass


def _set_paragraph_font_size(paragraph: Paragraph, pt: float = 12.0) -> None:
    """Выставить размер шрифта всем runs абзаца."""
    for run in paragraph.runs:
        _set_run_font(run, pt=pt)


def _fill_appendix_8_calculation(doc: Document, ctx: Dict[str, Any]) -> None:
    """Приложение № 8 — расчёт на прочность из calculation_data."""
    g = ctx["g"]
    calc = g("calculation_data", "calculationData", default=None)
    if calc in (None, MISSING, ""):
        calc = {}
    if not isinstance(calc, dict):
        calc = {"description": str(calc)}

    anchor = find_paragraph_containing(doc, "ПРИЛОЖЕНИЕ № 8")
    if anchor is None:
        anchor = find_paragraph_containing(doc, "Расчет на прочность")
    if anchor is None:
        return

    lines: List[str] = []
    residual = calc.get("residual_life_years") or calc.get("residual_life") or g(
        "residual_life_text", "residual_life_years", default=""
    )
    if residual and residual != MISSING:
        lines.append(f"Остаточный ресурс: {residual} лет.")
    description = calc.get("description") or calc.get("text") or calc.get("summary") or ""
    if description:
        lines.append(str(description))
    for key in (
        "allowable_stress",
        "design_pressure",
        "calc_thickness",
        "min_thickness",
        "corrosion_rate",
        "method",
        "conclusion",
    ):
        val = calc.get(key)
        if val not in (None, ""):
            labels = {
                "allowable_stress": "Допускаемое напряжение",
                "design_pressure": "Расчётное давление",
                "calc_thickness": "Расчётная толщина",
                "min_thickness": "Минимальная толщина",
                "corrosion_rate": "Скорость коррозии",
                "method": "Методика расчёта",
                "conclusion": "Заключение",
            }
            lines.append(f"{labels.get(key, key)}: {val}")

    # Таблица результатов, если есть rows
    rows = calc.get("rows") or calc.get("results") or []
    last = anchor
    if not lines and not rows:
        last = insert_paragraph_after(
            last,
            "Расчёт на прочность выполнен по результатам УЗТ. "
            "Значения толщин стенок удовлетворяют требованиям прочности.",
        )
        _set_paragraph_font_size(last, 12.0)
        _normalize_appendix_font(doc, "ПРИЛОЖЕНИЕ № 8", stop_markers=("ПРИЛОЖЕНИЕ № 9", "ПРИЛОЖЕНИЕ №9"))
        return

    for line in lines:
        last = insert_paragraph_after(last, line)
        _set_paragraph_font_size(last, 12.0)

    if isinstance(rows, list) and rows:
        last = insert_paragraph_after(last, "Результаты расчёта:")
        _set_paragraph_font_size(last, 12.0)
        for row in rows:
            if isinstance(row, dict):
                parts = [f"{k}: {v}" for k, v in row.items() if v not in (None, "")]
                last = insert_paragraph_after(last, "; ".join(parts))
            else:
                last = insert_paragraph_after(last, str(row))
            _set_paragraph_font_size(last, 12.0)

    _normalize_appendix_font(doc, "ПРИЛОЖЕНИЕ № 8", stop_markers=("ПРИЛОЖЕНИЕ № 9", "ПРИЛОЖЕНИЕ №9"))


def _normalize_appendix_font(
    doc: Document,
    start_marker: str,
    stop_markers: Sequence[str] = (),
    pt: float = 12.0,
) -> None:
    """Все абзацы приложения — 12 pt (как в остальном отчёте)."""
    started = False
    for p in _iter_all_paragraphs(doc):
        t = (p.text or "").strip()
        if not started:
            if start_marker.lower() in t.lower() or (
                "расчет на прочность" in t.lower() and "приложение" in t.lower()
            ):
                started = True
            elif t.upper().startswith("ПРИЛОЖЕНИЕ") and "8" in t:
                started = True
            else:
                continue
        else:
            up = t.upper()
            if any(m.upper() in up for m in stop_markers) or (
                up.startswith("ПРИЛОЖЕНИЕ") and "8" not in up[:20]
            ):
                break
        _set_paragraph_font_size(p, pt)

def _fill_appendix_9_hydraulic_act(doc: Document, ctx: Dict[str, Any]) -> None:
    """Приложение № 9 — копия акта гидравлического испытания (скан)."""
    data = ctx.get("data") or {}
    attachments = ctx.get("attachments") or {}
    find_image = ctx.get("find_image")
    paths = collect_hydraulic_act_paths(data, attachments)

    # Fallback: документ №9 журнала / или текст из истории испытаний
    if not paths:
        for key in ("9", "15", "17"):
            if key in attachments:
                paths.append(attachments[key])

    anchor = find_paragraph_containing(doc, "ПРИЛОЖЕНИЕ № 9")
    if anchor is None:
        anchor = find_paragraph_containing(doc, "гидравлического испытания")
    if anchor is None:
        return

    last = anchor
    inserted = 0
    for path in paths:
        resolved = resolve_image_path(path, find_image)
        if not resolved:
            continue
        if is_image_file(resolved):
            pic = add_picture_after_paragraph(
                last,
                resolved,
                width_inches=5.5,
                caption="Копия акта гидравлического испытания",
            )
            if pic is not None:
                last = pic
                inserted += 1
        else:
            last = insert_paragraph_after(
                last,
                f"Приложенный документ: {Path(resolved).name}",
            )
            inserted += 1

    if inserted == 0:
        g = ctx["g"]
        hist = g("hydraulic_test_history", "strength_tests", default=[])
        if isinstance(hist, list) and hist:
            last_rec = hist[-1] if isinstance(hist[-1], dict) else {}
            date = _fmt_date_ru(last_rec.get("date")) or last_rec.get("date") or "—"
            pressure = last_rec.get("pressure") or g("test_pressure", default="—")
            insert_paragraph_after(
                last,
                f"Акт гидравлического испытания от {date}, пробное давление {pressure}. "
                f"Скан акта не приложен к обследованию.",
            )
        else:
            insert_paragraph_after(
                last,
                "Скан акта гидравлического испытания не приложен к материалам обследования.",
            )


def _insert_schemes_and_photos(doc: Document, ctx: Dict[str, Any]) -> None:
    """Вставить схемы контроля (слои ВИК/УЗТ/ТК/УЗК/МПК) и фото измерений."""
    data = ctx.get("data") or {}
    attachments = ctx.get("attachments") or {}
    find_image = ctx.get("find_image")
    kind = str((ctx.get("g")("equipment_kind", default="") if callable(ctx.get("g")) else "") or "vessel")
    photos = collect_photo_paths(data, attachments)
    tmp_files: List[str] = []
    n_schemes = insert_ndt_layer_schemes(
        doc, data, attachments, find_image, kind=kind, tmp_files=tmp_files
    )
    logger.info("Вставлено схем: %s", n_schemes)

    uzt_photos = [p for p in photos if "УЗТ" in (p.get("label") or "")]
    other_photos = [p for p in photos if p not in uzt_photos]
    n_uzt = insert_media_block(
        doc, "Результаты контроля", uzt_photos, find_image=find_image, width_inches=4.5, max_items=15
    )
    vik_photos = [
        p
        for p in other_photos
        if "дефект" in (p.get("label") or "").lower() or "ВИК" in (p.get("label") or "")
    ]
    n_vik = insert_media_block(
        doc,
        "Результаты визуального и измерительного контроля",
        vik_photos,
        find_image=find_image,
        width_inches=4.5,
        max_items=10,
    )
    rest_photos = [p for p in other_photos if p not in vik_photos]
    n_rest = 0
    if rest_photos:
        n_rest = insert_media_block(
            doc, "ПРИЛОЖЕНИЕ № 9", rest_photos, find_image=find_image, width_inches=4.5, max_items=10
        )
    logger.info("Вставлено фото: УЗТ=%s ВИК=%s прочие=%s", n_uzt, n_vik, n_rest)
    for pth in tmp_files:
        try:
            Path(pth).unlink(missing_ok=True)
        except Exception:
            pass


def insert_ndt_layer_schemes(
    doc: Document,
    data: Dict[str, Any],
    attachments: Dict[str, str],
    find_image: Any,
    *,
    kind: str = "vessel",
    tmp_files: Optional[List[str]] = None,
) -> int:
    """Сгенерировать и вставить 5 слоёв карт контроля в приложения 3–7 (и generic)."""
    tmp_files = tmp_files if tmp_files is not None else []
    generated: List[Dict[str, Any]] = []
    try:
        generated = render_all_layer_pngs(data)
    except Exception:
        logger.exception("Не удалось сгенерировать слои схем НК")

    anchors = find_all_paragraphs_containing(doc, "Схема контроля")
    extra = find_all_paragraphs_containing(doc, "Схема оборудования")
    extra += find_all_paragraphs_containing(doc, "Карта контроля")
    seen_p = set()
    ordered: List[Paragraph] = []
    for p in anchors + extra:
        key = id(p._p)
        if key in seen_p:
            continue
        seen_p.add(key)
        ordered.append(p)

    n_schemes = 0
    if generated:
        for i, item in enumerate(generated):
            title = item.get("title") or layer_title(item.get("layer") or LAYER_ORDER[i], equipment_kind=kind)
            png = item.get("png")
            if not png:
                continue
            path = png_to_tempfile(png)
            tmp_files.append(path)
            extras = item.get("extra") or []
            extra_paths: List[Tuple[str, str]] = []
            for cap, extra_png in extras:
                ep = png_to_tempfile(extra_png)
                tmp_files.append(ep)
                extra_paths.append((cap, ep))
            if i < len(ordered):
                pic = _insert_scheme_landscape_block(
                    ordered[i],
                    title,
                    path,
                    extra_paths,
                    close_prev_as_portrait=(i == 0),
                )
                if pic is not None:
                    n_schemes += 1 + len(extra_paths)
            else:
                w = _fit_image_width_inches(path, _SCHEME_PAGE_MAX_W_IN, _SCHEME_PAGE_MAX_H_IN)
                n_schemes += insert_media_block(
                    doc,
                    title,
                    [{"path": path, "label": title}],
                    find_image=find_image,
                    width_inches=w,
                    max_items=1,
                )
        return n_schemes

    # Fallback: картинки из обследования / конструктора
    schemes = collect_scheme_paths(data, attachments)
    if not schemes:
        return 0
    for i, s in enumerate(schemes[:5]):
        if i < len(ordered):
            title = layer_title(LAYER_ORDER[i] if i < len(LAYER_ORDER) else "vik", equipment_kind=kind)
            path = resolve_image_path(s.get("path"), find_image)
            if path and is_image_file(path):
                pic = _insert_scheme_landscape_block(
                    ordered[i],
                    title,
                    path,
                    [],
                    close_prev_as_portrait=(i == 0),
                )
                if pic is not None:
                    n_schemes += 1
        else:
            path = resolve_image_path(s.get("path"), find_image)
            w = (
                _fit_image_width_inches(path, _SCHEME_PAGE_MAX_W_IN, _SCHEME_PAGE_MAX_H_IN)
                if path
                else 7.2
            )
            n_schemes += insert_media_block(
                doc, "Схема контроля", [s], find_image=find_image, width_inches=w, max_items=1
            )
    return n_schemes
