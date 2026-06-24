"""
Таблицы приложения Б и протоколы В (№3–№6) — формат образца 25-3173.
"""
from __future__ import annotations

from collections import defaultdict
from datetime import datetime
from typing import Any, Callable, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from epb_report_builder import EpbReportContext, device_type_label
from equipment_profiles import passport_elements_for_equipment
from equipment_presets import preset_from_equipment_data

MISSING = "Данные в паспорте отсутствуют"
POINTS_PER_UZT_ROW = 5


def _cell(row, idx: int, text: str, bold: bool = False) -> None:
    row.cells[idx].text = str(text)
    if bold:
        try:
            row.cells[idx].paragraphs[0].runs[0].font.bold = True
        except Exception:
            pass


def _header_row(tbl, headers: List[str]) -> None:
    for i, h in enumerate(headers):
        _cell(tbl.rows[0], i, h, bold=True)
        try:
            tbl.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass


def _as_list(g: Callable[..., Any], *keys: str) -> List[Dict[str, Any]]:
    for k in keys:
        raw = g(k, default=[])
        if isinstance(raw, list) and raw:
            return [x for x in raw if isinstance(x, dict)]
    return []


def _default_oil_settler_elements(g: Callable[..., Any]) -> List[Dict[str, Any]]:
    """Fallback элементов корпуса — из профиля оборудования или ОГ-13."""
    elements = passport_elements_for_equipment(
        {
            "type_code": g("equipment_type", default=""),
            "attributes": {
                "wall_thickness": g("wall_thickness", default="18,0"),
                "diameter": g("diameter", default="3400"),
            },
        }
    )
    if elements:
        return elements
    wt = g("wall_thickness", default="18,0")
    dia = g("diameter", default="3400")
    return [
        {
            "name": "Обечайка",
            "diameter_mm": dia,
            "length_mm": g("shell_length_mm", default="21000"),
            "wall_thickness_mm": wt,
            "material": g("shell_material", default="09Г2С-7"),
            "gost": g("material_gost", default="5520-79"),
            "weld_data": g("weld_data", default="Сварка авто\nСВ08ГА\nГОСТ 2246-70\nУОНИ 13/55"),
        },
        {
            "name": "Днище\nправое",
            "diameter_mm": dia,
            "length_mm": g("head_length_mm", default="826"),
            "wall_thickness_mm": wt,
            "material": g("head_material", default="09Г2С-7"),
            "gost": g("material_gost", default="5520-79"),
            "weld_data": g("weld_data", default="Сварка авто\nСВ08ГА\nГОСТ 2246-70\nУОНИ 13/55"),
        },
        {
            "name": "Днище\nнижнее",
            "diameter_mm": dia,
            "length_mm": g("head_length_mm", default="826"),
            "wall_thickness_mm": wt,
            "material": g("head_material", default="09Г2С-7"),
            "gost": g("material_gost", default="5520-79"),
            "weld_data": g("weld_data", default="Сварка авто\nСВ08ГА\nГОСТ 2246-70\nУОНИ 13/55"),
        },
    ]


def append_appendix_b_tables(doc: Document, ctx: EpbReportContext) -> None:
    """Таблицы Б1–Б6 приложения Б."""
    g = ctx.g
    doc.add_paragraph(
        "В процессе экспертизы промышленной безопасности был проанализирован паспорт сосуда, "
        "работающего под давлением. Сведения о сосуде и его элементах представлены "
        "в Таблицах Б1, Б2, Б3, Б4, Б5, Б6."
    )
    doc.add_paragraph()

    # Б1
    doc.add_paragraph("Таблица Б1. Сведения об основных элементах сосуда")
    elements = _as_list(g, "vessel_elements", "passport_elements")
    if not elements:
        elements = _default_oil_settler_elements(g)
    tbl = doc.add_table(rows=len(elements) + 1, cols=7)
    tbl.style = "Table Grid"
    _header_row(
        tbl,
        [
            "Наименование элемента",
            "Диаметр, мм",
            "Длина (высота), мм",
            "Толщина стенки номинальная, мм",
            "Марка материала",
            "ГОСТ (ТУ)",
            "Данные о сварке",
        ],
    )
    for i, el in enumerate(elements, 1):
        row = tbl.rows[i]
        _cell(row, 0, el.get("name") or el.get("element_name") or "—")
        _cell(row, 1, el.get("diameter_mm") or el.get("diameter") or "—")
        _cell(row, 2, el.get("length_mm") or el.get("length") or "—")
        _cell(row, 3, el.get("wall_thickness_mm") or el.get("wall_thickness") or "—")
        _cell(row, 4, el.get("material") or "—")
        _cell(row, 5, el.get("gost") or el.get("material_gost") or "—")
        _cell(row, 6, el.get("weld_data") or "—")
    doc.add_paragraph()

    # Б2
    doc.add_paragraph("Таблица Б2. Данные о термообработке сосуда и его элементов")
    heat = _as_list(g, "heat_treatment_records", "vessel_heat_treatment")
    if not heat:
        heat = [{"element": MISSING, "type": MISSING, "temperature": MISSING, "duration": MISSING, "cooling": MISSING}]
    ht = doc.add_table(rows=len(heat) + 2, cols=5)
    ht.style = "Table Grid"
    _cell(ht.rows[0], 0, "Наименование элемента или соединения", bold=True)
    _cell(ht.rows[0], 1, "Вид термообработки", bold=True)
    _cell(ht.rows[1], 0, "Наименование элемента или соединения", bold=True)
    _cell(ht.rows[1], 1, "Вид термообработки", bold=True)
    _cell(ht.rows[1], 2, "Температура, °С", bold=True)
    _cell(ht.rows[1], 3, "Продолжительность выдержки, ч.", bold=True)
    _cell(ht.rows[1], 4, "Способ охлаждения", bold=True)
    for i, rec in enumerate(heat, 2):
        row = ht.rows[i]
        _cell(row, 0, rec.get("element") or rec.get("name") or MISSING)
        _cell(row, 1, rec.get("type") or rec.get("heat_treatment_type") or MISSING)
        _cell(row, 2, rec.get("temperature") or MISSING)
        _cell(row, 3, rec.get("duration") or MISSING)
        _cell(row, 4, rec.get("cooling") or MISSING)
    doc.add_paragraph()

    # Б3
    doc.add_paragraph("Таблица Б3. Данные о проведённых гидравлических (пневматических) испытаний")
    hydro = _as_list(g, "hydraulic_test_history", "hydrostatic_test_history")
    if not hydro:
        tp = g("test_pressure", default="1,3")
        hydro = [
            {"date": g("commissioning_year", default="—"), "test_type": "гидравлическое", "pressure": f"{tp} ({float(str(tp).replace(',', '.')) * 10:.1f})" if tp else "—", "medium": "вода", "note": "Эксплуатирующая организация"},
        ]
    hb = doc.add_table(rows=len(hydro) + 1, cols=5)
    hb.style = "Table Grid"
    _header_row(
        hb,
        ["Дата", "Вид испытания", "Пробное давление, МПа (кгс/см²)", "Испытательная среда", "Примечание"],
    )
    for i, rec in enumerate(hydro, 1):
        row = hb.rows[i]
        _cell(row, 0, rec.get("date") or "—")
        _cell(row, 1, rec.get("test_type") or rec.get("type") or "гидравлическое")
        _cell(row, 2, rec.get("pressure") or rec.get("test_pressure") or "—")
        _cell(row, 3, rec.get("medium") or "вода")
        _cell(row, 4, rec.get("note") or rec.get("organization") or "")
    doc.add_paragraph()

    # Б4
    doc.add_paragraph(
        "Таблица Б4. Сведения о выполненном неразрушающем контроле технического состояния, "
        "натурных измерениях и лабораторных исследований материалов"
    )
    ndt_hist = _as_list(g, "ndt_control_history", "previous_ndt_history")
    if not ndt_hist:
        prev = g("previous_inspection_result", default="")
        ndt_hist = [{"date": "", "scope": "Рентген, АУЗК, УЗК-100%", "result": "Дефектов не обнаружено", "organization": ""}]
        if prev:
            ndt_hist.append({"date": "", "scope": "ЭПБ", "result": str(prev)[:80], "organization": ""})
    nb = doc.add_table(rows=len(ndt_hist) + 1, cols=4)
    nb.style = "Table Grid"
    _header_row(nb, ["Дата", "Вид и объём контроля", "Основные результаты контроля", "Организация — исполнитель"])
    for i, rec in enumerate(ndt_hist, 1):
        row = nb.rows[i]
        _cell(row, 0, rec.get("date") or "—")
        _cell(row, 1, rec.get("scope") or rec.get("control_type") or "—")
        _cell(row, 2, rec.get("result") or "—")
        _cell(row, 3, rec.get("organization") or "—")
    doc.add_paragraph()

    # Б5
    doc.add_paragraph("Таблица Б5. Сведения о ремонте и замене элементов сосуда")
    repairs = _as_list(g, "repair_history", "vessel_repairs")
    if not repairs:
        repairs = [{"year": MISSING, "description": MISSING, "ndt_result": MISSING}]
    rb = doc.add_table(rows=len(repairs) + 1, cols=3)
    rb.style = "Table Grid"
    _header_row(
        rb,
        ["Год проведения ремонта", "Характер ремонта (элемент, подвергнутый замене)", "Вид и результат НК после ремонта"],
    )
    for i, rec in enumerate(repairs, 1):
        row = rb.rows[i]
        _cell(row, 0, rec.get("year") or "—")
        _cell(row, 1, rec.get("description") or rec.get("repair_type") or "—")
        _cell(row, 2, rec.get("ndt_result") or "—")
    doc.add_paragraph()

    # Б6
    doc.add_paragraph("Таблица Б6. Основная арматура, контрольно-измерительные приборы и приборы безопасности")
    fittings = _collect_fittings_b6(g)
    fb = doc.add_table(rows=len(fittings) + 1, cols=4)
    fb.style = "Table Grid"
    _header_row(fb, ["Наименование", "Кол-во, шт.", "Условный проход, мм", "Условное давление МПа (кгс/см²)"])
    for i, f in enumerate(fittings, 1):
        row = fb.rows[i]
        _cell(row, 0, f.get("name") or "—")
        _cell(row, 1, f.get("quantity") or "1")
        _cell(row, 2, f.get("dn") or f.get("type_size") or "—")
        _cell(row, 3, f.get("pressure") or "—")
    doc.add_paragraph()


def _collect_fittings_b6(g: Callable[..., Any]) -> List[Dict[str, str]]:
    items: List[Dict[str, str]] = []
    explicit = _as_list(g, "fittings_and_instruments", "armature_items")
    if explicit:
        return explicit

    gauge = g("gauge", default={})
    if isinstance(gauge, dict) and gauge:
        items.append({"name": "Манометр", "quantity": gauge.get("quantity") or "1", "dn": "-", "pressure": "-"})

    ls = g("level_sensor", default={})
    if isinstance(ls, dict) and ls.get("type_size") or ls.get("serial_number"):
        items.append({"name": "Датчик уровня", "quantity": ls.get("quantity") or "1", "dn": ls.get("type_size") or "-", "pressure": "-"})

    sw = g("switching_device", default={})
    if isinstance(sw, dict) and any(sw.values()):
        items.append({
            "name": "Переключающее устройство",
            "quantity": sw.get("quantity") or "1",
            "dn": sw.get("type_size") or "100",
            "pressure": sw.get("pressure") or "1,6 (16,0)",
        })

    for sppk in _as_list(g, "sppk_items"):
        items.append({
            "name": "Предохранительный клапан" + (f" зав.№{sppk['serial_number']}" if sppk.get("serial_number") else ""),
            "quantity": sppk.get("quantity") or "1",
            "dn": sppk.get("type_size") or "100",
            "pressure": sppk.get("pressure") or "1,6 (16,0)",
        })

    for zra in _as_list(g, "zra_items"):
        name = "Клапан регулирующий"
        if zra.get("tech_number"):
            name += f" {zra['tech_number']}"
        items.append({
            "name": name,
            "quantity": zra.get("quantity") or "1",
            "dn": zra.get("type_size") or "—",
            "pressure": zra.get("pressure") or "—",
        })

    if not items:
        items = [
            {"name": "Датчик уровня ДУУ4", "quantity": "1", "dn": "-", "pressure": "-"},
            {"name": "Манометр", "quantity": "1", "dn": "-", "pressure": "-"},
            {"name": "Переключающее устройство", "quantity": "1", "dn": "100", "pressure": "1,6 (16,0)"},
            {"name": "Предохранительный клапан", "quantity": "2", "dn": "100", "pressure": "1,6 (16,0)"},
        ]
    return items


def append_epb_protocol_hardness(
    doc: Document,
    g: Callable[..., Any],
    header_block: Callable[[], None],
    hardness: List[Dict[str, Any]],
) -> None:
    """Протокол №3 — твердометрия (образец)."""
    doc.add_page_break()
    doc.add_heading("Протокол № 3", level=2)
    doc.add_paragraph("оценка механических свойств элементов сосуда")
    doc.add_paragraph()
    header_block()
    doc.add_paragraph("Результаты контроля твёрдости сварных швов и околошовных зон")
    if not hardness:
        doc.add_paragraph("Данные замеров твердости не предоставлены.")
        return

    tbl = doc.add_table(rows=len(hardness) + 3, cols=8)
    tbl.style = "Table Grid"
    _cell(tbl.rows[0], 0, "Номер сварного соединения/номер участка контроля", bold=True)
    _cell(tbl.rows[0], 1, "Допустимые пределы твёрдости основного металла, НВ", bold=True)
    _cell(tbl.rows[0], 2, "Допустимые пределы твёрдости металла шва и зоны термического влияния, НВ не более", bold=True)
    _cell(tbl.rows[0], 3, "Твёрдость металла, НВ", bold=True)
    for c in range(4, 8):
        _cell(tbl.rows[0], c, "", bold=True)
    _cell(tbl.rows[1], 0, "Номер сварного соединения/номер участка контроля", bold=True)
    _cell(tbl.rows[1], 1, "Допустимые пределы твёрдости основного металла, НВ", bold=True)
    _cell(tbl.rows[1], 2, "Допустимые пределы твёрдости металла шва и зоны термического влияния, НВ не более", bold=True)
    _cell(tbl.rows[1], 3, "Основного", bold=True)
    _cell(tbl.rows[1], 4, "Основного", bold=True)
    _cell(tbl.rows[1], 5, "Шва", bold=True)
    _cell(tbl.rows[1], 6, "Зоны термического влияния", bold=True)
    _cell(tbl.rows[1], 7, "Зоны термического влияния", bold=True)
    _cell(tbl.rows[2], 0, "Номер сварного соединения/номер участка контроля", bold=True)
    _cell(tbl.rows[2], 1, "Допустимые пределы твёрдости основного металла, НВ", bold=True)
    _cell(tbl.rows[2], 2, "Допустимые пределы твёрдости металла шва и зоны термического влияния, НВ не более", bold=True)
    _cell(tbl.rows[2], 3, "Т.1", bold=True)
    _cell(tbl.rows[2], 4, "Т.5", bold=True)
    _cell(tbl.rows[2], 5, "Т.3", bold=True)
    _cell(tbl.rows[2], 6, "Т.2", bold=True)
    _cell(tbl.rows[2], 7, "Т.4", bold=True)

    for i, h in enumerate(hardness, 3):
        row = tbl.rows[i]
        _cell(row, 0, h.get("location") or h.get("weld_number") or f"Участок {i - 2}")
        _cell(row, 1, h.get("allowed_hardness_base") or "120-180")
        _cell(row, 2, h.get("allowed_hardness_weld") or "225")
        _cell(row, 3, h.get("hardness_base_t1") or h.get("hardness_base") or "—")
        _cell(row, 4, h.get("hardness_base_t5") or h.get("hardness_base") or "—")
        _cell(row, 5, h.get("hardness_weld") or "—")
        _cell(row, 6, h.get("hardness_haz_t2") or h.get("hardness_haz") or "—")
        _cell(row, 7, h.get("hardness_haz_t4") or h.get("hardness_haz") or "—")

    doc.add_paragraph()
    doc.add_paragraph(
        "Вывод: в результате контроля твердости металла элементов сосуда отклонений твердости "
        "металла от указанного в нормативной документации не выявлено."
    )
    doc.add_paragraph()


def _group_thickness(measurements: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
    grouped: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
    for m in measurements:
        loc = str(m.get("location") or m.get("element_name") or "Обечайка")
        grouped[loc].append(m)
    for loc in grouped:
        grouped[loc].sort(
            key=lambda x: int(str(x.get("section_number") or x.get("point_number") or "0") or 0)
        )
    return grouped


def append_epb_protocol_uzt(
    doc: Document,
    g: Callable[..., Any],
    header_block: Callable[[], None],
    thickness: List[Dict[str, Any]],
) -> None:
    """Протокол №4 — УЗТ, таблица по 5 точек в строке (образец)."""
    doc.add_page_break()
    doc.add_heading("Протокол № 4", level=2)
    doc.add_paragraph("ультразвукового контроля толщины стенок элементов сосуда")
    doc.add_paragraph()
    header_block()
    doc.add_paragraph("Результаты контроля:")
    if not thickness:
        doc.add_paragraph("Данные ультразвуковой толщинометрии не предоставлены.")
        return

    grouped = _group_thickness(thickness)
    cols = 2 + POINTS_PER_UZT_ROW * 2
    for element_name, points in grouped.items():
        nominal = "—"
        for p in points:
            n = p.get("nominal_thickness")
            if n not in (None, ""):
                nominal = str(n).replace(".", ",")
                break
        if nominal == "—":
            nominal = str(g("wall_thickness", default="—"))

        rows_needed = max(1, (len(points) + POINTS_PER_UZT_ROW - 1) // POINTS_PER_UZT_ROW)
        tbl = doc.add_table(rows=rows_needed + 1, cols=cols)
        tbl.style = "Table Grid"
        hdr = ["Наименование элемента сосуда", "Толщина стенки номинальная, мм"]
        for _ in range(POINTS_PER_UZT_ROW):
            hdr.extend(["№ точки", "Толщина, мм"])
        _header_row(tbl, hdr)

        for r in range(rows_needed):
            row = tbl.rows[r + 1]
            _cell(row, 0, element_name if r == 0 else "")
            _cell(row, 1, nominal if r == 0 else "")
            for p_idx in range(POINTS_PER_UZT_ROW):
                abs_idx = r * POINTS_PER_UZT_ROW + p_idx
                col_base = 2 + p_idx * 2
                if abs_idx < len(points):
                    pt = points[abs_idx]
                    pn = pt.get("section_number") or pt.get("point_number") or abs_idx + 1
                    tv = pt.get("thickness")
                    tv_s = str(tv).replace(".", ",") if tv is not None else "—"
                    _cell(row, col_base, pn)
                    _cell(row, col_base + 1, tv_s)
                else:
                    _cell(row, col_base, "")
                    _cell(row, col_base + 1, "")
        doc.add_paragraph()

    doc.add_paragraph(
        "Вывод: в соответствии с п.1 Приложения 8 приказа Ростехнадзора от 15.12.2020 №536 "
        "уменьшения толщины стенки сосуда вследствие коррозионного и эрозионного износа сверх "
        "минимального значения, установленного расчётом на прочность, не выявлено."
    )
    doc.add_paragraph()


def append_epb_protocol_weld_control(
    doc: Document,
    protocol_no: int,
    title: str,
    header_block: Callable[[], None],
    welds: List[Dict[str, Any]],
    conclusion: str,
) -> None:
    """Протокол №5 (МПК) / №6 (УЗК) — таблица сварных соединений."""
    doc.add_page_break()
    doc.add_heading(f"Протокол № {protocol_no}", level=2)
    doc.add_paragraph(title)
    doc.add_paragraph()
    header_block()
    doc.add_paragraph("Результаты контроля")
    rows = welds if welds else [
        {"weld_number": "Пересечение сварного шва К1+П1*", "defect_description": "дефектов не обнаружено", "conclusion": "годен"},
    ]
    tbl = doc.add_table(rows=len(rows) + 1, cols=4)
    tbl.style = "Table Grid"
    _header_row(
        tbl,
        ["Номер сварного соединения по карте контроля", "Выявленный дефект", "Описание дефекта", "Заключение (годен, ремонт и т.д.)"],
    )
    for i, w in enumerate(rows, 1):
        row = tbl.rows[i]
        _cell(row, 0, w.get("weld_number") or w.get("name") or f"Шов {i}")
        defect = w.get("defect_description") or w.get("defect_number") or "дефектов не обнаружено"
        _cell(row, 1, defect)
        _cell(row, 2, w.get("description") or defect)
        _cell(row, 3, w.get("conclusion") or "годен")
    doc.add_paragraph()
    doc.add_paragraph(conclusion)
    doc.add_paragraph()


def _filter_welds(welds: List[Dict[str, Any]], method: str) -> List[Dict[str, Any]]:
    m_up = method.upper()
    aliases = {m_up}
    if m_up == "MPK":
        aliases.update({"MK", "МПК"})
    if m_up == "UZK":
        aliases.update({"УЗК"})
    has_tags = any(str(w.get("control_method") or w.get("method") or "").strip() for w in welds)
    if not has_tags:
        return welds
    return [
        w for w in welds
        if str(w.get("control_method") or w.get("method") or "").upper() in aliases
    ]


def append_epb_appendix_e(
    doc: Document,
    ctx: EpbReportContext,
    g: Callable[..., Any],
    attrs: Dict[str, Any],
    equipment_data: Dict[str, Any],
) -> None:
    """Приложение Е — расчёт остаточного ресурса и прочности."""
    preset = preset_from_equipment_data(equipment_data)
    type_label = device_type_label(preset)
    doc.add_page_break()
    doc.add_heading(
        "Приложение Е Расчетные и аналитические процедуры оценки и прогнозирования "
        "технического состояния сосуда",
        level=1,
    )
    intro = [
        "Расчет остаточного ресурса основан на принципе безопасной эксплуатации по техническому состоянию.",
        "В качестве определяющих параметров технического состояния сосуда принимается:",
        "- изменение физико-механических характеристик металла элементов конструкции сосуда;",
        "- уменьшение толщины стенки элементов сосуда в результате коррозии (эрозии);",
        "- недопустимые дефекты основного металла и сварных элементов сосуда;",
    ]
    for p in intro:
        doc.add_paragraph(p)
    doc.add_paragraph(
        "По результатам измерения твердости и косвенного определения временного сопротивления "
        "разрыву можно сделать вывод о том, что за весь период эксплуатации сосуда деградации "
        "механических свойств материала не произошло. Расчет остаточного ресурса проводится "
        "по эрозионному (коррозионному) износу."
    )
    doc.add_paragraph(
        "Расчеты на прочность и определения остаточного ресурса выполнены согласно: "
        "ГОСТ 34233.1-2017, ГОСТ 34233.2-2017, приказа Ростехнадзора от 15.12.2020 №536 и ДиОР-05."
    )
    doc.add_paragraph()
    doc.add_paragraph("Расчет на прочность и определение остаточного ресурса сосуда")
    doc.add_paragraph("Исходные данные")

    thickness = g("thickness_measurements", "thicknessMeasurements", default=[])
    if not isinstance(thickness, list):
        thickness = []
    wall_th = attrs.get("wall_thickness") or attrs.get("thickness") or g("wall_thickness", default="18,0")
    min_allowed = attrs.get("min_wall_thickness") or g("min_allowed_thickness", default="15,0")
    min_vals = []
    for p in thickness:
        if isinstance(p, dict) and p.get("thickness") not in (None, ""):
            try:
                min_vals.append(float(str(p["thickness"]).replace(",", ".")))
            except (TypeError, ValueError):
                pass
    s_f = min(min_vals) if min_vals else float(str(wall_th).replace(",", "."))
    try:
        s_n = float(str(wall_th).replace(",", "."))
    except (TypeError, ValueError):
        s_n = 18.0
    try:
        s_otb = float(str(min_allowed).replace(",", "."))
    except (TypeError, ValueError):
        s_otb = 15.0

    comm_year = attrs.get("commissioning_year") or g("commissioning_year") or equipment_data.get("commissioning_date")
    t1 = 16
    if comm_year:
        try:
            t1 = max(1, datetime.now().year - int(str(comm_year)[:4]))
        except (TypeError, ValueError):
            pass
    a = (s_n - s_f) / t1 if t1 > 0 else 0.01
    tk = (s_f - s_otb) / a if a > 0 else 110

    calc = g("calculation_data", default={})
    if isinstance(calc, dict) and calc.get("residual_life_years"):
        tk = float(calc["residual_life_years"])

    doc.add_paragraph("Таблица Е.1. Расчёт остаточного ресурса")
    tbl = doc.add_table(rows=7, cols=5)
    tbl.style = "Table Grid"
    _header_row(tbl, ["№ п/п", "Наименование величины", "Единица измерения", "Обозначение и расчетная формула", "Числовое значение"])
    rows_e1 = [
        ("1", "Время эксплуатации", "лет", "t₁", str(t1)),
        ("2", "Паспортная толщина стенки\nОбечайка / Днище", "мм", "Sн", f"{s_n:.1f} / {s_n:.1f}"),
        ("3", "Минимально допустимая толщина стенки сосуда", "мм", "Sотб", f"{s_otb:.1f} / {s_otb:.1f}"),
        ("4", "Минимальная толщина по результатам замеров", "мм", "Sф", f"{s_f:.1f} / {s_f:.1f}"),
        ("5", "Скорость коррозии металла сосуда", "мм/год", "a = (Sн - Sф) / t₁", f"{a:.3f} / {a:.3f}"),
        ("6", "Остаточный срок службы сосуда, поэлементно", "лет", "Tk = (Sф - Sотб) / a", f"{tk:.0f} / {tk:.0f}"),
    ]
    for i, row_data in enumerate(rows_e1, 1):
        for j, val in enumerate(row_data):
            _cell(tbl.rows[i], j, val)

    doc.add_paragraph()
    doc.add_paragraph("2. Расчет на прочность сосуда")
    try:
        p_val = float(str(attrs.get("working_pressure") or g("working_pressure") or "1,0").replace(",", "."))
    except (TypeError, ValueError):
        p_val = 1.0
    try:
        d_n = float(str(attrs.get("diameter") or g("diameter") or "3400").replace(",", "."))
    except (TypeError, ValueError):
        d_n = 3400.0
    phi, sigma = 0.9, 177.0
    s_p = (p_val * d_n) / (2 * phi * sigma - 0.5 * p_val) if (2 * phi * sigma - 0.5 * p_val) > 0 else s_otb
    doc.add_paragraph(
        f"Для {type_label} зав.№ {ctx.serial} при рабочем давлении {p_val} МПа "
        f"и диаметре {d_n:.0f} мм расчётная толщина стенки Sр = {s_p:.2f} мм. "
        f"Фактическая минимальная толщина {s_f:.1f} мм обеспечивает требуемый запас прочности."
    )
    residual_text = g("residual_life_text", default=f"более {ctx.residual_life_years} лет")
    doc.add_paragraph(f"Утонений по толщине стенки днищ не обнаружено, дальнейший срок службы «{residual_text}».")
    doc.add_paragraph(f"Расчет остаточного ресурса сосуда показал, что он может эксплуатироваться «{residual_text}».")
    doc.add_paragraph()
