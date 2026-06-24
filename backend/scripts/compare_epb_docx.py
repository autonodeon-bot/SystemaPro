"""
Сравнение сгенерированного заключения ЭПБ с образцом (ключевые маркеры разделов).

Использование:
  python scripts/compare_epb_docx.py <образец.docx> <сгенерированный.docx>
"""
from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import List, Set

from docx import Document

# Маркеры структуры образца 25-3173
EXPECTED_MARKERS: List[str] = [
    "ЗАКЛЮЧЕНИЕ ЭКСПЕРТИЗЫ ПРОМЫШЛЕННОЙ БЕЗОПАСНОСТИ",
    "1. Вводная часть",
    "2. Наименование объекта экспертизы",
    "6. Результаты проведенной экспертизы",
    "7.1. Анализ результатов",
    "Таблица Б1",
    "Таблица Б6",
    "Протокол № 3",
    "Протокол № 4",
    "Протокол № 5",
    "Протокол № 6",
    "Приложение Е",
    "Таблица Е.1",
    "Расчет остаточного ресурса",
]


def extract_text(doc_path: Path) -> str:
    doc = Document(str(doc_path))
    parts: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def find_markers(text: str, markers: List[str]) -> tuple[List[str], List[str]]:
    found: List[str] = []
    missing: List[str] = []
    norm = re.sub(r"\s+", " ", text.lower())
    for m in markers:
        key = re.sub(r"\s+", " ", m.lower())
        if key in norm:
            found.append(m)
        else:
            missing.append(m)
    return found, missing


def compare(sample_path: Path, generated_path: Path) -> int:
    sample_text = extract_text(sample_path)
    gen_text = extract_text(generated_path)

    print(f"Образец:       {sample_path}")
    print(f"Сгенерирован:  {generated_path}")
    print(f"Абзацев+ячеек образец: {len(sample_text.splitlines())}")
    print(f"Абзацев+ячеек сгенер.:   {len(gen_text.splitlines())}")
    print()

    _, sample_missing = find_markers(sample_text, EXPECTED_MARKERS)
    found_gen, missing_gen = find_markers(gen_text, EXPECTED_MARKERS)

    print("=== Маркеры в сгенерированном документе ===")
    for m in found_gen:
        print(f"  [OK] {m}")
    for m in missing_gen:
        print(f"  [--] {m}")

    if sample_missing:
        print("\n(В образце не найдены некоторые маркеры — возможно другая формулировка)")

    # Уникальные заголовки протоколов в образце
    sample_protocols = set(re.findall(r"Протокол №\s*\d+", sample_text, re.I))
    gen_protocols = set(re.findall(r"Протокол №\s*\d+", gen_text, re.I))
    print(f"\nПротоколы в образце: {sorted(sample_protocols)}")
    print(f"Протоколы в сгенер.: {sorted(gen_protocols)}")

    tables_sample = len(Document(str(sample_path)).tables)
    tables_gen = len(Document(str(generated_path)).tables)
    print(f"\nТаблиц в образце: {tables_sample}, в сгенер.: {tables_gen}")

    return len(missing_gen)


def main() -> None:
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(2)
    sample = Path(sys.argv[1])
    generated = Path(sys.argv[2])
    if not sample.is_file():
        print(f"Образец не найден: {sample}")
        sys.exit(1)
    if not generated.is_file():
        print(f"Файл не найден: {generated}")
        sys.exit(1)
    missing = compare(sample, generated)
    sys.exit(1 if missing else 0)


if __name__ == "__main__":
    main()
