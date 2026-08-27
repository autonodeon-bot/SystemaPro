"""Smoke: fill to-1 from latest TECHNICAL_REPORT inspection on server."""
import asyncio
import os
import sys
from pathlib import Path

sys.path.insert(0, "/app")
os.chdir("/app")

from sqlalchemy import select, text
from database import AsyncSessionLocal
from report_forms_registry import resolve_form_path, load_forms_catalog
from form_template_filler import fill_vessel_form_to1
from report_org_settings import load_report_org_settings


async def main():
    load_forms_catalog.cache_clear()
    p = resolve_form_path("to-1")
    print("template:", p, p.exists() if p else None)
    async with AsyncSessionLocal() as db:
        r = await db.execute(
            text(
                """
                SELECT i.id::text, i.data, e.name, e.id::text
                FROM inspections i
                JOIN equipment e ON e.id = i.equipment_id
                WHERE i.id = '45dcdf9a-aef7-44cc-9846-a2d6a7ef3acd'
                LIMIT 1
                """
            )
        )
        row = r.first()
        if not row:
            r = await db.execute(
                text(
                    """
                    SELECT i.id::text, i.data, e.name, e.id::text
                    FROM inspections i
                    JOIN equipment e ON e.id = i.equipment_id
                    ORDER BY COALESCE(i.updated_at, i.created_at) DESC NULLS LAST
                    LIMIT 1
                    """
                )
            )
            row = r.first()
        if not row:
            print("NO INSPECTION")
            return
        insp_id, data, eq_name, eq_id = row
        print("inspection:", insp_id, "equipment:", eq_name)
        er = await db.execute(
            text(
                """
                SELECT e.name, e.serial_number, et.code, et.name, e.attributes
                FROM equipment e
                LEFT JOIN equipment_types et ON et.id = e.type_id
                WHERE e.id = CAST(:eid AS uuid)
                """
            ),
            {"eid": eq_id},
        )
        eq = er.first()
        attrs = eq[4] if eq and isinstance(eq[4], dict) else {}
        equipment_data = {
            "id": eq_id,
            "name": eq[0] if eq else eq_name,
            "factory_number": (attrs or {}).get("factory_number") or (eq[1] if eq else None),
            "registration_number": (attrs or {}).get("registration_number"),
            "serial_number": eq[1] if eq else None,
            "type_code": eq[2] if eq else "VESSEL",
            "type_name": eq[3] if eq else None,
            "attributes": attrs or {},
        }
        inspection_data = {
            "id": insp_id,
            "data": data if isinstance(data, dict) else {},
            "report_form_id": "to-1",
        }
        out = Path("/app/reports/_smoke_to1_official.docx")
        fill_vessel_form_to1(
            inspection_data=inspection_data,
            equipment_data=equipment_data,
            output_path=str(out),
            org_settings=load_report_org_settings(),
        )
        print("OUT", out, out.stat().st_size)
        # quick structure check
        from docx import Document
        d = Document(str(out))
        texts = [p.text.strip() for p in d.paragraphs if p.text.strip()][:8]
        print("first paras:", texts)
        print("tables:", len(d.tables))


if __name__ == "__main__":
    asyncio.run(main())
