from docx import Document

d = Document("/app/reports/TECHNICAL_REPORT_45dcdf9a-aef7-44cc-9846-a2d6a7ef3acd_OFFICIAL.docx")
print("T2 name:", d.tables[2].rows[0].cells[1].text)
print("T2 manufacturer:", d.tables[2].rows[2].cells[1].text)
print("T2 material:", d.tables[2].rows[9].cells[1].text)
print("T3 R2:", [c.text for c in d.tables[3].rows[2].cells[:7]])
print("T3 R3:", [c.text for c in d.tables[3].rows[3].cells[:7]])
print("T4 pressure:", d.tables[4].rows[1].cells[2].text, d.tables[4].rows[1].cells[3].text)
print("T5 materials R2:", [c.text for c in d.tables[5].rows[2].cells[:3]])
