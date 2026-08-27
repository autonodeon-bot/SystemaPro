# -*- coding: utf-8 -*-
"""Inspect appendix/main tables content after fill."""
import asyncio
import sys
from pathlib import Path

sys.path.insert(0, "/app")

from docx import Document
from sqlalchemy import text
from database import AsyncSessionLocal
from form_template_filler import fill_vessel_form_to1
from report_org_settings import load_report_org_settings
from report_forms_registry import load_forms_catalog


async def main():
    load_forms_catalog.cache_clear()
    insp_id = "45dcdf9a-aef7-44cc-9846-a2d6a7ef3acd"
    async with AsyncSessionLocal() as db:
        r = await db.execute(
            text(
                """
                SELECT i.data, e.name, e.serial_number, e.attributes, e.id::text
                FROM inspections i JOIN equipment e ON e.id=i.equipment_id
                WHERE i.id=CAST(:iid AS uuid)
                """
            ),
            {"iid": insp_id},
        )
        data, eq_name, serial, attrs, eq_id = r.first()
        data = data if isinstance(data, dict) else {}
        out = Path("/app/reports/_diag_to1.docx")
        fill_vessel_form_to1(
            inspection_data={"id": insp_id, "data": data, "report_form_id": "to-1"},
            equipment_data={
                "id": eq_id,
                "name": eq_name,
                "serial_number": serial,
                "attributes": attrs if isinstance(attrs, dict) else {},
                "type_code": "VESSEL",
            },
            output_path=str(out),
            org_settings=load_report_org_settings(),
        )
    doc = Document(str(out))
    for idx in (0, 1, 2, 3, 4, 5, 7, 8, 15, 17, 21):
        if idx >= len(doc.tables):
            continue
        t = doc.tables[idx]
        print(f"\n=== TABLE {idx} rows={len(t.rows)} cols={len(t.columns)} ===")
        for ri, row in enumerate(t.rows[:6]):
            cells = [c.text.replace("\n", " | ")[:60] for c in row.cells[:5]]
            print(f"  r{ri}: {cells}")


if __name__ == "__main__":
    asyncio.run(main())
