import json, os, sys
sys.path.insert(0, "/app")
os.chdir("/app")
import asyncio
from sqlalchemy import text
from database import AsyncSessionLocal
from docx import Document

async def main():
    async with AsyncSessionLocal() as db:
        r = await db.execute(text("""
            SELECT i.id::text, i.questionnaire_id::text, i.data
            FROM inspections i WHERE i.id = CAST('45dcdf9a-aef7-44cc-9846-a2d6a7ef3acd' AS uuid)
        """))
        row = r.first()
        print("qid", row[1])
        # ndt methods
        r2 = await db.execute(text("""
            SELECT method_code, method_name, is_performed, performed_date,
                   results, defects, additional_data, photos, conclusion, inspector_name
            FROM ndt_methods
            WHERE inspection_id = CAST('45dcdf9a-aef7-44cc-9846-a2d6a7ef3acd' AS uuid)
               OR (CAST(:qid AS text) IS NOT NULL AND questionnaire_id = CAST(:qid AS uuid))
        """), {"qid": row[1]})
        rows = r2.fetchall()
        print("ndt_methods count", len(rows))
        for m in rows:
            print("---", m[0], m[1], "performed", m[2], "date", m[3], "inspector", m[9])
            for label, val in [("results", m[4]), ("defects", m[5]), ("additional_data", m[6]), ("photos", m[7])]:
                if val is None: continue
                s = json.dumps(val, ensure_ascii=False, default=str)
                print(f"  {label}: {s[:500]}")
        if row[1]:
            r3 = await db.execute(text("""
                SELECT document_number, file_name, file_path, file_size
                FROM questionnaire_document_files
                WHERE questionnaire_id = CAST(:qid AS uuid)
                ORDER BY document_number
            """), {"qid": row[1]})
            docs = r3.fetchall()
            print("qdocs", len(docs))
            for d in docs[:60]:
                print(d[0], d[1], d[2], d[3])
    # inspect latest UI-generated docx
    p = "/app/reports/TECHNICAL_REPORT_45dcdf9a-aef7-44cc-9846-a2d6a7ef3acd_20260710_112101.docx"
    d = Document(p)
    print("LATEST DOC tables", len(d.tables), "paras", len(d.paragraphs), "size", os.path.getsize(p))
    t2 = d.tables[2]
    for ri in range(min(8, len(t2.rows))):
        print("T2 R", ri, [c.text.strip()[:40] for c in t2.rows[ri].cells])
    # UZT table if present
    for ti in (17, 21, 25):
        t = d.tables[ti]
        print(f"T{ti} first data rows:")
        for ri in range(min(5, len(t.rows))):
            print(" ", ri, [c.text.strip()[:30] for c in t.rows[ri].cells[:6]])

asyncio.run(main())
