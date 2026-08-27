# -*- coding: utf-8 -*-
"""Smoke-тест заполнения to-1 после правок по замечаниям PDF."""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from form_template_filler import fill_vessel_form_to1, NOT_PROVIDED
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table

out = Path(__file__).resolve().parents[1] / "reports" / "_smoke_to1_remarks.docx"

inspection = {
    "date_performed": "2026-07-13T10:00:00",
    "data": {
        "report_form_id": "to-1",
        "protocol_number": "ТО-15/26",
        "vessel_name": "Сепаратор нефтегазовый",
        "serial_number": "1075",
        "reg_number": "477",
        "inventory_number": "—",
        "organization": 'ПАО "Сургутнефтегаз" / НГДУ «Лянторнефть» / Пункт подготовки и сбора нефти',
        "location": "",
        "manufacturer": "Снежнянский завод химического машиностроения",
        "manufacture_year": "1986",
        "commissioning_year": "1987",
        "working_pressure": "0,8",
        "design_pressure": "1,0",
        "test_pressure": "1,45",
        "design_temperature": "+100",
        "working_temperature": "-60",
        "working_medium": "нефть, газ",
        "hazard_class": "3",
        "explosion_hazard": "IIA",
        "fire_hazard": "В",
        "shell_material": "09Г2С",
        "volume": "100",
        "diameter": "3000",
        "corrosion_allowance": "2",
        "purpose": "сепарация",
        "service_life": "20",
        "calculation_result": "сепаратора нефтегазового при рабочих параметрах",
        "technical_state": "работоспособное, пригодно к дальнейшей эксплуатации",
        "documentation_conclusion": "в полном объёме",
        "documents": {
            "1": True,
            "2": True,
            "5": False,
            "10": False,
        },
        "documents_info": {
            "1": {"number": "1", "date": "2026-07-13", "pages": "3"},
            "2": {"number": "2", "date": "2026-07-13", "pages": "2"},
        },
        "vessel_elements": [
            {
                "name": "Обечайка",
                "quantity": "1",
                "diameter_mm": "3000",
                "length_mm": "16300",
                "wall_thickness_mm": "10,0",
                "calc_thickness": "8,0",
                "material": "09Г2С",
                "gost": "ГОСТ 19281",
            },
            {
                "name": "Днище левое",
                "quantity": "1",
                "diameter_mm": "3000",
                "length_mm": "825",
                "wall_thickness_mm": "12,0",
                "calc_thickness": "10,0",
                "material": "09Г2С",
            },
        ],
        "previous_inspections": [
            {
                "kind": "ТО",
                "date": "2020-05-01",
                "result": "Годен",
                "scope": "100%",
                "organization": "ООО Тест",
                "report_number": "ТО-1/20",
            }
        ],
        "thickness_measurements": [
            {"element": "Обечайка", "thickness": "9.8"},
            {"element": "Обечайка", "thickness": "9.5"},
            {"element": "Днище левое", "thickness": "11.2"},
        ],
        "hydraulic_test_history": [
            {"date": "2026-07-13", "type": "гидравлическое", "pressure": "1,45"}
        ],
    },
}
equipment = {
    "name": "Сепаратор нефтегазовый",
    "serial_number": "1075",
    "location": "",
    "attributes": {},
}
org = {
    "contractor": {
        "legal_name": "Общество с ограниченной ответственностью «ЮТАР»",
        "postal_address": "628285, ХМАО, г. Урай",
        "director_name": "Иванов И.И.",
        "phone": "+7",
        "email": "info@test.ru",
    },
    "customer": {"legal_name": 'ПАО "Сургутнефтегаз"', "director": "Петров П.П."},
    "ndt_lab": {"name": "ЛНК ООО «ЮТАР»", "certificate": "№ 1234"},
}

path = fill_vessel_form_to1(
    inspection_data=inspection,
    equipment_data=equipment,
    output_path=str(out),
    org_settings=org,
    specialist_docs=[{"full_name": "Коровин Александр Сергеевич", "certificate_number": "АА-11"}],
)
print("OUT", path)

d = Document(path)
# Main SDT checks
body = d.element.body
sdt = list(body.iterchildren())[0]
content = sdt.find(qn("w:sdtContent"))
tbls = [Table(t, d) for t in content.iter(qn("w:tbl"))]
print("main tables", len(tbls))
print("title name", tbls[0].rows[6].cells[1].text[:80])
print("title serial", tbls[0].rows[7].cells[1].text)
print("customer", tbls[2].rows[0].cells[1].text[:80])
print("tech name", tbls[7].rows[0].cells[1].text)
print("works caption fix:", end=" ")
for p in content.iter(qn("w:p")):
    tx = "".join(t.text or "" for t in p.iter(qn("w:t"))).strip()
    if tx.startswith("Таблица №"):
        # collect around section 10
        pass
# find captions after works
seen = False
caps = []
for p in content.iter(qn("w:p")):
    tx = "".join(t.text or "" for t in p.iter(qn("w:t"))).strip()
    if tx.startswith("10. Перечень работ, выполненных") and not tx[-1:].isdigit():
        seen = True
        continue
    if seen and tx.startswith("Таблица"):
        caps.append(tx)
        break
print(caps)

# appendix header location
print("app header org", d.tables[0].rows[0].cells[2].text[:80])
print("app header loc", d.tables[0].rows[2].cells[2].text[:80])
print("lab", d.tables[0].rows[4].cells[0].text[:80])
print("lab cert", d.tables[0].rows[6].cells[0].text[:80])

# docs not provided
print("doc5", d.tables[1].rows[5].cells[2].text)
print("doc10-ish", [d.tables[1].rows[i].cells[2].text for i in range(1, min(12, len(d.tables[1].rows)))])

# element material
print("el0 material", d.tables[3].rows[2].cells[6].text)
print("el0 calc", d.tables[3].rows[2].cells[5].text)

# section 15
for p in content.iter(qn("w:p")):
    tx = "".join(t.text or "" for t in p.iter(qn("w:t"))).strip()
    if tx.startswith("Фактическое значение"):
        print("sec15", tx[:160])
        break

assert "Пункт подготовки" in d.tables[0].rows[2].cells[2].text or "Лянтор" not in d.tables[0].rows[0].cells[2].text
assert NOT_PROVIDED in d.tables[1].rows[5].cells[2].text
assert "09Г2С" in d.tables[3].rows[2].cells[6].text
assert caps == ["Таблица № 7"]
print("OK")
