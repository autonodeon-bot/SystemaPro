# -*- coding: utf-8 -*-
"""Regenerate to-1 for known inspection and check key strings in output."""
import asyncio
import sys
from pathlib import Path

sys.path.insert(0, "/app")

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from sqlalchemy import text

from database import AsyncSessionLocal
from form_template_filler import fill_vessel_form_to1, _extract_specialists, _build_context
from report_org_settings import load_report_org_settings
from report_forms_registry import load_forms_catalog, resolve_form_path


def _all_text(doc: Document) -> str:
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                tx = cell.text.strip()
                if tx:
                    parts.append(tx)
    # SDT
    body = doc.element.body
    for sdt in body.iter(qn("w:sdt")):
        for t_el in sdt.iter(qn("w:t")):
            if t_el.text:
                parts.append(t_el.text)
    return "\n".join(parts)


async def main():
    load_forms_catalog.cache_clear()
    insp_id = "45dcdf9a-aef7-44cc-9846-a2d6a7ef3acd"
    async with AsyncSessionLocal() as db:
        r = await db.execute(
            text(
                """
                SELECT i.data, e.name, e.serial_number, e.attributes, e.id::text
                FROM inspections i
                JOIN equipment e ON e.id = i.equipment_id
                WHERE i.id = CAST(:iid AS uuid)
                """
            ),
            {"iid": insp_id},
        )
        row = r.first()
        data, eq_name, serial, attrs, eq_id = row
        data = data if isinstance(data, dict) else {}

        # verification equipment linked
        ve = []
        vr = await db.execute(
            text(
                """
                SELECT ve.name, ve.serial_number, ve.equipment_type,
                       ve.verification_certificate_number, ve.next_verification_date
                FROM inspection_equipment ie
                JOIN verification_equipment ve ON ve.id = ie.verification_equipment_id
                WHERE ie.inspection_id = CAST(:iid AS uuid)
                """
            ),
            {"iid": insp_id},
        )
        for v in vr.fetchall():
            ve.append(
                {
                    "name": v[0],
                    "serial_number": v[1],
                    "equipment_type": v[2],
                    "verification_certificate_number": v[3],
                    "next_verification_date": str(v[4]) if v[4] else None,
                }
            )
        print("VE count", len(ve), ve[:3])

        specs = _extract_specialists(data, [])
        print("SPECIALISTS", specs)

        equipment_data = {
            "id": eq_id,
            "name": eq_name,
            "serial_number": serial,
            "attributes": attrs if isinstance(attrs, dict) else {},
            "type_code": "VESSEL",
        }
        out = Path("/app/reports/_diag_to1.docx")
        fill_vessel_form_to1(
            inspection_data={"id": insp_id, "data": data, "report_form_id": "to-1"},
            equipment_data=equipment_data,
            output_path=str(out),
            org_settings=load_report_org_settings(),
            verification_equipment=ve,
            specialist_docs=[],
        )
        doc = Document(str(out))
        text_all = _all_text(doc)
        for needle in (
            "Коровин",
            "Федоров",
            "Соколов",
            "Попов",
            "Сепаратор",
            "1075",
            "Не предоставлено",
            "09Г2С",
            "1,45",
            "обечайка",
            "ВИК",
            "VIK",
        ):
            print(f"HAS[{needle}]", needle in text_all)

        # main SDT tables sample
        body = doc.element.body
        sdt = list(body.iterchildren())[0]
        content = sdt.find(qn("w:sdtContent"))
        tbls = [Table(t, doc) for t in content.iter(qn("w:tbl"))]
        print("main_tables", len(tbls))
        if tbls:
            print("title_r6", tbls[0].rows[6].cells[1].text[:80] if len(tbls[0].rows) > 6 else "?")
        if len(tbls) > 4:
            print("specs_r1", tbls[4].rows[1].cells[1].text if len(tbls[4].rows) > 1 else "?")
            print("specs_r2", tbls[4].rows[2].cells[1].text if len(tbls[4].rows) > 2 else "?")
        if len(tbls) > 7:
            print("tech_r0", tbls[7].rows[0].cells[1].text[:80] if len(tbls[7].rows) > 0 else "?")
        print("appendix_tables", len(doc.tables))
        if len(doc.tables) > 3:
            print("elements_sample", doc.tables[3].rows[1].cells[0].text if len(doc.tables[3].rows) > 1 else "?")


if __name__ == "__main__":
    asyncio.run(main())
