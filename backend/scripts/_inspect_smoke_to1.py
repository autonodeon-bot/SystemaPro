from docx import Document

d = Document("/app/reports/_smoke_to1_official.docx")
rels = list(d.part.rels.values())
imgs = [r for r in rels if "image" in str(getattr(r, "reltype", "")).lower()]
print("image_rels", len(imgs))
print("tables", len(d.tables), "paragraphs", len(d.paragraphs))
for i, p in enumerate(d.paragraphs):
    t = (p.text or "").strip()
    if not t:
        continue
    low = t.lower()
    if any(k in low for k in ("схем", "фото", "приложение № 7", "приложение №7", "неразрушающ", "гидравл")):
        print(i, t[:120])
