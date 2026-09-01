#!/usr/bin/env python3
"""Smoke: fill to-1 and verify landscape sectPr near schemes + section 15 fill."""
from __future__ import annotations

import json
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, "/app")
from form_template_filler import fill_vessel_form_to1
from scheme_ndt_overlays import render_all_layer_pngs


def _sect_orient(p_el):
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        return None
    sect = pPr.find(qn("w:sectPr"))
    if sect is None:
        return None
    pgSz = sect.find(qn("w:pgSz"))
    if pgSz is None:
        return None
    return pgSz.get(qn("w:orient")) or "portrait"


def main() -> int:
    data = {
        "report_form_id": "to-1",
        "vessel_name": "Аппарат тестовый А-1",
        "serial_number": "SN-100",
        "reg_number": "Р-55",
        "inventory_number": "ИНВ-9",
        "orientation": "vertical",
        "shell_count": 2,
        "thickness_measurements": [
            {"point_number": 1, "thickness": 8.2, "x_percent": 40, "y_percent": 50},
            {"point_number": 2, "thickness": 8.0, "x_percent": 55, "y_percent": 50},
        ],
        "geometry": {
            "kind": "vessel",
            "orientation": "vertical",
            "shell_count": 2,
            "heads": ["elliptical", "elliptical"],
            "nozzles": [
                {"label": "Штуцер 1", "dn": 50, "x_percent": 30, "y_percent": 40},
                {"label": "Штуцер 2", "dn": 80, "x_percent": 70, "y_percent": 60},
            ],
        },
    }
    inspection = {"data": data, "date_performed": "2026-08-28"}
    equipment = {
        "name": "Аппарат тестовый А-1",
        "serial_number": "SN-100",
        "attributes": {
            "registration_number": "Р-55",
            "inventory_number": "ИНВ-9",
        },
    }
    out = Path(tempfile.gettempdir()) / "_smoke_to1_landscape.docx"
    fill_vessel_form_to1(
        inspection_data=inspection,
        equipment_data=equipment,
        output_path=str(out),
        verification_equipment=[],
        org_settings={},
        specialist_docs=[],
        document_files=[],
        ndt_methods=[],
    )
    doc = Document(str(out))
    landscape_n = 0
    for p in doc.element.body.iter(qn("w:p")):
        if _sect_orient(p) == "landscape":
            landscape_n += 1
    body_sect = doc.element.body.find(qn("w:sectPr"))
    body_orient = "portrait"
    if body_sect is not None:
        pg = body_sect.find(qn("w:pgSz"))
        if pg is not None:
            body_orient = pg.get(qn("w:orient")) or "portrait"

    sec15 = ""
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("Фактическое значение"):
            sec15 = t
            break
    # also search SDT
    if not sec15:
        for p_el in doc.element.body.iter(qn("w:p")):
            t = "".join(x.text or "" for x in p_el.iter(qn("w:t"))).strip()
            if t.startswith("Фактическое значение"):
                sec15 = t
                break

    # picture widths
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    widths = []
    for rel in doc.part.rels.values():
        pass
    for p in doc.paragraphs:
        for r in p.runs:
            for d in r._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent"):
                cx = d.get("cx")
                if cx:
                    widths.append(round(int(cx) / 914400, 2))  # EMU -> inches

    print("OUT", out)
    print("LANDSCAPE_SECTIONS", landscape_n)
    print("BODY_ORIENT", body_orient)
    print("SEC15", sec15[:180] if sec15 else "MISSING")
    print("PIC_WIDTHS_IN", widths[:8])
    ok = landscape_n >= 1 and "SN-100" in (sec15 or "") and (not widths or max(widths) <= 9.5)
    print("OK" if ok else "FAIL")
    return 0 if ok else 1


if __name__ == "__main__":
    raise SystemExit(main())
