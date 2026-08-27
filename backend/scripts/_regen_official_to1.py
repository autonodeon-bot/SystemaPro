"""Regenerate TECHNICAL_REPORT via WordGenerator for known inspection."""
import asyncio
import os
import sys
from pathlib import Path

sys.path.insert(0, "/app")
os.chdir("/app")

from sqlalchemy import text
from database import AsyncSessionLocal
from report_forms_registry import load_forms_catalog
from word_generator import WordGenerator
from report_org_settings import load_report_org_settings
from docx import Document


async def main():
    load_forms_catalog.cache_clear()
    insp_id = "45dcdf9a-aef7-44cc-9846-a2d6a7ef3acd"
    async with AsyncSessionLocal() as db:
        r = await db.execute(
            text(
                """
                SELECT i.id::text, i.data, i.conclusion, i.status,
                       i.date_performed, e.id::text, e.name, e.serial_number,
                       et.code, et.name, e.attributes
                FROM inspections i
                JOIN equipment e ON e.id = i.equipment_id
                LEFT JOIN equipment_types et ON et.id = e.type_id
                WHERE i.id = CAST(:iid AS uuid)
                """
            ),
            {"iid": insp_id},
        )
        row = r.first()
        if not row:
            print("inspection not found")
            return
        data = row[1] if isinstance(row[1], dict) else {}
        data = dict(data)
        data["report_form_id"] = "to-1"
        attrs = row[10] if isinstance(row[10], dict) else {}
        inspection_data = {
            "id": row[0],
            "data": data,
            "conclusion": row[2],
            "status": row[3],
            "date_performed": str(row[4]) if row[4] else None,
            "report_form_id": "to-1",
        }
        equipment_data = {
            "id": row[5],
            "name": row[6],
            "serial_number": row[7],
            "factory_number": (attrs or {}).get("factory_number") or row[7],
            "registration_number": (attrs or {}).get("registration_number"),
            "type_code": row[8] or "VESSEL",
            "type_name": row[9],
            "attributes": attrs or {},
        }
        nr = await db.execute(
            text(
                """
                SELECT method_code, method_name, is_performed, performed_date,
                       results, defects, additional_data, photos, conclusion,
                       inspector_name, equipment, standard
                FROM ndt_methods
                WHERE inspection_id = CAST(:iid AS uuid)
                """
            ),
            {"iid": insp_id},
        )
        ndt_methods = []
        for m in nr.fetchall():
            ndt_methods.append(
                {
                    "method_code": m[0],
                    "method_name": m[1],
                    "is_performed": bool(m[2]),
                    "performed_date": str(m[3]) if m[3] else None,
                    "results": m[4],
                    "defects": m[5],
                    "additional_data": m[6] or {},
                    "photos": m[7] or [],
                    "conclusion": m[8],
                    "inspector_name": m[9],
                    "equipment": m[10],
                    "standard": m[11],
                }
            )
        out = Path("/app/reports/TECHNICAL_REPORT_45dcdf9a-aef7-44cc-9846-a2d6a7ef3acd_OFFICIAL.docx")
        wg = WordGenerator()
        result = wg.generate_report_word(
            inspection_data=inspection_data,
            equipment_data=equipment_data,
            ndt_methods=ndt_methods,
            output_path=str(out),
            report_type="TECHNICAL_REPORT",
            org_settings=load_report_org_settings(),
        )
        print("result", result)
        print("exists", out.exists(), out.stat().st_size if out.exists() else 0)
        d = Document(str(out))
        print("T2 name", d.tables[2].rows[0].cells[1].text)
        print("T2 material", d.tables[2].rows[9].cells[1].text)
        print("T3 R2", [c.text for c in d.tables[3].rows[2].cells[:7]])
        print("T3 R3", [c.text for c in d.tables[3].rows[3].cells[:7]])
        print("thickness points in data", len(data.get("thickness_measurements") or []))


if __name__ == "__main__":
    asyncio.run(main())
