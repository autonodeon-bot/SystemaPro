from docx import Document
from pathlib import Path

def dump_table(path, idxs, label):
    d = Document(path)
    print('===', label, 'tables', len(d.tables), '===')
    for ti in idxs:
        t = d.tables[ti]
        print(f'-- T{ti} {len(t.rows)}x{len(t.columns)} --')
        for ri, row in enumerate(t.rows[:14]):
            cells = []
            for ci, c in enumerate(row.cells):
                txt = (c.text or '').replace('\n',' / ').strip()[:60]
                cells.append(f'[{ci}]{txt}')
            print(f'R{ri}: ' + ' || '.join(cells))
        print()

official = '/app/report_forms/to-1.docx'
filled = '/app/reports/_smoke_to1_official.docx'
dump_table(official, [2,3,4], 'OFFICIAL')
dump_table(filled, [2,3,4], 'FILLED')
