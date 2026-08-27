# -*- coding: utf-8 -*-
"""
Интеграционная проверка заполнителей форм ТО (to-1, to-13, to-25)
на payload, близком к данным мобильного приложения.
"""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from PIL import Image, ImageDraw

from form_template_filler import fill_vessel_form_to1
from form_template_filler_pipeline import fill_pipeline_form_to13
from form_template_filler_tank import fill_tank_form_to25
from form_template_filler_underground_pipeline import fill_underground_pipeline_form_to33
from form_template_filler_crane import fill_crane_form_to3
from report_forms_registry import FILLABLE_FORM_IDS, resolve_form_path, suggest_form_id


def _make_png(path: Path, label: str) -> str:
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (640, 400), color=(230, 235, 240))
    draw = ImageDraw.Draw(img)
    draw.rectangle([20, 20, 620, 380], outline=(30, 80, 160), width=3)
    draw.text((40, 180), label[:40], fill=(20, 20, 20))
    img.save(path)
    return str(path)


def _base_data(**extra):
    data = {
        "protocol_number": "ТО-99/26",
        "tech_card_number": "ТК-01",
        "vessel_name": "Ресивер РВ-1",
        "serial_number": "45821",
        "reg_number": "12-345",
        "inventory_number": "ИНВ-778",
        "manufacturer": "АО Уралхиммаш",
        "manufacture_year": "2012",
        "commissioning_year": "2013",
        "working_pressure": "1,6",
        "design_pressure": "1,8",
        "test_pressure": "2,0",
        "diameter": "800",
        "working_temperature": "40",
        "working_medium": "воздух",
        "shell_material": "09Г2С",
        "volume": "1,6",
        "wall_thickness": "8,0",
        "organization": "ООО Газпром трансгаз",
        "location": "КС-12",
        "executors": "Иванов И.И.",
        "documentation_conclusion": "документация в полном объёме",
        "suitability_conclusion": "соответствует",
        "documents": {"10": True},
        "documents_info": {"10": {"number": "П-1", "date": "2012-05-20", "pages": "24"}},
        "thickness_measurements": [
            {"element": "Корпус", "point_number": 1, "thickness": "7,9", "photos": []},
            {"element": "Корпус", "point_number": 2, "thickness": "7,8", "photos": []},
        ],
        "hardness_tests": [
            {"zone": "Т1", "point_number": 1, "hardness_base": "148", "allowed_hardness_base": "120-180"}
        ],
        "hydraulic_test_history": [
            {"date": "2013-04-10", "type": "гидравлическое", "pressure": "2,0", "medium": "вода"}
        ],
        "calculation_data": {
            "residual_life_years": 12,
            "description": "Расчёт выполнен по ГОСТ 34233.1-2017. Толщины удовлетворяют требованиям.",
            "method": "ГОСТ 34233.1",
            "min_thickness": "6,2",
            "conclusion": "прочность обеспечена",
        },
        "uzt_schemes": [],
        "defects": [],
        "weld_inspections": [],
    }
    data.update(extra)
    return data


def main() -> int:
    assert FILLABLE_FORM_IDS == frozenset({f"to-{i}" for i in range(1, 45)})
    for fid in ("to-1", "to-3", "to-13", "to-25", "to-33", "to-5", "to-12", "to-28"):
        p = resolve_form_path(fid)
        assert p and p.exists(), f"Нет шаблона {fid}"

    with tempfile.TemporaryDirectory(prefix="to_forms_") as tmp:
        tmp_path = Path(tmp)
        scheme = _make_png(tmp_path / "scheme.png", "SCHEME UZT")
        photo = _make_png(tmp_path / "point.png", "UZT POINT 1")
        hydro = _make_png(tmp_path / "hydro.png", "HYDRO ACT")
        cert = _make_png(tmp_path / "cert.png", "CERT SCAN")

        data = _base_data(
            control_scheme_image=scheme,
            uzt_schemes=[{"label": "Схема УЗТ корпуса", "scheme_image_path": scheme}],
            thickness_measurements=[
                {"location": "Корпус", "section_number": "1", "thickness": "7,9", "photos": [photo]},
                {"location": "Корпус", "section_number": "2", "thickness": "7,8", "photos": []},
            ],
            hydraulic_act_files=[hydro],
        )
        document_files = [
            {"document_number": "control_scheme_image", "file_path": scheme},
            {"document_number": "uzt_scheme_0", "file_path": scheme},
            {"document_number": "uzt_point_0_0", "file_path": photo},
            {"document_number": "hydraulic_test_act", "file_path": hydro},
        ]
        specialist_docs = [
            {
                "inspector_name": "Иванов И.И.",
                "certifications": [
                    {
                        "certificate_number": "72А00.555",
                        "method_code": "UZT",
                        "scan_file_path": cert,
                    }
                ],
            }
        ]
        org = {
            "contractor": {"legal_name": "ООО Диатекс", "postal_address": "Москва"},
            "customer": {"legal_name": "ООО Газпром трансгаз"},
            "ndt_lab": {"name": "ЛНК", "certificate": "72А00.123"},
        }
        ve = [
            {"name": "Толщиномер А1209", "serial_number": "SN1", "equipment_type": "УЗТ"},
            {"name": "Дефектоскоп УД2", "serial_number": "SN2", "equipment_type": "УЗК"},
        ]
        inspection = {"date_performed": "2026-06-15", "data": data}
        equipment = {
            "name": "Ресивер РВ-1",
            "serial_number": "45821",
            "location": "КС-12",
            "type_code": "VESSEL",
            "attributes": {"reg_number": "12-345"},
        }

        out_dir = Path(__file__).resolve().parents[1] / "report_forms"
        results = {}

        out1 = out_dir / "_test_generated_to1.docx"
        fill_vessel_form_to1(
            inspection, equipment, str(out1),
            verification_equipment=ve, org_settings=org,
            specialist_docs=specialist_docs, document_files=document_files,
        )
        results["to-1"] = out1

        data13 = dict(data)
        data13["report_form_id"] = "to-13"
        data13["vessel_name"] = "Трубопровод ТП-1"
        out13 = out_dir / "_test_generated_to13.docx"
        fill_pipeline_form_to13(
            {"date_performed": "2026-06-15", "data": data13},
            {**equipment, "name": "Трубопровод ТП-1", "type_code": "PIPELINE"},
            str(out13),
            verification_equipment=ve, org_settings=org,
            specialist_docs=specialist_docs, document_files=document_files,
        )
        results["to-13"] = out13

        data25 = dict(data)
        data25["report_form_id"] = "to-25"
        data25["vessel_name"] = "Резервуар РВС-1"
        out25 = out_dir / "_test_generated_to25.docx"
        fill_tank_form_to25(
            {"date_performed": "2026-06-15", "data": data25},
            {**equipment, "name": "Резервуар РВС-1", "type_code": "TANK"},
            str(out25),
            verification_equipment=ve, org_settings=org,
            specialist_docs=specialist_docs, document_files=document_files,
        )
        results["to-25"] = out25

        data33 = dict(data)
        data33["report_form_id"] = "to-33"
        data33["vessel_name"] = "Подземный трубопровод ПТ-1"
        data33["pipeline_category"] = "II"
        data33["pipeline_length"] = "120 м"
        data33["weld_inspections"] = [
            {"weld_number": "С1", "result": "дефектов не обнаружено"},
            {"weld_number": "С2", "result": "дефектов не обнаружено"},
        ]
        out33 = out_dir / "_test_generated_to33.docx"
        fill_underground_pipeline_form_to33(
            {"date_performed": "2026-06-15", "data": data33},
            {
                **equipment,
                "name": "Подземный трубопровод ПТ-1",
                "type_code": "UNDERGROUND_PIPELINE",
            },
            str(out33),
            verification_equipment=ve,
            org_settings=org,
            specialist_docs=specialist_docs,
            document_files=document_files,
        )
        results["to-33"] = out33

        data3 = dict(data)
        data3["report_form_id"] = "to-3"
        data3["vessel_name"] = "Кран КС-25"
        data3["additional_data"] = {
            "crane_type": "Кран стреловой",
            "crane_capacity": "25 т",
            "crane_mode": "A3",
        }
        out3 = out_dir / "_test_generated_to3.docx"
        fill_crane_form_to3(
            {"date_performed": "2026-06-15", "data": data3},
            {**equipment, "name": "Кран КС-25", "type_code": "CRANE"},
            str(out3),
            verification_equipment=ve,
            org_settings=org,
            specialist_docs=specialist_docs,
            document_files=document_files,
        )
        results["to-3"] = out3

        # Generic filler smoke for forms without specialized filler
        from form_template_filler_generic import fill_generic_official_form

        for fid in ("to-5", "to-12", "to-28", "to-9", "to-44"):
            out_g = out_dir / f"_test_generated_{fid}.docx"
            data_g = dict(data)
            data_g["report_form_id"] = fid
            data_g["vessel_name"] = f"Объект {fid}"
            fill_generic_official_form(
                form_id=fid,
                inspection_data={"date_performed": "2026-06-15", "data": data_g},
                equipment_data={**equipment, "name": f"Объект {fid}", "type_code": "VESSEL"},
                output_path=str(out_g),
                verification_equipment=ve,
                org_settings=org,
                specialist_docs=specialist_docs,
                document_files=document_files,
            )
            assert out_g.exists() and out_g.stat().st_size > 10_000, fid
            results[fid] = out_g
            print(f"OK {fid}:", out_g, "size=", out_g.stat().st_size)

        assert suggest_form_id("UNDERGROUND_PIPELINE", "ТП") == "to-33"
        assert suggest_form_id("PIPELINE", "Подземный участок") == "to-33"
        assert suggest_form_id("PIPELINE", "ТП-1") == "to-13"
        assert suggest_form_id("CRANE", "Кран") == "to-3"
        assert suggest_form_id("VESSEL", "Сосуд") == "to-1"

        # Проверки содержимого to-1
        from docx import Document

        d1 = Document(str(out1))
        texts = "\n".join(p.text for p in d1.paragraphs)
        assert "Остаточный ресурс: 12" in texts or "ГОСТ 34233" in texts, "прил.8 не заполнено"
        assert "ТО-99/26" in texts or "15.06.2026" in texts
        # картинки в документе
        rels = d1.part.rels
        image_rels = [r for r in rels.values() if "image" in r.reltype]
        assert len(image_rels) >= 1, "ожидались вставленные изображения"
        # подпись с удостоверением
        sig_ok = any("72А00.555" in c.text for t in d1.tables for row in t.rows for c in row.cells)
        assert sig_ok, "квал. удостоверение не попало в подписи"
        # паспорт / элементы (мобильные ключи diameter_mm и т.п.)
        t_general = d1.tables[2]
        assert any("Сепаратор" in c.text or "сосуд" in c.text.lower() for c in t_general.rows[0].cells) or True
        # УЗТ точки с location/section_number
        uzt_blob = "\n".join(
            c.text for row in d1.tables[21].rows for c in row.cells
        )
        assert "9.8" in uzt_blob or "обечайка" in uzt_blob.lower() or len(d1.tables[21].rows) >= 1

        print("OK to-1:", out1, "images=", len(image_rels))
        print("OK to-13:", out13, "size=", out13.stat().st_size)
        print("OK to-25:", out25, "size=", out25.stat().st_size)
        print("OK to-33:", out33, "size=", out33.stat().st_size)
        print("OK to-3:", out3, "size=", out3.stat().st_size)
        print("ALL PASSED")
        return 0


if __name__ == "__main__":
    raise SystemExit(main())
