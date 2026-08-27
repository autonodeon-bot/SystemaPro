from docx import Document
from lxml import etree
from docx.oxml.ns import qn

d = Document('/app/report_forms/to-1.docx')
# count content controls / textboxes
root = d.element
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
sdts = root.findall('.//w:sdt', ns)
txbx = root.findall('.//w:txbxContent', ns)
print('sdt', len(sdts), 'txbx', len(txbx))
# check filled general table XML for R0 cell1
f = Document('/app/reports/TECHNICAL_REPORT_45dcdf9a-aef7-44cc-9846-a2d6a7ef3acd_20260710_112101.docx')
cell = f.tables[2].rows[0].cells[1]
print('cell text repr:', repr(cell.text))
print('cell xml snippet:', etree.tostring(cell._tc, encoding='unicode')[:800])
# vessel elements sample from data via filler mapping issue
print('T3 R2 cells:', [c.text for c in f.tables[3].rows[2].cells])
