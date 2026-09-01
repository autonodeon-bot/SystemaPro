#!/usr/bin/env python3
from __future__ import annotations

import sys
import tempfile
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, "/app")
from form_template_filler import (
    fill_vessel_form_to1,
    _build_context,
    _fill_main_report,
    _fill_paragraph_blanks,
    _iter_all_paragraphs,
    _paragraph_plain_text,
    _norm_ws,
    _set_paragraph_text,
)
from report_forms_registry import resolve_form_path


def dump_matches(doc, label):
    print("===", label, "===")
    n_all = 0
    n = 0
    for p in _iter_all_paragraphs(doc):
        n_all += 1
        t = _paragraph_plain_text(p)
        if not t.strip():
            continue
        low = t.lower()
        if (
            "фактическое" in low
            or "техническое состояние объекта" in low
            or ("зав" in low and "инв" in low and "____" in t)
            or "выводы по результатам" in low
        ):
            n += 1
            print(f"[{n}] runs={len(p.runs)} plain={t[:180]!r}")
            print(f"    p.text={((p.text or '')[:80])!r}")
    print("total paras", n_all, "matches", n)


def main():
    path = resolve_form_path("to-1")
    raw = Document(str(path))
    dump_matches(raw, "RAW")

    # count body.iter vs iter_all
    body_n = sum(1 for _ in raw.element.body.iter(qn("w:p")))
    iter_n = sum(1 for _ in _iter_all_paragraphs(raw))
    print("body.iter paragraphs", body_n, "iter_all", iter_n)

    data = {
        "report_form_id": "to-1",
        "vessel_name": "Аппарат тестовый А-1",
        "serial_number": "SN-100",
        "reg_number": "Р-55",
        "inventory_number": "ИНВ-9",
        "orientation": "vertical",
        "geometry": {"kind": "vessel", "orientation": "vertical", "shell_count": 2},
    }
    inspection = {"data": data, "date_performed": "2026-08-28"}
    equipment = {
        "name": "Аппарат тестовый А-1",
        "serial_number": "SN-100",
        "attributes": {"registration_number": "Р-55", "inventory_number": "ИНВ-9"},
    }
    ctx = _build_context(inspection, equipment, [], {}, [], {})
    print("CTX", ctx["device_name"], ctx["serial"], ctx["reg_no"], ctx["inv_no"])

    out1 = Path(tempfile.gettempdir()) / "_dbg2_main.docx"
    shutil.copy(path, out1)
    d1 = Document(str(out1))
    # manual match test
    hit = None
    for p in _iter_all_paragraphs(d1):
        t = _paragraph_plain_text(p)
        if "фактическое значение параметров" in t.lower():
            hit = p
            break
    print("MANUAL HIT", "YES" if hit else "NO")
    if hit:
        before = _paragraph_plain_text(hit)[:100]
        _set_paragraph_text(
            hit,
            f"Фактическое значение параметров, определяющих состояние эксплуатации сосуда "
            f"работающего под давлением – {ctx['device_name']}, зав.№ {ctx['serial']}, "
            f"рег.№ {ctx['reg_no']}, инв.№ {ctx['inv_no']}, удовлетворяют требованиям нормативных документов.",
        )
        after = _paragraph_plain_text(hit)[:160]
        print("BEFORE", before)
        print("AFTER", after)

    d1 = Document(str(out1))
    _fill_main_report(d1, ctx)
    dump_matches(d1, "AFTER MAIN")
    d1.save(str(out1))

    out = Path(tempfile.gettempdir()) / "_dbg2_full.docx"
    fill_vessel_form_to1(
        inspection_data=inspection,
        equipment_data=equipment,
        output_path=str(out),
        verification_equipment=[],
        org_settings={},
        specialist_docs=[],
        document_files=[],
        ndt_methods=[],
    )
    dump_matches(Document(str(out)), "AFTER FULL")


if __name__ == "__main__":
    main()
