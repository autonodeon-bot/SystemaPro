from docx import Document
# old programmatic
old = Document('/app/reports/TECHNICAL_REPORT_45dcdf9a-aef7-44cc-9846-a2d6a7ef3acd_20260710_110210.docx')
print('OLD tables', len(old.tables), 'paras', len(old.paragraphs))
texts = [p.text.strip() for p in old.paragraphs if p.text.strip()]
for t in texts:
    if any(k in t.lower() for k in ['толщин','узт','тверд','вик','сепаратор','1075','0,8','замеров']):
        print('P:', t[:120])
# sample some tables with data
for ti, t in enumerate(old.tables[:15]):
    nonempty = 0
    sample = []
    for ri, row in enumerate(t.rows[:6]):
        vals = [c.text.strip()[:40] for c in row.cells]
        if any(vals[1:] if len(vals)>1 else vals):
            nonempty += 1
            if len(sample)<2: sample.append((ri, vals))
    if nonempty:
        print(f'T{ti} nonempty~{nonempty}', sample[:1])
