"""Формирование DOCX для протоколов, созданных только в мобильном приложении."""

from __future__ import annotations

import io
import json
from typing import Any, Dict, Optional

from docx import Document


def _add_kv(doc: Document, label: str, value: Any) -> None:
    if value is None or value == "":
        return
    if isinstance(value, (dict, list)):
        value = json.dumps(value, ensure_ascii=False, indent=2)
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(str(value))


def _render_custom_template(doc: Document, payload: Dict[str, Any]) -> None:
    structure = payload.get("structure") or []
    values = payload.get("values") or {}
    if not isinstance(structure, list):
        return
    for block in structure:
        if not isinstance(block, dict):
            continue
        btype = block.get("block_type") or ""
        label = block.get("label") or ""
        key = block.get("field_key") or ""
        val = values.get(key) if key else None

        if btype == "section_header":
            doc.add_heading(str(label), level=2)
            continue
        if btype in ("text_field", "textarea", "date_field", "number_field", "instruments_field"):
            _add_kv(doc, label, val)
            continue
        if btype == "table":
            rows = val if isinstance(val, list) else []
            doc.add_paragraph().add_run(label).bold = True
            cols = block.get("columns") or []
            if not cols or not rows:
                doc.add_paragraph("—")
                continue
            table = doc.add_table(rows=1 + len(rows), cols=len(cols))
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for ci, col in enumerate(cols):
                hdr[ci].text = str(col.get("label") or col.get("key") or "")
            for ri, row in enumerate(rows):
                if not isinstance(row, dict):
                    continue
                for ci, col in enumerate(cols):
                    ck = str(col.get("key") or "")
                    table.rows[ri + 1].cells[ci].text = str(row.get(ck) or "")
            continue
        if btype == "photo_section":
            paths = val if isinstance(val, list) else []
            _add_kv(doc, label, f"файлов: {len(paths)} (приложите фото из мобильного архива)" if paths else "—")
            continue
        if btype == "checkbox_list":
            if isinstance(val, dict):
                parts = [f"{k}: {'да' if v else 'нет'}" for k, v in val.items()]
                _add_kv(doc, label, "; ".join(parts) if parts else "—")
            else:
                _add_kv(doc, label, val)
            continue
        if btype == "signature":
            _add_kv(doc, label, val or "—")
            continue
        _add_kv(doc, label, val)


def _render_flat_protocol(doc: Document, title: str, payload: Dict[str, Any]) -> None:
    doc.add_heading(title, level=1)
    order = [
        ("date", "Дата"),
        ("location", "Место проведения"),
        ("object_name", "Объект"),
        ("customer", "Заказчик"),
        ("executor", "Исполнитель"),
        ("devices", "Средства контроля"),
        ("norm_doc", "Нормативный документ"),
    ]
    for k, lab in order:
        if k in payload:
            _add_kv(doc, lab, payload.get(k))
    methods = payload.get("selected_methods")
    if methods:
        _add_kv(doc, "Методы НК", ", ".join(str(x) for x in methods))

    vik = payload.get("vik_defects")
    if isinstance(vik, list) and vik:
        doc.add_heading("Дефекты ВИК", level=2)
        for i, d in enumerate(vik, 1):
            if not isinstance(d, dict):
                continue
            doc.add_paragraph().add_run(f"{i}. ").bold = True
            inner = d.get("defect_type") or d.get("type")
            _add_kv(doc, "Тип", inner)
            _add_kv(doc, "Место", d.get("location"))
            _add_kv(doc, "Размер", d.get("size"))
            _add_kv(doc, "Описание", d.get("description"))

    uzt = payload.get("uzt_measurements")
    if isinstance(uzt, list) and uzt:
        doc.add_heading("Замеры УЗТ", level=2)
        table = doc.add_table(rows=1 + len(uzt), cols=5)
        table.style = "Table Grid"
        hdr = ["№ точки", "Место", "Ном., мм", "Факт., мм", "Коммент."]
        for i, h in enumerate(hdr):
            table.rows[0].cells[i].text = h
        for ri, m in enumerate(uzt):
            if not isinstance(m, dict):
                continue
            r = table.rows[ri + 1].cells
            r[0].text = str(m.get("section_number") or m.get("sectionNumber") or "")
            r[1].text = str(m.get("location") or "")
            r[2].text = str(m.get("nominal_thickness") or m.get("nominalThickness") or "")
            r[3].text = str(m.get("thickness") or m.get("measuredThickness") or m.get("measured_thickness") or "")
            r[4].text = str(m.get("comment") or m.get("notes") or "")

    photos = []
    for pk in ("vik_photos", "uzt_photos"):
        p = payload.get(pk)
        if isinstance(p, list):
            photos.extend(str(x) for x in p)
    if photos:
        _add_kv(doc, "Вложенные фото (локальные пути с устройства)", f"{len(photos)} файл(ов)")


def build_standalone_protocol_docx(
    *,
    title: str,
    kind: str,
    template_name: Optional[str],
    payload: Dict[str, Any],
) -> bytes:
    doc = Document()
    doc.add_paragraph("Протокол (мобильное приложение «Монитор»)").runs[0].bold = True
    sub = f"Вид: {kind}"
    if template_name:
        sub += f" — {template_name}"
    doc.add_paragraph(sub)

    if kind == "custom_template":
        doc.add_heading(title, level=1)
        _render_custom_template(doc, payload)
    else:
        _render_flat_protocol(doc, title, payload)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
