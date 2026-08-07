"""
Заполнение формы ТО to-25 «Обследование резервуаров (емкости)».

Шаблон содержит протоколы УЗТ (прил.5) и УЗК (прил.6) — заполняем их
данными измерений; схемы и фото вставляем в разделы «Схема контроля».
"""
from __future__ import annotations

import logging
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document

from form_media_helpers import (
    build_attachments_map,
    collect_photo_paths,
    collect_scheme_paths,
    insert_media_block,
)
from form_template_filler import (
    MISSING,
    _ensure_rows,
    _extract_specialists,
    _fill_signatures,
    _fmt_date_ru,
    _set,
    _set_paragraph_text,
)
from report_forms_registry import resolve_form_path
from report_org_settings import load_report_org_settings

logger = logging.getLogger(__name__)


def fill_tank_form_to25(
    inspection_data: Dict[str, Any],
    equipment_data: Dict[str, Any],
    output_path: str,
    verification_equipment: Optional[List[Dict[str, Any]]] = None,
    org_settings: Optional[Dict[str, Any]] = None,
    specialist_docs: Optional[List[Dict[str, Any]]] = None,
    document_files: Optional[List[Dict[str, Any]]] = None,
    find_image: Optional[Any] = None,
) -> str:
    template = resolve_form_path("to-25")
    if template is None or not template.exists():
        raise FileNotFoundError("Шаблон to-25 не найден в backend/report_forms/")

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, out)
    doc = Document(str(out))

    if org_settings is None:
        org_settings = load_report_org_settings()

    data = inspection_data.get("data") or {}
    if not isinstance(data, dict):
        data = {}
    attrs = equipment_data.get("attributes") or {}
    if not isinstance(attrs, dict):
        attrs = {}

    def g(*keys: str, default: Any = MISSING) -> Any:
        for k in keys:
            if k in data and data.get(k) not in (None, ""):
                return data.get(k)
            if k in attrs and attrs.get(k) not in (None, ""):
                return attrs.get(k)
        return default

    date_ru = _fmt_date_ru(inspection_data.get("date_performed")) or datetime.now().strftime("%d.%m.%Y")
    device = str(g("vessel_name", "equipment_device_name", default=equipment_data.get("name") or MISSING))
    serial = str(g("serial_number", default=equipment_data.get("serial_number") or MISSING))
    reg_no = str(g("reg_number", default=MISSING))
    inv_no = str(g("inventory_number", default=MISSING))
    location = str(g("location", default=equipment_data.get("location") or MISSING))
    org_name = str(
        g(
            "organization",
            default=(org_settings.get("customer") or {}).get("legal_name") or MISSING,
        )
    )
    contractor = org_settings.get("contractor") or {}
    contractor_name = contractor.get("legal_name") or ""
    lab = org_settings.get("ndt_lab") or {}
    protocol_no = str(g("protocol_number", "report_number", default="") or "")

    tables = doc.tables
    # Шапки протоколов: T1, T5 (8x3)
    for idx in (1, 5):
        if idx >= len(tables):
            continue
        t = tables[idx]
        if len(t.rows) < 7:
            continue
        _set(t, 0, 0, contractor_name)
        _set(t, 0, 2, org_name)
        _set(t, 2, 2, location)
        _set(t, 4, 0, lab.get("name") or contractor_name)
        _set(t, 4, 2, device)
        _set(t, 6, 0, lab.get("certificate") or "")
        _set(t, 6, 2, f"Зав.№ {serial}, рег.№ {reg_no}, инв.№ {inv_no}")

    # Приборы УЗТ T2, УЗК T6
    ve = verification_equipment or []
    for idx, keys in ((2, ("УЗТ", "UZT", "ТОЛЩИНОМЕР")), (6, ("УЗК", "UZK", "ДЕФЕКТОСКОП"))):
        if idx >= len(tables):
            continue
        matched = [
            eq
            for eq in ve
            if isinstance(eq, dict)
            and any(
                k in f"{eq.get('equipment_type','')} {eq.get('name','')}".upper()
                for k in keys
            )
        ]
        for i, eq in enumerate(matched[:3]):
            r = i + 1
            if r >= len(tables[idx].rows):
                break
            _set(tables[idx], r, 0, f"{i + 1}.")
            _set(tables[idx], r, 1, eq.get("name") or "")
            if len(tables[idx].rows[r].cells) > 2:
                _set(tables[idx], r, 2, eq.get("serial_number") or "")

    # Результаты УЗТ — T3 (широкая таблица)
    if len(tables) > 3:
        _fill_uzt_wide(tables[3], data)

    # Результаты УЗК — T7
    if len(tables) > 7:
        welds = data.get("weld_inspections") or data.get("uzk_results") or []
        if isinstance(welds, list) and welds:
            start = 1
            _ensure_rows(tables[7], start + len(welds))
            for i, w in enumerate(welds):
                if not isinstance(w, dict):
                    continue
                r = start + i
                vals = [
                    w.get("joint") or w.get("weld_number") or "",
                    w.get("defect_number") or "",
                    w.get("area") or "",
                    w.get("depth") or "",
                    w.get("length") or "",
                    w.get("form") or "",
                    w.get("location") or "",
                    w.get("conclusion") or "Годен",
                ]
                for c, v in enumerate(vals):
                    if c < len(tables[7].rows[r].cells):
                        _set(tables[7], r, c, v)
        elif len(tables[7].rows) > 1 and len(tables[7].rows[1].cells) > 7:
            _set(tables[7], 1, 7, "Дефектов не обнаружено")

    specs = _extract_specialists(data, specialist_docs or [])
    for idx in (0, 4, 8):
        if idx < len(tables):
            _fill_signatures(tables[idx], {"specialists": specs})

    for p in doc.paragraphs:
        text = p.text
        if not text:
            continue
        stripped = text.strip()
        if stripped.startswith("№") and "от" in stripped and "г." in stripped:
            no_part = protocol_no if protocol_no else "_________"
            _set_paragraph_text(p, f"№ {no_part} от {date_ru} г.")
        elif "утонений" in text.lower() and "______" in text:
            _set_paragraph_text(
                p,
                text.replace("___________________", device, 1).replace("______", device, 1),
            )

    attachments = build_attachments_map(document_files)
    insert_media_block(
        doc,
        "Схема контроля",
        collect_scheme_paths(data, attachments),
        find_image=find_image,
        max_items=6,
    )
    insert_media_block(
        doc,
        "Результаты контроля",
        collect_photo_paths(data, attachments),
        find_image=find_image,
        max_items=12,
    )

    doc.save(str(out))
    logger.info("Форма to-25 заполнена: %s", out)
    return str(out)


def _fill_uzt_wide(table, data: Dict[str, Any]) -> None:
    points = data.get("thickness_measurements") or data.get("thicknessMeasurements") or []
    if not isinstance(points, list) or not points:
        return
    start = 1
    _ensure_rows(table, start + len(points))
    for i, p in enumerate(points):
        if not isinstance(p, dict):
            continue
        r = start + i
        if r >= len(table.rows):
            break
        vals = [
            str(i + 1),
            p.get("element") or p.get("zone") or "",
            p.get("point_number") or p.get("number") or (i + 1),
            p.get("thickness") or p.get("value") or "",
        ]
        for c, v in enumerate(vals):
            if c < len(table.rows[r].cells):
                _set(table, r, c, v)
