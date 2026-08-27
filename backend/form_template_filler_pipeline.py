"""
Заполнение формы ТО to-13 «Обследование технологических трубопроводов».
"""
from __future__ import annotations

import logging
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.table import Table

from form_media_helpers import (
    build_attachments_map,
    collect_photo_paths,
    collect_scheme_paths,
    insert_media_block,
)
from form_template_filler import (
    MISSING,
    _fill_signatures,
    _fmt_date_ru,
    _replace_underscores,
    _set,
    _set_paragraph_text,
    _ensure_rows,
    finalize_official_form,
)
from report_forms_registry import resolve_form_path
from report_org_settings import load_report_org_settings

logger = logging.getLogger(__name__)
_BLANK_RE = re.compile(r"_+")


def fill_pipeline_form_to13(
    inspection_data: Dict[str, Any],
    equipment_data: Dict[str, Any],
    output_path: str,
    verification_equipment: Optional[List[Dict[str, Any]]] = None,
    org_settings: Optional[Dict[str, Any]] = None,
    specialist_docs: Optional[List[Dict[str, Any]]] = None,
    document_files: Optional[List[Dict[str, Any]]] = None,
    find_image: Optional[Any] = None,
) -> str:
    template = resolve_form_path("to-13")
    if template is None or not template.exists():
        raise FileNotFoundError("Шаблон to-13 не найден в backend/report_forms/")

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, out)
    doc = Document(str(out))

    if org_settings is None:
        org_settings = load_report_org_settings()

    data = inspection_data.get("data") or {}
    if not isinstance(data, dict):
        data = {}
    attrs = equipment_data.get("attributes") or {}
    if not isinstance(attrs, dict):
        attrs = {}

    def g(*keys: str, default: Any = MISSING) -> Any:
        for k in keys:
            if k in data and data.get(k) not in (None, ""):
                return data.get(k)
            if k in attrs and attrs.get(k) not in (None, ""):
                return attrs.get(k)
        return default

    date_ru = _fmt_date_ru(inspection_data.get("date_performed")) or datetime.now().strftime("%d.%m.%Y")
    device = str(g("vessel_name", "equipment_device_name", default=equipment_data.get("name") or MISSING))
    serial = str(g("serial_number", default=equipment_data.get("serial_number") or MISSING))
    reg_no = str(g("reg_number", default=MISSING))
    inv_no = str(g("inventory_number", "inv_number", default=MISSING))
    location = str(g("location", default=equipment_data.get("location") or MISSING))
    pressure = str(g("working_pressure", default=MISSING))
    org_name = str(
        g(
            "organization",
            "customer_name",
            default=(org_settings.get("customer") or {}).get("legal_name") or MISSING,
        )
    )
    contractor = (org_settings.get("contractor") or {})
    contractor_name = contractor.get("legal_name") or contractor.get("name") or ""

    tables = doc.tables

    # T1 — заказчик
    if len(tables) > 1:
        _set(tables[1], 0, 1, org_name)
        _set(tables[1], 1, 1, location)
    # T2 — исполнитель
    if len(tables) > 2:
        _set(tables[2], 0, 1, contractor_name)
        lab = org_settings.get("ndt_lab") or {}
        _set(tables[2], 1, 1, lab.get("name") or contractor_name)
        _set(tables[2], 2, 1, lab.get("certificate") or "")

    # T3 — специалисты
    if len(tables) > 3:
        _fill_specialists_table(tables[3], data, specialist_docs or [])

    # T4 — приборы
    if len(tables) > 4:
        _fill_instruments_generic(tables[4], verification_equipment or [])

    # T5 — объект
    if len(tables) > 5:
        _set(tables[5], 0, 1, device)
        if len(tables[5].rows) > 1:
            _set(tables[5], 1, 1, serial)
        if len(tables[5].rows) > 2:
            _set(tables[5], 2, 1, reg_no)
        if len(tables[5].rows) > 3:
            _set(tables[5], 3, 1, pressure)

    # T7 — параметры
    if len(tables) > 7:
        mapping = [
            (0, device),
            (1, serial),
            (2, reg_no),
            (3, inv_no),
            (4, pressure),
            (5, g("diameter", "pipe_diameter", default=MISSING)),
            (6, g("wall_thickness", "thickness", default=MISSING)),
            (7, g("working_medium", "medium", default=MISSING)),
            (8, g("shell_material", "material", default=MISSING)),
            (9, g("commissioning_year", default=MISSING)),
        ]
        for row, val in mapping:
            if row < len(tables[7].rows) and len(tables[7].rows[row].cells) > 1:
                _set(tables[7], row, 1, val)

    # УЗТ результаты — ищем таблицу с «толщина»
    _fill_uzt_in_pipeline(doc, data)
    _fill_paragraphs_pipeline(doc, date_ru, device, serial, reg_no, inv_no, pressure, org_name, g)

    attachments = build_attachments_map(document_files)
    schemes = collect_scheme_paths(data, attachments)
    photos = collect_photo_paths(data, attachments)
    insert_media_block(doc, "Схема контроля", schemes, find_image=find_image, max_items=8)
    insert_media_block(doc, "Результаты контроля", photos, find_image=find_image, max_items=12)

    # Подписи в типичных таблицах
    for t in tables:
        if t.rows and "Контроль провел" in (t.rows[0].cells[0].text or ""):
            ctx = {
                "specialists": _specs_from(data, specialist_docs or []),
            }
            _fill_signatures(t, ctx)

    finalize_official_form(doc, "to-13")

    doc.save(str(out))
    logger.info("Форма to-13 заполнена: %s", out)
    return str(out)


def _specs_from(data: Dict[str, Any], specialist_docs: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    from form_template_filler import _extract_specialists

    return _extract_specialists(data, specialist_docs)


def _fill_specialists_table(
    table: Table, data: Dict[str, Any], specialist_docs: List[Dict[str, Any]]
) -> None:
    specs = _specs_from(data, specialist_docs)
    start = 1
    _ensure_rows(table, start + max(len(specs), 1))
    for i, s in enumerate(specs):
        r = start + i
        if r >= len(table.rows):
            break
        cols = len(table.rows[r].cells)
        _set(table, r, 0, str(i + 1))
        if cols > 1:
            _set(table, r, 1, s.get("name") or "")
        if cols > 2:
            _set(table, r, 2, s.get("role") or "НК")
        if cols > 3:
            _set(table, r, 3, s.get("cert") or "")
        if cols > 4:
            _set(table, r, 4, "")


def _fill_instruments_generic(table: Table, ve: List[Dict[str, Any]]) -> None:
    start = 1
    items = ve[: max(0, len(table.rows) - 1)] or []
    for i, eq in enumerate(items):
        if not isinstance(eq, dict):
            continue
        r = start + i
        if r >= len(table.rows):
            break
        _set(table, r, 0, str(i + 1))
        if len(table.rows[r].cells) > 1:
            _set(table, r, 1, eq.get("name") or "")
        if len(table.rows[r].cells) > 2:
            _set(table, r, 2, eq.get("serial_number") or eq.get("factory_number") or "")
        if len(table.rows[r].cells) > 3:
            _set(table, r, 3, eq.get("verification_certificate_number") or "")
        if len(table.rows[r].cells) > 4:
            _set(table, r, 4, eq.get("equipment_type") or "")


def _fill_uzt_in_pipeline(doc: Document, data: Dict[str, Any]) -> None:
    points = data.get("thickness_measurements") or data.get("thicknessMeasurements") or []
    if not isinstance(points, list) or not points:
        return
    # Ищем таблицу с заголовком про толщину / № точки
    for table in doc.tables:
        if not table.rows:
            continue
        header = " ".join(c.text.lower() for c in table.rows[0].cells)
        if "толщин" not in header and "точки" not in header and "№ точки" not in header:
            continue
        start = 1
        _ensure_rows(table, start + len(points))
        for i, p in enumerate(points):
            if not isinstance(p, dict):
                continue
            r = start + i
            if r >= len(table.rows):
                break
            vals = [
                str(i + 1),
                p.get("element") or p.get("zone") or "",
                p.get("point_number") or p.get("number") or (i + 1),
                p.get("thickness") or p.get("value") or "",
                p.get("min_allowed_thickness") or "",
                p.get("nominal_thickness") or "",
            ]
            for c, v in enumerate(vals):
                if c < len(table.rows[r].cells):
                    _set(table, r, c, v)
        break


def _fill_paragraphs_pipeline(
    doc: Document,
    date_ru: str,
    device: str,
    serial: str,
    reg_no: str,
    inv_no: str,
    pressure: str,
    org_name: str,
    g,
) -> None:
    protocol_no = str(g("protocol_number", "report_number", default="") or "")
    for p in doc.paragraphs:
        text = p.text
        if not text:
            continue
        new_text = text
        stripped = text.strip()
        if stripped.startswith("№") and "от" in stripped and "г." in stripped:
            no_part = protocol_no if protocol_no else "_________"
            new_text = f"№ {no_part} от {date_ru} г."
        elif "Разрешенное давление" in text:
            new_text = re.sub(r"_+\s*Мпа", f"{pressure} Мпа", text, flags=re.IGNORECASE)
            if pressure != MISSING and "____" in new_text:
                new_text = new_text.replace("____", str(pressure), 1)
        elif "зав.№" in text and ("рег.№" in text or "инв.№" in text):
            new_text = _replace_underscores(text, [device, serial, reg_no, inv_no])
        elif "ВЫВОД:" in text:
            conclusion = str(g("suitability_conclusion", "conclusion", default="соответствует") or "")
            new_text = _replace_underscores(text, [device, serial, reg_no, inv_no, conclusion])
        elif "договором между" in text.lower():
            contract = str(g("contract_number", default="") or "")
            contract_date = _fmt_date_ru(g("contract_date", default="")) or ""
            if contract or org_name:
                new_text = _replace_underscores(
                    text,
                    [org_name, contract_date or "____", contract or "____"],
                )
        else:
            continue
        if new_text != text:
            _set_paragraph_text(p, new_text)
