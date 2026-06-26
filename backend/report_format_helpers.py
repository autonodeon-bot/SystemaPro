"""
Вспомогательные функции форматирования отчётов ТО/ЭПБ (таблицы, оглавление, подписи).
"""
from __future__ import annotations

from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from report_org_settings import TO_TOC_ITEMS


def append_technical_toc(doc: Document) -> None:
    doc.add_heading("СОДЕРЖАНИЕ", level=1)
    for item in TO_TOC_ITEMS:
        doc.add_paragraph(item)
    doc.add_page_break()


def customer_table_rows(
    org_settings: Dict[str, Any],
    g: Callable[..., Any],
    org: str,
    location: str,
) -> List[Tuple[str, str]]:
    customer = (org_settings or {}).get("customer") or {}
    return [
        (
            "Наименование организации",
            g("customer_legal_name", default=customer.get("legal_name") or org or "—"),
        ),
        (
            "Организационно-правовая форма",
            g("customer_legal_form", default=customer.get("legal_form") or "—"),
        ),
        (
            "Место нахождения",
            g("customer_address", default=customer.get("address") or location or "—"),
        ),
        (
            "Телефон / факс",
            g("customer_phone", default=customer.get("phone") or "—"),
        ),
        (
            "Руководитель",
            g("customer_director", default=customer.get("director") or "—"),
        ),
        (
            "Структурное подразделение",
            g("customer_department", default=customer.get("department") or "—"),
        ),
        (
            "Место нахождения подразделения",
            g("customer_department_address", default=customer.get("department_address") or location or "—"),
        ),
        (
            "Телефон / факс подразделения",
            g("customer_department_phone", default=customer.get("department_phone") or "—"),
        ),
        (
            "Руководитель подразделения",
            g("customer_department_head", default=customer.get("department_head") or "—"),
        ),
    ]


def contractor_table_rows(
    org_settings: Dict[str, Any],
    g: Callable[..., Any],
    contractor: str,
    director_title: str,
    director_name: str,
) -> List[Tuple[str, str]]:
    c = (org_settings or {}).get("contractor") or {}
    name = g("contractor_name", default=c.get("short_name") or c.get("name") or contractor or "—")
    return [
        ("Наименование организации:", str(name)),
        (
            "Организационно-правовая форма организации:",
            g("contractor_legal_form", default=c.get("legal_form") or "Общество с ограниченной ответственностью"),
        ),
        (
            "Юридический адрес:",
            g("contractor_address", default=c.get("address") or "—"),
        ),
        (
            "Лицензия:",
            g("contractor_license", default=c.get("license") or "—"),
        ),
        (
            "Свидетельство об аттестации лаборатории:",
            g("contractor_certificate", default=c.get("certificate") or "—"),
        ),
        (
            "Руководитель экспертной организации:",
            f"{g('director_title', default=director_title)} {g('director_name', default=director_name)}".strip(),
        ),
        (
            "Телефон / e-mail:",
            g(
                "contractor_contacts",
                default=" / ".join(
                    x for x in [c.get("phone"), c.get("email")] if x
                )
                or "—",
            ),
        ),
    ]


def normative_bullets(
    org_settings: Dict[str, Any],
    g: Callable[..., Any],
    fallback_text: str,
) -> List[str]:
    custom = g("normative_documents")
    if isinstance(custom, list) and custom:
        return [str(x) for x in custom if str(x).strip()]
    from_settings = (org_settings or {}).get("normative_documents")
    if isinstance(from_settings, list) and from_settings:
        return [str(x) for x in from_settings if str(x).strip()]
    explicit = g("normative_base")
    if explicit and str(explicit).strip():
        return [str(explicit).strip()]
    return [fallback_text]


def work_basis_text(org_settings: Dict[str, Any], g: Callable[..., Any]) -> str:
    return str(
        g(
            "basis",
            "work_basis",
            default=(org_settings or {}).get("work_basis")
            or "Работы по техническому диагностированию проведены согласно договору.",
        )
    )


def doc_meta_extended(
    num: str,
    docs_dict: Dict[str, Any],
    docs_info: Dict[str, Any],
) -> Tuple[Optional[bool], str, str, str]:
    """present, number, date, pages."""
    num_key = str(num)
    present = None
    doc_number = ""
    doc_date = ""
    pages = ""
    if isinstance(docs_dict, dict) and num_key in docs_dict:
        val = docs_dict.get(num_key)
        if isinstance(val, dict):
            present = val.get("present")
            if present is None:
                present = val.get("has") if val.get("has") is not None else val.get("value")
            doc_number = str(val.get("number") or val.get("doc_number") or "")
            doc_date = str(val.get("date") or val.get("doc_date") or "")
            pages = str(val.get("pages") or val.get("page_count") or val.get("volume") or "")
        else:
            if isinstance(val, str):
                present = val.strip().lower() in ("true", "1", "yes", "да")
            else:
                present = bool(val)
    if isinstance(docs_info, dict) and num_key in docs_info:
        info = docs_info.get(num_key) or {}
        if isinstance(info, dict):
            if present is None:
                present = info.get("present")
                if present is None:
                    present = info.get("has") if info.get("has") is not None else info.get("value")
            if not doc_number:
                doc_number = str(info.get("number") or info.get("doc_number") or "")
            if not doc_date:
                doc_date = str(info.get("date") or info.get("doc_date") or "")
            if not pages:
                pages = str(info.get("pages") or info.get("page_count") or info.get("volume") or "")
    return present, doc_number, doc_date, pages


def build_specialist_rows(
    inspection_engineers: List[Dict[str, Any]],
    ndt_methods: List[Dict[str, Any]],
    normalize_method: Callable[[str], str],
) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    seen = set()
    details: Dict[str, Dict[str, Any]] = {}

    if isinstance(inspection_engineers, list):
        for ie in inspection_engineers:
            if not isinstance(ie, dict):
                continue
            name = (ie.get("full_name") or "").strip()
            if not name:
                continue
            method = normalize_method(ie.get("method") or "")
            cert_num = (ie.get("certificate_number") or "").strip()
            valid_until = (ie.get("valid_until") or "").strip()
            area = (ie.get("role") or ie.get("expert_area") or ie.get("certification_area") or "").strip()
            md = details.setdefault(name, {"methods": [], "cert": "", "valid": "", "area": ""})
            if method and method not in md["methods"]:
                md["methods"].append(method)
            if cert_num:
                md["cert"] = cert_num
            if valid_until:
                md["valid"] = valid_until
            if area:
                md["area"] = area

    for m in ndt_methods or []:
        name = (m.get("inspector_name") or "").strip()
        if not name:
            continue
        md = details.setdefault(name, {"methods": [], "cert": "", "valid": "", "area": ""})
        method = normalize_method(m.get("method_code") or m.get("method_name") or "")
        if method and method not in md["methods"]:
            md["methods"].append(method)
        cert_num = (m.get("certificate_number") or m.get("certification_number") or "").strip()
        if cert_num and not md["cert"]:
            md["cert"] = cert_num
        level = (m.get("inspector_level") or "").strip()
        if level and not md["area"]:
            md["area"] = f"Уровень {level}"

    for name, md in details.items():
        if name in seen:
            continue
        seen.add(name)
        rows.append(
            {
                "name": name,
                "cert": md.get("cert") or "—",
                "valid_until": md.get("valid") or "—",
                "area": ", ".join(md.get("methods") or []) or md.get("area") or "—",
            }
        )
    return rows


def find_signature_image(
    specialist_docs: Optional[List[Dict[str, Any]]],
    inspector_name: str,
) -> Optional[str]:
    if not specialist_docs or not inspector_name:
        return None
    target = inspector_name.strip().lower()
    for block in specialist_docs:
        if not isinstance(block, dict):
            continue
        name = (block.get("inspector_name") or "").strip().lower()
        if name != target:
            continue
        for cert in block.get("certifications") or []:
            if not isinstance(cert, dict):
                continue
            path = cert.get("scan_file_path") or cert.get("signature_scan_path")
            if isinstance(path, str) and path.strip():
                return path.strip()
    return None


def add_protocol_header_from_settings(
    doc: Document,
    org_settings: Dict[str, Any],
    *,
    org: str,
    device_name: str,
    serial: str,
    location: str,
    date_perf_ru: str,
    normative_text: str,
) -> None:
    labels = (org_settings or {}).get("appendix_protocol_header") or {}
    ht = doc.add_table(rows=5, cols=2)
    ht.style = "Table Grid"
    rows = [
        (labels.get("customer_label") or "Заказчик:", str(org)),
        (labels.get("object_label") or "Объект контроля:", f"{device_name} зав.№ {serial}"),
        (labels.get("location_label") or "Место проведения контроля:", str(location)),
        (labels.get("date_label") or "Дата проведения контроля:", date_perf_ru),
        (labels.get("ntd_label") or "НТД, по которой выполнен контроль:", normative_text),
    ]
    for i, (lbl, val) in enumerate(rows):
        ht.rows[i].cells[0].text = lbl
        ht.rows[i].cells[1].text = val
        try:
            ht.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        except Exception:
            pass
    doc.add_paragraph()


def add_inspector_signature_block(
    doc: Document,
    inspector_name: str,
    specialist_docs: Optional[List[Dict[str, Any]]],
    resolve_image_path: Callable[[str], Optional[str]],
    *,
    position: str = "Дефектоскопист II уровня",
) -> None:
    name = inspector_name or "—"
    doc.add_paragraph()
    sig_tbl = doc.add_table(rows=3, cols=2)
    sig_tbl.style = "Table Grid"
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_tbl.rows[0].cells[0].text = "Контроль провел, заключение выдал:"
    sig_tbl.rows[0].cells[1].text = ""
    sig_tbl.rows[1].cells[0].text = position
    sig_tbl.rows[1].cells[1].text = str(name)
    sig_tbl.rows[2].cells[0].text = "Подпись"
    sig_cell = sig_tbl.rows[2].cells[1]
    sig_cell.text = ""
    image_path = find_signature_image(specialist_docs, name)
    if image_path:
        resolved = resolve_image_path(image_path) or image_path
        try:
            p = Path(resolved)
            if p.exists() and p.suffix.lower() in (".png", ".jpg", ".jpeg", ".webp", ".bmp"):
                sig_cell.paragraphs[0].add_run().add_picture(str(p), width=Inches(1.6))
            else:
                sig_cell.text = "__________________"
        except Exception:
            sig_cell.text = "__________________"
    else:
        sig_cell.text = "__________________"
    for row in sig_tbl.rows:
        try:
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        except Exception:
            pass
    doc.add_paragraph()


def tech_characteristic_rows(
    g: Callable[..., Any],
    device_name: str,
    purpose_default: Optional[str],
    equipment_data: Dict[str, Any],
) -> List[Tuple[str, str]]:
    commissioning = g("commissioning_year", default=str(equipment_data.get("commissioning_date") or "—"))
    rows = [
        ("Наименование объекта", device_name),
        ("Назначение", g("purpose", "vessel_purpose", default=purpose_default or "—")),
        (
            "Условное обозначение",
            g("designation", "conditional_designation", default="—"),
        ),
        ("Наименование завода-изготовителя", g("manufacturer", default="—")),
        ("Год изготовления", g("manufacturing_year", "manufacture_year", default="—")),
        ("Год ввода в эксплуатацию", commissioning),
        ("Рабочее давление, МПа", g("working_pressure", default="—")),
        ("Расчетное давление, МПа", g("design_pressure", default="—")),
        ("Пробное давление гидравлического испытания, МПа", g("test_pressure", "hydraulic_test_pressure", default="—")),
        ("Допустимая рабочая температура стенки, ℃", g("working_temperature", default="—")),
        ("Расчетная температура стенки, ℃", g("design_temperature", default="—")),
        ("Наименование рабочей среды", g("working_medium", default="—")),
        ("Характеристика рабочей среды", g("medium_characteristics", default="—")),
        ("Группа сосуда", g("vessel_group", default="—")),
        ("Группа рабочей среды", g("medium_group", default="—")),
        ("Прибавка для компенсации коррозии, мм", g("corrosion_allowance", default="—")),
    ]
    return rows


def previous_inspection_rows(g: Callable[..., Any]) -> List[Dict[str, str]]:
    prev = g("previous_inspections", default=[])
    if isinstance(prev, list) and prev:
        rows = []
        for item in prev:
            if not isinstance(item, dict):
                continue
            kind = str(item.get("kind") or item.get("type") or "Техническое диагностирование")
            date = str(item.get("date") or "")
            report = str(item.get("report") or item.get("report_number") or item.get("result") or "")
            result = str(item.get("result") or "")
            text = report
            if date and report:
                text = f"{report} от {date}"
            elif date:
                text = f"от {date}"
            if result and result not in text:
                text = f"{result} / {text}" if text else result
            rows.append({"kind": kind, "result": text or "—"})
        if rows:
            return rows
    legacy = g("previous_inspection_result", default="—")
    return [{"kind": "Техническое диагностирование", "result": str(legacy)}]


def format_work_result_row(
    idx: int,
    work_name: str,
    conclusion: str,
    method_code: str,
    appendix_map: Dict[str, int],
) -> str:
    appendix_no = appendix_map.get((method_code or "").upper())
    if appendix_no:
        return f"{conclusion} / Приложение № {appendix_no}"
    return f"{conclusion} / Протокол № {idx}"
