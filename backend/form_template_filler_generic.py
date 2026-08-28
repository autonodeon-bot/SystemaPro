"""
Универсальный заполнитель официальных форм ТО (to-2…to-44).

Специализированные fillers (to-1/3/13/25/33) остаются приоритетными.
Для остальных шаблонов заполняем типовые блоки по эвристикам:
заказчик / исполнитель / специалисты / приборы / паспорт объекта /
УЗТ / схемы и фото / даты в абзацах.
"""
from __future__ import annotations

import logging
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.table import Table

from form_media_helpers import (
    build_attachments_map,
    collect_photo_paths,
    collect_scheme_paths,
    insert_media_block,
)
from form_template_filler import (
    MISSING,
    _build_context,
    _ensure_rows,
    _enrich_inspection_data,
    _extract_specialists,
    _fill_hardness_steel_heading,
    _fill_signatures,
    _fmt_date_ru,
    _instrument_full_name,
    _merge_report_instruments,
    _replace_underscores,
    _set,
    apply_ndt_protocol_tables,
    finalize_official_form,
    insert_ndt_layer_schemes,
)
from form_template_filler_common import (
    fill_contractor_table,
    fill_customer_table,
    fill_instruments_table,
    fill_specialists_table,
)
from report_forms_registry import get_form, resolve_form_path
from report_org_settings import load_report_org_settings

logger = logging.getLogger(__name__)

_BLANK_RE = re.compile(r"_+")

# Метка строки (левая ячейка) → ключи данных
_LABEL_KEYS: List[Tuple[Tuple[str, ...], Tuple[str, ...]]] = [
    (("наименование объекта", "наименование сосуда", "наименование трубопровода",
      "наименование подъемного", "наименование подъёмного", "наименование резервуар",
      "наименование котл", "наименование трансформатор", "наименование агрегата",
      "наименование станции", "наименование установки", "наименование оборудования"),
     ("vessel_name", "equipment_device_name", "name")),
    (("заводской номер", "зав. номер", "зав номер", "зав.№", "заводской №"),
     ("serial_number",)),
    (("регистрационн", "учетный номер", "учётный номер", "рег. номер", "рег.№"),
     ("reg_number",)),
    (("инвентарн", "инв. номер", "инв.№", "технологический номер"),
     ("inventory_number", "inv_number")),
    (("местонахожд", "место установ", "адрес объект", "площадка"),
     ("location", "equipment_location")),
    (("изготовител", "завод-изготовител", "организация-изготовител"),
     ("manufacturer",)),
    (("год изготовлен", "год выпуска"), ("manufacture_year",)),
    (("год ввода", "ввод в эксплуатац", "год пуска"), ("commissioning_year",)),
    (("рабочее давление", "давление рабочее"), ("working_pressure",)),
    (("расчетное давление", "расчётное давление", "давление расчет"),
     ("design_pressure",)),
    (("пробное давление", "давление испыта"), ("test_pressure",)),
    (("диаметр",), ("diameter", "pipe_diameter")),
    (("толщина", "толщин"), ("wall_thickness", "thickness")),
    (("рабочая среда", "среда"), ("working_medium", "medium")),
    (("материал", "марка стали"), ("shell_material", "material")),
    (("объем", "объём", "вместимость"), ("volume",)),
    (("температур",), ("working_temperature",)),
    (("назначение",), ("purpose", "appointment")),
    (("грузоподъемн", "грузоподъёмн"), ("lifting_capacity", "capacity")),
    (("длина", "протяженн", "протяжённ"), ("length", "pipeline_length")),
]


def fill_generic_official_form(
    form_id: str,
    inspection_data: Dict[str, Any],
    equipment_data: Dict[str, Any],
    output_path: str,
    verification_equipment: Optional[List[Dict[str, Any]]] = None,
    org_settings: Optional[Dict[str, Any]] = None,
    specialist_docs: Optional[List[Dict[str, Any]]] = None,
    document_files: Optional[List[Dict[str, Any]]] = None,
    find_image: Optional[Any] = None,
    ndt_methods: Optional[List[Dict[str, Any]]] = None,
) -> str:
    """Заполнить любую официальную форму to-N по общим правилам."""
    fid = (form_id or "").strip().lower()
    template = resolve_form_path(fid)
    if template is None or not template.exists():
        raise FileNotFoundError(f"Шаблон {fid} не найден в backend/report_forms/")

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, out)
    doc = Document(str(out))

    if org_settings is None:
        org_settings = load_report_org_settings()

    # Те же нормализация мобильных ключей и подтягивание методов НК, что и в to-1
    inspection_data = dict(inspection_data or {})
    raw_data = inspection_data.get("data")
    if not isinstance(raw_data, dict):
        raw_data = {}
    data = _enrich_inspection_data(raw_data, ndt_methods or [])
    inspection_data["data"] = data

    attrs = equipment_data.get("attributes") or {}
    if not isinstance(attrs, dict):
        attrs = {}

    def g(*keys: str, default: Any = MISSING) -> Any:
        for k in keys:
            if k in data and data.get(k) not in (None, ""):
                return data.get(k)
            if k in attrs and attrs.get(k) not in (None, ""):
                return attrs.get(k)
            if k == "name" and equipment_data.get("name"):
                return equipment_data.get("name")
            if k == "serial_number" and equipment_data.get("serial_number"):
                return equipment_data.get("serial_number")
            if k == "location" and equipment_data.get("location"):
                return equipment_data.get("location")
        return default

    date_ru = _fmt_date_ru(inspection_data.get("date_performed")) or datetime.now().strftime(
        "%d.%m.%Y"
    )
    device = str(
        g("vessel_name", "equipment_device_name", "name", default=equipment_data.get("name") or MISSING)
    )
    serial = str(g("serial_number", default=equipment_data.get("serial_number") or MISSING))
    reg_no = str(g("reg_number", default=MISSING))
    inv_no = str(g("inventory_number", "inv_number", default=MISSING))
    location = str(g("location", "equipment_location", default=equipment_data.get("location") or MISSING))
    org_name = str(
        g(
            "organization",
            "customer_name",
            default=(org_settings.get("customer") or {}).get("legal_name")
            or (org_settings.get("customer") or {}).get("name")
            or MISSING,
        )
    )

    values = {
        "vessel_name": device,
        "equipment_device_name": device,
        "name": device,
        "serial_number": serial,
        "reg_number": reg_no,
        "inventory_number": inv_no,
        "inv_number": inv_no,
        "location": location,
        "equipment_location": location,
        "manufacturer": str(g("manufacturer", default=MISSING)),
        "manufacture_year": str(g("manufacture_year", default=MISSING)),
        "commissioning_year": str(g("commissioning_year", default=MISSING)),
        "working_pressure": str(g("working_pressure", default=MISSING)),
        "design_pressure": str(g("design_pressure", default=MISSING)),
        "test_pressure": str(g("test_pressure", default=MISSING)),
        "diameter": str(g("diameter", "pipe_diameter", default=MISSING)),
        "pipe_diameter": str(g("diameter", "pipe_diameter", default=MISSING)),
        "wall_thickness": str(g("wall_thickness", "thickness", default=MISSING)),
        "thickness": str(g("wall_thickness", "thickness", default=MISSING)),
        "working_medium": str(g("working_medium", "medium", default=MISSING)),
        "medium": str(g("working_medium", "medium", default=MISSING)),
        "shell_material": str(g("shell_material", "material", default=MISSING)),
        "material": str(g("shell_material", "material", default=MISSING)),
        "volume": str(g("volume", default=MISSING)),
        "working_temperature": str(g("working_temperature", default=MISSING)),
        "purpose": str(g("purpose", "appointment", default=MISSING)),
        "appointment": str(g("purpose", "appointment", default=MISSING)),
        "lifting_capacity": str(g("lifting_capacity", "capacity", default=MISSING)),
        "capacity": str(g("lifting_capacity", "capacity", default=MISSING)),
        "length": str(g("length", "pipeline_length", default=MISSING)),
        "pipeline_length": str(g("length", "pipeline_length", default=MISSING)),
        "organization": org_name,
        "customer_name": org_name,
    }

    meta = get_form(fid) or {}
    form_title = meta.get("title") or fid

    attachments = build_attachments_map(document_files)

    # --- таблицы ---
    _fill_org_blocks(doc, org_settings, location, org_name)
    _fill_people_and_instruments(doc, data, specialist_docs or [], verification_equipment, ndt_methods)
    _fill_kv_tables(doc, values)

    # Протоколы НК (ВИК/УЗТ/твёрдость/УЗК/МПК) — так же, как в форме сосудов,
    # но таблицы находятся по заголовкам, а не по фиксированным индексам.
    ndt_done: set = set()
    try:
        ndt_ctx = _build_context(
            inspection_data,
            equipment_data,
            _merge_report_instruments(verification_equipment, data, ndt_methods or []),
            org_settings,
            specialist_docs or [],
            attachments,
        )
        ndt_ctx["find_image"] = find_image
        ndt_ctx["ndt_methods"] = ndt_methods or []
        ndt_done = apply_ndt_protocol_tables(doc, ndt_ctx)
        _fill_hardness_steel_heading(doc, ndt_ctx)
    except Exception:
        logger.exception("generic: не удалось заполнить протокольные таблицы НК")

    # Резервный проход по «широким» таблицам толщинометрии, которых нет в to-1
    _fill_uzt_tables(doc, data, skip_ids=ndt_done)
    _fill_paragraph_blanks(doc, date_ru, device, org_name, form_title, g)

    # подписи
    for t in doc.tables:
        if not t.rows:
            continue
        head = (t.rows[0].cells[0].text or "") if t.rows[0].cells else ""
        if "Контроль провел" in head or "Специалист" in head:
            try:
                _fill_signatures(t, {"specialists": _extract_specialists(data, specialist_docs or [])})
            except Exception:
                pass

    tmp_files: List[str] = []
    finalize_official_form(doc, fid)
    try:
        insert_ndt_layer_schemes(
            doc, data, attachments, find_image, kind=str(data.get("equipment_kind") or fid), tmp_files=tmp_files
        )
    except Exception:
        logger.exception("generic: не удалось вставить слои схем")
        schemes = collect_scheme_paths(data, attachments)
        insert_media_block(doc, "Схема контроля", schemes, find_image=find_image, max_items=8)
        insert_media_block(doc, "Карта контроля", schemes, find_image=find_image, max_items=8)
    photos = collect_photo_paths(data, attachments)
    insert_media_block(doc, "Результаты контроля", photos, find_image=find_image, max_items=12)
    insert_media_block(doc, "Фотографии", photos, find_image=find_image, max_items=12)
    for pth in tmp_files:
        try:
            Path(pth).unlink(missing_ok=True)
        except Exception:
            pass

    doc.save(str(out))
    logger.info("Форма %s заполнена (generic): %s", fid, out)
    return str(out)


def _table_blob(table: Table, max_rows: int = 2) -> str:
    parts: List[str] = []
    for r in table.rows[:max_rows]:
        for c in r.cells:
            parts.append((c.text or "").lower())
    return " ".join(parts)


def _fill_org_blocks(
    doc: Document,
    org_settings: Dict[str, Any],
    location: str,
    org_name: str,
) -> None:
    customer_done = contractor_done = False
    for table in doc.tables:
        if len(table.columns) < 2 or len(table.rows) < 3:
            continue
        blob = _table_blob(table, 3)
        # шапки протоколов 8×3 — пропускаем (заполняются отдельно специализированными)
        if "предприятие-исполнитель" in blob and "заказчик" in blob:
            continue
        look = (table.rows[0].cells[0].text or "").lower()
        if "полное наименование организации" not in look and "наименование организации" not in look:
            continue
        if not customer_done:
            try:
                fill_customer_table(table, org_settings, location=location)
                # если пусто — хотя бы название
                if len(table.rows[0].cells) > 1 and not (table.rows[0].cells[1].text or "").strip():
                    _set(table, 0, 1, org_name)
            except Exception:
                if len(table.rows[0].cells) > 1:
                    _set(table, 0, 1, org_name)
            customer_done = True
            continue
        if not contractor_done:
            try:
                fill_contractor_table(table, org_settings)
            except Exception:
                contractor = org_settings.get("contractor") or {}
                if len(table.rows[0].cells) > 1:
                    _set(table, 0, 1, contractor.get("legal_name") or contractor.get("name") or MISSING)
            contractor_done = True


def _fill_people_and_instruments(
    doc: Document,
    data: Dict[str, Any],
    specialist_docs: List[Dict[str, Any]],
    verification_equipment: Optional[List[Dict[str, Any]]],
    ndt_methods: Optional[List[Dict[str, Any]]],
) -> None:
    specs_done = instr_done = False
    for table in doc.tables:
        if not table.rows:
            continue
        header = " ".join((c.text or "").lower() for c in table.rows[0].cells)
        if not specs_done and ("фамилия" in header or "удостоверен" in header) and "п/п" in header:
            try:
                fill_specialists_table(table, data, specialist_docs, start_row=1)
            except Exception:
                pass
            specs_done = True
            continue
        if not instr_done and (
            "наименование прибора" in header
            or ("заводской номер прибора" in header)
            or ("наименование" in header and "заводской" in header and "фамилия" not in header)
        ):
            # отличить от списка работ
            if "объем контроля" in header or "объём контроля" in header or "наименование работы" in header:
                continue
            if "наименование документа" in header:
                continue
            try:
                fill_instruments_table(
                    table,
                    verification_equipment,
                    data,
                    ndt_methods=ndt_methods,
                    start_row=1,
                )
            except Exception:
                # fallback короткий
                ve = _merge_report_instruments(verification_equipment, data, ndt_methods)
                for i, eq in enumerate(ve[: max(0, len(table.rows) - 1)]):
                    r = i + 1
                    if r >= len(table.rows):
                        break
                    _set(table, r, 0, f"{i + 1}.")
                    if len(table.rows[r].cells) > 1:
                        _set(table, r, 1, _instrument_full_name(eq))
                    if len(table.rows[r].cells) > 2:
                        _set(table, r, 2, eq.get("serial_number") or eq.get("factory_number") or "")
            instr_done = True


def _match_label(label: str) -> Optional[Tuple[str, ...]]:
    lab = (label or "").lower().strip()
    if not lab:
        return None
    # давление: отдельные правила
    if "давлен" in lab:
        if "расчет" in lab or "расчёт" in lab:
            return ("design_pressure",)
        if "пробн" in lab or "испыта" in lab:
            return ("test_pressure",)
        if "рабоч" in lab:
            return ("working_pressure",)
    for needles, keys in _LABEL_KEYS:
        if any(n in lab for n in needles):
            return keys
    return None


def _fill_kv_tables(doc: Document, values: Dict[str, Any]) -> None:
    """Заполнить таблицы вида «метка | значение» по подписям строк."""
    for table in doc.tables:
        if len(table.columns) < 2:
            continue
        header = " ".join((c.text or "").lower() for c in table.rows[0].cells) if table.rows else ""
        # пропускаем списки работ / документов / специалистов / приборов
        if any(
            x in header
            for x in (
                "наименование работы",
                "наименование документа",
                "фамилия",
                "наименование прибора",
                "объем контроля",
                "объём контроля",
            )
        ):
            continue
        for r_idx, row in enumerate(table.rows):
            if len(row.cells) < 2:
                continue
            label = (row.cells[0].text or "").strip()
            keys = _match_label(label)
            if not keys:
                continue
            # уже заполнено?
            cur = (row.cells[1].text or "").strip()
            if cur and cur not in ("", "—", "-", "–") and "_" not in cur:
                continue
            val = MISSING
            for k in keys:
                if values.get(k) not in (None, "", MISSING):
                    val = values[k]
                    break
            if val is MISSING:
                continue
            # если 3+ колонки и 2-я — подзаголовок (как давление/рабочее)
            target_col = 1
            if len(row.cells) >= 3:
                mid = (row.cells[1].text or "").lower()
                if mid.strip() in ("рабочее", "расчетное", "расчётное", "пробное"):
                    target_col = 2
            try:
                _set(table, r_idx, target_col, str(val))
            except Exception:
                pass


def _fill_uzt_tables(
    doc: Document, data: Dict[str, Any], skip_ids: Optional[set] = None
) -> None:
    points = data.get("thickness_measurements") or data.get("thicknessMeasurements") or []
    if not isinstance(points, list) or not points:
        return
    skip_ids = skip_ids or set()
    for table in doc.tables:
        if not table.rows or id(table._tbl) in skip_ids:
            continue
        header = " ".join((c.text or "").lower() for c in table.rows[0].cells)
        # Заголовок должен говорить именно о точках/зонах замера толщины:
        # одного слова «толщина» мало — оно есть и в таблицах элементов,
        # и в параметрах УЗК, которые эта эвристика раньше затирала.
        has_points = "№ точки" in header or "номер точки" in header or "точка" in header
        has_zone = "зоны контроля" in header or "зона контроля" in header
        has_snom = "s ном" in header or "sном" in header
        if not (has_zone or has_snom or (has_points and "толщин" in header)):
            continue
        if "сварного шва" in header and "толщин" not in header:
            continue
        start = 1
        # пропуск повторного заголовка
        if len(table.rows) > 1 and "зона" in " ".join(c.text.lower() for c in table.rows[1].cells):
            start = 2
        _ensure_rows(table, start + len(points))
        for i, p in enumerate(points):
            if not isinstance(p, dict):
                continue
            r = start + i
            if r >= len(table.rows):
                break
            cols = len(table.rows[r].cells)
            vals = [
                str(i + 1),
                str(p.get("element") or p.get("zone") or p.get("area") or ""),
                str(p.get("point_number") or p.get("point") or i + 1),
                str(p.get("nominal") or p.get("s_nom") or ""),
                str(p.get("thickness") or p.get("value") or ""),
            ]
            for c, v in enumerate(vals):
                if c < cols and v:
                    try:
                        _set(table, r, c, v)
                    except Exception:
                        pass


def _fill_paragraph_blanks(
    doc: Document,
    date_ru: str,
    device: str,
    org_name: str,
    form_title: str,
    g,
) -> None:
    contract = str(g("contract_number", "basis_document", default="") or "")
    for p in doc.paragraphs:
        text = p.text or ""
        if not text.strip():
            continue
        low = text.lower()
        new = text
        if "договор" in low and _BLANK_RE.search(text) and contract:
            new = _replace_underscores(new, [contract])
        if ("с «" in low or "с __" in low or "по «" in low or "202_" in low) and date_ru:
            if "____" in new or "202_" in new:
                new = new.replace("202__", date_ru[-2:] if len(date_ru) >= 2 else "26")
                new = new.replace("202_", date_ru[-2:] if len(date_ru) >= 2 else "26")
        if device and device != MISSING and "наименован" in low and _BLANK_RE.search(text):
            new = _replace_underscores(new, [device])
        if org_name and org_name != MISSING and "заказчик" in low and _BLANK_RE.search(text):
            new = _replace_underscores(new, [org_name])
        if new != text:
            try:
                p.text = new
            except Exception:
                pass
