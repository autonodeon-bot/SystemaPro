"""
Протокол анализа технической документации (Приложение № 1) — форма ТО
«Приложение № 1. Обследование сосудов и аппаратов».
"""
from __future__ import annotations

from dataclasses import dataclass
from typing import Any, Callable, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from equipment_profiles import passport_elements_for_equipment
from equipment_presets import preset_from_equipment_data
from epb_report_builder import device_type_label

MISSING = "—"

TO_DOCUMENT_NAMES: Dict[str, str] = {
    "1": "Лицензия на осуществление эксплуатации",
    "2": "Свидетельство о регистрации опасных производственных объектов",
    "3": "Договор обязательного страхования гражданской ответственности",
    "4": "Страховой полис",
    "5": "Положение о производственном контроле",
    "6": "Приказ об организации производственного контроля",
    "7": "План мероприятий по локализации и ликвидации последствий аварий",
    "8": "Предписания надзорных органов",
    "9": "Журнал учета аварий и инцидентов, происшедших на опасных производственных объектах",
    "10": "Технический паспорт сосуда",
    "11": "Инструкция по монтажу и эксплуатации",
    "12": "Паспорта на предохранительные клапаны",
    "13": "Паспорта на запорную арматуру",
    "14": "Документация на контрольно-измерительные приборы",
    "15": "Ремонтная (исполнительная) документация",
    "16": "Заключение экспертизы промышленной безопасности",
    "17": "Акты проведения УЗТ",
}


@dataclass
class TechnicalReportContext:
    g: Callable[..., Any]
    doc_meta_fn: Callable[[str], Tuple[bool, str, str]]
    device_name: str
    serial: str
    reg_no: str
    inv_no: str
    org: str
    location: str
    date_perf_ru: str
    equipment_data: Dict[str, Any]


def _header_row(tbl, headers: List[str]) -> None:
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
        try:
            tbl.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
            tbl.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass


def _doc_meta_rows(
    docs_dict: Dict[str, Any],
    docs_info: Dict[str, Any],
    doc_meta_fn: Callable[[str], Tuple],
) -> List[Tuple[str, str, str, str]]:
    keys = set()
    if isinstance(docs_dict, dict):
        keys.update(str(k) for k in docs_dict.keys())
    if isinstance(docs_info, dict):
        keys.update(str(k) for k in docs_info.keys())
    keys.update(TO_DOCUMENT_NAMES.keys())
    ordered = sorted(keys, key=lambda x: int(x) if str(x).isdigit() else 999)
    rows: List[Tuple[str, str, str, str]] = []
    for num in ordered:
        if not str(num).isdigit():
            continue
        name = TO_DOCUMENT_NAMES.get(str(num), f"Документ {num}")
        meta = doc_meta_fn(str(num))
        if len(meta) >= 4:
            _present, doc_number, doc_date, pages = meta[:4]
        else:
            _present, doc_number, doc_date = meta[:3]
            pages = ""
        ident = doc_number or ""
        if doc_date:
            ident = f"{ident} от {doc_date}".strip() if ident else f"от {doc_date}"
        rows.append((str(num), name, ident or MISSING, pages or MISSING))
    return rows


def append_technical_protocol_doc_analysis(doc: Document, ctx: TechnicalReportContext) -> None:
    """Приложение № 1: протокол анализа технической документации (таблицы № 1–9)."""
    g = ctx.g
    preset = preset_from_equipment_data(ctx.equipment_data)
    type_label = device_type_label(preset)

    doc.add_paragraph("Протокол анализа технической документации")
    doc.add_paragraph(f"№ _________ от {ctx.date_perf_ru} г.")
    doc.add_paragraph()

    docs_dict = g("documents", default={})
    docs_info = g("documents_info", default={})
    if not isinstance(docs_dict, dict):
        docs_dict = {}
    if not isinstance(docs_info, dict):
        docs_info = {}

    doc.add_paragraph(
        "1. Сведения о рассмотренных в процессе технического диагностирования документах"
    )
    doc.add_paragraph("Таблица № 1")
    doc.add_paragraph()
    doc_rows = _doc_meta_rows(docs_dict, docs_info, ctx.doc_meta_fn)
    t1 = doc.add_table(rows=len(doc_rows) + 1, cols=4)
    t1.style = "Table Grid"
    _header_row(
        t1,
        [
            "№ п/п",
            "Наименование документа",
            "Идентификационные сведения (шифр, номер, марка и/или другая информация)",
            "Объём документа, листов",
        ],
    )
    for i, (num, name, ident, volume) in enumerate(doc_rows, 1):
        t1.rows[i].cells[0].text = num
        t1.rows[i].cells[1].text = name
        t1.rows[i].cells[2].text = ident
        t1.rows[i].cells[3].text = volume
    doc.add_paragraph()

    doc.add_paragraph("2. Общие данные.")
    doc.add_paragraph("Таблица № 2")
    doc.add_paragraph()
    general = [
        ("Наименование объекта", g("vessel_name", "device_name", default=ctx.device_name)),
        ("Условное обозначение", g("designation", "conditional_designation", default=MISSING)),
        ("Наименование завода-изготовителя", g("manufacturer", default=MISSING)),
        ("Год изготовления", g("manufacture_year", "year_of_manufacture", default=MISSING)),
        ("Год ввода в эксплуатацию", g("commissioning_year", default=MISSING)),
        ("Рабочее давление, МПа", g("working_pressure", default=MISSING)),
        ("Номинальный наружный или внутренний диаметр, мм", g("diameter", "inner_diameter", default=MISSING)),
        ("Рабочая температура среды, ℃", g("working_temperature", default=MISSING)),
        ("Наименование рабочей среды", g("working_medium", "medium", default=MISSING)),
        ("Марка материала корпуса", g("shell_material", "material", default=MISSING)),
        ("Вместимость, м³", g("volume", "capacity", default=MISSING)),
        ("Схема подключения сосуда в установку", g("connection_scheme", default=MISSING)),
        ("Климатическое исполнение", g("climate_design", default=MISSING)),
    ]
    t2 = doc.add_table(rows=len(general), cols=2)
    t2.style = "Table Grid"
    for i, (lbl, val) in enumerate(general):
        t2.rows[i].cells[0].text = lbl
        t2.rows[i].cells[1].text = str(val) if val not in (None, "") else MISSING
    doc.add_paragraph()

    doc.add_paragraph("3. Сведения об основных элементах сосуда.")
    doc.add_paragraph("Таблица № 3")
    doc.add_paragraph()
    elements = passport_elements_for_equipment(ctx.equipment_data)
    if not elements:
        raw = g("vessel_elements", "passport_elements", default=[])
        elements = raw if isinstance(raw, list) else []
    if not elements:
        elements = [{"name": "Корпус"}, {"name": "Верхнее днище"}, {"name": "Нижнее днище"}]
    t3 = doc.add_table(rows=len(elements) + 1, cols=11)
    t3.style = "Table Grid"
    _header_row(
        t3,
        [
            "Наименование элемента сосуда",
            "Количество",
            "Диаметр внутренний",
            "Длина или высота",
            "Толщина стенки номинальная",
            "Расчет толщины стенки",
            "Марка стали",
            "ГОСТ",
            "Вид сварки",
            "Электроды, сварочная проволока",
            "Метод неразрушающего контроля",
        ],
    )
    for i, el in enumerate(elements, 1):
        if not isinstance(el, dict):
            continue
        row = t3.rows[i]
        row.cells[0].text = str(el.get("name") or el.get("element_name") or "—")
        row.cells[1].text = str(el.get("quantity") or "1")
        row.cells[2].text = str(el.get("diameter_mm") or el.get("diameter") or "—")
        row.cells[3].text = str(el.get("length_mm") or el.get("length") or "—")
        row.cells[4].text = str(el.get("wall_thickness_mm") or el.get("wall_thickness") or "—")
        row.cells[5].text = str(el.get("calculated_thickness") or "—")
        row.cells[6].text = str(el.get("material") or "—")
        row.cells[7].text = str(el.get("gost") or el.get("material_gost") or "—")
        weld = str(el.get("weld_data") or el.get("weld_type") or "—")
        row.cells[8].text = weld.split("\n")[0] if weld else "—"
        row.cells[9].text = weld
        row.cells[10].text = str(el.get("ndt_method") or "—")
    doc.add_paragraph()

    doc.add_paragraph("4. Технические характеристики и параметры.")
    doc.add_paragraph("Таблица № 4")
    doc.add_paragraph()
    params = [
        ("Давление, кгс/см²", "рабочее", g("working_pressure", default=MISSING)),
        ("Давление, кгс/см²", "расчетное", g("design_pressure", default=MISSING)),
        ("Давление, кгс/см²", "пробное пневматическое испытание", g("pneumatic_test_pressure", default=MISSING)),
        ("Давление, кгс/см²", "пробное гидравлическое испытание", g("hydraulic_test_pressure", default=MISSING)),
        ("Температура, °С", "расчетная стенки", g("design_wall_temperature", default=MISSING)),
        ("Температура, °С", "рабочей среды", g("working_temperature", default=MISSING)),
        ("Температура, °С", "стенки, минимально допустимая отрицательная", g("min_wall_temperature", default=MISSING)),
        ("Рабочая среда", "состав", g("medium_composition", default=MISSING)),
        ("Рабочая среда", "класс опасности по ГОСТ 12.1.007", g("hazard_class", default=MISSING)),
        ("Рабочая среда", "взрывоопасность", g("explosiveness", default=MISSING)),
        ("Рабочая среда", "пожароопасность", g("fire_hazard", default=MISSING)),
        ("Объем, м³", "Объем, м³", g("volume", default=MISSING)),
        ("Масса порожнего сосуда, кг", "Масса порожнего сосуда, кг", g("empty_mass", default=MISSING)),
        ("Прибавка к толщине для компенсации коррозии (эрозии), мм", "Прибавка к толщине для компенсации коррозии (эрозии), мм", g("corrosion_allowance", default=MISSING)),
        ("Число циклов нагружения (пусков-остановов)", "Число циклов нагружения (пусков-остановов)", g("load_cycles", default=MISSING)),
        ("Срок службы, лет", "Срок службы, лет", g("service_life_years", default=MISSING)),
    ]
    t4 = doc.add_table(rows=len(params) + 1, cols=5)
    t4.style = "Table Grid"
    _header_row(t4, ["Наименование показателей", "Наименование показателей", "Проектные", "Фактические", "Примечание"])
    for i, (g1, g2, val) in enumerate(params, 1):
        row = t4.rows[i]
        row.cells[0].text = g1
        row.cells[1].text = g2
        row.cells[2].text = str(val) if val not in (None, "") else MISSING
        row.cells[3].text = MISSING
        row.cells[4].text = MISSING
    doc.add_paragraph()

    doc.add_paragraph("5. Материалы элементов сосуда.")
    doc.add_paragraph("Таблица № 5")
    doc.add_paragraph()
    materials = g("material_certificates", "passport_materials", default=[])
    if not isinstance(materials, list) or not materials:
        materials = elements if elements else [{"name": "Корпус", "material": g("shell_material", default=MISSING)}]
    t5 = doc.add_table(rows=len(materials) + 1, cols=10)
    t5.style = "Table Grid"
    _header_row(
        t5,
        [
            "Наименование элемента",
            "Марка материала",
            "ГОСТ, ТУ",
            "Предел текучести",
            "Временное сопротивление",
            "Относительное удлинение, %",
            "Относительное сужение, %",
            "Ударная вязкость",
            "Температура, °С",
            "Тип образца",
        ],
    )
    for i, m in enumerate(materials, 1):
        if not isinstance(m, dict):
            continue
        row = t5.rows[i]
        row.cells[0].text = str(m.get("name") or m.get("element") or "—")
        row.cells[1].text = str(m.get("material") or "—")
        row.cells[2].text = str(m.get("gost") or m.get("material_gost") or "—")
        row.cells[3].text = str(m.get("yield_strength") or MISSING)
        row.cells[4].text = str(m.get("tensile_strength") or MISSING)
        row.cells[5].text = str(m.get("elongation") or MISSING)
        row.cells[6].text = str(m.get("reduction") or MISSING)
        row.cells[7].text = str(m.get("impact_toughness") or MISSING)
        row.cells[8].text = str(m.get("temperature") or MISSING)
        row.cells[9].text = str(m.get("specimen_type") or MISSING)
    doc.add_paragraph()

    doc.add_paragraph("6. Термообработка.")
    doc.add_paragraph("Таблица № 6")
    doc.add_paragraph()
    heat = g("heat_treatment_records", "vessel_heat_treatment", default=[])
    if not isinstance(heat, list) or not heat:
        heat = [{"element": MISSING, "type": MISSING, "mode": MISSING, "temperature": MISSING, "duration": MISSING, "cooling": MISSING}]
    t6 = doc.add_table(rows=len(heat) + 1, cols=6)
    t6.style = "Table Grid"
    _header_row(
        t6,
        [
            "Наименование элемента или соединения",
            "Вид термообработки",
            "Режим термообработки",
            "Температура, °С",
            "Продолжительность выдержки, ч.",
            "Способ охлаждения",
        ],
    )
    for i, rec in enumerate(heat, 1):
        if not isinstance(rec, dict):
            continue
        row = t6.rows[i]
        row.cells[0].text = str(rec.get("element") or rec.get("name") or MISSING)
        row.cells[1].text = str(rec.get("type") or rec.get("treatment_type") or MISSING)
        row.cells[2].text = str(rec.get("mode") or rec.get("regime") or MISSING)
        row.cells[3].text = str(rec.get("temperature") or MISSING)
        row.cells[4].text = str(rec.get("duration") or MISSING)
        row.cells[5].text = str(rec.get("cooling") or MISSING)
    doc.add_paragraph()

    doc.add_paragraph("7. Испытание на прочность.")
    doc.add_paragraph("Таблица № 7")
    doc.add_paragraph()
    tests = g("hydraulic_test_history", "strength_test_records", "hydraulic_tests", default=[])
    if not isinstance(tests, list) or not tests:
        tests = [
            {
                "test_type": "Гидравлическое",
                "pressure": g("hydraulic_test_pressure", default=MISSING),
                "duration": MISSING,
                "medium": "вода",
                "temperature": MISSING,
                "result": g("hydraulic_test_result", default="испытание выдержано"),
                "date": g("hydraulic_test_date", default=MISSING),
            }
        ]
    t7 = doc.add_table(rows=len(tests) + 1, cols=7)
    t7.style = "Table Grid"
    _header_row(
        t7,
        [
            "Дата",
            "Вид испытания",
            "Давление, МПа (кгс/см²)",
            "Среда",
            "Температура, °С",
            "Продолжительность, ч",
            "Результат / примечание",
        ],
    )
    for i, rec in enumerate(tests, 1):
        if not isinstance(rec, dict):
            continue
        row = t7.rows[i]
        row.cells[0].text = str(rec.get("date") or MISSING)
        row.cells[1].text = str(
            rec.get("test_type") or rec.get("type") or rec.get("kind") or "Гидравлическое"
        )
        row.cells[2].text = str(rec.get("pressure") or MISSING)
        row.cells[3].text = str(rec.get("medium") or MISSING)
        row.cells[4].text = str(rec.get("temperature") or MISSING)
        row.cells[5].text = str(rec.get("duration") or MISSING)
        row.cells[6].text = str(rec.get("note") or rec.get("result") or MISSING)
    doc.add_paragraph()

    doc.add_paragraph("8. Сведения о предыдущих обследованиях")
    doc.add_paragraph("Таблица № 8")
    doc.add_paragraph()
    prev = g("previous_inspections", default=[])
    if not isinstance(prev, list) or not prev:
        prev = [
            {
                "kind": "Техническое диагностирование",
                "result": g("previous_inspection_result", default=MISSING),
            }
        ]
    t8 = doc.add_table(rows=len(prev) + 1, cols=3)
    t8.style = "Table Grid"
    _header_row(t8, ["№ п/п", "Вид обследования", "Результаты контроля / отчётная документация"])
    for i, rec in enumerate(prev, 1):
        if not isinstance(rec, dict):
            continue
        t8.rows[i].cells[0].text = str(i)
        t8.rows[i].cells[1].text = str(rec.get("kind") or rec.get("type") or "Техническое диагностирование")
        t8.rows[i].cells[2].text = str(rec.get("result") or rec.get("report") or MISSING)
    doc.add_paragraph()

    doc.add_paragraph("9. Дополнительные данные о сосуде.")
    doc.add_paragraph("Таблица № 9")
    doc.add_paragraph()
    extra = g("additional_vessel_data", default=MISSING)
    if isinstance(extra, dict):
        extra_rows = list(extra.items())
    else:
        extra_rows = [
            ("Дополнительные сведения", str(extra) if extra not in (None, "") else MISSING),
        ]
    t9 = doc.add_table(rows=len(extra_rows) + 1, cols=2)
    t9.style = "Table Grid"
    _header_row(t9, ["Показатель", "Значение"])
    for i, (k, v) in enumerate(extra_rows, 1):
        t9.rows[i].cells[0].text = str(k)
        t9.rows[i].cells[1].text = str(v) if v not in (None, "") else MISSING
    doc.add_paragraph()

    doc.add_paragraph("10. Результаты анализа технической документации.")
    analysis = g(
        "doc_analysis_result",
        default="При анализе технической документации установлено соответствие представленных документов требованиям НТД.",
    )
    doc.add_paragraph(str(analysis))
    doc.add_paragraph()
    doc.add_paragraph(
        f"ВЫВОД: Представленная техническая документация на {type_label.lower()}, "
        f"работающий под давлением – {ctx.device_name} зав. № {ctx.serial}, "
        f"рег. № {ctx.reg_no}, инв. № {ctx.inv_no} ведется "
        f"в соответствии с требованиями действующей нормативно-технической документации."
    )
    doc.add_paragraph()
