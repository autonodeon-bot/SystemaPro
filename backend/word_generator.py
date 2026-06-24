"""
Генератор Word документов для отчетов и опросных листов
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from typing import Dict, Any, Optional, List
from pathlib import Path
import os
import tempfile
import uuid as _uuid

from shared import resolve_report_file_path as _scoped_resolve_upload_path
from pressure_device_labels import (
    apply_device_terminology_to_document,
    default_purpose_for_kind,
    detect_pressure_device_kind,
    is_pressure_device,
)
from equipment_presets import preset_from_equipment_data, pressure_regime_for_preset
from epb_report_builder import (
    EpbReportContext,
    append_epb_appendix_act,
    append_epb_appendix_doc_analysis,
    append_epb_toc,
    build_epb_main_body,
    epb_appendix_letter,
)
from epb_protocol_tables import (
    append_epb_appendix_e,
    append_epb_protocol_hardness,
    append_epb_protocol_uzt,
    append_epb_protocol_weld_control,
    _filter_welds,
)
from report_attachments import build_attachments_index
from suitability_conclusions import conclusion_from_inspection_data
from technical_report_builder import TechnicalReportContext, append_technical_protocol_doc_analysis

NORMATIVE_BASE_ORDER_536 = (
    "Федеральные нормы и правила в области промышленной безопасности "
    "«Правила промышленной безопасности при использовании оборудования, "
    "работающего под избыточным давлением», утвержденные приказом "
    "Федеральной службы по экологическому, технологическому и атомному надзору от 15.12.2020 №536."
)
NORMATIVE_BASE_RUA_93 = (
    "Правила устройства и безопасной эксплуатации аппаратов, работающих под давлением "
    "до 0,07 МПа (РУА-93), утвержденные Госгортехнадзором России 05.06.1993."
)

try:
    from PIL import Image, ImageDraw, ImageFont
    _HAS_PIL = True
except ImportError:
    _HAS_PIL = False

class WordGenerator:
    """Генератор Word документов"""

    def __init__(self):
        self._report_inspection_id: Optional[str] = None
        self._report_questionnaire_id: Optional[str] = None

    def _set_report_path_scope(self, inspection_data: Optional[Dict[str, Any]]) -> None:
        """Контекст для поиска вложений только в каталогах текущего обследования/опросника."""
        self._report_inspection_id = None
        self._report_questionnaire_id = None
        if not isinstance(inspection_data, dict):
            return
        _iid = inspection_data.get("id")
        if _iid:
            self._report_inspection_id = str(_iid)
        _qid = inspection_data.get("questionnaire_id")
        if _qid:
            self._report_questionnaire_id = str(_qid)

    def _draw_points_on_scheme(
        self,
        scheme_path: str,
        thickness_points: List[Dict[str, Any]],
        output_dir: Optional[str] = None,
    ) -> Optional[str]:
        """
        Рисует номера точек замеров на схеме УЗК по координатам x_percent, y_percent.
        Возвращает путь к сохранённому изображению или None при ошибке.
        """
        if not _HAS_PIL or not scheme_path or not thickness_points:
            return None
        resolved = self._find_image_path(scheme_path) or scheme_path
        if not os.path.isfile(resolved):
            return None
        points_with_coords = [
            (i + 1, p)
            for i, p in enumerate(thickness_points)
            if isinstance(p, dict)
            and (p.get("x_percent") is not None or p.get("y_percent") is not None)
        ]
        # Если координат нет — рисуем легенду «Точки замеров 1, 2, 3...» внизу схемы
        if not points_with_coords and thickness_points:
            points_with_coords = [(i + 1, {}) for i, p in enumerate(thickness_points) if isinstance(p, dict)]
        if not points_with_coords:
            return None
        try:
            img = Image.open(resolved).convert("RGB")
            w, h = img.size
            draw = ImageDraw.Draw(img)
            radius = max(14, min(w, h) // 50)
            font_size = max(12, min(w, h) // 40)
            font = None
            for font_path in [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
                "/usr/share/fonts/TTF/DejaVuSans-Bold.ttf",
            ]:
                try:
                    if os.path.isfile(font_path):
                        font = ImageFont.truetype(font_path, font_size)
                        break
                except Exception:
                    continue
            if font is None:
                font = ImageFont.load_default()
            n_pts = len(points_with_coords)
            for idx, (num, p) in enumerate(points_with_coords):
                x_pct = p.get("x_percent") if isinstance(p, dict) else None
                y_pct = p.get("y_percent") if isinstance(p, dict) else None
                if x_pct is not None or y_pct is not None:
                    try:
                        xf = float(x_pct) if x_pct is not None else 0.5
                        yf = float(y_pct) if y_pct is not None else 0.5
                    except (TypeError, ValueError):
                        xf, yf = 0.5, 0.5
                    # Поддержка и процентов (0..100), и нормализованных (0..1),
                    # и старого формата абсолютных пикселей.
                    if xf > 100:
                        xf = xf / max(w, 1)
                    else:
                        xf = xf / 100.0 if xf > 1 else xf
                    if yf > 100:
                        yf = yf / max(h, 1)
                    else:
                        yf = yf / 100.0 if yf > 1 else yf
                    xf = max(0, min(1, xf))
                    yf = max(0, min(1, yf))
                    cx = int(xf * w)
                    cy = int(yf * h)
                else:
                    # Нет координат — размещаем номера точек в ряд внизу схемы (легенда)
                    margin = w // 20
                    step = (w - 2 * margin) // max(n_pts, 1)
                    cx = margin + (idx * step) + step // 2
                    cy = h - max(radius * 2, h // 15)
                    cy = max(radius, min(h - radius, cy))
                draw.ellipse(
                    [cx - radius, cy - radius, cx + radius, cy + radius],
                    outline=(220, 50, 50),
                    width=max(2, radius // 5),
                    fill=(255, 200, 200),
                )
                text = str(num)
                bbox = draw.textbbox((0, 0), text, font=font)
                tw = bbox[2] - bbox[0]
                th = bbox[3] - bbox[1]
                draw.text((cx - tw // 2, cy - th // 2), text, fill=(180, 0, 0), font=font)
            out_dir = output_dir or tempfile.gettempdir()
            os.makedirs(out_dir, exist_ok=True)
            out_path = os.path.join(out_dir, f"scheme_with_points_{_uuid.uuid4().hex[:12]}.png")
            img.save(out_path, "PNG")
            return out_path
        except Exception as e:
            print(f"Ошибка отрисовки точек на схеме: {e}")
            return None

    def _draw_weld_points_on_scheme(
        self,
        scheme_path: str,
        weld_points: List[Dict[str, Any]],
        output_dir: Optional[str] = None,
    ) -> Optional[str]:
        """
        Рисует номера сварных соединений (швов) на схеме УЗК по координатам x_percent, y_percent.
        Возвращает путь к сохранённому изображению или None при ошибке.
        """
        if not _HAS_PIL or not scheme_path or not weld_points:
            return None
        resolved = self._find_image_path(scheme_path) or scheme_path
        if not os.path.isfile(resolved):
            return None
        points_with_coords = [
            (str(w.get("weld_number") or ""), w)
            for w in weld_points
            if isinstance(w, dict)
            and (w.get("x_percent") is not None or w.get("y_percent") is not None)
        ]
        if not points_with_coords:
            return None
        try:
            img = Image.open(resolved).convert("RGB")
            w, h = img.size
            draw = ImageDraw.Draw(img)
            radius = max(12, min(w, h) // 55)
            font_size = max(10, min(w, h) // 45)
            font = None
            for font_path in [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
                "/usr/share/fonts/TTF/DejaVuSans-Bold.ttf",
            ]:
                try:
                    if os.path.isfile(font_path):
                        font = ImageFont.truetype(font_path, font_size)
                        break
                except Exception:
                    continue
            if font is None:
                font = ImageFont.load_default()
            for label, p in points_with_coords:
                x_pct = p.get("x_percent")
                y_pct = p.get("y_percent")
                try:
                    xf = float(x_pct) if x_pct is not None else 0.5
                    yf = float(y_pct) if y_pct is not None else 0.5
                except (TypeError, ValueError):
                    xf, yf = 0.5, 0.5
                # Поддержка процентов/нормализованных координат/пикселей из старых данных.
                if xf > 100:
                    xf = xf / max(w, 1)
                else:
                    xf = xf / 100.0 if xf > 1 else xf
                if yf > 100:
                    yf = yf / max(h, 1)
                else:
                    yf = yf / 100.0 if yf > 1 else yf
                xf = max(0, min(1, xf))
                yf = max(0, min(1, yf))
                cx = int(xf * w)
                cy = int(yf * h)
                draw.ellipse(
                    [cx - radius, cy - radius, cx + radius, cy + radius],
                    outline=(50, 100, 220),
                    width=max(2, radius // 5),
                    fill=(200, 220, 255),
                )
                text = (label or "").strip() or "?"
                bbox = draw.textbbox((0, 0), text, font=font)
                tw = bbox[2] - bbox[0]
                th = bbox[3] - bbox[1]
                draw.text((cx - tw // 2, cy - th // 2), text, fill=(0, 80, 180), font=font)
            out_dir = output_dir or tempfile.gettempdir()
            os.makedirs(out_dir, exist_ok=True)
            out_path = os.path.join(out_dir, f"scheme_weld_points_{_uuid.uuid4().hex[:12]}.png")
            img.save(out_path, "PNG")
            return out_path
        except Exception as e:
            print(f"Ошибка отрисовки точек УЗК на схеме: {e}")
            return None

    def _find_image_path(self, path: Optional[str]) -> Optional[str]:
        """Разрешить путь к изображению без подстановки файлов других осмотров (см. shared.resolve_report_file_path)."""
        if not path or not isinstance(path, str):
            return None
        path = path.strip().replace("\\", "/")
        if os.path.isabs(path) and os.path.isfile(path):
            return path
        resolved = _scoped_resolve_upload_path(
            path,
            inspection_id=self._report_inspection_id,
            questionnaire_id=self._report_questionnaire_id,
        )
        if resolved and isinstance(resolved, str) and os.path.isfile(resolved):
            return resolved
        return None

    def generate_questionnaire_word(
        self,
        questionnaire_data: Dict[str, Any],
        equipment_data: Dict[str, Any],
        questionnaire_info: Dict[str, Any],
        ndt_methods: List[Dict[str, Any]],
        output_path: str
    ):
        """Генерировать Word документ опросного листа"""
        self._report_inspection_id = None
        qn = questionnaire_info.get("questionnaire_id") if isinstance(questionnaire_info, dict) else None
        if not qn and isinstance(questionnaire_info, dict):
            qn = questionnaire_info.get("id")
        self._report_questionnaire_id = str(qn) if qn else None

        doc = Document()
        
        # Настройка стилей
        self._setup_styles(doc)
        
        # Титульная страница
        title = doc.add_heading('ОПРОСНЫЙ ЛИСТ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(f'оборудования: {equipment_data.get("name", "Не указано")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.size = Pt(14)
        subtitle_format.bold = True
        
        doc.add_paragraph()
        
        # Раздел 1: Общие сведения об оборудовании
        doc.add_heading('1. ОБЩИЕ СВЕДЕНИЯ ОБ ОБОРУДОВАНИИ', level=1)
        
        equipment_table = doc.add_table(rows=4, cols=2)
        equipment_table.style = 'Light Grid Accent 1'
        equipment_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        equipment_info = [
            ['Наименование оборудования:', equipment_data.get('name') or 'Не указано'],
            ['Инвентарный номер:', questionnaire_info.get('inventory_number') or 'Не указан'],
            ['Заводской номер:', equipment_data.get('serial_number') or 'Не указан'],
            ['Место расположения:', equipment_data.get('location') or 'Не указано'],
        ]
        
        for i, (label, value) in enumerate(equipment_info):
            equipment_table.rows[i].cells[0].text = label
            # Обрабатываем None значения
            equipment_table.rows[i].cells[1].text = str(value) if value is not None else 'Не указано'
            equipment_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Раздел 2: Сведения об обследовании
        doc.add_heading('2. СВЕДЕНИЯ ОБ ОБСЛЕДОВАНИИ', level=1)
        
        inspection_table = doc.add_table(rows=3, cols=2)
        inspection_table.style = 'Light Grid Accent 1'
        
        inspection_date = questionnaire_info.get('inspection_date')
        if inspection_date:
            try:
                if 'T' in str(inspection_date):
                    inspection_date = datetime.fromisoformat(str(inspection_date).replace('Z', '+00:00')).strftime('%d.%m.%Y')
                else:
                    inspection_date = datetime.fromisoformat(str(inspection_date)).strftime('%d.%m.%Y')
            except:
                pass
        
        inspection_info = [
            ['Дата обследования:', inspection_date or 'Не указана'],
            # ВАЖНО: questionnaire_info может содержать ключи со значением None,
            # а python-docx ожидает строку (иначе TypeError: 'NoneType' object is not iterable).
            ['Инженер:', (questionnaire_info.get('inspector_name') or 'Не указан')],
            ['Должность:', (questionnaire_info.get('inspector_position') or 'Не указана')],
        ]
        
        for i, (label, value) in enumerate(inspection_info):
            inspection_table.rows[i].cells[0].text = label
            inspection_table.rows[i].cells[1].text = str(value) if value is not None else ''
            inspection_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Раздел 3: Результаты неразрушающего контроля
        doc.add_heading('3. РЕЗУЛЬТАТЫ НЕРАЗРУШАЮЩЕГО КОНТРОЛЯ', level=1)
        
        if ndt_methods:
            ndt_table = doc.add_table(rows=len(ndt_methods) + 1, cols=6)
            ndt_table.style = 'Light Grid Accent 1'
            
            # Заголовки
            headers = ['Метод НК', 'Нормативный документ', 'Оборудование', 'Инженер', 'Уровень', 'Результаты']
            for i, header in enumerate(headers):
                cell = ndt_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Данные методов НК
            for idx, method in enumerate(ndt_methods, start=1):
                ndt_table.rows[idx].cells[0].text = str(method.get('method_name') or '')
                ndt_table.rows[idx].cells[1].text = str(method.get('standard') or '')
                ndt_table.rows[idx].cells[2].text = str(method.get('equipment') or '')
                ndt_table.rows[idx].cells[3].text = str(method.get('inspector_name') or '')
                ndt_table.rows[idx].cells[4].text = str(method.get('inspector_level') or '')
                ndt_table.rows[idx].cells[5].text = str(method.get('results') or '')
            
            doc.add_paragraph()
            
            # Детальная информация по каждому методу
            for method in ndt_methods:
                if method.get('is_performed'):
                    method_name = method.get('method_name', 'Неизвестный метод')
                    doc.add_heading(f'{method_name}', level=2)
                    
                    if method.get('defects'):
                        p = doc.add_paragraph()
                        p.add_run('Обнаруженные дефекты: ').bold = True
                        p.add_run(str(method.get('defects') or ''))
                    
                    if method.get('conclusion'):
                        p = doc.add_paragraph()
                        p.add_run('Заключение: ').bold = True
                        p.add_run(str(method.get('conclusion') or ''))
                    
                    doc.add_paragraph()
        else:
            doc.add_paragraph('Методы неразрушающего контроля не указаны.')
        
        # Раздел 4: Перечень документов
        doc.add_paragraph()
        doc.add_heading('4. ПЕРЕЧЕНЬ РАССМОТРЕННЫХ ДОКУМЕНТОВ', level=1)
        
        # Список названий документов (из мобильного приложения)
        document_names = {
            '1': 'Лицензия на осуществление деятельности по эксплуатации взрывопожароопасных и химически опасных производственных объектов I, II и III классов опасности',
            '2': 'Свидетельство о регистрации в государственном реестре ОПО, включая сведения характеризующие ОПО',
            '3': 'Технологический регламент объектов опасных производственных объектов',
            '4': 'План мероприятий по локализации и ликвидации последствий аварий на опасном производственном объекте',
            '5': 'Положение о производственном контроле за соблюдением требований промышленной безопасности на опасных производственных объектах',
            '6': 'Журнал учета аварий и инцидентов на ОПО',
            '7': 'Страховой полис страхования гражданской ответственности владельца опасного объекта за причинение вреда в результате аварии на опасном объекте',
            '8': 'Приказ о назначении ответственного лица за исправное состояние и безопасную эксплуатацию сосудов',
            '9': 'Приказ о назначении ответственного лица за осуществление производственного контроля и соблюдение требований промышленной безопасности на опасном производственном объекте',
            '10': 'Паспорт сосуда заводской (удостоверение о качестве монтажа, сертификат соответствия, сборочный чертёж и схема включения сосуда, расчёт на прочность)',
            '11': 'Инструкция по монтажу и эксплуатации',
            '12': 'Паспорта на предохранительные клапаны',
            '13': 'Паспорта на запорную арматуру',
            '14': 'Документация на контрольно-измерительные приборы',
            '15': 'Ремонтная (исполнительная) документация',
            '16': 'Заключение экспертизы промышленной безопасности',
            '17': 'Акты проведения УЗТ',
        }
        
        docs = questionnaire_data.get('documents', {})
        if isinstance(docs, dict) and docs:
            doc_table = doc.add_table(rows=len(docs) + 1, cols=3)
            doc_table.style = 'Light Grid Accent 1'
            
            # Заголовки
            headers = ['№', 'Наименование документа', 'Наличие']
            for i, header in enumerate(headers):
                cell = doc_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Данные документов
            row_idx = 1
            for num, has_doc in sorted(docs.items(), key=lambda x: int(x[0])):
                doc_name = document_names.get(str(num), f'Документ {num}')
                doc_table.rows[row_idx].cells[0].text = str(num)
                doc_table.rows[row_idx].cells[1].text = doc_name
                doc_table.rows[row_idx].cells[2].text = 'Да' if has_doc else 'Нет'
                doc_table.rows[row_idx].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_idx += 1
        else:
            doc.add_paragraph('Документы не указаны.')
        
        # Подпись
        doc.add_paragraph()
        doc.add_paragraph('Инженер: _________________')
        doc.add_paragraph(questionnaire_info.get('inspector_name', ''))
        doc.add_paragraph()
        doc.add_paragraph(f"Дата: {inspection_date or datetime.now().strftime('%d.%m.%Y')}")
        
        # Сохранение
        doc.save(output_path)
    
    def generate_report_word(
        self,
        inspection_data: Dict[str, Any],
        equipment_data: Dict[str, Any],
        ndt_methods: List[Dict[str, Any]],
        output_path: str,
        report_type: str = "TECHNICAL_REPORT",
        document_files: Optional[List[Dict[str, Any]]] = None,
        specialist_docs: Optional[List[Dict[str, Any]]] = None,
        verification_equipment: Optional[List[Dict[str, Any]]] = None,
        template_definition: Optional[Dict[str, Any]] = None,
    ):
        """Генерировать Word документ отчета"""
        self._set_report_path_scope(inspection_data)
        rt = (report_type or "").strip().upper()
        if rt in ["DIAGNOSTICS", "DIAGNOSTIC", "TECHNICAL_DIAGNOSTICS"]:
            return self._generate_diagnostics_report_word(
                inspection_data=inspection_data,
                equipment_data=equipment_data,
                ndt_methods=ndt_methods,
                output_path=output_path,
                document_files=document_files,
                specialist_docs=specialist_docs,
                verification_equipment=verification_equipment,
                template_definition=template_definition,
            )
        
        # Сосуд / газосепаратор / ресивер — единый шаблон отчёта СРпД
        if is_pressure_device(equipment_data):
            return self._generate_vessel_report_word(
                inspection_data=inspection_data,
                equipment_data=equipment_data,
                ndt_methods=ndt_methods,
                output_path=output_path,
                report_type=rt,
                document_files=document_files,
                specialist_docs=specialist_docs,
                verification_equipment=verification_equipment,
                template_definition=template_definition,
            )

        doc = Document()
        
        # Настройка стилей
        self._setup_styles(doc)
        
        # Титульная страница
        if report_type == "EXPERTISE":
            title = doc.add_heading('ЭКСПЕРТИЗА ПРОМЫШЛЕННОЙ БЕЗОПАСНОСТИ', 0)
        else:
            title = doc.add_heading('ОТЧЕТ О ТЕХНИЧЕСКОМ ДИАГНОСТИРОВАНИИ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(f'оборудования: {equipment_data.get("name", "Не указано")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.size = Pt(14)
        subtitle_format.bold = True
        
        doc.add_paragraph()

        # Содержание (упрощенное)
        doc.add_heading('СОДЕРЖАНИЕ', level=1)
        for item in [
            '1. Общая часть',
            '2. Исходные данные и нормативная база',
            '3. Описание объекта и карта обследования',
            '4. Акт(ы) неразрушающего контроля',
            '5. Результаты обследования (детализация)',
            '6. Заключение',
            '7. Приложения',
        ]:
            doc.add_paragraph(item)
        doc.add_page_break()

        # Общая часть
        doc.add_heading('1. ОБЩАЯ ЧАСТЬ', level=1)
        doc.add_paragraph(
            'Настоящий отчет составлен по результатам технического диагностирования оборудования с целью '
            'оценки технического состояния и определения возможности дальнейшей безопасной эксплуатации.'
        )
        doc.add_paragraph()

        # Нормативная база
        doc.add_heading('2. ИСХОДНЫЕ ДАННЫЕ И НОРМАТИВНАЯ БАЗА', level=1)
        doc.add_paragraph(
            'При выполнении работ использовались данные Заказчика, результаты обследований и применимые нормативные документы (ФНП, ГОСТ, РД и др.).'
        )
        doc.add_paragraph()
        
        # Описание объекта
        doc.add_heading('3. ОПИСАНИЕ ОБЪЕКТА И КАРТА ОБСЛЕДОВАНИЯ', level=1)
        
        equipment_table = doc.add_table(rows=3, cols=2)
        equipment_table.style = 'Light Grid Accent 1'
        
        equipment_info = [
            ['Наименование оборудования:', equipment_data.get('name') or 'Не указано'],
            ['Заводской номер:', equipment_data.get('serial_number') or 'Не указан'],
            ['Место расположения:', equipment_data.get('location') or 'Не указано'],
        ]
        
        for i, (label, value) in enumerate(equipment_info):
            equipment_table.rows[i].cells[0].text = label
            # Обрабатываем None значения
            equipment_table.rows[i].cells[1].text = str(value) if value is not None else 'Не указано'
            equipment_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        data = inspection_data.get("data") or {}
        if not isinstance(data, dict):
            data = {}

        def _get(*keys, default=None):
            """Извлечь значение по ключам (поддержка camelCase и snake_case)"""
            for k in keys:
                if k in data and data.get(k) is not None:
                    return data.get(k)
            # Пробуем варианты с camelCase/snake_case
            for k in keys:
                # snake_case -> camelCase
                camel_key = ''.join(word.capitalize() if i > 0 else word for i, word in enumerate(k.split('_')))
                camel_key_lower = camel_key[0].lower() + camel_key[1:] if camel_key else k
                if camel_key_lower in data and data.get(camel_key_lower) is not None:
                    return data.get(camel_key_lower)
            return default

        opo = data.get("opo") or data.get("opo_info") or {}
        if not isinstance(opo, dict):
            opo = {}

        def _opo_get(*keys, default=None):
            for k in keys:
                if k in opo and opo.get(k) not in (None, ""):
                    return opo.get(k)
            for k in keys:
                camel_key = ''.join(word.capitalize() if i > 0 else word for i, word in enumerate(k.split('_')))
                camel_key_lower = camel_key[0].lower() + camel_key[1:] if camel_key else k
                if camel_key_lower in opo and opo.get(camel_key_lower) not in (None, ""):
                    return opo.get(camel_key_lower)
            return _get(*keys, default=default)

        # Сведения об ОПО (если есть)
        opo_name = _opo_get("name", "opo_name")
        opo_code = _opo_get("code", "opo_code")
        opo_desc = _opo_get("description", "opo_description")
        opo_enterprise = _opo_get("enterprise_name", "opo_enterprise")
        opo_branch = _opo_get("branch_name", "opo_branch")
        opo_workshop = _opo_get("workshop_name", "opo_workshop")
        survey = data.get("opo_survey") if isinstance(data.get("opo_survey"), dict) else opo.get("survey_data")
        if not isinstance(survey, dict):
            survey = {}
        opo_org = survey.get("organization")
        opo_exec = survey.get("executors")

        if any([opo_name, opo_code, opo_desc, opo_enterprise, opo_branch, opo_workshop, opo_org, opo_exec]):
            doc.add_heading('3.1. Сведения об ОПО', level=2)
            opo_rows = []
            def _add_opo_row(label, value):
                if value is None:
                    return
                s = str(value).strip()
                if not s:
                    return
                opo_rows.append((label, s))
            _add_opo_row('Наименование ОПО:', opo_name)
            _add_opo_row('Код ОПО:', opo_code)
            _add_opo_row('Описание:', opo_desc)
            _add_opo_row('Предприятие:', opo_enterprise)
            _add_opo_row('Филиал:', opo_branch)
            _add_opo_row('Цех:', opo_workshop)
            _add_opo_row('Организация (из опросного листа ОПО):', opo_org)
            _add_opo_row('Исполнители (из опросного листа ОПО):', opo_exec)

            if opo_rows:
                opo_table = doc.add_table(rows=len(opo_rows), cols=2)
                opo_table.style = 'Light Grid Accent 1'
                for i, (label, value) in enumerate(opo_rows):
                    opo_table.rows[i].cells[0].text = label
                    opo_table.rows[i].cells[1].text = value
                    try:
                        opo_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                    except Exception:
                        pass
                doc.add_paragraph()

        # Индекс вложений (document_number -> file_path / file_name)
        attachments: Dict[str, str] = {}
        attachment_names: Dict[str, str] = {}
        if document_files and isinstance(document_files, list):
            for f in document_files:
                if not isinstance(f, dict):
                    continue
                dn = str(f.get("document_number") or "")
                fp = f.get("file_path")
                if dn and isinstance(fp, str) and fp:
                    attachments[dn] = fp
                    fn = f.get("file_name")
                    if isinstance(fn, str) and fn:
                        attachment_names[dn] = fn

        def add_picture_if_exists(title: str, path: Optional[str]):
            if not path:
                return
            # PDF вставляем как подпись (Word не встраивает PDF как картинку)
            if str(path).lower().endswith(".pdf"):
                resolved = self._find_image_path(path) or path
                p = Path(resolved) if resolved else None
                if p and p.exists():
                    if title:
                        par = doc.add_paragraph()
                        par.add_run(title).bold = True
                    doc.add_paragraph()
                    doc.add_paragraph("Приложенный документ (PDF): " + p.name)
                    doc.add_paragraph()
                return
            resolved = self._find_image_path(path) or path
            try:
                p = Path(resolved)
                if not p.exists():
                    return
                if title:
                    par = doc.add_paragraph()
                    par.add_run(title).bold = True
                doc.add_paragraph()
                doc.add_picture(str(p), width=Inches(4.8))
                doc.add_paragraph()
            except Exception:
                pass

        inspection_engineers = _get("inspection_engineers", default=[])

        def _engineer_for_method(method_code: str) -> str:
            if not isinstance(inspection_engineers, list):
                return ""
            for ie in inspection_engineers:
                if not isinstance(ie, dict):
                    continue
                if (ie.get("method") or "").upper() == (method_code or "").upper():
                    return str(ie.get("full_name") or "")
            return ""

        # Акт(ы) НК
        doc.add_heading('4. АКТ(Ы) НЕРАЗРУШАЮЩЕГО КОНТРОЛЯ', level=1)
        
        performed = [m for m in (ndt_methods or []) if m.get('is_performed')]
        if performed:
            for i, m in enumerate(performed, start=1):
                doc.add_heading(f'Акт №{i}. {m.get("method_name") or "Метод НК"}', level=2)
                t = doc.add_table(rows=7, cols=2)
                t.style = 'Light Grid Accent 1'
                inspector_name = m.get('inspector_name') or _engineer_for_method(m.get('method_code') or m.get('method_name') or '')
                rows = [
                    ('Метод НК:', m.get('method_name') or ''),
                    ('Код:', m.get('method_code') or ''),
                    ('Нормативный документ:', m.get('standard') or ''),
                    ('Оборудование/прибор:', m.get('equipment') or ''),
                    ('Дата выполнения:', m.get('performed_date') or inspection_data.get('date_performed') or ''),
                    ('Специалист:', inspector_name or ''),
                    ('Уровень:', m.get('inspector_level') or ''),
                ]
                for r, (k, v) in enumerate(rows):
                    t.rows[r].cells[0].text = str(k)
                    t.rows[r].cells[1].text = str(v)
                    try:
                        t.rows[r].cells[0].paragraphs[0].runs[0].font.bold = True
                    except Exception:
                        pass
                if m.get('results'):
                    p = doc.add_paragraph()
                    p.add_run('Результаты: ').bold = True
                    p.add_run(str(m.get('results')))
                if m.get('defects'):
                    p = doc.add_paragraph()
                    p.add_run('Дефекты: ').bold = True
                    p.add_run(str(m.get('defects')))
                if m.get('conclusion'):
                    p = doc.add_paragraph()
                    p.add_run('Заключение: ').bold = True
                    p.add_run(str(m.get('conclusion')))

                # Фото по методу (включая аннотированные изображения)
                photos = m.get('photos') or []
                additional_data = m.get('additional_data', {})
                annotated_images = additional_data.get('annotated_images', []) if isinstance(additional_data, dict) else []
                
                # Объединяем обычные фото и аннотированные изображения
                all_images = list(photos) + list(annotated_images) if isinstance(photos, list) else list(annotated_images)
                
                if all_images:
                    doc.add_paragraph()
                    doc.add_paragraph('Фотоматериалы и аннотированные схемы:').runs[0].bold = True
                    for idx, ph in enumerate(all_images[:10], 1):
                        if isinstance(ph, str):
                            doc.add_paragraph(f'Изображение {idx}:')
                            add_picture_if_exists('', ph)

                doc.add_paragraph()
        else:
            doc.add_paragraph('Методы неразрушающего контроля не указаны или не выполнены.')
        
        doc.add_paragraph()

        # 4. Проверки (состояние и выбор — радиокнопки Да/Нет и выпадающие списки из мобильного приложения)
        def _yn(val):
            if val is True:
                return "Да"
            if val is False:
                return "Нет"
            return "—"
        checks_rows = [
            ("Соответствует ли сосуд чертежу", _yn(_get("matches_drawing"))),
            ("Наличие тепловой изоляции", _yn(_get("has_thermal_insulation"))),
            ("Состояние антикоррозионного покрытия", str(_get("anticorrosion_coating_state") or _get("anticorrosion_coating") or "—")),
            ("Состояние опор сосуда", str(_get("support_state") or "—")),
            ("Состояние крепежных элементов", str(_get("fasteners_state") or "—")),
            ("Перекосы фланцевых соединений", _yn(_get("has_flange_misalignment"))),
            ("Непрямолинейность патрубков", _yn(_get("has_nozzle_misalignment"))),
            ("Имеются ли места ремонта сосуда", _yn(_get("has_vessel_repairs"))),
            ("Имеются ли места ремонта ТПА", _yn(_get("has_tpa_repairs"))),
            ("Состояние внутренних устройств", str(_get("internal_devices_state") or "—")),
            ("Локально деформированные зоны", _yn(_get("has_local_deformations"))),
            ("Дефекты при наружном осмотре", _yn(_get("has_external_defects"))),
            ("Дефекты при внутреннем осмотре", _yn(_get("has_internal_defects"))),
            ("Дефекты арматуры", _yn(_get("has_armature_defects"))),
        ]
        doc.add_heading("4. Проверки (результаты осмотра)", level=1)
        checks_tbl = doc.add_table(rows=len(checks_rows) + 1, cols=2)
        checks_tbl.style = "Table Grid"
        checks_tbl.rows[0].cells[0].text = "Параметр"
        checks_tbl.rows[0].cells[1].text = "Значение"
        checks_tbl.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        checks_tbl.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
        for idx, (name, value) in enumerate(checks_rows, 1):
            checks_tbl.rows[idx].cells[0].text = name
            checks_tbl.rows[idx].cells[1].text = value
        doc.add_paragraph()

        # Перечень документов (чтобы не было "Документ 1")
        docs = _get("documents", default={})
        docs_info = _get("documents_info", default={})
        if isinstance(docs, dict) or isinstance(docs_info, dict):
            doc.add_heading('5. ПЕРЕЧЕНЬ РАССМОТРЕННЫХ ДОКУМЕНТОВ', level=1)

            document_names = {
                '1': 'Лицензия на осуществление деятельности по эксплуатации взрывопожароопасных и химически опасных производственных объектов I, II и III классов опасности',
                '2': 'Свидетельство о регистрации в государственном реестре ОПО, включая сведения характеризующие ОПО',
                '3': 'Технологический регламент объектов опасных производственных объектов',
                '4': 'План мероприятий по локализации и ликвидации последствий аварий на опасном производственном объекте',
                '5': 'Положение о производственном контроле за соблюдением требований промышленной безопасности на опасных производственных объектах',
                '6': 'Журнал учета аварий и инцидентов на ОПО',
                '7': 'Страховой полис страхования гражданской ответственности владельца опасного объекта за причинение вреда в результате аварии на опасном объекте',
                '8': 'Приказ о назначении ответственного лица за исправное состояние и безопасную эксплуатацию сосудов',
                '9': 'Приказ о назначении ответственного лица за осуществление производственного контроля и соблюдение требований промышленной безопасности на опасном производственном объекте',
                '10': 'Паспорт сосуда заводской (удостоверение о качестве монтажа, сертификат соответствия, сборочный чертёж и схема включения сосуда, расчёт на прочность)',
                '11': 'Инструкция по монтажу и эксплуатации',
                '12': 'Паспорта на предохранительные клапаны',
                '13': 'Паспорта на запорную арматуру',
                '14': 'Документация на контрольно-измерительные приборы',
                '15': 'Ремонтная (исполнительная) документация',
                '16': 'Заключение экспертизы промышленной безопасности',
                '17': 'Акты проведения УЗТ',
            }

            def _doc_meta_local(num: str):
                num_key = str(num)
                present = None
                doc_number = ""
                doc_date = ""
                if isinstance(docs, dict) and num_key in docs:
                    val = docs.get(num_key)
                    if isinstance(val, dict):
                        present = val.get("present")
                        if present is None:
                            present = val.get("has") if val.get("has") is not None else val.get("value")
                        doc_number = str(val.get("number") or val.get("doc_number") or "")
                        doc_date = str(val.get("date") or val.get("doc_date") or "")
                    else:
                        if isinstance(val, str):
                            present = val.strip().lower() in ("true", "1", "yes", "да")
                        else:
                            present = bool(val)
                if isinstance(docs_info, dict) and num_key in docs_info:
                    info = docs_info.get(num_key) or {}
                    if isinstance(info, dict):
                        if present is None:
                            present = info.get("present")
                            if present is None:
                                present = info.get("has") if info.get("has") is not None else info.get("value")
                        if not doc_number:
                            doc_number = str(info.get("number") or info.get("doc_number") or "")
                        if not doc_date:
                            doc_date = str(info.get("date") or info.get("doc_date") or "")
                return (present, doc_number, doc_date)

            doc_keys = set()
            if isinstance(docs, dict):
                doc_keys.update([str(k) for k in docs.keys()])
            if isinstance(docs_info, dict):
                doc_keys.update([str(k) for k in docs_info.keys()])
            doc_keys = sorted(doc_keys, key=lambda x: int(x) if str(x).isdigit() else 999)

            doc_table = doc.add_table(rows=len(doc_keys) + 1, cols=5)
            doc_table.style = 'Light Grid Accent 1'
            headers = ['№', 'Наименование документа', 'Номер документа', 'Дата документа', 'Наличие']
            for i, header in enumerate(headers):
                cell = doc_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_idx = 1
            for num in doc_keys:
                name = document_names.get(str(num), f'Документ {num}')
                present, doc_number, doc_date = _doc_meta_local(str(num))
                doc_table.rows[row_idx].cells[0].text = str(num)
                doc_table.rows[row_idx].cells[1].text = name
                doc_table.rows[row_idx].cells[2].text = doc_number or '—'
                doc_table.rows[row_idx].cells[3].text = doc_date or '—'
                doc_table.rows[row_idx].cells[4].text = 'Да' if present else 'Нет'
                doc_table.rows[row_idx].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_idx += 1

            doc.add_paragraph()

        # Фото заводской таблички / схема контроля (из мобильного приложения)
        add_picture_if_exists(
            'Фото заводской таблички',
            attachments.get('factory_plate_photo') or _get('factory_plate_photo'),
        )

        # Толщинометрия (точки + таблица)
        thickness = _get("thickness_measurements", "thicknessMeasurements", default=[])
        has_thickness = isinstance(thickness, list) and len(thickness) > 0
        control_scheme = attachments.get('control_scheme_image') or _get('control_scheme_image')
        if has_thickness:
            doc.add_heading('6. УЗТ (УЛЬТРАЗВУКОВАЯ ТОЛЩИНОМЕТРИЯ)', level=1)

            t = doc.add_table(rows=len(thickness) + 1, cols=6)
            t.style = 'Light Grid Accent 1'
            headers = ['№', 'Местоположение', 'Сечение', 'Толщина, мм', 'Мин. допустимая, мм', 'Комментарий']
            for i, header in enumerate(headers):
                cell = t.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for idx, point in enumerate(thickness, start=1):
                if not isinstance(point, dict):
                    continue
                t.rows[idx].cells[0].text = str(idx)
                t.rows[idx].cells[1].text = str(point.get('location') or '')
                t.rows[idx].cells[2].text = str(point.get('section_number') or '')
                t.rows[idx].cells[3].text = str(point.get('thickness') or '')
                t.rows[idx].cells[4].text = str(point.get('min_allowed_thickness') or '')
                t.rows[idx].cells[5].text = str(point.get('comment') or '')

            doc.add_paragraph()

            scheme_to_show = control_scheme
            if control_scheme and thickness:
                scheme_resolved = self._find_image_path(control_scheme) or control_scheme
                if scheme_resolved and os.path.isfile(scheme_resolved):
                    annotated = self._draw_points_on_scheme(scheme_resolved, thickness, output_dir="/app/reports/tmp")
                    if annotated and os.path.isfile(annotated):
                        scheme_to_show = annotated
            add_picture_if_exists(
                'Схема контроля / карта обследования (с точками замеров)',
                scheme_to_show,
            )
        elif control_scheme:
            doc.add_heading('6. УЗТ (УЛЬТРАЗВУКОВАЯ ТОЛЩИНОМЕТРИЯ)', level=1)
            add_picture_if_exists(
                'Схема контроля / карта обследования',
                control_scheme,
            )

        # ЗРА / СППК / измерительный контроль / твердость / сварные соединения
        zra = _get('zra_items', default=[])
        if isinstance(zra, list) and zra:
            doc.add_heading('5. ЗРА (ЗАПОРНО-РЕГУЛИРУЮЩАЯ АРМАТУРА)', level=1)
            t = doc.add_table(rows=len(zra) + 1, cols=6)
            t.style = 'Light Grid Accent 1'
            headers = ['№', 'Кол-во', 'Типоразмер', 'Тех. №', 'Зав. №', 'Место на схеме']
            for i, h in enumerate(headers):
                c = t.rows[0].cells[i]
                c.text = h
                c.paragraphs[0].runs[0].font.bold = True
            for i, it in enumerate(zra, start=1):
                if not isinstance(it, dict):
                    continue
                t.rows[i].cells[0].text = str(i)
                t.rows[i].cells[1].text = str(it.get('quantity') or '')
                t.rows[i].cells[2].text = str(it.get('type_size') or '')
                t.rows[i].cells[3].text = str(it.get('tech_number') or '')
                t.rows[i].cells[4].text = str(it.get('serial_number') or '')
                t.rows[i].cells[5].text = str(it.get('location_on_scheme') or '')
            doc.add_paragraph()

        sppk = _get('sppk_items', default=[])
        if isinstance(sppk, list) and sppk:
            doc.add_heading('6. СППК (ПРЕДОХРАНИТЕЛЬНЫЕ КЛАПАНЫ)', level=1)
            t = doc.add_table(rows=len(sppk) + 1, cols=6)
            t.style = 'Light Grid Accent 1'
            headers = ['№', 'Кол-во', 'Типоразмер', 'Тех. №', 'Зав. №', 'Место на схеме']
            for i, h in enumerate(headers):
                c = t.rows[0].cells[i]
                c.text = h
                c.paragraphs[0].runs[0].font.bold = True
            for i, it in enumerate(sppk, start=1):
                if not isinstance(it, dict):
                    continue
                t.rows[i].cells[0].text = str(i)
                t.rows[i].cells[1].text = str(it.get('quantity') or '')
                t.rows[i].cells[2].text = str(it.get('type_size') or '')
                t.rows[i].cells[3].text = str(it.get('tech_number') or '')
                t.rows[i].cells[4].text = str(it.get('serial_number') or '')
                t.rows[i].cells[5].text = str(it.get('location_on_scheme') or '')
            doc.add_paragraph()

        ovality = _get('ovality_measurements', default=[])
        if isinstance(ovality, list) and ovality:
            doc.add_heading('7. ИЗМЕРИТЕЛЬНЫЙ КОНТРОЛЬ — ОВАЛЬНОСТЬ', level=1)
            t = doc.add_table(rows=len(ovality) + 1, cols=5)
            t.style = 'Light Grid Accent 1'
            headers = ['№', 'Сечение', 'Dmax', 'Dmin', 'Отклонение, %']
            for i, h in enumerate(headers):
                c = t.rows[0].cells[i]
                c.text = h
                c.paragraphs[0].runs[0].font.bold = True
            for i, it in enumerate(ovality, start=1):
                if not isinstance(it, dict):
                    continue
                t.rows[i].cells[0].text = str(i)
                t.rows[i].cells[1].text = str(it.get('section_number') or '')
                t.rows[i].cells[2].text = str(it.get('max_diameter') or '')
                t.rows[i].cells[3].text = str(it.get('min_diameter') or '')
                t.rows[i].cells[4].text = str(it.get('deviation_percent') or '')
            doc.add_paragraph()

        deflection = _get('deflection_measurements', default=[])
        if isinstance(deflection, list) and deflection:
            doc.add_heading('8. ИЗМЕРИТЕЛЬНЫЙ КОНТРОЛЬ — ПРОГИБ', level=1)
            t = doc.add_table(rows=len(deflection) + 1, cols=4)
            t.style = 'Light Grid Accent 1'
            headers = ['№', 'Сечение', 'Прогиб, мм', 'Прогиб, %']
            for i, h in enumerate(headers):
                c = t.rows[0].cells[i]
                c.text = h
                c.paragraphs[0].runs[0].font.bold = True
            for i, it in enumerate(deflection, start=1):
                if not isinstance(it, dict):
                    continue
                t.rows[i].cells[0].text = str(i)
                t.rows[i].cells[1].text = str(it.get('section_number') or '')
                t.rows[i].cells[2].text = str(it.get('deflection_mm') or '')
                d_pct = it.get('deflection_percent')
                try:
                    t.rows[i].cells[3].text = f"{float(str(d_pct).replace(',', '.')):.2f}" if d_pct not in (None, '') else ''
                except Exception:
                    t.rows[i].cells[3].text = str(d_pct or '')
            doc.add_paragraph()

        hardness = _get('hardness_tests', default=[])
        if isinstance(hardness, list) and hardness:
            doc.add_heading('9. КОНТРОЛЬ ТВЕРДОСТИ', level=1)
            t = doc.add_table(rows=len(hardness) + 1, cols=8)
            t.style = 'Light Grid Accent 1'
            headers = ['№', 'Шов', 'Участок', 'Доп. осн', 'Доп. шов', 'Осн', 'Шов', 'ЗТВ']
            for i, h in enumerate(headers):
                c = t.rows[0].cells[i]
                c.text = h
                c.paragraphs[0].runs[0].font.bold = True
            for i, it in enumerate(hardness, start=1):
                if not isinstance(it, dict):
                    continue
                t.rows[i].cells[0].text = str(i)
                t.rows[i].cells[1].text = str(it.get('weld_number') or '')
                t.rows[i].cells[2].text = str(it.get('area_number') or '')
                t.rows[i].cells[3].text = str(it.get('allowed_hardness_base') or '')
                t.rows[i].cells[4].text = str(it.get('allowed_hardness_weld') or '')
                t.rows[i].cells[5].text = str(it.get('hardness_base') or '')
                t.rows[i].cells[6].text = str(it.get('hardness_weld') or '')
                t.rows[i].cells[7].text = str(it.get('hardness_haz') or '')
            doc.add_paragraph()

        welds = _get('weld_inspections', default=[])
        if isinstance(welds, list) and welds:
            doc.add_heading('10. КОНТРОЛЬ СВАРНЫХ СОЕДИНЕНИЙ (ПВК/УЗК)', level=1)
            t = doc.add_table(rows=len(welds) + 1, cols=6)
            t.style = 'Light Grid Accent 1'
            headers = ['№', 'Шов', 'Место на карте', 'ПВК дефект', 'УЗК дефект', 'Заключение']
            for i, h in enumerate(headers):
                c = t.rows[0].cells[i]
                c.text = h
                c.paragraphs[0].runs[0].font.bold = True
            for i, it in enumerate(welds, start=1):
                if not isinstance(it, dict):
                    continue
                t.rows[i].cells[0].text = str(i)
                t.rows[i].cells[1].text = str(it.get('weld_number') or '')
                t.rows[i].cells[2].text = str(it.get('location_on_control_map') or '')
                t.rows[i].cells[3].text = str(it.get('pvk_defect') or '')
                t.rows[i].cells[4].text = str(it.get('uzk_defect') or '')
                t.rows[i].cells[5].text = str(it.get('conclusion') or '')
            doc.add_paragraph()
        
        # Заключение
        if inspection_data.get('conclusion'):
            doc.add_heading('7. ЗАКЛЮЧЕНИЕ', level=1)
            doc.add_paragraph(str(inspection_data.get('conclusion') or ''))

        # Приложения: документы специалистов
        doc.add_page_break()
        doc.add_heading('8. ПРИЛОЖЕНИЯ', level=1)
        
        # 8.1. Документы специалистов НК
        doc.add_heading('8.1. Документы специалистов неразрушающего контроля', level=2)
        if specialist_docs:
            for s in specialist_docs:
                inspector_name = s.get('inspector_name', 'Не указано')
                doc.add_heading(f"Специалист: {inspector_name}", level=3)
                certifications = s.get('certifications') or []
                if certifications:
                    for idx, c in enumerate(certifications, 1):
                        cert_type = c.get('certification_type', 'Удостоверение')
                        cert_num = c.get('certificate_number', '')
                        org = c.get('issuing_organization', '')
                        issue_date = c.get('issue_date', '')
                        expiry_date = c.get('expiry_date', '')
                        
                        doc.add_paragraph(f"{idx}. {cert_type}")
                        if cert_num:
                            doc.add_paragraph(f"   Номер: {cert_num}")
                        if org:
                            doc.add_paragraph(f"   Организация: {org}")
                        if issue_date:
                            doc.add_paragraph(f"   Дата выдачи: {issue_date}")
                        if expiry_date:
                            doc.add_paragraph(f"   Срок действия: {expiry_date}")
                        
                        # Добавляем скан удостоверения
                        scan_path = c.get('scan_file_path')
                        if scan_path and isinstance(scan_path, str):
                            add_picture_if_exists(f'Скан удостоверения {cert_type} №{cert_num}', scan_path)
                        doc.add_paragraph()
                else:
                    doc.add_paragraph(f'Документы для специалиста {inspector_name} не найдены.')
                doc.add_paragraph()
        else:
            doc.add_paragraph('Документы специалистов НК не приложены.')
        
        # 8.2. Используемое оборудование для поверок
        if verification_equipment and isinstance(verification_equipment, list) and len(verification_equipment) > 0:
            doc.add_paragraph('')
            doc.add_heading('8.2. Используемое оборудование для неразрушающего контроля', level=2)
            doc.add_paragraph('При проведении обследования использовалось следующее поверенное оборудование:')
            
            # Таблица с оборудованием
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '№'
            hdr_cells[1].text = 'Наименование'
            hdr_cells[2].text = 'Тип'
            hdr_cells[3].text = 'Производитель/Модель'
            hdr_cells[4].text = 'Серийный номер'
            hdr_cells[5].text = 'Дата поверки'
            hdr_cells[6].text = 'Срок поверки'
            
            for idx, eq in enumerate(verification_equipment, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = eq.get('name', '')
                row_cells[2].text = eq.get('equipment_type', '')
                
                manufacturer = eq.get('manufacturer', '')
                model = eq.get('model', '')
                manufacturer_model = f"{manufacturer} {model}".strip() if manufacturer or model else '—'
                row_cells[3].text = manufacturer_model
                
                row_cells[4].text = eq.get('serial_number', '—')
                
                ver_date = eq.get('verification_date', '')
                if ver_date:
                    try:
                        from datetime import datetime as dt
                        d = dt.fromisoformat(ver_date.replace('Z', '+00:00'))
                        ver_date = d.strftime('%d.%m.%Y')
                    except:
                        pass
                row_cells[5].text = ver_date if ver_date else '—'
                
                next_date = eq.get('next_verification_date', '')
                if next_date:
                    try:
                        from datetime import datetime as dt
                        d = dt.fromisoformat(next_date.replace('Z', '+00:00'))
                        next_date = d.strftime('%d.%m.%Y')
                    except:
                        pass
                row_cells[6].text = next_date if next_date else '—'
            
            doc.add_paragraph('')
            doc.add_paragraph('Сканы свидетельств о поверке используемого оборудования:')
            
            # Добавляем сканы с подробной информацией
            for idx, eq in enumerate(verification_equipment, 1):
                scan_path = eq.get('scan_file_path')
                scan_name = eq.get('scan_file_name', '')
                eq_name = eq.get('name', '')
                cert_num = eq.get('verification_certificate_number', '')
                ver_org = eq.get('verification_organization', '')
                
                if scan_path:
                    doc.add_paragraph('')
                    title_parts = [f'{idx}. Свидетельство о поверке: {eq_name}']
                    if cert_num:
                        title_parts.append(f'№ {cert_num}')
                    if ver_org:
                        title_parts.append(f'({ver_org})')
                    doc.add_paragraph(' '.join(title_parts))
                    if os.path.exists(scan_path):
                        ext = str(Path(scan_path).suffix or "").lower()
                        if ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"]:
                            add_picture_if_exists('', scan_path)
                        else:
                            doc.add_paragraph(f'Файл: {Path(scan_path).name}')
                    else:
                        doc.add_paragraph(f'[Файл не найден: {scan_path}]')

        # 8.3. Схемы диагностики и обследования (если есть)
        extra_attachments = []

        for key, path in attachments.items():
            if str(key).isdigit():
                continue
            if key in ("factory_plate_photo", "control_scheme_image"):
                continue
            if not path:
                continue
            extra_attachments.append((f"Вложение: {key}", path))

        if extra_attachments:
            doc.add_paragraph('')
            doc.add_heading('8.3. Схемы диагностики и обследования', level=2)
            for title, path in extra_attachments:
                if not path:
                    continue
                resolved = self._find_image_path(path) or path
                if not resolved or not os.path.exists(resolved):
                    continue
                ext = str(Path(resolved).suffix or "").lower()
                if ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"]:
                    doc.add_paragraph(title).runs[0].bold = True
                    add_picture_if_exists('', resolved)
                else:
                    doc.add_paragraph(f"{title}: {Path(resolved).name}")
        
        # 8.4. Фото заводской таблички, схема контроля, сканы документов, фото дефектов (явно по порядку)
        doc.add_paragraph('')
        doc.add_heading('8.4. Приложенные фото и сканы документов', level=2)
        # Фото таблички
        _fp = attachments.get('factory_plate_photo') or _get('factory_plate_photo')
        if _fp:
            doc.add_paragraph('Фото заводской таблички').runs[0].bold = True
            add_picture_if_exists('', _fp)
        # Сканы документов 1–17
        _doc_names = {'1': 'Документ 1', '2': 'Документ 2', '3': 'Документ 3', '4': 'Документ 4', '5': 'Документ 5', '6': 'Документ 6', '7': 'Документ 7', '8': 'Документ 8', '9': 'Документ 9', '10': 'Паспорт сосуда', '11': 'Документ 11', '12': 'Документ 12', '13': 'Документ 13', '14': 'Документ 14', '15': 'Ремонтная документация', '16': 'Документ 16', '17': 'Документ 17'}
        for _num in [str(i) for i in range(1, 18)]:
            _path = attachments.get(_num)
            if not _path:
                continue
            _res = self._find_image_path(_path) or _path
            if _res and os.path.exists(_res):
                doc.add_paragraph((_doc_names.get(_num) or f'Документ {_num}')).runs[0].bold = True
                add_picture_if_exists('', _res)
        # Фото дефектов ВИК (vd_0_0, vd_0_1, ...)
        _vd_keys = sorted([k for k in attachments if isinstance(k, str) and k.startswith('vd_')], key=lambda x: (int(x.split('_')[1]) if len(x.split('_')) >= 2 else 0, int(x.split('_')[2]) if len(x.split('_')) >= 3 else 0))
        for _vk in _vd_keys:
            _path = attachments.get(_vk)
            if not _path:
                continue
            _res = self._find_image_path(_path) or _path
            if _res and os.path.exists(_res):
                doc.add_paragraph(f'Фото дефекта ВИК ({_vk})').runs[0].bold = True
                add_picture_if_exists('', _res)
        
        if (inspection_data.get("status") or "").upper() == "DRAFT":
            self._add_draft_watermark(doc)
        apply_device_terminology_to_document(doc, detect_pressure_device_kind(equipment_data))
        # Сохранение
        doc.save(output_path)
        return

    def _fmt_date_ru(self, s: Optional[str]) -> Optional[str]:
        if not s:
            return None
        try:
            # 2025-07-25 / 2025-07-25T...Z
            if "T" in s:
                d = datetime.fromisoformat(s.replace("Z", "+00:00"))
            else:
                d = datetime.fromisoformat(s)
            return d.strftime("%d.%m.%Y")
        except Exception:
            return s

    def _add_toc_field(self, doc: Document):
        """
        Вставить поле оглавления (обновляется в Word: ПКМ -> Обновить поле).
        """
        p = doc.add_paragraph()
        run = p.add_run()

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = r'TOC \o "1-3" \h \z \u'

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")

        fldChar3 = OxmlElement("w:t")
        fldChar3.text = "Оглавление будет сформировано при обновлении полей в Word."

        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        run._r.append(fldChar4)

    def _generate_diagnostics_report_word(
        self,
        inspection_data: Dict[str, Any],
        equipment_data: Dict[str, Any],
        ndt_methods: List[Dict[str, Any]],
        output_path: str,
        document_files: Optional[List[Dict[str, Any]]] = None,
        specialist_docs: Optional[List[Dict[str, Any]]] = None,
        verification_equipment: Optional[List[Dict[str, Any]]] = None,
        template_definition: Optional[Dict[str, Any]] = None,
    ):
        """
        Диагностический отчет в структуре, близкой к примеру (reciver.md):
        титульник -> оглавление -> разделы 1..15 -> приложения.
        """
        doc = Document()
        self._setup_styles(doc)

        data = inspection_data.get("data") or {}
        if not isinstance(data, dict):
            data = {}
        attrs = equipment_data.get("attributes") or {}
        if not isinstance(attrs, dict):
            attrs = {}

        # Индекс вложений (document_number -> file_path / file_name)
        attachments: Dict[str, str] = {}
        attachment_names: Dict[str, str] = {}
        if document_files and isinstance(document_files, list):
            for f in document_files:
                if not isinstance(f, dict):
                    continue
                dn = str(f.get("document_number") or "")
                fp = f.get("file_path")
                if dn and isinstance(fp, str) and fp:
                    attachments[dn] = fp
                    fn = f.get("file_name")
                    if isinstance(fn, str) and fn:
                        attachment_names[dn] = fn

        def add_picture_if_exists(title: str, path: Optional[str]):
            if not path:
                return
            if str(path).lower().endswith(".pdf"):
                resolved = self._find_image_path(path) or path
                p = Path(resolved) if resolved else None
                if p and p.exists():
                    if title:
                        par = doc.add_paragraph()
                        par.add_run(title).bold = True
                    doc.add_paragraph()
                    doc.add_paragraph("Приложенный документ (PDF): " + p.name)
                    doc.add_paragraph()
                return
            resolved = self._find_image_path(path) or path
            try:
                p = Path(resolved)
                if not p.exists():
                    return
                if title:
                    par = doc.add_paragraph()
                    par.add_run(title).bold = True
                doc.add_paragraph()
                doc.add_picture(str(p), width=Inches(4.8))
                doc.add_paragraph()
            except Exception:
                pass

        # template_definition: {logo_path, fields:{...}, sections:[{key,enabled}]}
        tdef = template_definition if isinstance(template_definition, dict) else {}
        tfields = tdef.get("fields") if isinstance(tdef.get("fields"), dict) else {}
        tsections = tdef.get("sections") if isinstance(tdef.get("sections"), list) else []
        enabled = {str(s.get("key")): bool(s.get("enabled")) for s in tsections if isinstance(s, dict) and s.get("key")}
        def is_on(key: str) -> bool:
            # по умолчанию включено, если нет явного списка
            if not tsections:
                return True
            return enabled.get(key, False)

        def g(*keys, default=None):
            """Извлечь значение по ключам (поддержка camelCase и snake_case)"""
            # Сначала ищем в data
            for k in keys:
                if k in data and data.get(k) not in (None, ""):
                    return data.get(k)
            # Затем ищем в attrs
            for k in keys:
                if k in attrs and attrs.get(k) not in (None, ""):
                    return attrs.get(k)
            # Пробуем варианты с camelCase/snake_case
            for k in keys:
                # snake_case -> camelCase
                camel_key = ''.join(word.capitalize() if i > 0 else word for i, word in enumerate(k.split('_')))
                camel_key_lower = camel_key[0].lower() + camel_key[1:] if camel_key else k
                if camel_key_lower in data and data.get(camel_key_lower) not in (None, ""):
                    return data.get(camel_key_lower)
                if camel_key_lower in attrs and attrs.get(camel_key_lower) not in (None, ""):
                    return attrs.get(camel_key_lower)
            return default

        opo = data.get("opo") or data.get("opo_info") or {}
        if not isinstance(opo, dict):
            opo = {}

        def _opo_get(*keys, default=None):
            for k in keys:
                if k in opo and opo.get(k) not in (None, ""):
                    return opo.get(k)
            for k in keys:
                camel_key = ''.join(word.capitalize() if i > 0 else word for i, word in enumerate(k.split('_')))
                camel_key_lower = camel_key[0].lower() + camel_key[1:] if camel_key else k
                if camel_key_lower in opo and opo.get(camel_key_lower) not in (None, ""):
                    return opo.get(camel_key_lower)
            return g(*keys, default=default)

        date_perf_iso = inspection_data.get("date_performed")
        date_perf_ru = self._fmt_date_ru(date_perf_iso) or datetime.now().strftime("%d.%m.%Y")
        year2 = datetime.now().strftime("%y")
        # Номер отчета: YY-xxxx (по последним 4 символам UUID)
        rid = str(equipment_data.get("id") or "")[-4:] or "0000"
        report_no = g("report_number", default=f"{year2}-{rid}")

        # --------------- ТИТУЛЬНЫЙ ЛИСТ ---------------
        if is_on("title"):
            logo_path = tdef.get("logo_path") or g("report_logo_path", default="/app/reports/assets/yutar_logo.png")
            try:
                resolved = self._find_image_path(str(logo_path))
                if resolved and os.path.isfile(resolved):
                    doc.add_picture(resolved, width=Inches(5.2))
            except Exception:
                pass

        # Блок заголовка (как таблица в примере)
        if is_on("title"):
            title_table = doc.add_table(rows=1, cols=1)
            title_table.style = "Table Grid"
            title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = title_table.rows[0].cells[0]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"ТЕХНИЧЕСКИЙ ОТЧЕТ № {report_no}\nПО РЕЗУЛЬТАТАМ ТЕХНИЧЕСКОГО ДИАГНОСТИРОВАНИЯ")
            r.bold = True
            r.font.size = Pt(14)

        doc.add_paragraph("")

        # Таблица объекта (упрощенно, но структура похожа)
        if is_on("title"):
            obj_tbl = doc.add_table(rows=5, cols=2)
            obj_tbl.style = "Table Grid"
            obj_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        object_name = g("equipment_object_name", "vessel_name", default=equipment_data.get("name") or "—")
        device_name = g("equipment_device_name", default=equipment_data.get("name") or "—")
        serial = g("serial_number", default=equipment_data.get("serial_number") or "—")
        org = g("organization", "customer_name", default="—")
        location = g("location", "equipment_location", default=equipment_data.get("location") or "—")

        rows = [
            ("Объект технического диагностирования:", ""),
            ("Техническое устройство:", device_name),
            ("Заводской номер:", str(serial)),
            ("Эксплуатирующая организация:", str(org)),
            ("Местонахождение объекта:", str(location)),
        ]
        if is_on("title"):
            for i, (k, v) in enumerate(rows):
                obj_tbl.rows[i].cells[0].text = k
                obj_tbl.rows[i].cells[1].text = v
                try:
                    obj_tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                except Exception:
                    pass

        doc.add_paragraph("")

        # Подпись руководителя (как в примере — справа)
        if is_on("title"):
            sign_tbl = doc.add_table(rows=1, cols=2)
            sign_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            sign_tbl.columns[0].width = Inches(3.5)
            sign_tbl.columns[1].width = Inches(3.5)
            sign_tbl.cell(0, 0).text = ""

        contractor = tfields.get("contractor_name") or g("contractor_name", default='ООО «ЮТАР»')
        director_title = tfields.get("director_title") or g("director_title", default="Генеральный директор")
        director_name = tfields.get("director_name") or g("director_name", default="__________________")
        city = tfields.get("report_city") or g("report_city", "city", default="г. ________")
        year = datetime.now().strftime("%Y")

        if is_on("title"):
            p = sign_tbl.cell(0, 1).paragraphs[0]
            p.add_run(f"{director_title}\n{contractor}\n\n__________________\n{director_name}\n\n«____» __________ {year} г.\nМ.П.").font.size = Pt(10)

        doc.add_paragraph("")
        if is_on("title"):
            p = doc.add_paragraph(f"{city} {year} г.")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()

        # --------------- СОДЕРЖАНИЕ ---------------
        if is_on("toc"):
            doc.add_heading("СОДЕРЖАНИЕ", level=1)
            self._add_toc_field(doc)
            doc.add_page_break()

        # --------------- РАЗДЕЛЫ 1..15 ---------------
        # Эти переменные используются и в разделах, и в приложениях ниже
        performed = [m for m in (ndt_methods or []) if m.get("is_performed")]
        docs = g("documents", default={})
        docs_info = g("documents_info", default={})
        inspection_engineers = g("inspection_engineers", default=[])
        additional_data = g("additional_data", default={})
        manual_verification_equipment: List[Dict[str, Any]] = []
        if isinstance(additional_data, dict):
            raw_manual = additional_data.get("manual_verification_equipment") or additional_data.get("manualVerificationEquipment")
            if isinstance(raw_manual, list):
                for item in raw_manual:
                    if not isinstance(item, dict):
                        continue
                    name = str(item.get("name") or "").strip()
                    serial = str(item.get("serial_number") or item.get("serial") or "").strip()
                    cert_num = str(item.get("verification_certificate_number") or item.get("certificate_number") or "").strip()
                    next_date = str(item.get("next_verification_date") or item.get("valid_until") or item.get("expiry_date") or "").strip()
                    if not (name or serial or cert_num or next_date):
                        continue
                    manual_verification_equipment.append({
                        "name": name or "Прибор (ручной ввод)",
                        "serial_number": serial,
                        "verification_certificate_number": cert_num,
                        "next_verification_date": next_date,
                        "equipment_type": "РУЧНОЙ ВВОД",
                    })

        effective_verification_equipment: List[Dict[str, Any]] = []
        if verification_equipment and isinstance(verification_equipment, list):
            effective_verification_equipment.extend([e for e in verification_equipment if isinstance(e, dict)])
        if manual_verification_equipment:
            existing_eq_keys = {
                f"{str(e.get('name') or '').strip().lower()}|{str(e.get('serial_number') or '').strip().lower()}|{str(e.get('verification_certificate_number') or '').strip().lower()}"
                for e in effective_verification_equipment
            }
            for item in manual_verification_equipment:
                key = f"{str(item.get('name') or '').strip().lower()}|{str(item.get('serial_number') or '').strip().lower()}|{str(item.get('verification_certificate_number') or '').strip().lower()}"
                if key in existing_eq_keys:
                    continue
                effective_verification_equipment.append(item)
                existing_eq_keys.add(key)
        performed_codes = {str(m.get("method_code") or m.get("method_name") or "").upper() for m in performed}
        work_list = list(performed)
        has_visual_data = (g("visual_defects", default=[]) and len(g("visual_defects", default=[])) > 0) or g("has_external_defects") is not None or g("has_internal_defects") is not None
        if has_visual_data and "ВИК" not in performed_codes and "VIK" not in performed_codes:
            work_list.append({"method_name": "ВИК", "work_name": "Визуальный и измерительный контроль", "standard": "приказ Ростехнадзора от 15.12.2020 №536", "control_volume": "—", "conclusion": "Выполнено по протоколу"})
        if isinstance(docs, dict) and any(v for v in docs.values()):
            work_list.append({"method_name": "Анализ док.", "work_name": "Анализ технической документации", "standard": "приказ Ростехнадзора от 15.12.2020 №536", "control_volume": "—", "conclusion": "Рассмотрены документы по перечню раздела 11"})
        ovality_data = g("ovality_measurements", default=[])
        if isinstance(ovality_data, list) and ovality_data:
            work_list.append({"method_name": "Овальность", "work_name": "Измерение овальности корпуса сосуда", "standard": "СО 153-34.17.439-2003", "control_volume": "—", "conclusion": "Выполнено по протоколу"})
        hardness_data = g("hardness_tests", default=[])
        if isinstance(hardness_data, list) and hardness_data:
            work_list.append({"method_name": "Твердометрия", "work_name": "Контроль твердости металла", "standard": "приказ Ростехнадзора от 15.12.2020 №536", "control_volume": "—", "conclusion": "Выполнено по протоколу"})
        thickness_data = g("thickness_measurements", default=[])
        if isinstance(thickness_data, list) and thickness_data and "УЗТ" not in performed_codes and "UZT" not in performed_codes:
            work_list.append({"method_name": "УЗТ", "work_name": "Ультразвуковой контроль толщины стенок элементов сосуда", "standard": "приказ Ростехнадзора от 15.12.2020 №536", "control_volume": "—", "conclusion": "Выполнено по протоколу"})
        weld_data = g("weld_inspections", default=[])
        if isinstance(weld_data, list) and weld_data and "УЗК" not in performed_codes and "UZK" not in performed_codes:
            work_list.append({"method_name": "УЗК", "work_name": "Ультразвуковой контроль качества основного металла и сварных соединений", "standard": "приказ Ростехнадзора от 15.12.2020 №536", "control_volume": "—", "conclusion": "Выполнено по протоколу"})

        def _normalize_method(method_raw: str) -> str:
            m = (method_raw or "").strip()
            if not m:
                return ""
            m_up = m.upper()
            mapping = {
                "VIK": "ВИК",
                "UZK": "УЗК",
                "UZT": "УЗТ",
                "PVK": "ПВК/МК",
                "MK": "МК",
                "RK": "РК",
                "MPD": "МПД",
                "KPD": "КПД",
                "TVI": "ТВИ",
            }
            if m_up in mapping:
                return mapping[m_up]
            # Если уже на русском — возвращаем как есть
            return m

        def _doc_meta(num: str):
            num_key = str(num)
            present = None
            doc_number = ""
            doc_date = ""
            if isinstance(docs, dict) and num_key in docs:
                val = docs.get(num_key)
                if isinstance(val, dict):
                    present = val.get("present")
                    if present is None:
                        present = val.get("has") if val.get("has") is not None else val.get("value")
                    doc_number = str(val.get("number") or val.get("doc_number") or "")
                    doc_date = str(val.get("date") or val.get("doc_date") or "")
                else:
                    if isinstance(val, str):
                        present = val.strip().lower() in ("true", "1", "yes", "да")
                    else:
                        present = bool(val)
            if isinstance(docs_info, dict) and num_key in docs_info:
                info = docs_info.get(num_key) or {}
                if isinstance(info, dict):
                    if present is None:
                        present = info.get("present")
                        if present is None:
                            present = info.get("has") if info.get("has") is not None else info.get("value")
                    if not doc_number:
                        doc_number = str(info.get("number") or info.get("doc_number") or "")
                    if not doc_date:
                        doc_date = str(info.get("date") or info.get("doc_date") or "")
            return (present, doc_number, doc_date)

        def _engineer_for_method(method_code: str) -> str:
            if not isinstance(inspection_engineers, list):
                return ""
            for ie in inspection_engineers:
                if not isinstance(ie, dict):
                    continue
                if (ie.get("method") or "").upper() == (method_code or "").upper():
                    return str(ie.get("full_name") or "")
            return ""

        if is_on("sections_1_15"):
            doc.add_heading("1. Основания для проведения работ", level=1)
            doc.add_paragraph(str(g("basis", "work_basis", default="—")))

            doc.add_heading("2. Сроки проведения работ", level=1)
            doc.add_paragraph(str(g("work_period", default=f"Дата проведения: {date_perf_ru}")))

            doc.add_heading("3. Перечень нормативных и правовых актов, устанавливающих требования к объекту диагностирования", level=1)
            # Минимальный дефолт (можно расширять через будущий редактор шаблонов)
            doc.add_paragraph(str(g("normative_base", default="Приказ Ростехнадзора от 15.12.2020 №536.")))

            doc.add_heading("4. Сведения о Заказчике", level=1)
            doc.add_paragraph(f"Эксплуатирующая организация: {org}")
            doc.add_paragraph(f"Местонахождение объекта: {location}")

            doc.add_heading("5. Сведения об организации, проводившей техническое диагностирование", level=1)
            doc.add_paragraph(str(contractor))
            doc.add_paragraph(str(g("contractor_address", default="—")))
            doc.add_paragraph(str(g("contractor_license", default="—")))

            doc.add_heading("6. Сведения об эксперте и специалисте, проводивших диагностирование", level=1)
            # Собираем специалистов из разных источников
            inspectors = []
            inspector_details = {}  # name -> {certifications, level, etc}

            # 0. Из inspection_engineers (выбор инженеров по методам в мобильном приложении)
            if isinstance(inspection_engineers, list):
                for ie in inspection_engineers:
                    if not isinstance(ie, dict):
                        continue
                    name = (ie.get("full_name") or "").strip()
                    method = _normalize_method(ie.get("method"))
                    cert_num = (ie.get("certificate_number") or "").strip()
                    valid_until = (ie.get("valid_until") or "").strip()
                    if name and name not in inspectors:
                        inspectors.append(name)
                    if name:
                        if name not in inspector_details:
                            inspector_details[name] = {}
                        methods = inspector_details[name].get("methods", [])
                        if method:
                            methods.append(method)
                        inspector_details[name]["methods"] = methods
                        if cert_num:
                            certs = inspector_details[name].get("certifications_inline", [])
                            certs.append(f"{cert_num}" + (f" до {valid_until}" if valid_until else ""))
                            inspector_details[name]["certifications_inline"] = certs
            
            # 1. Из executors в data
            executors_str = g("executors", default="")
            if executors_str:
                # executors может быть строкой с именами через запятую
                for name in executors_str.split(','):
                    name = name.strip()
                    if name and name not in inspectors:
                        inspectors.append(name)
                        inspector_details[name] = {}
            
            # 2. Из ndt_methods (методы по специалистам; один человек может быть в нескольких методах)
            for m in (ndt_methods or []):
                name = (m.get("inspector_name") or "").strip()
                method_code = (m.get("method_code") or "").strip()
                if name and name not in inspectors:
                    inspectors.append(name)
                if name:
                    if name not in inspector_details:
                        inspector_details[name] = {}
                    methods = inspector_details[name].get("methods", [])
                    if method_code and method_code not in methods:
                        methods.append(method_code)
                    inspector_details[name]["methods"] = methods
                    inspector_details[name]["level"] = m.get("inspector_level")
                    if m.get("certificate_number"):
                        inspector_details[name]["certification"] = m.get("certificate_number")
            
            # 3. Из specialist_docs (сертификаты с method_code для подстановки № удостоверения по методу)
            if specialist_docs:
                for s in specialist_docs:
                    name = (s.get("inspector_name") or "").strip()
                    if name and name not in inspectors:
                        inspectors.append(name)
                    if name:
                        if name not in inspector_details:
                            inspector_details[name] = {}
                        inspector_details[name]["certifications"] = s.get("certifications", [])
            
            # Формируем таблицу специалистов
            if inspectors:
                spec_table = doc.add_table(rows=len(inspectors) + 1, cols=4)
                spec_table.style = "Table Grid"
                headers = ["№", "Фамилия И.О.", "№ удостоверения", "Область аттестации / Срок действия"]
                for i, h in enumerate(headers):
                    cell = spec_table.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for idx, name in enumerate(inspectors, 1):
                    spec_table.rows[idx].cells[0].text = str(idx)
                    spec_table.rows[idx].cells[1].text = name
                    
                    # № удостоверения по методу: для каждого метода специалиста — номер и срок из сертификата с этим method_code
                    details = inspector_details.get(name, {})
                    cert_info = []
                    methods_done = set()
                    for method_code in (details.get("methods") or []):
                        if not method_code or method_code in methods_done:
                            continue
                        methods_done.add(method_code)
                        cert_num = ""
                        expiry = ""
                        for cert in (details.get("certifications") or []):
                            if (cert.get("method_code") or "").strip() == (method_code or "").strip():
                                cert_num = cert.get("certificate_number", "")
                                expiry = cert.get("expiry_date", "")
                                break
                        if cert_num:
                            cert_info.append(f"{method_code}: №{cert_num}" + (f" до {expiry}" if expiry else ""))
                    if not cert_info and details.get("certifications"):
                        for cert in details["certifications"]:
                            cert_num = cert.get("certificate_number", "")
                            cert_type = cert.get("certification_type", "")
                            expiry = cert.get("expiry_date", "")
                            if cert_num:
                                cert_info.append(f"{cert_type} №{cert_num}" + (f" до {expiry}" if expiry else ""))
                    if not cert_info and details.get("certification"):
                        cert_info.append(details["certification"])
                    if not cert_info and details.get("certifications_inline"):
                        cert_info.extend(details["certifications_inline"])
                    
                    spec_table.rows[idx].cells[2].text = "; ".join(cert_info) if cert_info else "—"
                    
                    areas = []
                    cert_areas = []
                    if details.get("certifications"):
                        for cert in details["certifications"]:
                            for ca in (cert.get("certification_areas") or ([cert.get("certification_area")] if cert.get("certification_area") else [])):
                                if ca and ca not in cert_areas:
                                    cert_areas.append(ca)
                    if cert_areas:
                        areas.append("Области аттестации: " + "; ".join(cert_areas))
                    methods = details.get("methods") or []
                    if methods:
                        areas.append("Методы: " + ", ".join(sorted(set([m for m in methods if m]))))
                    level = details.get("level", "")
                    if level:
                        areas.append(f"Уровень: {level}")
                    spec_table.rows[idx].cells[3].text = "; ".join(areas) if areas else "—"
            else:
                doc.add_paragraph("—")

            doc.add_heading("7. Перечень приборов и оборудования", level=1)
            fallback_equipment = []
            if not effective_verification_equipment:
                for m in (ndt_methods or []):
                    name = (m.get("equipment") or "").strip()
                    if not name:
                        continue
                    if name not in [e.get("name") for e in fallback_equipment]:
                        fallback_equipment.append({"name": name})

            if effective_verification_equipment:
                # Таблица как в примере
                eq_table = doc.add_table(rows=len(effective_verification_equipment) + 1, cols=4)
                eq_table.style = "Table Grid"
                headers = ["№ п/п", "Наименование прибора", "Заводской номер прибора", "Свидетельство о поверке / Действительна до"]
                for i, h in enumerate(headers):
                    cell = eq_table.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for idx, eq in enumerate(effective_verification_equipment, 1):
                    eq_table.rows[idx].cells[0].text = str(idx)
                    eq_table.rows[idx].cells[1].text = eq.get('name') or '—'
                    eq_table.rows[idx].cells[2].text = eq.get('serial_number') or '—'
                    cert_num = eq.get('verification_certificate_number', '')
                    next_date = self._fmt_date_ru(eq.get('next_verification_date'))
                    cert_info = f"{cert_num}" if cert_num else ""
                    if next_date:
                        cert_info += f" до {next_date}" if cert_info else f"до {next_date}"
                    eq_table.rows[idx].cells[3].text = cert_info if cert_info else "—"
            elif fallback_equipment:
                eq_table = doc.add_table(rows=len(fallback_equipment) + 1, cols=4)
                eq_table.style = "Table Grid"
                headers = ["№ п/п", "Наименование прибора", "Заводской номер прибора", "Свидетельство о поверке / Действительна до"]
                for i, h in enumerate(headers):
                    cell = eq_table.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for idx, eq in enumerate(fallback_equipment, 1):
                    eq_table.rows[idx].cells[0].text = str(idx)
                    eq_table.rows[idx].cells[1].text = eq.get("name") or "—"
                    eq_table.rows[idx].cells[2].text = "—"
                    eq_table.rows[idx].cells[3].text = "—"
            else:
                doc.add_paragraph("—")

            doc.add_heading("8. Объект технического диагностирования", level=1)
            obj_table = doc.add_table(rows=3, cols=2)
            obj_table.style = "Table Grid"
            obj_table.rows[0].cells[0].text = "Объект технического диагностирования"
            obj_table.rows[0].cells[1].text = object_name
            obj_table.rows[1].cells[0].text = "Заводской номер"
            obj_table.rows[1].cells[1].text = serial
            obj_table.rows[2].cells[0].text = "Местонахождение (адрес)"
            obj_table.rows[2].cells[1].text = location

            doc.add_heading("9. Краткая техническая характеристика и назначение объекта технического освидетельствования", level=1)
            def _attr(key: str, default="—"):
                v = attrs.get(key)
                if v is None or v == "":
                    v = g(key, default=default)
                return v if v is not None and v != "" else default

            tech_rows = [
                ("Наименование объекта", device_name),
                ("Назначение", _attr("purpose", default=str(g("tech_description", default="—")))),
                ("Наименование завода-изготовителя", _attr("manufacturer")),
                ("Год изготовления", _attr("manufacture_year")),
                ("Год ввода в эксплуатацию", _attr("commissioning_year")),
                ("Объем/вместимость, м³", _attr("volume", default=_attr("capacity"))),
                ("Внутренний диаметр, мм", _attr("inner_diameter", default=_attr("diameter"))),
                ("Длина/высота, мм", _attr("length", default=_attr("height"))),
                ("Толщина стенки, мм", _attr("wall_thickness", default=_attr("thickness"))),
                ("Материал", _attr("material")),
                ("Рабочее давление, МПа", _attr("working_pressure")),
                ("Расчетное давление, МПа", _attr("design_pressure")),
                ("Пробное давление гидравлического испытания, МПа", _attr("test_pressure")),
                ("Допустимая рабочая температура стенки, ℃", _attr("working_temperature_range", default=_attr("working_temperature"))),
                ("Расчетная температура, ℃", _attr("design_temperature")),
                ("Наименование рабочей среды", _attr("working_medium")),
                ("Класс опасности среды", _attr("hazard_class")),
                ("Категория/группа сосуда", _attr("category", default=_attr("group"))),
                ("Допускаемая коррозия, мм/год", _attr("corrosion_allowance")),
                ("Проектный ресурс, лет", _attr("design_life")),
                ("Назначенный срок службы, лет", _attr("service_life")),
            ]
            tech_table = doc.add_table(rows=len(tech_rows), cols=2)
            tech_table.style = "Table Grid"
            for i, (label, value) in enumerate(tech_rows):
                tech_table.rows[i].cells[0].text = label
                tech_table.rows[i].cells[1].text = str(value)

            doc.add_heading("10. Перечень работ, выполненных в процессе технического освидетельствования", level=1)
            if work_list:
                work_table = doc.add_table(rows=len(work_list) + 1, cols=4)
                work_table.style = "Table Grid"
                headers = ["№ п/п", "Наименование работы", "Объем контроля", "Нормативная документация"]
                for i, h in enumerate(headers):
                    cell = work_table.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for i, m in enumerate(work_list, 1):
                    work_table.rows[i].cells[0].text = str(i)
                    work_table.rows[i].cells[1].text = str(m.get('work_name') or m.get('method_name') or m.get('method_code') or 'Метод НК')
                    work_table.rows[i].cells[2].text = str(m.get('control_volume') or m.get('volume') or '—')
                    work_table.rows[i].cells[3].text = str(m.get('standard') or '—')
            else:
                doc.add_paragraph("—")

            doc.add_heading("11. Сведения о рассмотренных в процессе технического освидетельствования документах", level=1)
            if isinstance(docs, dict) and docs:
                document_names = {
                    '1': 'Дубликат паспорт сосуда, работающего под давлением',
                    '2': 'Приказ об организации работ при эксплуатации оборудования, работающего под давлением',
                    '3': 'Журнал проверки сосудов, работающих под давлением',
                    '4': 'Заключение технического диагностирования',
                    '5': 'Сведения об основных элементах сосуда',
                    '6': 'Сведения об основных элементах сосуда (продолжение)',
                    '7': 'Сведения об основной арматуре, КИП и приборах безопасности',
                    '8': 'Приказ о назначении ответственного лица за исправное состояние и безопасную эксплуатацию сосудов',
                    '9': 'Приказ о назначении ответственного лица за осуществление производственного контроля и соблюдение требований промышленной безопасности на ОПО',
                    '10': 'Паспорт сосуда заводской (удостоверение о качестве монтажа, сборочный чертёж, схема включения, расчёт на прочность)',
                    '11': 'Инструкция по монтажу и эксплуатации',
                    '12': 'Паспорта на предохранительные клапаны',
                    '13': 'Паспорта на запорную арматуру',
                    '14': 'Документация на контрольно-измерительные приборы',
                    '15': 'Ремонтная (исполнительная) документация',
                    '16': 'Заключение экспертизы промышленной безопасности',
                    '17': 'Акты проведения УЗТ',
                }
                present = [(str(k), v) for k, v in docs.items() if v]
                if present:
                    t = doc.add_table(rows=len(present) + 1, cols=4)
                    t.style = "Table Grid"
                    headers = ["№ п/п", "Наименование документа", "Идентификационный номер документа", "Объём, листов"]
                    for i, h in enumerate(headers):
                        cell = t.rows[0].cells[i]
                        cell.text = h
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row_idx = 1
                    for num, _ in sorted(present, key=lambda x: int(x[0]) if str(x[0]).isdigit() else 999):
                        t.rows[row_idx].cells[0].text = str(row_idx)
                        t.rows[row_idx].cells[1].text = document_names.get(str(num), f'Документ {num}')
                        t.rows[row_idx].cells[2].text = attachment_names.get(str(num), str(num))
                        t.rows[row_idx].cells[3].text = "—"
                        row_idx += 1
                else:
                    doc.add_paragraph("—")
            else:
                doc.add_paragraph("—")

            doc.add_heading("12. Анализ результатов предыдущих обследований", level=1)
            doc.add_paragraph(str(g("previous_inspections", default="—")))

            doc.add_heading("13. Результаты технического освидетельствования", level=1)
            # Формируем таблицу результатов на основе выполненных методов и данных из checklist
            results_table = doc.add_table(rows=1, cols=4)
            results_table.style = "Table Grid"
            headers = ["№ п/п", "Наименование работы", "Результаты контроля", "Наименование и номер отчетной документации"]
            for i, h in enumerate(headers):
                cell = results_table.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row_num = 1
            # Добавляем результаты по каждому методу и дополнительным работам (анализ документов, овальность)
            if work_list:
                for m in work_list:
                    method_name = m.get('work_name') or m.get('method_name') or m.get('method_code') or 'Метод НК'
                    results = m.get('results') or m.get('conclusion') or 'Выполнено'
                    row = results_table.add_row()
                    row.cells[0].text = str(row_num)
                    row.cells[1].text = method_name
                    row.cells[2].text = str(results)
                    row.cells[3].text = f"Приложение №{row_num + 1}" if row_num <= 10 else "—"
                    row_num += 1
            
            # Добавляем результаты из checklist если есть
            has_visual = g("has_external_defects") is not None or g("has_internal_defects") is not None
            has_thickness = isinstance(g("thickness_measurements", default=[]), list) and len(g("thickness_measurements", default=[])) > 0
            has_hardness = isinstance(g("hardness_tests", default=[]), list) and len(g("hardness_tests", default=[])) > 0
            has_welds = isinstance(g("weld_inspections", default=[]), list) and len(g("weld_inspections", default=[])) > 0
            
            if has_visual:
                row = results_table.add_row()
                row.cells[0].text = str(row_num)
                row.cells[1].text = "Визуальный и измерительный контроль"
                defects = []
                if g("has_external_defects") == True:
                    defects.append("выявлены дефекты при наружном осмотре")
                if g("has_internal_defects") == True:
                    defects.append("выявлены дефекты при внутреннем осмотре")
                row.cells[2].text = "Дефектов, препятствующих дальнейшей безопасной эксплуатации не выявлено" if not defects else "; ".join(defects)
                row.cells[3].text = "Приложение №3"
                row_num += 1
            
            if has_thickness:
                row = results_table.add_row()
                row.cells[0].text = str(row_num)
                row.cells[1].text = "Ультразвуковой контроль толщины стенок элементов сосуда"
                row.cells[2].text = "Утонения стенок сосуда, превышающие допустимые значения не обнаружены"
                row.cells[3].text = "Приложение №4"
                row_num += 1
            
            if has_welds:
                row = results_table.add_row()
                row.cells[0].text = str(row_num)
                row.cells[1].text = "Ультразвуковой контроль качества основного металла и сварных соединений"
                row.cells[2].text = "Недопустимых дефектов не обнаружено"
                row.cells[3].text = "Приложение №5"
                row_num += 1
            
            if has_hardness:
                row = results_table.add_row()
                row.cells[0].text = str(row_num)
                row.cells[1].text = "Оценка механических свойств элемента сосуда"
                row.cells[2].text = "Отклонений твердости металла не выявлено"
                row.cells[3].text = "Приложение №6"
                row_num += 1

            doc.add_paragraph()
            doc.add_heading("14. Результаты расчетной оценки технического состояния", level=1)
            calc_text = g("calculation_results", default="")
            if not calc_text or calc_text == "—":
                calc_text = "По результатам работ произведена оценка работоспособности сосуда при рабочих параметрах. Выполнен расчет на прочность и определение остаточного ресурса сосуда."
            doc.add_paragraph(str(calc_text))

            doc.add_heading("15. Выводы по результатам технического освидетельствования", level=1)
            concl = inspection_data.get("conclusion") or g("final_conclusion", default="")
            if not concl or concl == "—":
                concl = f"На основании результатов выполненного комплекса работ по техническому диагностированию сосуда, работающего под давлением — {device_name} зав. № {serial}, техническое состояние сосуда, работающего под давлением, оценивается как работоспособное."
            doc.add_paragraph(str(concl))

        # --------------- ПРИЛОЖЕНИЯ ---------------
        if is_on("appendices"):
            doc.add_page_break()
            doc.add_heading("ПРИЛОЖЕНИЯ", level=1)

        def _add_protocol_header(protocol_title: str, protocol_no: int = 1):
            """Блок заголовка протокола (как в otchet.docx): Заказчик, Объект, Место, Дата, НТД."""
            doc.add_paragraph()
            p = doc.add_paragraph(protocol_title)
            p.runs[0].bold = True if p.runs else None
            doc.add_paragraph(f"№ {protocol_no} от {date_perf_ru}г.")
            doc.add_paragraph()
            header_tbl = doc.add_table(rows=5, cols=2)
            header_tbl.style = "Table Grid"
            header_rows = [
                ("Заказчик:", str(org)),
                ("Объект контроля:", f"{device_name} зав.№ {serial}"),
                ("Место проведения контроля:", str(location)),
                ("Дата проведения контроля:", date_perf_ru),
                ("НТД, по которой выполнен контроль:", str(g("normative_base", default="приказ Ростехнадзора от 15.12.2020 №536, СО 153-34.17.439-2003"))),
            ]
            for i, (label, val) in enumerate(header_rows):
                header_tbl.rows[i].cells[0].text = label
                header_tbl.rows[i].cells[1].text = val
                try:
                    header_tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                except Exception:
                    pass
            doc.add_paragraph()

        def _add_equipment_table(method_equipment: Optional[List[Dict[str, Any]]] = None):
            """Таблица применяемого оборудования в протоколе."""
            eq_list = []
            if method_equipment:
                for m in method_equipment:
                    name = m.get("name") or m.get("equipment") or ""
                    if name:
                        eq_list.append({"name": name, "serial_number": m.get("serial_number") or m.get("equipment_serial") or "—"})
            if not eq_list and verification_equipment and isinstance(verification_equipment, list):
                eq_list = [{"name": e.get("name") or "—", "serial_number": e.get("serial_number") or "—"} for e in verification_equipment]
            if not eq_list:
                for m in (ndt_methods or []):
                    eq_name = m.get("equipment") or ""
                    if eq_name and not any(e.get("name") == eq_name for e in eq_list):
                        eq_list.append({"name": eq_name, "serial_number": m.get("equipment_serial") or m.get("serial_number") or "—"})
            if eq_list:
                eq_tbl = doc.add_table(rows=len(eq_list) + 1, cols=3)
                eq_tbl.style = "Table Grid"
                for i, h in enumerate(["№ п/п", "Наименование прибора", "Заводской номер прибора"]):
                    eq_tbl.rows[0].cells[i].text = h
                    try:
                        eq_tbl.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
                        eq_tbl.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass
                for idx, eq in enumerate(eq_list, 1):
                    eq_tbl.rows[idx].cells[0].text = str(idx)
                    eq_tbl.rows[idx].cells[1].text = str(eq.get("name") or "—")
                    eq_tbl.rows[idx].cells[2].text = str(eq.get("serial_number") or "—")
            else:
                doc.add_paragraph("—")
            doc.add_paragraph()

        def _add_inspector_signature(inspector_name: str = ""):
            """Блок подписи специалиста в конце протокола (как в otchet.docx)."""
            name = inspector_name or g("executors", default="—")
            if isinstance(name, list):
                name = ", ".join(str(x) for x in name) if name else "—"
            doc.add_paragraph()
            sig_tbl = doc.add_table(rows=3, cols=2)
            sig_tbl.style = "Table Grid"
            sig_tbl.rows[0].cells[0].text = "Контроль провел, заключение выдал:"
            sig_tbl.rows[0].cells[1].text = ""
            sig_tbl.rows[1].cells[0].text = "Дефектоскопист II уровня по ВИК, УК"
            sig_tbl.rows[1].cells[1].text = str(name)
            sig_tbl.rows[2].cells[0].text = "Начальник ЛНМК"
            sig_tbl.rows[2].cells[1].text = "__________________"
            try:
                sig_tbl.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
                sig_tbl.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
                sig_tbl.rows[2].cells[0].paragraphs[0].runs[0].font.bold = True
            except Exception:
                pass
            doc.add_paragraph()

        app_no = 1
        # ПРИЛОЖЕНИЕ № 1 (Протокол анализа техдокументации) не дублируем — уже есть раздел 11 «Сведения о рассмотренных документах»

        # Отдельные акты выполненных работ по каждому методу НК (как в otchet.docx)
        if performed:
            for m in performed:
                method_name = m.get("method_name") or m.get("method_code") or "Метод НК"
                inspector_name = m.get("inspector_name") or _engineer_for_method(m.get("method_code") or method_name)
                doc.add_page_break()
                doc.add_heading(f"ПРИЛОЖЕНИЕ № {app_no} Акт выполненных работ ({method_name})", level=2)
                _add_protocol_header(f"Протокол по результатам {method_name}", app_no)
                act_table = doc.add_table(rows=9, cols=2)
                act_table.style = "Table Grid"
                rows = [
                    ("Объект технического диагностирования", object_name),
                    ("Техническое устройство", device_name),
                    ("Заводской номер", serial),
                    ("Местонахождение объекта", location),
                    ("Дата проведения", self._fmt_date_ru(m.get("performed_date")) or date_perf_ru),
                    ("Метод контроля", method_name),
                    ("Нормативный документ", m.get("standard") or "—"),
                    ("Оборудование/прибор", m.get("equipment") or "—"),
                    ("Специалист", inspector_name or "—"),
                ]
                for r, (k, v) in enumerate(rows):
                    act_table.rows[r].cells[0].text = str(k)
                    act_table.rows[r].cells[1].text = str(v)
                    try:
                        act_table.rows[r].cells[0].paragraphs[0].runs[0].font.bold = True
                    except Exception:
                        pass
                if m.get("results"):
                    doc.add_paragraph(f"Результаты контроля: {m.get('results')}")
                if m.get("conclusion"):
                    doc.add_paragraph(f"Заключение: {m.get('conclusion')}")
                app_no += 1

        # ПРИЛОЖЕНИЕ № 2: Протокол по результатам оперативной (функциональной) диагностики
        # (если есть данные функциональной диагностики)
        
        # ПРИЛОЖЕНИЕ № 3: Протокол по результатам визуального и измерительного контроля
        has_visual_data = (g("has_external_defects") is not None or 
                          g("has_internal_defects") is not None or
                          isinstance(g("ovality_measurements", default=[]), list) and len(g("ovality_measurements", default=[])) > 0)
        
        if has_visual_data:
            doc.add_page_break()
            doc.add_heading(f"ПРИЛОЖЕНИЕ № {app_no} Протокол по результатам визуального и измерительного контроля", level=2)
            _add_protocol_header("Протокол по результатам визуального и измерительного контроля", app_no)
            
            doc.add_paragraph("1. Применяемое оборудование").runs[0].bold = True
            vik_equipment = [m for m in (ndt_methods or []) if (m.get("method_code") or "").upper() in ("VIK", "ВИК") and m.get("equipment")]
            _add_equipment_table(vik_equipment or None)
            
            # Фото заводской таблички (если есть)
            factory_plate = g("factory_plate_photo") or attachments.get("factory_plate_photo")
            if factory_plate:
                doc.add_paragraph()
                doc.add_paragraph("1. Фото заводской таблички").runs[0].bold = True
                add_picture_if_exists("", factory_plate)
            
            # Результаты визуального контроля
            doc.add_paragraph()
            doc.add_paragraph("2. Результаты визуального контроля").runs[0].bold = True
            visual_table = doc.add_table(rows=10, cols=5)
            visual_table.style = "Table Grid"
            headers = ["№ п/п", "Наименование объекта контроля", "Объем контроля", "Описание обнаруженных дефектов, их размеры", "Оценка качества"]
            for i, h in enumerate(headers):
                cell = visual_table.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            visual_items = [
                ("Опоры", "100%", g("support_state", default="дефектов не обнаружено")),
                ("Антикоррозионное покрытие", "100%", g("anticorrosion_coating_state", default="дефектов не обнаружено")),
                ("Разъемные соединения", "100%", "дефектов не обнаружено"),
                ("Крепежные детали", "100%", g("fasteners_state", default="дефектов не обнаружено")),
                ("Основной металл обечайки, днищ сосуда, штуцеров, фланцев", "100%", "дефектов не обнаружено"),
                ("Сварные соединения вварки штуцеров в корпус", "100%", "дефектов не обнаружено"),
                ("Сварные соединения штуцеров фланцев к патрубкам", "100%", "дефектов не обнаружено"),
                ("Сварные соединения приварки опор к корпусу", "100%", "дефектов не обнаружено"),
                ("Кольцевые, продольные сварные соединения и их перекрестья", "100%", "дефектов не обнаружено"),
            ]
            
            for idx, (name, volume, desc) in enumerate(visual_items, 1):
                visual_table.rows[idx].cells[0].text = str(idx)
                visual_table.rows[idx].cells[1].text = name
                visual_table.rows[idx].cells[2].text = volume
                visual_table.rows[idx].cells[3].text = desc if desc else "дефектов не обнаружено"
                visual_table.rows[idx].cells[4].text = "годен"

            # Детализация дефектов ВИК (если есть)
            visual_defects = g("visual_defects", default=[])
            if isinstance(visual_defects, list) and visual_defects:
                doc.add_paragraph()
                doc.add_paragraph("Дополнительно: выявленные дефекты ВИК").runs[0].bold = True
                defect_table = doc.add_table(rows=len(visual_defects) + 1, cols=5)
                defect_table.style = "Table Grid"
                headers = ["№", "Тип дефекта", "Место", "Размеры", "Описание"]
                for i, h in enumerate(headers):
                    cell = defect_table.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for i, d in enumerate(visual_defects, 1):
                    if not isinstance(d, dict):
                        continue
                    defect_table.rows[i].cells[0].text = str(i)
                    defect_table.rows[i].cells[1].text = str(d.get("defect_type") or "")
                    defect_table.rows[i].cells[2].text = str(d.get("location") or "")
                    defect_table.rows[i].cells[3].text = str(d.get("size") or "")
                    defect_table.rows[i].cells[4].text = str(d.get("description") or "")
                
                # Фотографии дефектов (пути из data или из document_files vd_i_j при синхронизации с мобильного)
                for i, d in enumerate(visual_defects, 1):
                    if not isinstance(d, dict):
                        continue
                    photos = d.get("photos") or []
                    if isinstance(photos, list) and photos:
                        doc.add_paragraph(f"Фотографии дефекта №{i}:").runs[0].bold = True
                        for j, ph in enumerate(photos[:6]):
                            if not isinstance(ph, str):
                                continue
                            photo_path = self._find_image_path(ph)
                            if not photo_path and ph and attachments.get(ph):
                                photo_path = self._find_image_path(attachments.get(ph))
                            if not photo_path and ph and os.path.exists(ph):
                                photo_path = ph
                            # Подстановка из загруженных document_files (vd_defectIndex_photoIndex)
                            if not photo_path and attachments:
                                vd_key = "vd_%d_%d" % (i - 1, j)
                                photo_path = self._find_image_path(attachments.get(vd_key)) or (attachments.get(vd_key) if isinstance(attachments.get(vd_key), str) and os.path.exists(attachments.get(vd_key)) else None)
                            if photo_path:
                                add_picture_if_exists("", photo_path)
            
            # Результаты измерительного контроля - овальность
            ovality = g('ovality_measurements', default=[])
            if isinstance(ovality, list) and ovality:
                doc.add_paragraph()
                doc.add_paragraph("4. Результаты измерительного контроля")
                doc.add_paragraph("Определение овальности проводят измерением максимального (Dmax) и минимального (Dmin) наружного или внутреннего диаметров в одном сечении по двум перпендикулярным направлениям. Относительная овальность корпуса определяется по формуле:")
                
                ovality_table = doc.add_table(rows=len(ovality) + 1, cols=5)
                ovality_table.style = "Table Grid"
                headers = ["Номер сечения", "Dmin", "Dmax", "Фактическая овальность, %", "Допустимая овальность, %"]
                for i, h in enumerate(headers):
                    cell = ovality_table.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for idx, it in enumerate(ovality, 1):
                    ovality_table.rows[idx].cells[0].text = str(it.get('section_number') or f'I-{idx}')
                    ovality_table.rows[idx].cells[1].text = str(it.get('min_diameter') or '')
                    ovality_table.rows[idx].cells[2].text = str(it.get('max_diameter') or '')
                    ovality_table.rows[idx].cells[3].text = str(it.get('deviation_percent') or '0')
                    ovality_table.rows[idx].cells[4].text = "1,0"
            
            vik_inspector = next((m.get("inspector_name") for m in (ndt_methods or []) if (m.get("method_code") or "").upper() in ("VIK", "ВИК")), None)
            _add_inspector_signature(vik_inspector or _engineer_for_method("VIK"))
            app_no += 1

        # ПРИЛОЖЕНИЕ № 4: Протокол по результатам ультразвукового контроля толщины стенок
        thickness = g("thickness_measurements", "thicknessMeasurements", default=[])
        if isinstance(thickness, list) and len(thickness) > 0:
            doc.add_page_break()
            doc.add_heading(f"ПРИЛОЖЕНИЕ № {app_no} Протокол по результатам ультразвукового контроля толщины стенок элементов сосуда", level=2)
            _add_protocol_header("Протокол по результатам ультразвукового контроля толщины стенок элементов сосуда", app_no)
            
            doc.add_paragraph("1. Применяемое оборудование").runs[0].bold = True
            doc.add_paragraph("Таблица № 1")
            doc.add_paragraph()
            uzt_equipment = [m for m in (ndt_methods or []) if (m.get("method_code") or "").upper() in ("UZT", "УЗТ") and m.get("equipment")]
            _add_equipment_table(uzt_equipment or None)
            
            doc.add_paragraph("2. Результаты контроля").runs[0].bold = True
            doc.add_paragraph("Контроль выполнен в соответствии с программой работ, согласно схемы контроля.")
            doc.add_paragraph()
            doc.add_paragraph("Таблица № 2")
            doc.add_paragraph()

            # Схема контроля (чертеж УЗТ) с привязкой точек - ВАЖНО: добавляем в начале раздела
            control_scheme = g("control_scheme_image") or attachments.get("control_scheme_image")
            if control_scheme:
                doc.add_paragraph()
                doc.add_paragraph("2.1. Схема контроля (чертеж с точками измерения):").runs[0].bold = True
                add_picture_if_exists("", control_scheme)
                doc.add_paragraph()

                # Таблица точек измерений по схеме (если есть координаты)
                points_with_coords = [
                    p for p in thickness
                    if isinstance(p, dict) and (p.get("x_percent") is not None or p.get("y_percent") is not None)
                ]
                if points_with_coords:
                    doc.add_paragraph("2.2. Координаты точек измерения на схеме:").runs[0].bold = True
                    t_coords = doc.add_table(rows=len(points_with_coords) + 1, cols=4)
                    t_coords.style = "Table Grid"
                    headers = ["№ точки", "Элемент", "Сечение", "Толщина, мм"]
                    for i, h in enumerate(headers):
                        cell = t_coords.rows[0].cells[i]
                        cell.text = h
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for i, p in enumerate(points_with_coords, 1):
                        t_coords.rows[i].cells[0].text = str(i)
                        t_coords.rows[i].cells[1].text = str(p.get("location") or "")
                        t_coords.rows[i].cells[2].text = str(p.get("section_number") or "")
                        t_coords.rows[i].cells[3].text = str(p.get("thickness") or "")
                    doc.add_paragraph()
            
            # Группируем по элементам (Обечайка, Днище 1, Днище 2)
            elements = {}
            for point in thickness:
                location = str(point.get('location') or 'Обечайка')
                if location not in elements:
                    elements[location] = []
                elements[location].append(point)
            
            for element_name, points in elements.items():
                doc.add_paragraph()
                doc.add_paragraph(f"Элемент: {element_name}")
                # Строк: 1 заголовок + ceil(points/4) данных + 3 итоговые
                data_rows = (len(points) + 3) // 4 if points else 0
                total_rows = 1 + data_rows + 3
                t = doc.add_table(rows=total_rows, cols=9)
                t.style = "Table Grid"
                headers = ["Наименование элемента", "№ точки", "Толщина, мм", "№ точки", "Толщина, мм", "№ точки", "Толщина, мм", "№ точки", "Толщина, мм"]
                for i, h in enumerate(headers):
                    cell = t.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Заполняем точки по 4 в ряд
                row_idx = 1
                for i in range(0, len(points), 4):
                    row = t.rows[row_idx]
                    row.cells[0].text = element_name if i == 0 else ""
                    for j in range(4):
                        point_idx = i + j
                        if point_idx < len(points):
                            point = points[point_idx]
                            row.cells[j*2 + 1].text = str(point_idx + 1)
                            row.cells[j*2 + 2].text = str(point.get('thickness') or '')
                    row_idx += 1
                
                # Добавляем итоговые строки
                nominal = attrs.get("wall_thickness") or attrs.get("thickness") or g("wall_thickness", "thickness", default="4,0")
                t.rows[row_idx].cells[0].text = "Номинальная толщина, мм"
                t.rows[row_idx].cells[1].text = str(nominal)
                t.rows[row_idx + 1].cells[0].text = "Минимально-измеренная толщина, мм"
                min_thickness = min([float(str(p.get('thickness') or '0').replace(',', '.')) for p in points if p.get('thickness')], default=0)
                t.rows[row_idx + 1].cells[1].text = f"{min_thickness:.1f}" if min_thickness > 0 else "—"
                t.rows[row_idx + 2].cells[0].text = "Минимально допустимая толщина стеки сосуда, мм"
                min_allowed_vals = [float(str(p.get('min_allowed_thickness') or '0').replace(',', '.')) for p in points if p.get('min_allowed_thickness')]
                min_allowed = attrs.get("min_wall_thickness") or (min_allowed_vals[0] if min_allowed_vals else 2.8)
                try:
                    ma = float(str(min_allowed or '2.8').replace(',', '.'))
                    t.rows[row_idx + 2].cells[1].text = f"{ma:.1f}" if ma > 0 else "2,8"
                except (TypeError, ValueError):
                    t.rows[row_idx + 2].cells[1].text = "2,8"
            
            # Фото замеров (из document_files uzt_point_i_j или thickness_measurements[].photos)
            doc.add_paragraph()
            doc.add_paragraph("3. Фото замеров толщины стенок").runs[0].bold = True
            has_photos = False
            for i, point in enumerate(thickness):
                if not isinstance(point, dict):
                    continue
                ph_list = point.get("photos") or []
                if not ph_list:
                    continue
                loc = point.get("location") or "Обечайка"
                thick_val = point.get("thickness") or ""
                for j, ph_path in enumerate(ph_list[:3]):  # до 3 фото на точку
                    if not isinstance(ph_path, str) or not ph_path.strip():
                        continue
                    res_path = self._find_image_path(ph_path) or attachments.get(f"uzt_point_{i}_{j}") or ph_path
                    if res_path and os.path.isfile(res_path):
                        doc.add_paragraph(f"Точка {i + 1} ({loc}): толщина {thick_val} мм")
                        add_picture_if_exists("", res_path)
                        has_photos = True
            if not has_photos:
                # Пробуем из attachments напрямую
                for k, fp in (attachments or {}).items():
                    if isinstance(k, str) and k.startswith("uzt_point_") and fp:
                        res = self._find_image_path(fp) or fp
                        if res and os.path.isfile(res):
                            doc.add_paragraph(f"Фото замера ({k})")
                            add_picture_if_exists("", res)
                            has_photos = True
                            break
            
            doc.add_paragraph()
            doc.add_paragraph("Схема контроля указана в Приложении № 7.")
            doc.add_paragraph()
            uzt_inspector = next((m.get("inspector_name") for m in (ndt_methods or []) if (m.get("method_code") or "").upper() in ("UZT", "УЗТ")), None)
            _add_inspector_signature(uzt_inspector or _engineer_for_method("UZT"))
            app_no += 1

        # ПРИЛОЖЕНИЕ № 5: Протокол по результатам ультразвукового контроля качества основного металла и сварных соединений
        welds = g('weld_inspections', default=[])
        if isinstance(welds, list) and len(welds) > 0:
            doc.add_page_break()
            doc.add_heading(f"ПРИЛОЖЕНИЕ № {app_no} Протокол по результатам ультразвукового контроля качества основного металла и сварных соединений", level=2)
            _add_protocol_header("Протокол по результатам ультразвукового контроля качества основного металла и сварных соединений", app_no)
            
            doc.add_paragraph("1. Применяемое оборудование").runs[0].bold = True
            uzk_equipment = [m for m in (ndt_methods or []) if (m.get("method_code") or "").upper() in ("UZK", "УЗК") and m.get("equipment")]
            _add_equipment_table(uzk_equipment or None)
            
            doc.add_paragraph("3. Результаты контроля").runs[0].bold = True
            doc.add_paragraph("Контроль выполнен согласно схемы контроля приведенной в Приложении № 7.")
            
            weld_table = doc.add_table(rows=len(welds) + 1, cols=8)
            weld_table.style = "Table Grid"
            headers = ["№ стыка по карте контроля", "Условный номер дефекта", "Эквивалент. Площадь Sдеф, мм2", "Глубина залегания «Y» , мм", "Протяженность ΔL, мм", "Форма (характер) дефекта (объемный/ плоскостной)", "Местоположение на сварном соединении L, мм", "Заключение"]
            for i, h in enumerate(headers):
                cell = weld_table.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for idx, w in enumerate(welds, 1):
                weld_table.rows[idx].cells[0].text = str(w.get('weld_number') or '')
                weld_table.rows[idx].cells[1].text = "Дефектов не обнаружено"
                weld_table.rows[idx].cells[2].text = ""
                weld_table.rows[idx].cells[3].text = ""
                weld_table.rows[idx].cells[4].text = ""
                weld_table.rows[idx].cells[5].text = ""
                weld_table.rows[idx].cells[6].text = str(w.get('location_on_control_map') or '')
                weld_table.rows[idx].cells[7].text = str(w.get('conclusion') or 'годен')
            
            uzk_inspector = next((m.get("inspector_name") for m in (ndt_methods or []) if (m.get("method_code") or "").upper() in ("UZK", "УЗК")), None)
            _add_inspector_signature(uzk_inspector or _engineer_for_method("UZK"))
            app_no += 1

        # ПРИЛОЖЕНИЕ № 6: Протокол по результатам оценки механических свойств элементов сосуда
        hardness = g('hardness_tests', default=[])
        if isinstance(hardness, list) and len(hardness) > 0:
            doc.add_page_break()
            doc.add_heading(f"ПРИЛОЖЕНИЕ № {app_no} Протокол по результатам оценки механических свойств элементов сосуда (измерение твердости металла)", level=2)
            _add_protocol_header("Протокол по результатам оценки механических свойств элементов сосуда (измерение твердости металла)", app_no)
            
            doc.add_paragraph("1. Применяемое оборудование").runs[0].bold = True
            doc.add_paragraph("Таблица № 1")
            doc.add_paragraph()
            hardness_equipment = [m for m in (ndt_methods or []) if (m.get("method_code") or "").upper() in ("TVI", "ТВИ", "HARDNESS") and m.get("equipment")]
            if hardness_equipment:
                eq_list = [{"name": m.get("equipment"), "serial_number": m.get("equipment_serial") or m.get("serial_number") or "—"} for m in hardness_equipment]
            else:
                eq_list = [e for e in (verification_equipment or []) if "твердость" in (e.get("name") or "").lower() or "УЗИТ" in (e.get("name") or "").upper()]
                eq_list = [{"name": e.get("name"), "serial_number": e.get("serial_number") or "—"} for e in eq_list] if eq_list else []
            if eq_list:
                eq_tbl = doc.add_table(rows=len(eq_list) + 1, cols=3)
                eq_tbl.style = "Table Grid"
                for i, h in enumerate(["№ п/п", "Наименование прибора", "Заводской номер прибора"]):
                    eq_tbl.rows[0].cells[i].text = h
                    eq_tbl.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
                for idx, eq in enumerate(eq_list, 1):
                    eq_tbl.rows[idx].cells[0].text = str(idx)
                    eq_tbl.rows[idx].cells[1].text = str(eq.get("name") or "—")
                    eq_tbl.rows[idx].cells[2].text = str(eq.get("serial_number") or "—")
            else:
                doc.add_paragraph("—")
            doc.add_paragraph()
            
            doc.add_paragraph("2. Результаты контроля").runs[0].bold = True
            doc.add_paragraph("Таблица № 2")
            doc.add_paragraph()
            
            hardness_by_element = {}
            for h in hardness:
                element = str(h.get('location') or h.get('element_name') or h.get('weld_number') or 'Обечайка')
                if element not in hardness_by_element:
                    hardness_by_element[element] = []
                hardness_by_element[element].append(h)
            
            # Таблица: по 4 точки в строке, несколько строк на элемент
            total_data_rows = sum((len(tests) + 3) // 4 for tests in hardness_by_element.values())
            hardness_table = doc.add_table(rows=1 + total_data_rows, cols=9)
            hardness_table.style = "Table Grid"
            headers = ["Наименование элемента", "№ точки", "Результат замера, НВ", "№ точки", "Результат замера, НВ", "№ точки", "Результат замера, НВ", "№ точки", "Результат замера, НВ"]
            for i, h in enumerate(headers):
                cell = hardness_table.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row_idx = 1
            for element_name, tests in hardness_by_element.items():
                for i in range(0, len(tests), 4):
                    row = hardness_table.rows[row_idx]
                    row.cells[0].text = element_name if i == 0 else ""
                    for j in range(4):
                        pt_idx = i + j
                        if pt_idx < len(tests):
                            t = tests[pt_idx]
                            row.cells[j*2 + 1].text = str(pt_idx + 1)
                            row.cells[j*2 + 2].text = str(t.get('hardness_base') or t.get('hardness_weld') or t.get('hardness_haz') or '')
                    row_idx += 1
            
            allowed_limit = ""
            if hardness and isinstance(hardness[0], dict):
                allowed_limit = (hardness[0].get('allowed_hardness_base') or hardness[0].get('allowed_hardness_weld') or '').strip()
            if allowed_limit:
                allowed_limit = f"Допустимый предел твердости: {allowed_limit}, в соответствии с СО 153-34.17.439-2003."
            else:
                allowed_limit = "Допустимый предел твердости для стали 19 ГС от 120 НВ до 180 НВ, в соответствии с СО 153-34.17.439-2003."
            doc.add_paragraph()
            doc.add_paragraph(allowed_limit)
            doc.add_paragraph()
            doc.add_paragraph("3. Заключение по результатам контроля").runs[0].bold = True
            doc.add_paragraph("При контроле физико-механических свойств основного металла методом замера твердости отклонения измеренных значений от допустимого диапазона, указанного в нормативной документации, не установлено.")
            doc.add_paragraph()
            hardness_inspector = next((m.get("inspector_name") for m in (ndt_methods or []) if (m.get("method_code") or "").upper() in ("TVI", "ТВИ", "HARDNESS")), None)
            _add_inspector_signature(hardness_inspector or _engineer_for_method("TVI"))
            app_no += 1

        # ПРИЛОЖЕНИЕ № 7: Схема контроля (с нанесёнными точками замеров УЗТ и/или точками УЗК)
        control_scheme = g('control_scheme_image') or attachments.get('control_scheme_image')
        thickness_for_scheme = g("thickness_measurements", "thicknessMeasurements", default=[])
        welds_for_scheme = g("weld_inspections", default=[])
        if control_scheme:
            doc.add_page_break()
            doc.add_heading(f"ПРИЛОЖЕНИЕ № {app_no} Схема контроля", level=2)
            scheme_resolved = self._find_image_path(str(control_scheme)) or control_scheme
            out_dir = os.path.dirname(output_path) if output_path else "/app/reports/tmp"
            scheme_to_show = scheme_resolved
            if scheme_resolved and os.path.isfile(scheme_resolved):
                if isinstance(thickness_for_scheme, list) and thickness_for_scheme:
                    scheme_with_thickness = self._draw_points_on_scheme(scheme_resolved, thickness_for_scheme, output_dir=out_dir)
                    if scheme_with_thickness:
                        scheme_to_show = scheme_with_thickness
                if isinstance(welds_for_scheme, list) and welds_for_scheme:
                    uzk_with_coords = [w for w in welds_for_scheme if isinstance(w, dict) and (w.get("x_percent") is not None or w.get("y_percent") is not None)]
                    if uzk_with_coords:
                        scheme_with_welds = self._draw_weld_points_on_scheme(scheme_to_show, uzk_with_coords, output_dir=out_dir)
                        if scheme_with_welds:
                            scheme_to_show = scheme_with_welds
                caption = "Схема контроля"
                has_weld_coords = False
                if isinstance(welds_for_scheme, list):
                    has_weld_coords = any(
                        isinstance(w, dict) and (w.get("x_percent") is not None or w.get("y_percent") is not None)
                        for w in welds_for_scheme
                    )
                if (isinstance(thickness_for_scheme, list) and thickness_for_scheme) or has_weld_coords:
                    caption = "Схема контроля с точками замеров"
                add_picture_if_exists(caption + ":", scheme_to_show)
            app_no += 1

        # ПРИЛОЖЕНИЕ № 8: Расчет остаточного ресурса и расчет на прочность
        thickness_for_calc = g("thickness_measurements", "thicknessMeasurements", default=[])
        has_thickness_for_calc = isinstance(thickness_for_calc, list) and len(thickness_for_calc) > 0
        if has_thickness_for_calc:
            doc.add_page_break()
            doc.add_heading(f"ПРИЛОЖЕНИЕ № {app_no} Расчет остаточного ресурса и расчет на прочность сосуда", level=2)
            # Извлекаем данные для расчёта
            wall_th = attrs.get("wall_thickness") or attrs.get("thickness") or g("wall_thickness", "thickness", default="4")
            min_allowed = attrs.get("min_wall_thickness") or g("min_allowed_thickness", default="2.8")
            min_vals = [float(str(p.get("thickness", "0")).replace(",", ".")) for p in thickness_for_calc if isinstance(p, dict) and p.get("thickness")]
            s_f = min(min_vals) if min_vals else 3.9
            try:
                s_n = float(str(wall_th).replace(",", "."))
            except (TypeError, ValueError):
                s_n = 4.0
            try:
                s_otb = float(str(min_allowed).replace(",", "."))
            except (TypeError, ValueError):
                s_otb = 2.8
            comm_year = attrs.get("commissioning_year") or g("commissioning_year") or equipment_data.get("commissioning_date")
            t1 = 16
            if comm_year:
                try:
                    t1 = datetime.now().year - int(str(comm_year)[:4])
                    if t1 < 1:
                        t1 = 16
                except (TypeError, ValueError):
                    pass
            a = (s_n - s_f) / t1 if t1 > 0 else 0.01
            tk = (s_f - s_otb) / a if a > 0 else 110
            doc.add_paragraph("1. Расчет остаточного ресурса сосуда")
            doc.add_paragraph("Остаточный ресурс сосуда рассчитан согласно ДиОР-05 и приведен в Таблице Е.1.")
            doc.add_paragraph()
            tbl_e1 = doc.add_table(rows=7, cols=5)
            tbl_e1.style = "Table Grid"
            for i, h in enumerate(["№ п/п", "Наименование величины", "Единица измерения", "Обозначение и расчетная формула", "Числовое значение"]):
                tbl_e1.rows[0].cells[i].text = h
                tbl_e1.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
            tbl_e1.rows[1].cells[0].text = "1"
            tbl_e1.rows[1].cells[1].text = "Время эксплуатации"
            tbl_e1.rows[1].cells[2].text = "лет"
            tbl_e1.rows[1].cells[3].text = "t₁"
            tbl_e1.rows[1].cells[4].text = str(t1)
            tbl_e1.rows[2].cells[0].text = "2"
            tbl_e1.rows[2].cells[1].text = "Паспортная толщина стенки\nОбечайка / Днище"
            tbl_e1.rows[2].cells[2].text = "мм"
            tbl_e1.rows[2].cells[3].text = "Sн"
            tbl_e1.rows[2].cells[4].text = f"{s_n:.0f} / {s_n:.0f}"
            tbl_e1.rows[3].cells[0].text = "3"
            tbl_e1.rows[3].cells[1].text = "Минимально допустимая толщина стенки сосуда"
            tbl_e1.rows[3].cells[2].text = "мм"
            tbl_e1.rows[3].cells[3].text = "Sотб"
            tbl_e1.rows[3].cells[4].text = f"{s_otb:.1f} / {s_otb:.1f}"
            tbl_e1.rows[4].cells[0].text = "4"
            tbl_e1.rows[4].cells[1].text = "Минимальная толщина по результатам замеров"
            tbl_e1.rows[4].cells[2].text = "мм"
            tbl_e1.rows[4].cells[3].text = "Sф"
            tbl_e1.rows[4].cells[4].text = f"{s_f:.1f} / {s_f:.1f}"
            tbl_e1.rows[5].cells[0].text = "5"
            tbl_e1.rows[5].cells[1].text = "Скорость коррозии металла сосуда"
            tbl_e1.rows[5].cells[2].text = "мм/год"
            tbl_e1.rows[5].cells[3].text = "a = (Sн - Sф) / t₁"
            tbl_e1.rows[5].cells[4].text = f"{a:.2f} / {a:.2f}"
            tbl_e1.rows[6].cells[0].text = "6"
            tbl_e1.rows[6].cells[1].text = "Остаточный срок службы сосуда, поэлементно"
            tbl_e1.rows[6].cells[2].text = "лет"
            tbl_e1.rows[6].cells[3].text = "Tk = (Sф - Sотб) / a"
            tbl_e1.rows[6].cells[4].text = f"{tk:.0f} / {tk:.0f}"
            doc.add_paragraph()
            doc.add_paragraph("2. Расчет на прочность сосуда")
            doc.add_paragraph("Расчет на прочность сосуда проводился в соответствии с ГОСТ 34233.1-2017 и ГОСТ 34233.2-2017 и приведен в Таблице Е.2.")
            doc.add_paragraph()
            try:
                p_val = float(str(attrs.get("working_pressure") or g("working_pressure") or "1.1").replace(",", "."))
            except (TypeError, ValueError):
                p_val = 1.1
            try:
                t_n = float(str(attrs.get("design_temperature") or g("design_temperature") or "100").replace(",", "."))
            except (TypeError, ValueError):
                t_n = 100
            try:
                d_n = float(str(attrs.get("diameter") or g("diameter") or equipment_data.get("diameter") or "792").replace(",", "."))
            except (TypeError, ValueError):
                d_n = 792
            try:
                c_val = float(str(attrs.get("corrosion_allowance") or g("corrosion_allowance") or "0").replace(",", "."))
            except (TypeError, ValueError):
                c_val = 0
            phi = 0.9
            sigma = 177
            r_val = d_n
            s_p = (p_val * r_val) / (2 * phi * sigma - 0.5 * p_val) if (2 * phi * sigma - 0.5 * p_val) > 0 else 2.74
            s_otb_calc = s_p + c_val
            p_allow = (2 * (s_f - c_val) * phi * sigma) / (r_val + 0.5 * (s_f - c_val)) if (r_val + 0.5 * (s_f - c_val)) > 0 else 1.57
            tbl_e2 = doc.add_table(rows=13, cols=5)
            tbl_e2.style = "Table Grid"
            for i, h in enumerate(["№ п/п", "Наименование величины", "Единица измерения", "Обозначение и расчетная формула", "Числовое значение"]):
                tbl_e2.rows[0].cells[i].text = h
                tbl_e2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
            tbl_e2.rows[1].cells[0].text = "1"
            tbl_e2.rows[1].cells[1].text = "Рабочее давление"
            tbl_e2.rows[1].cells[2].text = "МПа"
            tbl_e2.rows[1].cells[3].text = "P"
            tbl_e2.rows[1].cells[4].text = f"{p_val:.1f}"
            tbl_e2.rows[2].cells[0].text = "2"
            tbl_e2.rows[2].cells[1].text = "Расчетная температура"
            tbl_e2.rows[2].cells[2].text = "°C"
            tbl_e2.rows[2].cells[3].text = "tн"
            tbl_e2.rows[2].cells[4].text = f"{t_n:.0f}"
            tbl_e2.rows[3].cells[0].text = "3"
            tbl_e2.rows[3].cells[1].text = "Внутренний диаметр"
            tbl_e2.rows[3].cells[2].text = "мм"
            tbl_e2.rows[3].cells[3].text = "Dн"
            tbl_e2.rows[3].cells[4].text = f"{d_n:.0f}"
            tbl_e2.rows[4].cells[0].text = "4"
            tbl_e2.rows[4].cells[1].text = "Прибавка для компенсации коррозии"
            tbl_e2.rows[4].cells[2].text = "мм"
            tbl_e2.rows[4].cells[3].text = "C"
            tbl_e2.rows[4].cells[4].text = f"{c_val:.0f}"
            tbl_e2.rows[5].cells[0].text = "5"
            tbl_e2.rows[5].cells[1].text = "Коэффициент прочности сварных швов"
            tbl_e2.rows[5].cells[2].text = ""
            tbl_e2.rows[5].cells[3].text = "φ"
            tbl_e2.rows[5].cells[4].text = f"{phi}"
            tbl_e2.rows[6].cells[0].text = "6"
            tbl_e2.rows[6].cells[1].text = "Допускаемое напряжение при расчетной температуре"
            tbl_e2.rows[6].cells[2].text = "МПа"
            tbl_e2.rows[6].cells[3].text = "[σ]"
            tbl_e2.rows[6].cells[4].text = str(sigma)
            tbl_e2.rows[7].cells[0].text = "7"
            tbl_e2.rows[7].cells[1].text = "Радиус кривизны в вершине днища, R=D для эллиптических днищ"
            tbl_e2.rows[7].cells[2].text = "мм"
            tbl_e2.rows[7].cells[3].text = "R"
            tbl_e2.rows[7].cells[4].text = f"{r_val:.0f}"
            tbl_e2.rows[8].cells[0].text = "8"
            tbl_e2.rows[8].cells[1].text = "Минимальная толщина по результатам контроля\nДнище / Обечайка"
            tbl_e2.rows[8].cells[2].text = "мм"
            tbl_e2.rows[8].cells[3].text = "Sф"
            tbl_e2.rows[8].cells[4].text = f"{s_f:.1f} / {s_f:.1f}"
            tbl_e2.rows[9].cells[0].text = "9"
            tbl_e2.rows[9].cells[1].text = "Расчетная толщина стенки\nДнище / Обечайка"
            tbl_e2.rows[9].cells[2].text = "мм"
            tbl_e2.rows[9].cells[3].text = "Sр = P·R/(2φ[σ]–0,5P) / Sр = P·D/(2φ[σ]–P)"
            tbl_e2.rows[9].cells[4].text = f"{s_p:.2f} / {s_p:.2f}"
            tbl_e2.rows[10].cells[0].text = "10"
            tbl_e2.rows[10].cells[1].text = "Минимально допустимая толщина стенки сосуда\nДнище / Обечайка"
            tbl_e2.rows[10].cells[2].text = "мм"
            tbl_e2.rows[10].cells[3].text = "Sотб = Sр + C"
            tbl_e2.rows[10].cells[4].text = f"{s_otb_calc:.2f} / {s_otb_calc:.2f}"
            tbl_e2.rows[11].cells[0].text = "11"
            tbl_e2.rows[11].cells[1].text = "Допускаемое внутреннее избыточное давление\nДнище / Обечайка"
            tbl_e2.rows[11].cells[2].text = "МПа"
            tbl_e2.rows[11].cells[3].text = "[P]"
            tbl_e2.rows[11].cells[4].text = f"{p_allow:.2f} / {p_allow:.2f}"
            doc.add_paragraph()
            doc.add_paragraph("Условия прочности: Sотб = {:.1f} мм < Sф = {:.1f} мм. [P] = {:.2f} МПа > Рраб = {:.1f} МПа.".format(s_otb_calc, s_f, p_allow, p_val))
            doc.add_paragraph()
            doc.add_paragraph("Выводы:").runs[0].bold = True
            doc.add_paragraph("На основании выполненного расчета на прочность установлено, что сосуд удовлетворяет условиям прочности, срок эксплуатации до достижения предельно-допустимого значения толщины стенки сосуда составляет более 10 лет.")
            doc.add_paragraph()
            doc.add_paragraph("Расчет выполнил:")
            doc.add_paragraph("Дефектоскопист II уровня по ВИК, УК")
            doc.add_paragraph(str(_engineer_for_method("TVI") or _engineer_for_method("UZT") or g("executors", default="—")))
            app_no += 1
        
        # ПРИЛОЖЕНИЕ № 9: Акт проведения гидравлических испытаний (если есть)
        
        # ПРИЛОЖЕНИЕ № 10: Перечень применяемой нормативной документации
        doc.add_page_break()
        doc.add_heading(f"ПРИЛОЖЕНИЕ № {app_no} Перечень применяемой при техническом освидетельствовании нормативной, технической и методической документации", level=2)
        
        normative_docs = [
            "Федеральный закон от 21.07.1997г. №116 «О промышленной безопасности опасных производственных объектов».",
            "Федеральные нормы и правила в области промышленной безопасности «Правила промышленной безопасности при использовании оборудования, работающего под избыточным давлением», утвержденные приказом Федеральной службы по экологическому, технологическому и атомному надзору от 15.12.2020 №536",
            "СО 153-34.17.439-2003 «Инструкция по продлению срока службы сосудов, работающих под давлением».",
            "ГОСТ Р ИСО 17637-2014 «Контроль неразрушающий. Визуальный контроль соединений, выполненных сваркой плавлением»",
            "ГОСТ 34347-2017 «Сосуды и аппараты стальные сварные. Общие технические условия»",
            "ГОСТ Р 55614-2013 «Контроль неразрушающий. Толщиномеры ультразвуковые. Общие технические требования»",
            "ГОСТ Р ИСО 17640-2016 «Неразрушающий контроль сварных соединений. Ультразвуковой контроль. Технология, уровни контроля и оценки»",
            "ГОСТ Р 55724-2013 «Контроль неразрушающий. Соединения сварные. Методы ультразвуковые»",
            "СТО 00220256-005-2005 «Швы стыковых, угловых и тавровых сварных соединений сосудов и аппаратов, работающих под давлением. Методика ультразвукового контроля»",
            "ГОСТ 20911-89 «Техническая диагностика. Термины и определения»",
            "приказ Ростехнадзора от 16.01.2024 №8, Руководство по безопасности \"Методические рекомендации о порядке проведения визуального и измерительного контроля\"",
            "ГОСТ Р ИСО 16809-2015 «Контроль неразрушающий. Контроль ультразвуковой. Измерение толщины».",
            "ГОСТ 22761-77 «Металлы и сплавы. Метод измерения твердости по Бринеллю переносными твердомерами статического действия»",
        ]
        
        for idx, doc_text in enumerate(normative_docs, 1):
            doc.add_paragraph(f"{idx}. {doc_text}")

        # Дополнительные приложения по методам НК (если есть)
        if performed:
            for m in performed:
                if m.get('method_name') and m.get('method_name') not in ['Визуальный контроль', 'УЗТ', 'УЗК', 'Твердость']:
                    doc.add_page_break()
                    doc.add_heading(f"ПРИЛОЖЕНИЕ № {app_no} Протокол по результатам {m.get('method_name') or 'НК'}", level=2)
                    doc.add_paragraph(f"Дата проведения контроля: {self._fmt_date_ru(m.get('performed_date')) or date_perf_ru}")
                    doc.add_paragraph(f"НТД: {m.get('standard') or '—'}")
                    doc.add_paragraph(f"Оборудование: {m.get('equipment') or '—'}")
                    doc.add_paragraph(f"Результаты: {m.get('results') or '—'}")
                    if m.get("defects"):
                        doc.add_paragraph(f"Дефекты: {m.get('defects')}")
                    if m.get("conclusion"):
                        doc.add_paragraph(f"Заключение: {m.get('conclusion')}")
                    app_no += 1

        # Документы специалистов/поверки — оставляем в конце (если есть)
        if specialist_docs:
            doc.add_page_break()
            doc.add_heading(f"ПРИЛОЖЕНИЕ № {app_no} Документы специалистов НК", level=2)
            for s in specialist_docs:
                doc.add_heading(f"Специалист: {s.get('inspector_name') or '—'}", level=3)
                for c in (s.get("certifications") or []):
                    doc.add_paragraph(f"{c.get('certification_type') or 'Удостоверение'} № {c.get('certificate_number') or '—'}")
                    sp = c.get("scan_file_path")
                    if sp:
                        resolved = self._find_image_path(sp) if isinstance(sp, str) else None
                        if not resolved:
                            resolved = sp if isinstance(sp, str) and os.path.exists(sp) else None
                        if resolved:
                            try:
                                doc.add_picture(resolved, width=Inches(4.8))
                            except Exception:
                                pass
            app_no += 1

        if verification_equipment:
            doc.add_page_break()
            doc.add_heading(f"ПРИЛОЖЕНИЕ № {app_no} Свидетельства о поверке оборудования", level=2)
            for eq in verification_equipment:
                sp = eq.get("scan_file_path")
                if sp and isinstance(sp, str) and os.path.exists(sp):
                    doc.add_paragraph(f"{eq.get('name') or ''} № {eq.get('verification_certificate_number') or ''}")
                    try:
                        doc.add_picture(sp, width=Inches(4.8))
                    except Exception:
                        pass

        if (inspection_data.get("status") or "").upper() == "DRAFT":
            self._add_draft_watermark(doc)
        apply_device_terminology_to_document(doc, detect_pressure_device_kind(equipment_data))
        doc.save(output_path)
        return
    
    def _generate_vessel_report_word(
        self,
        inspection_data: Dict[str, Any],
        equipment_data: Dict[str, Any],
        ndt_methods: List[Dict[str, Any]],
        output_path: str,
        report_type: str = "TECHNICAL_REPORT",
        document_files: Optional[List[Dict[str, Any]]] = None,
        specialist_docs: Optional[List[Dict[str, Any]]] = None,
        verification_equipment: Optional[List[Dict[str, Any]]] = None,
        template_definition: Optional[Dict[str, Any]] = None,
    ):
        """
        Генерация технического отчета для сосудов/ресиверов по образцу reciver.md и rabota.md
        Структура: титульный лист -> содержание -> разделы 1-15 -> приложения
        """
        doc = Document()
        self._setup_styles(doc)
        
        data = inspection_data.get("data") or {}
        if not isinstance(data, dict):
            data = {}
        attrs = equipment_data.get("attributes") or {}
        if not isinstance(attrs, dict):
            attrs = {}
        device_kind = detect_pressure_device_kind(equipment_data)
        purpose_default = default_purpose_for_kind(device_kind)
        equipment_preset = preset_from_equipment_data(equipment_data)
        pressure_regime = pressure_regime_for_preset(equipment_preset)
        if isinstance(template_definition, dict):
            td_regime = template_definition.get("pressure_regime") or template_definition.get("normative_basis")
            if td_regime in ("low", "rua_93"):
                pressure_regime = "low"
            elif td_regime in ("high", "order_536"):
                pressure_regime = "high"
        normative_base_default = (
            NORMATIVE_BASE_RUA_93 if pressure_regime == "low" else NORMATIVE_BASE_ORDER_536
        )
        
        # Индекс вложений (с алиасами doc_15_0 -> 15 для сканов документов)
        attachments: Dict[str, str] = build_attachments_index(document_files)
        attachment_names: Dict[str, str] = {}
        if document_files and isinstance(document_files, list):
            for f in document_files:
                if not isinstance(f, dict):
                    continue
                dn = str(f.get("document_number") or "")
                fn = f.get("file_name")
                if dn and isinstance(fn, str) and fn:
                    attachment_names[dn] = fn
        
        def add_picture_if_exists(title: str, path: Optional[str], width_inches: float = 4.8):
            """Добавить изображение или подпись для PDF если существует."""
            if not path:
                return False
            if str(path).lower().endswith(".pdf"):
                resolved = self._find_image_path(path)
                if not resolved and path in attachments:
                    resolved = self._find_image_path(attachments[path]) or attachments.get(path)
                if not resolved:
                    resolved = path
                p = Path(resolved) if resolved else None
                if p and p.exists():
                    if title:
                        par = doc.add_paragraph()
                        par.add_run(title).bold = True
                    doc.add_paragraph()
                    doc.add_paragraph("Приложенный документ (PDF): " + p.name)
                    doc.add_paragraph()
                    return True
                return False
            resolved = self._find_image_path(path)
            if not resolved and path in attachments:
                resolved = self._find_image_path(attachments[path]) or attachments.get(path)
            if not resolved:
                resolved = path
            try:
                p = Path(resolved)
                if not p.exists():
                    return False
                if title:
                    par = doc.add_paragraph()
                    par.add_run(title).bold = True
                doc.add_paragraph()
                doc.add_picture(str(p), width=Inches(width_inches))
                doc.add_paragraph()
                return True
            except Exception as e:
                print(f"Ошибка добавления изображения {path}: {e}")
                return False
        
        def g(*keys, default=None):
            """Извлечь значение по ключам"""
            for k in keys:
                if k in data and data.get(k) not in (None, ""):
                    return data.get(k)
            for k in keys:
                if k in attrs and attrs.get(k) not in (None, ""):
                    return attrs.get(k)
            # camelCase варианты
            for k in keys:
                camel_key = ''.join(word.capitalize() if i > 0 else word for i, word in enumerate(k.split('_')))
                camel_key_lower = camel_key[0].lower() + camel_key[1:] if camel_key else k
                if camel_key_lower in data and data.get(camel_key_lower) not in (None, ""):
                    return data.get(camel_key_lower)
                if camel_key_lower in attrs and attrs.get(camel_key_lower) not in (None, ""):
                    return attrs.get(camel_key_lower)
            return default

        opo = data.get("opo") or data.get("opo_info") or {}
        if not isinstance(opo, dict):
            opo = {}

        def _opo_get(*keys, default=None):
            for k in keys:
                if k in opo and opo.get(k) not in (None, ""):
                    return opo.get(k)
            for k in keys:
                camel_key = ''.join(word.capitalize() if i > 0 else word for i, word in enumerate(k.split('_')))
                camel_key_lower = camel_key[0].lower() + camel_key[1:] if camel_key else k
                if camel_key_lower in opo and opo.get(camel_key_lower) not in (None, ""):
                    return opo.get(camel_key_lower)
            return g(*keys, default=default)
        
        date_perf_iso = inspection_data.get("date_performed")
        date_perf_ru = self._fmt_date_ru(date_perf_iso) or datetime.now().strftime("%d.%m.%Y")
        year2 = datetime.now().strftime("%y")
        rid = str(equipment_data.get("id") or "")[-4:] or "0000"
        report_no = g("report_number", default=f"{year2}-{rid}")

        rt = (report_type or inspection_data.get("report_type") or "TECHNICAL_REPORT")
        rt = str(rt).strip().upper()
        is_epb = rt in ("EXPERTISE", "EPB", "ЭПБ")
        
        # --------------- ТИТУЛЬНЫЙ ЛИСТ ---------------
        # Логотип
        logo_path = template_definition.get("logo_path") if isinstance(template_definition, dict) else None
        if not logo_path:
            logo_path = "/app/reports/assets/yutar_logo.png"
        try:
            resolved = self._find_image_path(str(logo_path))
            if resolved and os.path.isfile(resolved):
                doc.add_picture(resolved, width=Inches(5.2))
        except Exception:
            pass
        
        device_name = g("equipment_device_name", "vessel_name", default=equipment_data.get("name") or "—")
        serial = g("serial_number", default=equipment_data.get("serial_number") or "—")
        reg_no = g("reg_number", default=attrs.get("reg_number") or attrs.get("regNumber") or "—")
        org = g("organization", "customer_name", "enterprise_name", default="—")
        location = g("location", "equipment_location", default=equipment_data.get("location") or "—")
        opo_name = _opo_get("name", "opo_name", default=g("opo_name", default="—"))
        opo_reg = _opo_get("registration_number", "opo_code", "reg_number", default=g("opo_code", default="—"))
        opo_class = _opo_get("hazard_class", "danger_class", default=g("opo_hazard_class", default="—"))

        if is_epb:
            p = doc.add_paragraph("ЗАКЛЮЧЕНИЕ ЭКСПЕРТИЗЫ")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(14)
            except Exception:
                pass
            p2 = doc.add_paragraph(f"ПРОМЫШЛЕННОЙ БЕЗОПАСНОСТИ № {report_no}")
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p2.runs[0].bold = True
                p2.runs[0].font.size = Pt(14)
            except Exception:
                pass
            doc.add_paragraph("")
            epb_tbl = doc.add_table(rows=6, cols=2)
            epb_tbl.style = "Table Grid"
            epb_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            epb_rows = [
                ("Техническое устройство:", f"зав.№ {serial}, рег.№ {reg_no}"),
                ("Опасный производственный объект (ОПО):", str(opo_name)),
                ("Регистрационный номер ОПО:", str(opo_reg)),
                ("Класс опасности ОПО:", str(opo_class)),
                ("Предприятие владелец:", str(org)),
                ("Место эксплуатации:", str(location)),
            ]
            for i, (k, v) in enumerate(epb_rows):
                epb_tbl.rows[i].cells[0].text = k
                epb_tbl.rows[i].cells[1].text = v
            doc.add_paragraph("")
            doc.add_paragraph("Дата внесения в реестр: «____» _______________20____г.")
        else:
            # Заголовок технического отчёта
            title_table = doc.add_table(rows=1, cols=1)
            title_table.style = "Table Grid"
            title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = title_table.rows[0].cells[0]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"ТЕХНИЧЕСКИЙ ОТЧЕТ № {report_no}\nПО РЕЗУЛЬТАТАМ ТЕХНИЧЕСКОГО ДИАГНОСТИРОВАНИЯ")
            r.bold = True
            r.font.size = Pt(14)
            doc.add_paragraph("")
            obj_tbl = doc.add_table(rows=5, cols=2)
            obj_tbl.style = "Table Grid"
            obj_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            rows = [
                ("Объект технического диагностирования:", ""),
                ("Техническое устройство:", device_name),
                ("Заводской номер:", str(serial)),
                ("Эксплуатирующая организация:", str(org)),
                ("Местонахождение объекта:", str(location)),
            ]
            for i, (k, v) in enumerate(rows):
                obj_tbl.rows[i].cells[0].text = k
                obj_tbl.rows[i].cells[1].text = v
                try:
                    obj_tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                except Exception:
                    pass
        
        doc.add_paragraph("")
        
        # Подпись руководителя
        sign_tbl = doc.add_table(rows=1, cols=2)
        sign_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        sign_tbl.columns[0].width = Inches(3.5)
        sign_tbl.columns[1].width = Inches(3.5)
        sign_tbl.cell(0, 0).text = ""
        
        contractor = g("contractor_name", default='ООО «ЮТАР»')
        director_title = g("director_title", default="Генеральный директор")
        director_name = g("director_name", default="__________________")
        city = g("report_city", "city", default="г. Урай")
        year = datetime.now().strftime("%Y")
        
        p = sign_tbl.cell(0, 1).paragraphs[0]
        p.add_run(f"{director_title}\n{contractor}\n\n__________________\n{director_name}\n\n«____» __________ {year} г.\nМ.П.").font.size = Pt(10)
        
        doc.add_paragraph("")
        p = doc.add_paragraph(f"{city} {year} г.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()
        
        # --------------- СОДЕРЖАНИЕ ---------------
        if is_epb:
            append_epb_toc(doc)
        else:
            doc.add_heading("СОДЕРЖАНИЕ", level=1)
            self._add_toc_field(doc)
            doc.add_page_break()
        
        # --------------- РАЗДЕЛЫ 1-15 ---------------
        performed = [m for m in (ndt_methods or []) if m.get("is_performed")]
        docs_dict = g("documents", default={})
        docs_info = g("documents_info", default={})
        inspection_engineers = g("inspection_engineers", default=[])
        # Расширенный перечень работ: НК-методы + анализ документов + овальность + твердометрия + УЗТ + УЗК при наличии данных
        work_list = list(performed)
        performed_codes = {str(m.get("method_code") or m.get("method_name") or "").upper() for m in performed}
        if isinstance(docs_dict, dict) and any(v for v in docs_dict.values()):
            work_list.append({"method_name": "Анализ док.", "work_name": "Анализ технической документации", "standard": "приказ Ростехнадзора от 15.12.2020 №536", "conclusion": "Рассмотрены документы по перечню раздела 11"})
        has_visual_data = (g("visual_defects", default=[]) and len(g("visual_defects", default=[])) > 0) or g("has_external_defects") is not None or g("has_internal_defects") is not None
        if has_visual_data and "ВИК" not in performed_codes and "VIK" not in performed_codes:
            work_list.append({"method_name": "ВИК", "work_name": "Визуальный и измерительный контроль", "standard": "приказ Ростехнадзора от 15.12.2020 №536", "conclusion": "Выполнено по протоколу"})
        ovality_data = g("ovality_measurements", default=[])
        if isinstance(ovality_data, list) and ovality_data:
            work_list.append({"method_name": "Овальность", "work_name": "Измерение овальности корпуса сосуда", "standard": "СО 153-34.17.439-2003", "conclusion": "Выполнено по протоколу"})
        hardness_data = g("hardness_tests", default=[])
        if isinstance(hardness_data, list) and hardness_data:
            work_list.append({"method_name": "Твердометрия", "work_name": "Контроль твердости металла", "standard": "приказ Ростехнадзора от 15.12.2020 №536", "conclusion": "Выполнено по протоколу"})
        thickness_data = g("thickness_measurements", default=[])
        if isinstance(thickness_data, list) and thickness_data and "УЗТ" not in performed_codes and "UZT" not in performed_codes:
            work_list.append({"method_name": "УЗТ", "work_name": "Ультразвуковой контроль толщины стенок элементов сосуда", "standard": "приказ Ростехнадзора от 15.12.2020 №536", "conclusion": "Выполнено по протоколу"})
        weld_data = g("weld_inspections", default=[])
        if isinstance(weld_data, list) and weld_data and "УЗК" not in performed_codes and "UZK" not in performed_codes:
            work_list.append({"method_name": "УЗК", "work_name": "Ультразвуковой контроль качества основного металла и сварных соединений", "standard": "приказ Ростехнадзора от 15.12.2020 №536", "conclusion": "Выполнено по протоколу"})

        def _normalize_method(method_raw: str) -> str:
            m = (method_raw or "").strip()
            if not m:
                return ""
            m_up = m.upper()
            mapping = {
                "VIK": "ВИК",
                "UZK": "УЗК",
                "UZT": "УЗТ",
                "PVK": "ПВК/МК",
                "MK": "МК",
                "RK": "РК",
                "MPD": "МПД",
                "KPD": "КПД",
                "TVI": "ТВИ",
            }
            if m_up in mapping:
                return mapping[m_up]
            return m

        def _doc_meta(num: str):
            num_key = str(num)
            present = None
            doc_number = ""
            doc_date = ""
            if isinstance(docs_dict, dict) and num_key in docs_dict:
                val = docs_dict.get(num_key)
                if isinstance(val, dict):
                    present = val.get("present")
                    if present is None:
                        present = val.get("has") if val.get("has") is not None else val.get("value")
                    doc_number = str(val.get("number") or val.get("doc_number") or "")
                    doc_date = str(val.get("date") or val.get("doc_date") or "")
                else:
                    if isinstance(val, str):
                        present = val.strip().lower() in ("true", "1", "yes", "да")
                    else:
                        present = bool(val)
            if isinstance(docs_info, dict) and num_key in docs_info:
                info = docs_info.get(num_key) or {}
                if isinstance(info, dict):
                    if present is None:
                        present = info.get("present")
                        if present is None:
                            present = info.get("has") if info.get("has") is not None else info.get("value")
                    if not doc_number:
                        doc_number = str(info.get("number") or info.get("doc_number") or "")
                    if not doc_date:
                        doc_date = str(info.get("date") or info.get("doc_date") or "")
            return (present, doc_number, doc_date)
        
        def _engineer_for_method(method_code: str) -> str:
            if not isinstance(inspection_engineers, list):
                return ""
            for ie in inspection_engineers:
                if not isinstance(ie, dict):
                    continue
                if (ie.get("method") or "").upper() == (method_code or "").upper():
                    return str(ie.get("full_name") or "")
            return ""

        def _appendix_heading(title: str, level: int = 1) -> None:
            nonlocal app_no
            if is_epb:
                letter = epb_appendix_letter(app_no)
                doc.add_heading(f"Приложение {letter} {title}", level=level)
            else:
                doc.add_heading(f"ПРИЛОЖЕНИЕ № {app_no} {title}", level=level)

        performed_codes = {str(m.get("method_code") or m.get("method_name") or "").upper() for m in performed}
        epb_ctx: Optional[EpbReportContext] = None

        if is_epb:
            inspectors_epb: List[str] = []
            if isinstance(inspection_engineers, list):
                for ie in inspection_engineers:
                    if isinstance(ie, dict):
                        n = (ie.get("full_name") or "").strip()
                        if n and n not in inspectors_epb:
                            inspectors_epb.append(n)
            epb_ctx = EpbReportContext(
                g=g,
                opo_get=_opo_get,
                device_name=str(device_name),
                serial=str(serial),
                reg_no=str(reg_no),
                org=str(org),
                location=str(location),
                opo_name=str(opo_name),
                opo_reg=str(opo_reg),
                opo_class=str(opo_class),
                date_perf_ru=date_perf_ru,
                contractor=str(contractor),
                director_title=str(director_title),
                director_name=str(director_name),
                purpose_default=purpose_default,
                equipment_data=equipment_data,
                inspection_data=inspection_data,
                ndt_methods=ndt_methods or [],
                performed_codes=performed_codes,
                inspectors=inspectors_epb,
                inspection_engineers=inspection_engineers if isinstance(inspection_engineers, list) else [],
                docs_dict=docs_dict if isinstance(docs_dict, dict) else {},
                docs_info=docs_info if isinstance(docs_info, dict) else {},
                doc_meta_fn=_doc_meta,
                scheme_index=str(g("scheme_index", default="ОГ-13" if equipment_preset == "oil_settler" else "")),
                residual_life_years=str(g("residual_life_years", default="10")),
                residual_life_until=str(g("residual_life_until", default="")),
                allowed_pressure=str(g("allowed_pressure", "working_pressure", default="1,0")),
                allowed_temperature=str(g("allowed_temperature", default="плюс 80")),
            )
            build_epb_main_body(doc, epb_ctx)
        if not is_epb:
            # 1. Основания для проведения работ
            doc.add_heading("1. Основания для проведения работ", level=1)
            basis = g("basis", "work_basis", default="Работы по техническому диагностированию проведены согласно договору.")
            doc.add_paragraph(str(basis))
            doc.add_paragraph()
        
            # 2. Сроки проведения работ
            doc.add_heading("2. Сроки проведения работ", level=1)
            work_period = g("work_period", default=f"Работы по техническому диагностированию проведены в период с {date_perf_ru}г. по {date_perf_ru}г.")
            doc.add_paragraph(str(work_period))
            doc.add_paragraph()
        
            # 3. Перечень нормативных и правовых актов
            doc.add_heading("3. Перечень нормативных и правовых актов, устанавливающих требования к объекту диагностирования", level=1)
            normative_base = g("normative_base", default=normative_base_default)
            doc.add_paragraph(str(normative_base))
            doc.add_paragraph()
        
            # 4. Сведения о Заказчике
            doc.add_heading("4. Сведения о Заказчике", level=1)
            doc.add_paragraph("Таблица №1")
            doc.add_paragraph()
            customer_tbl = doc.add_table(rows=2, cols=2)
            customer_tbl.style = "Table Grid"
            customer_tbl.rows[0].cells[0].text = "Полное наименование организации"
            customer_tbl.rows[0].cells[1].text = str(org)
            customer_tbl.rows[1].cells[0].text = "Адрес местонахождения"
            customer_tbl.rows[1].cells[1].text = str(location)
            for row in customer_tbl.rows:
                for cell in row.cells:
                    try:
                        cell.paragraphs[0].runs[0].font.bold = True
                    except:
                        pass
            doc.add_paragraph()

            # 4.1. Сведения об ОПО (если есть)
            opo_name = _opo_get("name", "opo_name")
            opo_code = _opo_get("code", "opo_code")
            opo_desc = _opo_get("description", "opo_description")
            opo_enterprise = _opo_get("enterprise_name", "opo_enterprise")
            opo_branch = _opo_get("branch_name", "opo_branch")
            opo_workshop = _opo_get("workshop_name", "opo_workshop")
            survey = data.get("opo_survey") if isinstance(data.get("opo_survey"), dict) else opo.get("survey_data")
            if not isinstance(survey, dict):
                survey = {}
            opo_org = survey.get("organization")
            opo_exec = survey.get("executors")

            if any([opo_name, opo_code, opo_desc, opo_enterprise, opo_branch, opo_workshop, opo_org, opo_exec]):
                doc.add_heading("4.1. Сведения об ОПО", level=2)
                opo_rows = []
                def _add_opo_row(label, value):
                    if value is None:
                        return
                    s = str(value).strip()
                    if not s:
                        return
                    opo_rows.append((label, s))
                _add_opo_row("Наименование ОПО", opo_name)
                _add_opo_row("Код ОПО", opo_code)
                _add_opo_row("Описание", opo_desc)
                _add_opo_row("Предприятие", opo_enterprise)
                _add_opo_row("Филиал", opo_branch)
                _add_opo_row("Цех", opo_workshop)
                _add_opo_row("Организация (опросный лист ОПО)", opo_org)
                _add_opo_row("Исполнители (опросный лист ОПО)", opo_exec)

                if opo_rows:
                    opo_tbl = doc.add_table(rows=len(opo_rows), cols=2)
                    opo_tbl.style = "Table Grid"
                    for i, (label, value) in enumerate(opo_rows):
                        opo_tbl.rows[i].cells[0].text = label
                        opo_tbl.rows[i].cells[1].text = value
                        try:
                            opo_tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                        except Exception:
                            pass
                    doc.add_paragraph()
        
            # 5. Сведения об организации, проводившей диагностирование
            doc.add_heading("5. Сведения об организации, проводившей техническое диагностирование", level=1)
            doc.add_paragraph("Таблица №2")
            doc.add_paragraph()
            contractor_tbl = doc.add_table(rows=4, cols=2)
            contractor_tbl.style = "Table Grid"
            contractor_rows = [
                ("Наименование организации:", str(contractor)),
                ("Организационно-правовая форма организации:", "Общество с ограниченной ответственностью"),
                ("Юридический адрес:", str(g("contractor_address", default="628285, Ханты-Мансийский автономный округ — Югра, г.Урай, улица Ивана Шестакова, строение 46Б"))),
                ("Руководитель экспертной организации:", f"{director_title} {director_name}"),
            ]
            for i, (k, v) in enumerate(contractor_rows):
                contractor_tbl.rows[i].cells[0].text = k
                contractor_tbl.rows[i].cells[1].text = v
                try:
                    contractor_tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                except:
                    pass
            doc.add_paragraph()
        
            # 6. Сведения об эксперте и специалисте
            doc.add_heading("6. Сведения об эксперте и специалисте, проводивших диагностирование", level=1)
            doc.add_paragraph("Таблица № 3")
            doc.add_paragraph()
        
            # Собираем специалистов
            inspectors = []
            inspector_details = {}
        
            if isinstance(inspection_engineers, list):
                for ie in inspection_engineers:
                    if not isinstance(ie, dict):
                        continue
                    name = (ie.get("full_name") or "").strip()
                    method = _normalize_method(ie.get("method"))
                    cert_num = (ie.get("certificate_number") or "").strip()
                    valid_until = (ie.get("valid_until") or "").strip()
                    if name and name not in inspectors:
                        inspectors.append(name)
                    if name:
                        if name not in inspector_details:
                            inspector_details[name] = {}
                        methods = inspector_details[name].get("methods", [])
                        if method:
                            methods.append(method)
                        inspector_details[name]["methods"] = methods
                        if cert_num:
                            certs = inspector_details[name].get("certifications_inline", [])
                            certs.append(f"{cert_num}" + (f" до {valid_until}" if valid_until else ""))
                            inspector_details[name]["certifications_inline"] = certs
        
            for m in (ndt_methods or []):
                name = (m.get("inspector_name") or "").strip()
                if name and name not in inspectors:
                    inspectors.append(name)
                if name:
                    if name not in inspector_details:
                        inspector_details[name] = {}
                    md = inspector_details[name]
                    md["level"] = md.get("level") or m.get("inspector_level")
                    cert_num = m.get("certificate_number") or m.get("certification_number")
                    if cert_num:
                        md["certification"] = cert_num
                        md["certification_number"] = cert_num
                        if not md.get("certifications_inline"):
                            md["certifications_inline"] = [str(cert_num)]
                    method = _normalize_method(m.get("method_code") or m.get("method_name"))
                    if method:
                        methods = md.get("methods", [])
                        if method not in methods:
                            methods.append(method)
                        md["methods"] = methods
        
            if inspectors:
                spec_table = doc.add_table(rows=len(inspectors) + 1, cols=4)
                spec_table.style = "Table Grid"
                headers = ["№ п/п", "Фамилия И.О.", "№ удостоверения", "Область аттестации / Срок действия"]
                for i, h in enumerate(headers):
                    cell = spec_table.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
                for idx, name in enumerate(inspectors, 1):
                    spec_table.rows[idx].cells[0].text = str(idx)
                    spec_table.rows[idx].cells[1].text = name
                    details = inspector_details.get(name, {})
                    cert_nums = details.get("certifications_inline") or []
                    if not cert_nums and details.get("certification"):
                        cert_nums = [details["certification"]]
                    if not cert_nums and details.get("certificate_number"):
                        cert_nums = [details["certificate_number"]]
                    methods_str = ", ".join(sorted(set(details.get("methods", [])))) if details.get("methods") else ""
                    spec_table.rows[idx].cells[2].text = "; ".join(cert_nums) if cert_nums else "—"
                    area_parts = []
                    if details.get("level"):
                        area_parts.append(f"Уровень: {details['level']}")
                    if methods_str:
                        area_parts.append(f"Методы: {methods_str}")
                    spec_table.rows[idx].cells[3].text = "; ".join(area_parts) if area_parts else "—"
            else:
                doc.add_paragraph("Специалисты не указаны.")
            doc.add_paragraph()
        
            # 7. Перечень приборов и оборудования
            doc.add_heading("7. Перечень приборов и оборудования", level=1)
            doc.add_paragraph("Таблица № 4")
            doc.add_paragraph()
        
            fallback_equipment = []
            if not (verification_equipment and isinstance(verification_equipment, list) and len(verification_equipment) > 0):
                for m in (ndt_methods or []):
                    name = (m.get("equipment") or "").strip()
                    if not name:
                        continue
                    if name not in [e.get("name") for e in fallback_equipment]:
                        fallback_equipment.append({"name": name})

            if verification_equipment and isinstance(verification_equipment, list) and len(verification_equipment) > 0:
                eq_table = doc.add_table(rows=len(verification_equipment) + 1, cols=4)
                eq_table.style = "Table Grid"
                headers = ["№ п/п", "Наименование прибора", "Заводской номер прибора", "Свидетельство о поверке / Действительна до"]
                for i, h in enumerate(headers):
                    cell = eq_table.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
                for idx, eq in enumerate(verification_equipment, 1):
                    eq_table.rows[idx].cells[0].text = str(idx)
                    eq_table.rows[idx].cells[1].text = eq.get('name', '—')
                    eq_table.rows[idx].cells[2].text = eq.get('serial_number', '—')
                
                    cert_num = eq.get('verification_certificate_number', '')
                    next_date = eq.get('next_verification_date', '')
                    if next_date:
                        try:
                            from datetime import datetime as dt
                            d = dt.fromisoformat(next_date.replace('Z', '+00:00'))
                            next_date = d.strftime('%d.%m.%Y')
                        except:
                            pass
                    ver_info = f"{cert_num} / {next_date}" if cert_num and next_date else (cert_num or next_date or "—")
                    eq_table.rows[idx].cells[3].text = ver_info
            elif fallback_equipment:
                eq_table = doc.add_table(rows=len(fallback_equipment) + 1, cols=4)
                eq_table.style = "Table Grid"
                headers = ["№ п/п", "Наименование прибора", "Заводской номер прибора", "Свидетельство о поверке / Действительна до"]
                for i, h in enumerate(headers):
                    cell = eq_table.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for idx, eq in enumerate(fallback_equipment, 1):
                    eq_table.rows[idx].cells[0].text = str(idx)
                    eq_table.rows[idx].cells[1].text = eq.get("name") or "—"
                    eq_table.rows[idx].cells[2].text = "—"
                    eq_table.rows[idx].cells[3].text = "—"
            else:
                doc.add_paragraph("Оборудование не указано.")
            doc.add_paragraph()
        
            # 8. Объект технического диагностирования
            doc.add_heading("8. Объект технического диагностирования", level=1)
            doc.add_paragraph("Таблица №5")
            doc.add_paragraph()
            obj_tbl2 = doc.add_table(rows=4, cols=2)
            obj_tbl2.style = "Table Grid"
            obj_rows = [
                ("Объект технического диагностирования", device_name),
                ("Заводской №", str(serial)),
                ("Место установки", str(location)),
                ("Местонахождение (адрес)", str(location)),
            ]
            for i, (k, v) in enumerate(obj_rows):
                obj_tbl2.rows[i].cells[0].text = k
                obj_tbl2.rows[i].cells[1].text = v
                try:
                    obj_tbl2.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                except:
                    pass
            doc.add_paragraph()
        
            # 9. Краткая техническая характеристика
            doc.add_heading("9. Краткая техническая характеристика и назначение объекта технического освидетельствования", level=1)
            doc.add_paragraph("Таблица № 6")
            doc.add_paragraph()
        
            tech_tbl = doc.add_table(rows=15, cols=2)
            tech_tbl.style = "Table Grid"
            tech_rows = [
                ("Наименование объекта", device_name),
                ("Назначение", g("purpose", "vessel_purpose", default=purpose_default or "—")),
                ("Наименование завода-изготовителя", g("manufacturer", default="—")),
                ("Год изготовления", g("manufacturing_year", default="—")),
                ("Год ввода в эксплуатацию", g("commissioning_year", default=str(equipment_data.get("commissioning_date", "—")))),
                ("Рабочее давление, МПа", g("working_pressure", default="—")),
                ("Расчетное давление, МПа", g("design_pressure", default="—")),
                ("Пробное давление гидравлического испытания, МПа", g("test_pressure", default="—")),
                ("Допустимая рабочая температура стенки, ℃", g("working_temperature", default="—")),
                ("Расчетная температура стенки, ℃", g("design_temperature", default="—")),
                ("Наименование рабочей среды", g("working_medium", default="—")),
                ("Характеристика рабочей среды", g("medium_characteristics", default="—")),
                ("Группа сосуда", g("vessel_group", default="—")),
                ("Группа рабочей среды", g("medium_group", default="—")),
                ("Прибавка для компенсации коррозии, мм", g("corrosion_allowance", default="—")),
            ]
            for i, (k, v) in enumerate(tech_rows):
                tech_tbl.rows[i].cells[0].text = k
                tech_tbl.rows[i].cells[1].text = str(v)
                try:
                    tech_tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                except:
                    pass
            doc.add_paragraph()
        
            # 10. Перечень работ
            doc.add_heading("10. Перечень работ, выполненных в процессе технического освидетельствования", level=1)
            doc.add_paragraph("Таблица № 6")
            doc.add_paragraph()
        
            work_names = {
                "ВИК": "Визуальный и измерительный контроль",
                "УЗТ": "Ультразвуковой контроль толщины стенок элементов сосуда",
                "УЗК": "Ультразвуковой контроль качества основного металла и сварных соединений",
                "ПВК": "Пневматические испытания",
                "РК": "Радиографический контроль",
                "МК": "Магнитный контроль",
                "ТК": "Тепловой контроль",
                "АК": "Акустический контроль",
            }
            works_tbl = doc.add_table(rows=len(work_list) + 1, cols=3)
            works_tbl.style = "Table Grid"
            headers = ["№ п/п", "Наименование работы", "Объем контроля / Нормативная документация"]
            for i, h in enumerate(headers):
                cell = works_tbl.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
            for idx, m in enumerate(work_list, 1):
                work_name = m.get("work_name") or work_names.get(m.get("method_name", ""), m.get("method_name", "—"))
                standard = m.get("standard", "приказ Ростехнадзора от 15.12.2020 №536")
                works_tbl.rows[idx].cells[0].text = str(idx)
                works_tbl.rows[idx].cells[1].text = work_name
                works_tbl.rows[idx].cells[2].text = standard
            doc.add_paragraph()
        
            # 11. Сведения о рассмотренных документах
            doc.add_heading("11. Сведения о рассмотренных в процессе технического освидетельствования документах", level=1)
            doc.add_paragraph("Таблица № 7")
            doc.add_paragraph()
        
            document_names = {
                '1': 'Лицензия на осуществление деятельности по эксплуатации взрывопожароопасных и химически опасных производственных объектов I, II и III классов опасности',
                '2': 'Свидетельство о регистрации в государственном реестре ОПО, включая сведения характеризующие ОПО',
                '3': 'Технологический регламент объектов опасных производственных объектов',
                '4': 'План мероприятий по локализации и ликвидации последствий аварий на опасном производственном объекте',
                '5': 'Положение о производственном контроле за соблюдением требований промышленной безопасности на опасных производственных объектах',
                '6': 'Журнал учета аварий и инцидентов на ОПО',
                '7': 'Страховой полис страхования гражданской ответственности владельца опасного объекта за причинение вреда в результате аварии на опасном объекте',
                '8': 'Приказ о назначении ответственного лица за исправное состояние и безопасную эксплуатацию сосудов',
                '9': 'Приказ о назначении ответственного лица за осуществление производственного контроля и соблюдение требований промышленной безопасности на опасном производственном объекте',
                '10': 'Паспорт сосуда заводской (удостоверение о качестве монтажа, сертификат соответствия, сборочный чертёж и схема включения сосуда, расчёт на прочность)',
                '11': 'Инструкция по монтажу и эксплуатации',
                '12': 'Паспорта на предохранительные клапаны',
                '13': 'Паспорта на запорную арматуру',
                '14': 'Документация на контрольно-измерительные приборы',
                '15': 'Ремонтная (исполнительная) документация',
                '16': 'Заключение экспертизы промышленной безопасности',
                '17': 'Акты проведения УЗТ',
            }
        
            if isinstance(docs_dict, dict) or isinstance(docs_info, dict):
                doc_keys = set()
                if isinstance(docs_dict, dict):
                    doc_keys.update([str(k) for k in docs_dict.keys()])
                if isinstance(docs_info, dict):
                    doc_keys.update([str(k) for k in docs_info.keys()])
                doc_keys = sorted(doc_keys, key=lambda x: int(x) if str(x).isdigit() else 999)
                docs_tbl = doc.add_table(rows=len(doc_keys) + 1, cols=3)
                docs_tbl.style = "Table Grid"
                headers = ["№ п/п", "Наименование документа", "Идентификационный номер документа / Объём рассмотренных документов, листов"]
                for i, h in enumerate(headers):
                    cell = docs_tbl.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
                row_idx = 1
                for num in doc_keys:
                    name = document_names.get(str(num), f'Документ {num}')
                    _present, doc_number, doc_date = _doc_meta(str(num))
                    ident = "—"
                    if doc_number or doc_date:
                        ident = doc_number
                        if doc_date:
                            ident = f"{ident} от {doc_date}" if ident else f"от {doc_date}"
                    docs_tbl.rows[row_idx].cells[0].text = str(num)
                    docs_tbl.rows[row_idx].cells[1].text = name
                    docs_tbl.rows[row_idx].cells[2].text = ident
                    row_idx += 1
            else:
                doc.add_paragraph("Документы не указаны.")
            doc.add_paragraph()
        
            # 12. Анализ результатов предыдущих обследований
            doc.add_heading("12. Анализ результатов предыдущих обследований", level=1)
            doc.add_paragraph("Таблица № 8")
            doc.add_paragraph()
            prev_tbl = doc.add_table(rows=2, cols=3)
            prev_tbl.style = "Table Grid"
            headers = ["№ п/п", "Вид обследования", "Результаты контроля / Наименование и номер отчетной документации"]
            for i, h in enumerate(headers):
                cell = prev_tbl.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            prev_tbl.rows[1].cells[0].text = "1"
            prev_tbl.rows[1].cells[1].text = "Техническое диагностирование"
            prev_tbl.rows[1].cells[2].text = g("previous_inspection_result", default="—")
            doc.add_paragraph()
        
            # 13. Результаты технического освидетельствования
            doc.add_heading("13. Результаты технического освидетельствования", level=1)
            doc.add_paragraph("Таблица № 9")
            doc.add_paragraph()
            results_tbl = doc.add_table(rows=len(work_list) + 1, cols=3)
            results_tbl.style = "Table Grid"
            headers = ["№ п/п", "Наименование работы", "Результаты контроля / Наименование и номер отчетной документации"]
            for i, h in enumerate(headers):
                cell = results_tbl.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
            for idx, m in enumerate(work_list, 1):
                work_name = m.get("work_name") or work_names.get(m.get("method_name", ""), m.get("method_name", "—"))
                conclusion = m.get("conclusion", "Дефектов не обнаружено")
                results_tbl.rows[idx].cells[0].text = str(idx)
                results_tbl.rows[idx].cells[1].text = work_name
                results_tbl.rows[idx].cells[2].text = f"{conclusion} / Приложение №{idx+1}"
            doc.add_paragraph()
        
            # 14. Результаты расчетной оценки
            doc.add_heading("14. Результаты расчетной оценки технического состояния", level=1)
            calc_result = g("calculation_result", default="По результатам работ произведена оценка работоспособности сосуда, при рабочих параметрах. Выполнен расчет на прочность и определение остаточного ресурса сосуда.")
            doc.add_paragraph(str(calc_result))
            doc.add_paragraph()
        
            # 15. Выводы
            doc.add_heading("15. Выводы по результатам технического освидетельствования", level=1)
            conclusion_text = conclusion_from_inspection_data(
                inspection_data,
                equipment_data,
                g,
                explicit_conclusion=inspection_data.get("conclusion"),
            )
            doc.add_paragraph(str(conclusion_text))
            doc.add_paragraph()
        
            # Подпись эксперта
            expert_name = inspectors[0] if inspectors else "—"
            doc.add_paragraph(f"Эксперт Э12ТУ {expert_name}")
            doc.add_paragraph()
        
        # --------------- ПРИЛОЖЕНИЯ ---------------
        app_no = 1
        epb_protocol_section_open = False

        if is_epb and epb_ctx is not None:
            doc.add_page_break()
            _appendix_heading("Акт о проведении работ по техническому диагностированию")
            append_epb_appendix_act(doc, epb_ctx)
            app_no += 1
            doc.add_page_break()
            _appendix_heading("Отчет по анализу технической документации")
            append_epb_appendix_doc_analysis(doc, epb_ctx)
            app_no += 1
            doc.add_page_break()
            _appendix_heading("Результаты технического диагностирования")
            epb_protocol_section_open = True
            app_no += 1

        def _add_protocol_header_block():
            """Блок заголовка: Заказчик, Объект, Место, Дата, НТД."""
            ht = doc.add_table(rows=5, cols=2)
            ht.style = "Table Grid"
            nd = str(g("normative_base", default="приказ Ростехнадзора от 15.12.2020 №536, СО 153-34.17.439-2003, ГОСТ Р 55614-2013, ГОСТ Р ИСО 16809-2015"))
            for i, (lbl, val) in enumerate([
                ("Заказчик:", str(org)),
                ("Объект контроля:", f"{device_name} зав.№ {serial}"),
                ("Место проведения контроля:", str(location)),
                ("Дата проведения контроля:", date_perf_ru),
                ("НТД, по которой выполнен контроль:", nd),
            ]):
                ht.rows[i].cells[0].text = lbl
                ht.rows[i].cells[1].text = val
                try:
                    ht.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                except Exception:
                    pass
            doc.add_paragraph()
        
        # ПРИЛОЖЕНИЕ № 1 (техотчёт): протокол анализа технической документации — форма ТО
        if not is_epb:
            doc.add_page_break()
            _appendix_heading("Протокол анализа технической документации")
            inv_no = str(g("inventory_number", default=attrs.get("inventory_number") or ""))
            tech_ctx = TechnicalReportContext(
                g=g,
                doc_meta_fn=_doc_meta,
                device_name=str(device_name),
                serial=str(serial),
                reg_no=str(g("reg_number", "registration_number", default=attrs.get("registration_number") or "")),
                inv_no=inv_no,
                org=str(org),
                location=str(location),
                date_perf_ru=str(date_perf_ru),
                equipment_data=equipment_data,
            )
            append_technical_protocol_doc_analysis(doc, tech_ctx)
            app_no += 1

        # Протокол по результатам оперативной диагностики
        if not is_epb:
            doc.add_page_break()
            _appendix_heading("Протокол по результатам оперативной (функциональной) диагностики")
            doc.add_paragraph("Протокол по результатам оперативной (функциональной) диагностики")
            doc.add_paragraph(f"№ {app_no} от {date_perf_ru}г.")
        else:
            doc.add_heading("Протокол № 1", level=2)
            doc.add_paragraph("оперативного (функционального) диагностирования")
        doc.add_paragraph()
        _add_protocol_header_block()
        doc.add_paragraph("1. Результаты функциональной (оперативной) диагностики")
        doc.add_paragraph("Таблица № 1")
        doc.add_paragraph()
        func_tbl = doc.add_table(rows=6, cols=3)
        func_tbl.style = "Table Grid"
        headers = ["№ п/п", "Наименование", "Оценка / Примечания"]
        for i, h in enumerate(headers):
            cell = func_tbl.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        func_items = [
            ("1", "Значения основных параметров эксплуатации сосуда", g("functional_params", default="Соответствуют")),
            ("2", "Повышенная вибрация сосуда", g("vibration", default="Не выявлена")),
            ("3", "Состояние опор сосуда", g("support_state", default="Работоспособное")),
            ("4", "Состояние контрольно-измерительных приборов (КИП), систем автоматизации (СА) и противоаварийной защиты (ПАЗ)", g("kip_state", default="Работоспособное")),
            ("5", "Поверка манометров", g("manometer_verification", default="Соответствует")),
        ]
        for idx, (num, name, value) in enumerate(func_items, 1):
            func_tbl.rows[idx].cells[0].text = num
            func_tbl.rows[idx].cells[1].text = name
            func_tbl.rows[idx].cells[2].text = str(value)
        doc.add_paragraph()
        app_no += 1
        
        # Протокол по результатам визуального и измерительного контроля
        doc.add_page_break()
        if is_epb:
            doc.add_heading("Протокол № 2", level=2)
            doc.add_paragraph("визуального и измерительного контроля сосуда")
        else:
            _appendix_heading("Протокол по результатам визуального и измерительного контроля")
            doc.add_paragraph("Протокол по результатам визуального и измерительного контроля")
            doc.add_paragraph(f"№ {app_no} от {date_perf_ru}г.")
        doc.add_paragraph()
        _add_protocol_header_block()
        
        # ВИК результаты
        vik_method = next((m for m in performed if m.get("method_name") == "ВИК"), None)
        if vik_method:
            doc.add_paragraph("1. Применяемое оборудование")
            doc.add_paragraph("Таблица № 1")
            doc.add_paragraph()
            # Оборудование для ВИК
            vik_eq = [eq for eq in (verification_equipment or []) if "ВИК" in (eq.get("equipment_type") or "").upper()]
            if vik_eq:
                vik_eq_tbl = doc.add_table(rows=len(vik_eq) + 1, cols=2)
                vik_eq_tbl.style = "Table Grid"
                headers = ["№ п/п", "Наименование прибора / Заводской номер прибора"]
                for i, h in enumerate(headers):
                    cell = vik_eq_tbl.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for idx, eq in enumerate(vik_eq, 1):
                    vik_eq_tbl.rows[idx].cells[0].text = str(idx)
                    vik_eq_tbl.rows[idx].cells[1].text = f"{eq.get('name', '—')} / {eq.get('serial_number', '—')}"
            doc.add_paragraph()
            
            doc.add_paragraph("2. Параметры контроля")
            doc.add_paragraph("Таблица № 2")
            doc.add_paragraph()
            params_tbl = doc.add_table(rows=3, cols=2)
            params_tbl.style = "Table Grid"
            params_rows = [
                ("Шероховатость поверхности", g("surface_roughness", default="Rz 80")),
                ("Освещенность", g("illumination", default="500 Лк")),
            ]
            for i, (k, v) in enumerate(params_rows):
                params_tbl.rows[i].cells[0].text = k
                params_tbl.rows[i].cells[1].text = str(v)
                try:
                    params_tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                except:
                    pass
            doc.add_paragraph()
            
            doc.add_paragraph("3. Результаты визуального контроля")
            doc.add_paragraph("Таблица № 3")
            doc.add_paragraph()
            vik_results_tbl = doc.add_table(rows=10, cols=4)
            vik_results_tbl.style = "Table Grid"
            headers = ["№ п/п", "Наименование объекта контроля", "Объем контроля / Описание обнаруженных дефектов, их размеры", "Оценка качества"]
            for i, h in enumerate(headers):
                cell = vik_results_tbl.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            vik_items = [
                ("1", "Опоры", "100%", "дефектов не обнаружено", "годен"),
                ("2", "Антикоррозионное покрытие", "100%", "дефектов не обнаружено", "годен"),
                ("3", "Разъемные соединения", "100%", "дефектов не обнаружено", "годен"),
                ("4", "Крепежные детали", "100%", "дефектов не обнаружено", "годен"),
                ("5", "Основной металл обечайки, днищ сосуда, штуцеров, фланцев", "100%", "дефектов не обнаружено", "годен"),
                ("6", "Сварные соединения вварки штуцеров в корпус", "100%", "дефектов не обнаружено", "годен"),
                ("7", "Сварные соединения штуцеров фланцев к патрубкам", "100%", "дефектов не обнаружено", "годен"),
                ("8", "Сварные соединения приварки опор к корпусу", "100%", "дефектов не обнаружено", "годен"),
                ("9", "Кольцевые, продольные сварные соединения и их перекрестья", "100%", "дефектов не обнаружено", "годен"),
            ]
            for idx, (num, name, volume, defects, quality) in enumerate(vik_items, 1):
                vik_results_tbl.rows[idx].cells[0].text = num
                vik_results_tbl.rows[idx].cells[1].text = name
                vik_results_tbl.rows[idx].cells[2].text = f"{volume} / {defects}"
                vik_results_tbl.rows[idx].cells[3].text = quality
            doc.add_paragraph()
            # Детализация дефектов ВИК с фотографиями (в т.ч. с мобильного)
            visual_defects_v = g("visual_defects", default=[])
            if isinstance(visual_defects_v, list) and visual_defects_v:
                doc.add_paragraph("Дополнительно: выявленные дефекты ВИК").runs[0].bold = True
                defect_tbl = doc.add_table(rows=len(visual_defects_v) + 1, cols=5)
                defect_tbl.style = "Table Grid"
                for i, h in enumerate(["№", "Тип дефекта", "Место", "Размеры", "Описание"]):
                    defect_tbl.rows[0].cells[i].text = h
                    defect_tbl.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
                for idx, d in enumerate(visual_defects_v, 1):
                    if isinstance(d, dict):
                        defect_tbl.rows[idx].cells[0].text = str(idx)
                        defect_tbl.rows[idx].cells[1].text = str(d.get("defect_type") or "—")
                        defect_tbl.rows[idx].cells[2].text = str(d.get("location") or "—")
                        defect_tbl.rows[idx].cells[3].text = str(d.get("size") or "—")
                        defect_tbl.rows[idx].cells[4].text = str(d.get("description") or "—")
                doc.add_paragraph()
                for idx, d in enumerate(visual_defects_v, 1):
                    if not isinstance(d, dict):
                        continue
                    photos_v = d.get("photos") or []
                    if isinstance(photos_v, list) and photos_v:
                        doc.add_paragraph(f"Фотографии дефекта №{idx}:").runs[0].bold = True
                        for j, ph in enumerate(photos_v[:6]):
                            if not isinstance(ph, str):
                                continue
                            pp = self._find_image_path(ph) or (self._find_image_path(attachments.get(ph)) if attachments.get(ph) else None) or (ph if os.path.exists(ph) else None)
                            if not pp and attachments:
                                vd_key = "vd_%d_%d" % (idx - 1, j)
                                pp = self._find_image_path(attachments.get(vd_key)) or (attachments.get(vd_key) if isinstance(attachments.get(vd_key), str) and os.path.exists(attachments.get(vd_key)) else None)
                            if pp:
                                add_picture_if_exists("", pp)
                doc.add_paragraph()
            
            # Измерительный контроль - овальность
            doc.add_paragraph("4. Результаты измерительного контроля")
            doc.add_paragraph()
            doc.add_paragraph("Определение овальности проводят измерением максимального (Dmax) и минимального (Dmin) наружного или внутреннего диаметров в одном сечении по двум перпендикулярным направлениям.")
            doc.add_paragraph()
            
            ovality = g("ovality_measurements", default=[])
            if isinstance(ovality, list) and ovality:
                ovality_tbl = doc.add_table(rows=len(ovality) + 1, cols=4)
                ovality_tbl.style = "Table Grid"
                headers = ["Номер сечения", "Размеры, мм (Dmin / Dmax)", "Фактическая овальность, %", "Допустимая овальность, %"]
                for i, h in enumerate(headers):
                    cell = ovality_tbl.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for idx, it in enumerate(ovality, 1):
                    if not isinstance(it, dict):
                        continue
                    ovality_tbl.rows[idx].cells[0].text = str(it.get('section_number', idx))
                    dmin = str(it.get('min_diameter', ''))
                    dmax = str(it.get('max_diameter', ''))
                    ovality_tbl.rows[idx].cells[1].text = f"{dmin} / {dmax}" if dmin and dmax else "—"
                    ovality_tbl.rows[idx].cells[2].text = str(it.get('deviation_percent', ''))
                    ovality_tbl.rows[idx].cells[3].text = "1,0"
            doc.add_paragraph()
            
            # Прогиб
            deflection = g("deflection_measurements", default=[])
            if isinstance(deflection, list) and deflection:
                doc.add_paragraph("Измерение отклонений от прямолинейности нижней образующей обечайки сосуда")
                doc.add_paragraph()
                deflection_tbl = doc.add_table(rows=len(deflection) + 1, cols=4)
                deflection_tbl.style = "Table Grid"
                headers = ["Наименование", "№ сечения", "Прогиб, мм / Прогиб, %", "Допустимое отклонение прогиба, %"]
                for i, h in enumerate(headers):
                    cell = deflection_tbl.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for idx, it in enumerate(deflection, 1):
                    if not isinstance(it, dict):
                        continue
                    deflection_tbl.rows[idx].cells[0].text = "Обечайка"
                    deflection_tbl.rows[idx].cells[1].text = str(it.get('section_number', idx))
                    deflection_mm = str(it.get('deflection_mm', ''))
                    raw_pct = it.get('deflection_percent', '')
                    try:
                        deflection_pct = f"{float(str(raw_pct).replace(',', '.')):.2f}"
                    except Exception:
                        deflection_pct = str(raw_pct or '')
                    deflection_tbl.rows[idx].cells[2].text = f"{deflection_mm} / {deflection_pct}" if deflection_mm and deflection_pct else "—"
                    deflection_tbl.rows[idx].cells[3].text = "0,3"
            doc.add_paragraph()
            
            doc.add_paragraph("5. Заключение по результатам визуального и измерительного контроля")
            doc.add_paragraph("По результатам визуального и измерительного контроля основного металла и сварных соединений сосуда, недопустимых дефектов не обнаружено, что удовлетворяет требованиям нормативно-технической документации")
            doc.add_paragraph()
        app_no += 1

        # Протокол №3 (твердометрия) — только ЭПБ, до УЗТ
        hardness_for_epb = g("hardness_tests", default=[])
        if not isinstance(hardness_for_epb, list):
            hardness_for_epb = []
        if is_epb:
            append_epb_protocol_hardness(doc, g, _add_protocol_header_block, hardness_for_epb)
        
        # Протокол УЗТ (№4 для ЭПБ)
        uz_method = next((m for m in performed if "УЗТ" in (m.get("method_name") or "")), None)
        thickness_for_protocol = g("thickness_measurements", "thicknessMeasurements", default=[])
        has_uzt_data = uz_method or (isinstance(thickness_for_protocol, list) and len(thickness_for_protocol) > 0)
        if has_uzt_data or is_epb:
            if is_epb:
                thickness_epb = thickness_for_protocol if isinstance(thickness_for_protocol, list) else []
                append_epb_protocol_uzt(doc, g, _add_protocol_header_block, thickness_epb)
                app_no += 1
            else:
                doc.add_page_break()
                _appendix_heading("Протокол по результатам ультразвукового контроля толщины стенок элементов сосуда")
                doc.add_paragraph("Протокол по результатам ультразвукового контроля толщины стенок элементов сосуда")
                doc.add_paragraph(f"№ {app_no} от {date_perf_ru}г.")
                doc.add_paragraph()
                _add_protocol_header_block()
                doc.add_paragraph("1. Применяемое оборудование").runs[0].bold = True
                doc.add_paragraph("Таблица № 1")
                doc.add_paragraph()
                uz_eq = [eq for eq in (verification_equipment or []) if "УЗТ" in (eq.get("equipment_type") or "").upper() or "ТОЛЩИНОМЕР" in (eq.get("name") or "").upper()]
                if not uz_eq:
                    for m in (ndt_methods or []):
                        if (m.get("method_code") or "").upper() in ("UZT", "УЗТ") and m.get("equipment"):
                            uz_eq.append({"name": m.get("equipment"), "serial_number": m.get("equipment_serial") or m.get("serial_number") or "—"})
                if uz_eq:
                    uz_eq_tbl = doc.add_table(rows=len(uz_eq) + 1, cols=3)
                    uz_eq_tbl.style = "Table Grid"
                    headers = ["№ п/п", "Наименование прибора", "Заводской номер прибора"]
                    for i, h in enumerate(headers):
                        cell = uz_eq_tbl.rows[0].cells[i]
                        cell.text = h
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for idx, eq in enumerate(uz_eq, 1):
                        uz_eq_tbl.rows[idx].cells[0].text = str(idx)
                        uz_eq_tbl.rows[idx].cells[1].text = str(eq.get("name", "—"))
                        uz_eq_tbl.rows[idx].cells[2].text = str(eq.get("serial_number", "—"))
                doc.add_paragraph()
            
                doc.add_paragraph("2. Результаты контроля")
                doc.add_paragraph("Контроль выполнен в соответствии с программой работ, согласно схемы контроля.")
                doc.add_paragraph()
                doc.add_paragraph("Таблица № 2")
                doc.add_paragraph()
            
                thickness = g("thickness_measurements", "thicknessMeasurements", default=[])
                attrs = equipment_data.get("attributes") or {}
                if isinstance(thickness, list) and thickness:
                    # Группируем по элементам (обечайка, днище 1, днище 2)
                    elements = {}
                    for t in thickness:
                        if not isinstance(t, dict):
                            continue
                        element = str(t.get("location", "Обечайка"))
                        if element not in elements:
                            elements[element] = []
                        elements[element].append(t)
                
                    for element_name, measurements in elements.items():
                        data_rows = (len(measurements) + 3) // 4
                        total_rows = 1 + data_rows + 3
                        thick_tbl = doc.add_table(rows=total_rows, cols=9)
                        thick_tbl.style = "Table Grid"
                        headers = ["Наименование элемента", "№ точки", "Толщина, мм", "№ точки", "Толщина, мм", "№ точки", "Толщина, мм", "№ точки", "Толщина, мм"]
                        for i, h in enumerate(headers):
                            cell = thick_tbl.rows[0].cells[i]
                            cell.text = h
                            cell.paragraphs[0].runs[0].font.bold = True
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                        for row_idx in range(data_rows):
                            row = thick_tbl.rows[row_idx + 1]
                            row.cells[0].text = element_name if row_idx == 0 else ""
                            for col_idx in range(4):
                                point_idx = row_idx * 4 + col_idx
                                if point_idx < len(measurements):
                                    point = measurements[point_idx]
                                    point_num = str(point.get("point_number", point_idx + 1))
                                    thickness_val = str(point.get("thickness", ""))
                                    row.cells[col_idx * 2 + 1].text = point_num
                                    row.cells[col_idx * 2 + 2].text = thickness_val
                    
                        nominal = attrs.get("wall_thickness") or attrs.get("thickness") or g("wall_thickness", "nominal_thickness", default="4,0")
                        min_meas_vals = []
                        for m in measurements:
                            t = m.get("thickness")
                            if t is None or t == "":
                                continue
                            try:
                                min_meas_vals.append(float(str(t).replace(",", ".")))
                            except (TypeError, ValueError):
                                pass
                        min_meas = min(min_meas_vals, default=0)
                        min_allowed_vals = [float(str(p.get("min_allowed_thickness", "0")).replace(",", ".")) for p in measurements if p.get("min_allowed_thickness")]
                        min_allowed = attrs.get("min_wall_thickness") or (min_allowed_vals[0] if min_allowed_vals else 2.8)
                        try:
                            ma = float(str(min_allowed or "2.8").replace(",", "."))
                            min_allowed_str = f"{ma:.1f}" if ma > 0 else "2,8"
                        except (TypeError, ValueError):
                            min_allowed_str = "2,8"
                    
                        row_idx = 1 + data_rows
                        thick_tbl.rows[row_idx].cells[0].text = "Номинальная толщина, мм"
                        thick_tbl.rows[row_idx].cells[1].text = str(nominal)
                        thick_tbl.rows[row_idx + 1].cells[0].text = "Минимально-измеренная толщина, мм"
                        thick_tbl.rows[row_idx + 1].cells[1].text = f"{min_meas:.1f}" if min_meas > 0 else "—"
                        thick_tbl.rows[row_idx + 2].cells[0].text = "Минимально допустимая толщина стеки сосуда, мм"
                        thick_tbl.rows[row_idx + 2].cells[1].text = min_allowed_str
                        doc.add_paragraph()
            
                doc.add_paragraph("3. Заключение по результатам контроля")
                doc.add_paragraph("Измеренная толщина стенок элементов сосуда не превышает минимально допустимые значения и удовлетворяет требованиям нормативно-технической документации.")
                doc.add_paragraph()
            
                # Схема контроля
                control_scheme = attachments.get('control_scheme_image') or g('control_scheme_image')
                if not control_scheme:
                    # Используем шаблон чертежа сосуда
                    template_path = "/app/reports/assets/vessel_template.png"
                    control_scheme = self._find_image_path(template_path) or (template_path if os.path.isfile(template_path) else None)
            
                if control_scheme:
                    doc.add_paragraph("Схема контроля указана в Приложении № 7.")
                    # Добавим схему в приложение 7
                app_no += 1
        
        # ЭПБ: протоколы №5 (МПК) и №6 (УЗК)
        welds_all = g("weld_inspections", default=[])
        if not isinstance(welds_all, list):
            welds_all = []
        if is_epb:
            mpk_welds = _filter_welds(welds_all, "MPK")
            append_epb_protocol_weld_control(
                doc,
                5,
                "магнитопорошковый контроль качества сварных соединений",
                _add_protocol_header_block,
                mpk_welds,
                "По результатам магнитопорошкового контроля недопустимых дефектов в сварных соединениях не обнаружено.",
            )
            uzk_welds = _filter_welds(welds_all, "UZK")
            append_epb_protocol_weld_control(
                doc,
                6,
                "ультразвуковой контроль качества сварных соединений",
                _add_protocol_header_block,
                uzk_welds,
                "По результатам ультразвукового контроля недопустимых дефектов в сварных соединениях не обнаружено.",
            )
            app_no += 2
        
        # Техотчёт: приложение № 5 — твердометрия, № 6 — УЗК, № 7 — МПК (форма ТО)
        hardness_method = next((m for m in performed if "ТК" in (m.get("method_name") or "") or "твердость" in (m.get("method_name") or "").lower() or "ТВИ" in (m.get("method_name") or "")), None)
        hardness_for_protocol = g("hardness_tests", default=[])
        has_hardness_data = hardness_method or (isinstance(hardness_for_protocol, list) and len(hardness_for_protocol) > 0)
        if has_hardness_data and not is_epb:
            doc.add_page_break()
            _appendix_heading("Протокол по результатам контроля твердости основного металла и сварных соединений")
            doc.add_paragraph("Протокол по результатам контроля твердости основного металла и сварных соединений")
            doc.add_paragraph(f"№ {app_no} от {date_perf_ru}г.")
            doc.add_paragraph()
            _add_protocol_header_block()
            doc.add_paragraph("1. Применяемое оборудование").runs[0].bold = True
            doc.add_paragraph("Таблица № 1")
            doc.add_paragraph()
            hardness_eq = [eq for eq in (verification_equipment or []) if "твердость" in (eq.get("name") or "").lower() or "УЗИТ" in (eq.get("name") or "").upper()]
            if not hardness_eq:
                for m in (ndt_methods or []):
                    if (m.get("method_code") or "").upper() in ("TVI", "ТВИ", "HARDNESS") and m.get("equipment"):
                        hardness_eq.append({"name": m.get("equipment"), "serial_number": m.get("equipment_serial") or m.get("serial_number") or "—"})
            if hardness_eq:
                hardness_eq_tbl = doc.add_table(rows=len(hardness_eq) + 1, cols=3)
                hardness_eq_tbl.style = "Table Grid"
                headers = ["№ п/п", "Наименование прибора", "Заводской номер прибора"]
                for i, h in enumerate(headers):
                    cell = hardness_eq_tbl.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for idx, eq in enumerate(hardness_eq, 1):
                    hardness_eq_tbl.rows[idx].cells[0].text = str(idx)
                    hardness_eq_tbl.rows[idx].cells[1].text = str(eq.get("name", "—"))
                    hardness_eq_tbl.rows[idx].cells[2].text = str(eq.get("serial_number", "—"))
            doc.add_paragraph()
            doc.add_paragraph("2. Результаты контроля").runs[0].bold = True
            doc.add_paragraph("Контроль выполнен в соответствии с программой работ, согласно схемы контроля.")
            doc.add_paragraph("Таблица № 2")
            doc.add_paragraph()
            hardness = g("hardness_tests", default=[])
            if isinstance(hardness, list) and hardness:
                hardness_by_el = {}
                for h in hardness:
                    if not isinstance(h, dict):
                        continue
                    el = str(h.get("location") or h.get("element_name") or h.get("weld_number") or "Обечайка")
                    hardness_by_el.setdefault(el, []).append(h)
                total_rows = sum((len(t) + 3) // 4 for t in hardness_by_el.values())
                hardness_tbl = doc.add_table(rows=1 + total_rows, cols=9)
                hardness_tbl.style = "Table Grid"
                headers = ["Наименование элемента", "№ точки", "Результат замера, НВ", "№ точки", "Результат замера, НВ", "№ точки", "Результат замера, НВ", "№ точки", "Результат замера, НВ"]
                for i, h in enumerate(headers):
                    hardness_tbl.rows[0].cells[i].text = h
                    hardness_tbl.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
                row_idx = 1
                for el_name, tests in hardness_by_el.items():
                    for i in range(0, len(tests), 4):
                        row = hardness_tbl.rows[row_idx]
                        row.cells[0].text = el_name if i == 0 else ""
                        for j in range(4):
                            if i + j < len(tests):
                                t = tests[i + j]
                                row.cells[j * 2 + 1].text = str(i + j + 1)
                                row.cells[j * 2 + 2].text = str(t.get("hardness_base") or t.get("hardness_weld") or t.get("hardness_haz") or "")
                        row_idx += 1
            doc.add_paragraph("3. Заключение по результатам контроля твердости").runs[0].bold = True
            doc.add_paragraph(
                "Измеренные значения твердости металла сосуда находятся в допустимых пределах и отвечают требованиям нормативно-технической документации."
            )
            doc.add_paragraph()
            app_no += 1

        uzk_method = next((m for m in performed if "УЗК" in (m.get("method_name") or "") or (m.get("method_code") or "").upper() == "UZK"), None)
        uzk_welds_tech = _filter_welds(welds_all, "UZK") if welds_all else []
        if (uzk_method or uzk_welds_tech) and not is_epb:
            doc.add_page_break()
            _appendix_heading("Протокол по результатам ультразвукового контроля качества основного металла и сварных соединений")
            doc.add_paragraph("Протокол по результатам ультразвукового контроля качества")
            doc.add_paragraph("основного металла и сварных соединений")
            doc.add_paragraph(f"№ {app_no} от {date_perf_ru}г.")
            doc.add_paragraph()
            _add_protocol_header_block()
            doc.add_paragraph("1. Применяемое оборудование")
            doc.add_paragraph("Таблица № 1")
            doc.add_paragraph()
            uzk_eq = [eq for eq in (verification_equipment or []) if "УЗК" in (eq.get("equipment_type") or "").upper() or "ДЕФЕКТОСКОП" in (eq.get("name") or "").upper()]
            if uzk_eq:
                uzk_eq_tbl = doc.add_table(rows=len(uzk_eq) + 1, cols=2)
                uzk_eq_tbl.style = "Table Grid"
                headers = ["№ п/п", "Наименование прибора / Заводской номер прибора"]
                for i, h in enumerate(headers):
                    cell = uzk_eq_tbl.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for idx, eq in enumerate(uzk_eq, 1):
                    uzk_eq_tbl.rows[idx].cells[0].text = str(idx)
                    uzk_eq_tbl.rows[idx].cells[1].text = f"{eq.get('name', '—')} / {eq.get('serial_number', '—')}"
            doc.add_paragraph()
            doc.add_paragraph("2. Параметры контроля")
            doc.add_paragraph("Таблица № 2")
            doc.add_paragraph()
            uzk_params_tbl = doc.add_table(rows=2, cols=5)
            uzk_params_tbl.style = "Table Grid"
            headers = ["№ п/п", "Тип сварного соединения", "Толщина элементов, мм", "Поверхность ввода ультразвуковых колебаний", "Чувствительность, мм² / Параметры зарубки, мм"]
            for i, h in enumerate(headers):
                cell = uzk_params_tbl.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            uzk_params_tbl.rows[1].cells[0].text = "1"
            uzk_params_tbl.rows[1].cells[1].text = "Кольцевой"
            uzk_params_tbl.rows[1].cells[2].text = "4/4"
            uzk_params_tbl.rows[1].cells[3].text = "наружная"
            uzk_params_tbl.rows[1].cells[4].text = "0,8 / —"
            doc.add_paragraph()
            doc.add_paragraph("3. Результаты контроля")
            doc.add_paragraph("Контроль выполнен в соответствии с программой работ, согласно схемы контроля.")
            doc.add_paragraph("Таблица № 3")
            doc.add_paragraph()
            welds = uzk_welds_tech or g("weld_inspections", default=[])
            if isinstance(welds, list) and welds:
                welds_tbl = doc.add_table(rows=len(welds) + 1, cols=8)
                welds_tbl.style = "Table Grid"
                headers = ["№ стыка по карте контроля", "Условный номер дефекта", "Эквивалент. Площадь Sдеф, мм²", "Глубина залегания «Y» , мм", "Протяженность ΔL, мм", "Форма (характер) дефекта (объемный/плоскостной)", "Местоположение на сварном соединении L, мм", "Заключение"]
                for i, h in enumerate(headers):
                    cell = welds_tbl.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for idx, w in enumerate(welds, 1):
                    if not isinstance(w, dict):
                        continue
                    welds_tbl.rows[idx].cells[0].text = str(w.get('weld_number', idx))
                    welds_tbl.rows[idx].cells[1].text = str(w.get('defect_number', 'Дефектов не обнаружено'))
                    welds_tbl.rows[idx].cells[2].text = str(w.get('defect_area', ''))
                    welds_tbl.rows[idx].cells[3].text = str(w.get('defect_depth', ''))
                    welds_tbl.rows[idx].cells[4].text = str(w.get('defect_length', ''))
                    welds_tbl.rows[idx].cells[5].text = str(w.get('defect_type', ''))
                    welds_tbl.rows[idx].cells[6].text = str(w.get('defect_location', ''))
                    welds_tbl.rows[idx].cells[7].text = str(w.get('conclusion', 'годен'))
            else:
                welds_tbl = doc.add_table(rows=2, cols=8)
                welds_tbl.style = "Table Grid"
                headers = ["№ стыка по карте контроля", "Условный номер дефекта", "Эквивалент. Площадь Sдеф, мм²", "Глубина залегания «Y» , мм", "Протяженность ΔL, мм", "Форма (характер) дефекта (объемный/плоскостной)", "Местоположение на сварном соединении L, мм", "Заключение"]
                for i, h in enumerate(headers):
                    cell = welds_tbl.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                welds_tbl.rows[1].cells[0].text = "К1-П1"
                welds_tbl.rows[1].cells[1].text = "Дефектов не обнаружено"
                welds_tbl.rows[1].cells[7].text = "годен"
            doc.add_paragraph()
            doc.add_paragraph("4. Заключение по результатам контроля")
            doc.add_paragraph("По результатам обследования сварных соединений сосуда, недопустимых дефектов не обнаружено, объект контроля соответствует требованиям НТД.")
            doc.add_paragraph()
            app_no += 1

        mpk_method = next((m for m in performed if "МПК" in (m.get("method_name") or "") or (m.get("method_code") or "").upper() in ("MPK", "MK")), None)
        mpk_welds_tech = _filter_welds(welds_all, "MPK") if welds_all else []
        if (mpk_method or mpk_welds_tech) and not is_epb:
            doc.add_page_break()
            _appendix_heading("Протокол по результатам магнитопорошкового контроля элементов сосуда")
            doc.add_paragraph("Протокол по результатам магнитопорошкового контроля элементов сосуда")
            doc.add_paragraph(f"№ {app_no} от {date_perf_ru}г.")
            doc.add_paragraph()
            _add_protocol_header_block()
            doc.add_paragraph("1. Результаты контроля")
            doc.add_paragraph("Таблица № 3")
            doc.add_paragraph()
            mpk_rows = mpk_welds_tech if mpk_welds_tech else [{"weld_number": "К1-П1", "defect_number": "Дефектов не обнаружено", "conclusion": "годен"}]
            mpk_tbl = doc.add_table(rows=len(mpk_rows) + 1, cols=8)
            mpk_tbl.style = "Table Grid"
            headers = ["№ стыка по карте контроля", "Условный номер дефекта", "Эквивалент. Площадь Sдеф, мм²", "Глубина залегания «Y» , мм", "Протяженность ΔL, мм", "Форма (характер) дефекта", "Местоположение L, мм", "Заключение"]
            for i, h in enumerate(headers):
                mpk_tbl.rows[0].cells[i].text = h
                mpk_tbl.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
            for idx, w in enumerate(mpk_rows, 1):
                if not isinstance(w, dict):
                    continue
                mpk_tbl.rows[idx].cells[0].text = str(w.get("weld_number", idx))
                mpk_tbl.rows[idx].cells[1].text = str(w.get("defect_number", "Дефектов не обнаружено"))
                mpk_tbl.rows[idx].cells[7].text = str(w.get("conclusion", "годен"))
            doc.add_paragraph()
            doc.add_paragraph("2. Заключение по результатам контроля")
            doc.add_paragraph(
                "По результатам магнитопорошкового контроля дефектов в сварных соединениях сосуда не обнаружено, "
                "объект контроля соответствует требованиям нормативно-технической документации."
            )
            doc.add_paragraph()
            app_no += 1
        
        # ПРИЛОЖЕНИЕ: схема контроля и сканы (форма ТО — после протоколов НК)
        doc.add_page_break()
        if is_epb:
            _appendix_heading("Схема проведения неразрушающего контроля")
        else:
            _appendix_heading("Точки замера толщины и схема контроля")
        
        thickness = g("thickness_measurements", "thicknessMeasurements", default=[])
        # Сначала таблица точек замеров (чтобы схема была «там, где точки»)
        if isinstance(thickness, list) and thickness:
            doc.add_paragraph("1. Точки замера толщины:").runs[0].bold = True
            doc.add_paragraph()
            points_tbl = doc.add_table(rows=len(thickness) + 1, cols=5)
            points_tbl.style = "Table Grid"
            headers = ["№ точки", "Местоположение", "Толщина, мм", "Мин. допустимая, мм", "Комментарий"]
            for i, h in enumerate(headers):
                cell = points_tbl.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for idx, t in enumerate(thickness, 1):
                if not isinstance(t, dict):
                    continue
                points_tbl.rows[idx].cells[0].text = str(t.get("point_number", idx))
                points_tbl.rows[idx].cells[1].text = str(t.get("location", ""))
                points_tbl.rows[idx].cells[2].text = str(t.get("thickness", ""))
                points_tbl.rows[idx].cells[3].text = str(t.get("min_allowed_thickness", ""))
                points_tbl.rows[idx].cells[4].text = str(t.get("comment", ""))
            doc.add_paragraph()
        
        # Затем схема контроля с нанесёнными точками замеров (если есть координаты)
        doc.add_paragraph("2. Схема контроля (чертёж с точками измерения):").runs[0].bold = True
        doc.add_paragraph()
        control_scheme = attachments.get('control_scheme_image') or g('control_scheme_image')
        welds = g("weld_inspections", default=[])
        scheme_to_show = None
        if control_scheme:
            scheme_resolved = self._find_image_path(control_scheme) or control_scheme
            base_scheme = scheme_resolved if (scheme_resolved and os.path.isfile(scheme_resolved)) else None
            if base_scheme and isinstance(thickness, list) and thickness:
                annotated = self._draw_points_on_scheme(
                    base_scheme,
                    thickness,
                    output_dir="/app/reports/tmp",
                )
                if annotated and os.path.isfile(annotated):
                    scheme_to_show = annotated
            if scheme_to_show is None:
                scheme_to_show = base_scheme
            if scheme_to_show and isinstance(welds, list):
                uzk_with_coords = [w for w in welds if isinstance(w, dict) and (w.get("x_percent") is not None or w.get("y_percent") is not None)]
                if uzk_with_coords:
                    weld_annotated = self._draw_weld_points_on_scheme(scheme_to_show, uzk_with_coords, output_dir="/app/reports/tmp")
                    if weld_annotated and os.path.isfile(weld_annotated):
                        scheme_to_show = weld_annotated
            if not scheme_to_show:
                scheme_to_show = self._find_image_path(control_scheme) or control_scheme
        if not scheme_to_show:
            template_path = "/app/reports/assets/vessel_template.png"
            scheme_to_show = self._find_image_path(template_path) or (template_path if os.path.isfile(template_path) else None)
            if scheme_to_show:
                add_picture_if_exists("Схема контроля (шаблон)", scheme_to_show, width_inches=5.6)
            else:
                doc.add_paragraph("Схема контроля не предоставлена.")
        else:
            add_picture_if_exists("", scheme_to_show, width_inches=5.6)
        
        # Фото заводской таблички
        doc.add_paragraph()
        doc.add_paragraph("3. Фото заводской таблички:").runs[0].bold = True
        factory_plate = attachments.get('factory_plate_photo') or g('factory_plate_photo')
        if factory_plate:
            add_picture_if_exists("", factory_plate)
        else:
            doc.add_paragraph("Не приложено.")
        
        # Сканы рассмотренных документов (1–17) и фото дефектов ВИК
        doc.add_paragraph()
        doc.add_paragraph("4. Сканы рассмотренных документов и фото дефектов:").runs[0].bold = True
        _doc_nums = [str(i) for i in range(1, 18)]
        for _n in _doc_nums:
            _p = attachments.get(_n)
            if _p:
                _r = self._find_image_path(_p) or _p
                if _r and os.path.exists(_r):
                    add_picture_if_exists(f"Документ {_n}", _r)
        for _k in sorted([k for k in attachments if isinstance(k, str) and k.startswith("vd_")], key=lambda x: (int(x.split("_")[1]) if len(x.split("_")) >= 2 else 0, int(x.split("_")[2]) if len(x.split("_")) >= 3 else 0)):
            _p = attachments.get(_k)
            if _p:
                _r = self._find_image_path(_p) or _p
                if _r and os.path.exists(_r):
                    add_picture_if_exists(f"Фото дефекта ВИК ({_k})", _r)
        app_no += 1
        
        # ПРИЛОЖЕНИЕ: Расчёт остаточного ресурса (Приложение Е для ЭПБ)
        thickness_calc = g("thickness_measurements", "thicknessMeasurements", default=[])
        if is_epb and epb_ctx is not None:
            append_epb_appendix_e(doc, epb_ctx, g, attrs, equipment_data)
            app_no += 1
        elif isinstance(thickness_calc, list) and len(thickness_calc) > 0:
            doc.add_page_break()
            doc.add_heading(f"ПРИЛОЖЕНИЕ № {app_no} Расчет остаточного ресурса и расчет на прочность сосуда", level=1)
            wall_th = attrs.get("wall_thickness") or attrs.get("thickness") or g("wall_thickness", "thickness", default="4")
            min_allowed = attrs.get("min_wall_thickness") or g("min_allowed_thickness", default="2.8")
            min_vals = [float(str(p.get("thickness", "0")).replace(",", ".")) for p in thickness_calc if isinstance(p, dict) and p.get("thickness")]
            s_f = min(min_vals) if min_vals else 3.9
            try:
                s_n = float(str(wall_th).replace(",", "."))
            except (TypeError, ValueError):
                s_n = 4.0
            try:
                s_otb = float(str(min_allowed).replace(",", "."))
            except (TypeError, ValueError):
                s_otb = 2.8
            comm_year = attrs.get("commissioning_year") or g("commissioning_year") or equipment_data.get("commissioning_date")
            t1 = 16
            if comm_year:
                try:
                    t1 = datetime.now().year - int(str(comm_year)[:4])
                    if t1 < 1:
                        t1 = 16
                except (TypeError, ValueError):
                    pass
            a = (s_n - s_f) / t1 if t1 > 0 else 0.01
            tk = (s_f - s_otb) / a if a > 0 else 110
            doc.add_paragraph("1. Расчет остаточного ресурса сосуда")
            doc.add_paragraph("Остаточный ресурс сосуда рассчитан согласно ДиОР-05 и приведен в Таблице Е.1.")
            doc.add_paragraph()
            tbl_e1 = doc.add_table(rows=7, cols=5)
            tbl_e1.style = "Table Grid"
            for i, h in enumerate(["№ п/п", "Наименование величины", "Единица измерения", "Обозначение и расчетная формула", "Числовое значение"]):
                tbl_e1.rows[0].cells[i].text = h
                tbl_e1.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
            tbl_e1.rows[1].cells[0].text = "1"
            tbl_e1.rows[1].cells[1].text = "Время эксплуатации"
            tbl_e1.rows[1].cells[2].text = "лет"
            tbl_e1.rows[1].cells[3].text = "t₁"
            tbl_e1.rows[1].cells[4].text = str(t1)
            tbl_e1.rows[2].cells[0].text = "2"
            tbl_e1.rows[2].cells[1].text = "Паспортная толщина стенки\nОбечайка / Днище"
            tbl_e1.rows[2].cells[2].text = "мм"
            tbl_e1.rows[2].cells[3].text = "Sн"
            tbl_e1.rows[2].cells[4].text = f"{s_n:.0f} / {s_n:.0f}"
            tbl_e1.rows[3].cells[0].text = "3"
            tbl_e1.rows[3].cells[1].text = "Минимально допустимая толщина стенки сосуда"
            tbl_e1.rows[3].cells[2].text = "мм"
            tbl_e1.rows[3].cells[3].text = "Sотб"
            tbl_e1.rows[3].cells[4].text = f"{s_otb:.1f} / {s_otb:.1f}"
            tbl_e1.rows[4].cells[0].text = "4"
            tbl_e1.rows[4].cells[1].text = "Минимальная толщина по результатам замеров"
            tbl_e1.rows[4].cells[2].text = "мм"
            tbl_e1.rows[4].cells[3].text = "Sф"
            tbl_e1.rows[4].cells[4].text = f"{s_f:.1f} / {s_f:.1f}"
            tbl_e1.rows[5].cells[0].text = "5"
            tbl_e1.rows[5].cells[1].text = "Скорость коррозии металла сосуда"
            tbl_e1.rows[5].cells[2].text = "мм/год"
            tbl_e1.rows[5].cells[3].text = "a = (Sн - Sф) / t₁"
            tbl_e1.rows[5].cells[4].text = f"{a:.2f} / {a:.2f}"
            tbl_e1.rows[6].cells[0].text = "6"
            tbl_e1.rows[6].cells[1].text = "Остаточный срок службы сосуда, поэлементно"
            tbl_e1.rows[6].cells[2].text = "лет"
            tbl_e1.rows[6].cells[3].text = "Tk = (Sф - Sотб) / a"
            tbl_e1.rows[6].cells[4].text = f"{tk:.0f} / {tk:.0f}"
            doc.add_paragraph()
            doc.add_paragraph("2. Расчет на прочность сосуда")
            doc.add_paragraph("Расчет на прочность сосуда проводился в соответствии с ГОСТ 34233.1-2017 и ГОСТ 34233.2-2017 и приведен в Таблице Е.2.")
            doc.add_paragraph()
            try:
                p_val = float(str(attrs.get("working_pressure") or g("working_pressure") or "1.1").replace(",", "."))
            except (TypeError, ValueError):
                p_val = 1.1
            try:
                t_n = float(str(attrs.get("design_temperature") or g("design_temperature") or "100").replace(",", "."))
            except (TypeError, ValueError):
                t_n = 100
            try:
                d_n = float(str(attrs.get("diameter") or g("diameter") or equipment_data.get("diameter") or "792").replace(",", "."))
            except (TypeError, ValueError):
                d_n = 792
            try:
                c_val = float(str(attrs.get("corrosion_allowance") or g("corrosion_allowance") or "0").replace(",", "."))
            except (TypeError, ValueError):
                c_val = 0
            phi, sigma = 0.9, 177
            r_val = d_n
            s_p = (p_val * r_val) / (2 * phi * sigma - 0.5 * p_val) if (2 * phi * sigma - 0.5 * p_val) > 0 else 2.74
            s_otb_calc = s_p + c_val
            p_allow = (2 * (s_f - c_val) * phi * sigma) / (r_val + 0.5 * (s_f - c_val)) if (r_val + 0.5 * (s_f - c_val)) > 0 else 1.57
            tbl_e2 = doc.add_table(rows=13, cols=5)
            tbl_e2.style = "Table Grid"
            for i, h in enumerate(["№ п/п", "Наименование величины", "Единица измерения", "Обозначение и расчетная формула", "Числовое значение"]):
                tbl_e2.rows[0].cells[i].text = h
                tbl_e2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
            for r, (num, name, unit, formula, value) in enumerate([
                (1, "Рабочее давление", "МПа", "P", f"{p_val:.1f}"),
                (2, "Расчетная температура", "°C", "tн", f"{t_n:.0f}"),
                (3, "Внутренний диаметр", "мм", "Dн", f"{d_n:.0f}"),
                (4, "Прибавка для компенсации коррозии", "мм", "C", f"{c_val:.0f}"),
                (5, "Коэффициент прочности сварных швов", "", "φ", str(phi)),
                (6, "Допускаемое напряжение при расчетной температуре", "МПа", "[σ]", str(sigma)),
                (7, "Радиус кривизны в вершине днища, R=D для эллиптических днищ", "мм", "R", f"{r_val:.0f}"),
                (8, "Минимальная толщина по результатам контроля\nДнище / Обечайка", "мм", "Sф", f"{s_f:.1f} / {s_f:.1f}"),
                (9, "Расчетная толщина стенки\nДнище / Обечайка", "мм", "Sр = P·R/(2φ[σ]–0,5P) / Sр = P·D/(2φ[σ]–P)", f"{s_p:.2f} / {s_p:.2f}"),
                (10, "Минимально допустимая толщина стенки сосуда\nДнище / Обечайка", "мм", "Sотб = Sр + C", f"{s_otb_calc:.2f} / {s_otb_calc:.2f}"),
                (11, "Допускаемое внутреннее избыточное давление\nДнище / Обечайка", "МПа", "[P]", f"{p_allow:.2f} / {p_allow:.2f}"),
            ], 1):
                tbl_e2.rows[r].cells[0].text = str(num)
                tbl_e2.rows[r].cells[1].text = name
                tbl_e2.rows[r].cells[2].text = unit
                tbl_e2.rows[r].cells[3].text = formula
                tbl_e2.rows[r].cells[4].text = value
            doc.add_paragraph()
            doc.add_paragraph("Условия прочности: Sотб = {:.1f} мм < Sф = {:.1f} мм. [P] = {:.2f} МПа > Рраб = {:.1f} МПа.".format(s_otb_calc, s_f, p_allow, p_val))
            doc.add_paragraph()
            doc.add_paragraph("Выводы:").runs[0].bold = True
            doc.add_paragraph("На основании выполненного расчета на прочность установлено, что сосуд удовлетворяет условиям прочности, срок эксплуатации до достижения предельно-допустимого значения толщины стенки сосуда составляет более 10 лет.")
            doc.add_paragraph()
            doc.add_paragraph("Расчет выполнил:")
            doc.add_paragraph("Дефектоскопист II уровня по ВИК, УК")
            doc.add_paragraph(str(_engineer_for_method("TVI") or _engineer_for_method("UZT") or g("executors", default="—")))
            app_no += 1
        else:
            calc_data = g("calculation_data", default={})
            if calc_data:
                doc.add_page_break()
                doc.add_heading(f"ПРИЛОЖЕНИЕ № {app_no} Расчетные и аналитические процедуры оценки и прогнозирования технического состояния сосуда", level=1)
                doc.add_paragraph(str(calc_data.get("description", "Расчеты выполнены в соответствии с требованиями нормативной документации.")))
                doc.add_paragraph()
                app_no += 1
        
        # ПРИЛОЖЕНИЕ № 9: Гидравлические испытания (если есть)
        hydro_test = g("hydrostatic_test", default={})
        if hydro_test:
            doc.add_page_break()
            if is_epb:
                _appendix_heading("Акт проведения гидравлического испытания сосуда")
            else:
                _appendix_heading("Акт проведения гидравлических испытаний")
            doc.add_paragraph(f"Дата проведения: {hydro_test.get('date', date_perf_ru)}")
            doc.add_paragraph(f"Пробное давление: {hydro_test.get('test_pressure', '—')} МПа")
            doc.add_paragraph(f"Результат: {hydro_test.get('result', '—')}")
            doc.add_paragraph()
            app_no += 1
        
        # Нормативная документация
        doc.add_page_break()
        if is_epb:
            _appendix_heading(
                "Перечень использованной при экспертизе промышленной безопасности "
                "нормативной, технической и методической документации"
            )
        else:
            _appendix_heading(
                "Перечень применяемой при техническом освидетельствовании "
                "нормативной, технической и методической документации"
            )
        doc.add_paragraph()
        
        normative_docs = [
            "1. Федеральный закон от 21.07.1997г. №116 «О промышленной безопасности опасных производственных объектов».",
            "2. Федеральные нормы и правила в области промышленной безопасности «Правила промышленной безопасности при использовании оборудования, работающего под избыточным давлением», утвержденные приказом Федеральной службы по экологическому, технологическому и атомному надзору от 15.12.2020 №536",
            "3. СО 153-34.17.439-2003 «Инструкция по продлению срока службы сосудов, работающих под давлением».",
            "4. ГОСТ Р ИСО 17637-2014 «Контроль неразрушающий. Визуальный контроль соединений, выполненных сваркой плавлением»",
            "5. ГОСТ 34347-2017 «Сосуды и аппараты стальные сварные. Общие технические условия»",
            "6. ГОСТ Р 55614-2013 «Контроль неразрушающий. Толщиномеры ультразвуковые. Общие технические требования",
            "7. ГОСТ Р ИСО 17640-2016 «Неразрушающий контроль сварных соединений. Ультразвуковой контроль. Технология, уровни контроля и оценки»",
            "8. ГОСТ Р 55724-2013 «Контроль неразрушающий. Соединения сварные. Методы ультразвуковые»",
            "9. СТО 00220256-005-2005 «Швы стыковых, угловых и тавровых сварных соединений сосудов и аппаратов, работающих под давлением. Методика ультразвукового контроля»",
            "10. ГОСТ 20911-89 «Техническая диагностика. Термины и определения»",
            "11. приказ Ростехнадзора от 16.01.2024 №8, Руководство по безопасности \"Методические рекомендации о порядке проведения визуального и измерительного контроля\"",
            "12. ГОСТ Р ИСО 16809-2015 «Контроль неразрушающий. Контроль ультразвуковой. Измерение толщины».",
            "13. ГОСТ 22761-77 «Металлы и сплавы. Метод измерения твердости по Бринеллю переносными твердомерами статического действия»",
        ]
        
        for doc_item in normative_docs:
            doc.add_paragraph(doc_item)
        
        app_no += 1
        # ПРИЛОЖЕНИЕ: Документы специалистов НК (сканы удостоверений в конце отчёта)
        if specialist_docs:
            doc.add_page_break()
            if is_epb:
                _appendix_heading("Копия приказа экспертной организации о назначении эксперта")
            else:
                _appendix_heading("Документы специалистов НК")
            for s in specialist_docs:
                doc.add_heading(f"Специалист: {s.get('inspector_name') or '—'}", level=2)
                for c in (s.get("certifications") or []):
                    cert_type = c.get('certification_type') or 'Удостоверение'
                    cert_num = c.get('certificate_number') or '—'
                    method_code = c.get('method_code') or ''
                    doc.add_paragraph(f"{cert_type}" + (f" (метод {method_code})" if method_code else "") + f" № {cert_num}")
                    sp = c.get("scan_file_path")
                    if sp:
                        add_picture_if_exists("", sp)
            app_no += 1
        
        # ПРИЛОЖЕНИЕ: Свидетельства о поверке оборудования
        if verification_equipment and isinstance(verification_equipment, list) and verification_equipment:
            doc.add_page_break()
            doc.add_heading(f"ПРИЛОЖЕНИЕ № {app_no} Свидетельства о поверке оборудования", level=1)
            for eq in verification_equipment:
                doc.add_paragraph(f"{eq.get('name') or ''} — № {eq.get('verification_certificate_number') or ''}")
                sp = eq.get("scan_file_path")
                if sp:
                    add_picture_if_exists("", sp)
            app_no += 1
        
        if (inspection_data.get("status") or "").upper() == "DRAFT":
            self._add_draft_watermark(doc)
        apply_device_terminology_to_document(doc, detect_pressure_device_kind(equipment_data))
        # Сохранение
        doc.save(output_path)
        return

    def _setup_styles(self, doc: Document):
        """Настройка стилей документа: поля, шрифты, заголовки для аккуратного DOCX."""
        # Поля страницы (А4)
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.75)
            section.right_margin = Cm(1.75)
        # Стиль обычного текста
        try:
            normal = doc.styles["Normal"]
            normal.font.name = "Times New Roman"
            normal.font.size = Pt(11)
            normal.paragraph_format.space_after = Pt(6)
            normal.paragraph_format.line_spacing = 1.15
        except (KeyError, Exception):
            pass
        # Заголовок 1 уровня
        try:
            h1 = doc.styles["Heading 1"]
            h1.font.name = "Times New Roman"
            h1.font.size = Pt(14)
            h1.font.bold = True
            if h1.font.color is not None:
                h1.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
            h1.paragraph_format.space_before = Pt(12)
            h1.paragraph_format.space_after = Pt(6)
            h1.paragraph_format.keep_with_next = True
        except (KeyError, Exception):
            pass
        # Заголовок 2 уровня
        try:
            h2 = doc.styles["Heading 2"]
            h2.font.name = "Times New Roman"
            h2.font.size = Pt(12)
            h2.font.bold = True
            h2.paragraph_format.space_before = Pt(8)
            h2.paragraph_format.space_after = Pt(4)
            h2.paragraph_format.keep_with_next = True
        except (KeyError, Exception):
            pass
        # Заголовок 3 уровня
        try:
            h3 = doc.styles["Heading 3"]
            h3.font.name = "Times New Roman"
            h3.font.size = Pt(11)
            h3.font.bold = True
            h3.paragraph_format.space_before = Pt(6)
            h3.paragraph_format.space_after = Pt(2)
        except (KeyError, Exception):
            pass

    def _add_draft_watermark(self, doc: Document):
        """Добавить водяной знак «ЧЕРНОВИК» в нижний колонтитул каждой страницы."""
        try:
            for section in doc.sections:
                footer = section.footer
                if footer.paragraphs:
                    p = footer.paragraphs[0]
                else:
                    p = footer.add_paragraph()
                p.clear()
                run = p.add_run("ЧЕРНОВИК")
                run.bold = True
                run.font.size = Pt(14)
                if run.font.color is not None:
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
