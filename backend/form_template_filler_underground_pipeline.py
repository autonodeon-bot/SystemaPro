"""
Заполнение формы ТО to-33 «Обследование подземных трубопроводов» (полная глубина).
"""
from __future__ import annotations

import logging
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document

from form_media_helpers import (
    build_attachments_map,
    collect_photo_paths,
    collect_scheme_paths,
    insert_media_block,
)
from form_template_filler import (
    MISSING,
    _fill_signatures,
    _fmt_date_ru,
    _replace_underscores,
    _set,
    _set_paragraph_text,
    _extract_specialists,
    finalize_official_form,
)
from form_template_filler_common import (
    fill_contractor_table,
    fill_customer_table,
    fill_instruments_table,
    fill_kv_table,
    fill_specialists_table,
    g_data,
)
from form_template_filler_ndt_shared import (
    fill_documents_by_name,
    fill_ehz_points,
    fill_ehz_station,
    fill_geometry_table,
    fill_hardness_pipeline_table,
    fill_mpk_table,
    fill_pipeline_calc,
    fill_pipeline_life,
    fill_protocol_header_ids,
    fill_uzt_wide_table,
    fill_vik_element_table,
    fill_vik_weld_table,
    fill_vtk_table,
)
from report_forms_registry import resolve_form_path
from report_org_settings import load_report_org_settings

logger = logging.getLogger(__name__)


def fill_underground_pipeline_form_to33(
    inspection_data: Dict[str, Any],
    equipment_data: Dict[str, Any],
    output_path: str,
    verification_equipment: Optional[List[Dict[str, Any]]] = None,
    org_settings: Optional[Dict[str, Any]] = None,
    specialist_docs: Optional[List[Dict[str, Any]]] = None,
    document_files: Optional[List[Dict[str, Any]]] = None,
    find_image: Optional[Any] = None,
    ndt_methods: Optional[List[Dict[str, Any]]] = None,
) -> str:
    template = resolve_form_path("to-33")
    if template is None or not template.exists():
        raise FileNotFoundError("Шаблон to-33 не найден в backend/report_forms/")

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, out)
    doc = Document(str(out))

    if org_settings is None:
        org_settings = load_report_org_settings()

    data = inspection_data.get("data") or {}
    if not isinstance(data, dict):
        data = {}
    attrs = equipment_data.get("attributes") or {}
    if not isinstance(attrs, dict):
        attrs = {}

    def g(*keys: str, default: Any = MISSING) -> Any:
        return g_data(data, attrs, *keys, default=default)

    date_ru = _fmt_date_ru(inspection_data.get("date_performed")) or datetime.now().strftime(
        "%d.%m.%Y"
    )
    device = str(
        g(
            "vessel_name",
            "pipeline_name",
            "equipment_device_name",
            default=equipment_data.get("name") or MISSING,
        )
    )
    serial = str(g("serial_number", default=equipment_data.get("serial_number") or MISSING))
    reg_no = str(g("reg_number", default=MISSING))
    inv_no = str(g("inventory_number", "inv_number", default=MISSING))
    location = str(
        g("location", "equipment_location", default=equipment_data.get("location") or MISSING)
    )
    pressure = str(g("working_pressure", "design_pressure", default=MISSING))
    diameter = str(g("diameter", "pipe_diameter", "nominal_diameter", default=MISSING))
    thickness = str(g("wall_thickness", "thickness", default=MISSING))
    medium = str(g("working_medium", "medium", default=MISSING))
    purpose = str(g("purpose", "pipeline_purpose", default=MISSING))
    category = str(g("pipeline_category", "category", default=MISSING))
    length = str(g("pipeline_length", "length", "total_length", default=MISSING))
    material = str(g("shell_material", "material", "pipe_material", default=MISSING))
    temp = str(g("working_temperature", "working_medium_temperature", default=MISSING))
    year = str(g("commissioning_year", "manufacture_year", default=MISSING))

    cust = data.get("customer_info") if isinstance(data.get("customer_info"), dict) else {}
    org_name = str(
        cust.get("legal_name")
        or g(
            "organization",
            "customer_name",
            default=(org_settings.get("customer") or {}).get("legal_name") or MISSING,
        )
    )
    contractor = (org_settings.get("contractor") or {})
    contractor_name = contractor.get("legal_name") or contractor.get("name") or ""

    tables = doc.tables
    ve = verification_equipment or []

    # --- Основной отчёт ---
    if len(tables) > 1:
        fill_customer_table(
            tables[1],
            {"customer": {**(org_settings.get("customer") or {}), **cust}},
            location if location != MISSING else "",
        )
    if len(tables) > 2:
        fill_contractor_table(tables[2], org_settings)
    if len(tables) > 3:
        fill_specialists_table(tables[3], data, specialist_docs or [])
    if len(tables) > 4:
        fill_instruments_table(tables[4], ve, data, ndt_methods)
    if len(tables) > 5:
        fill_kv_table(tables[5], [device, inv_no, location, reg_no, org_name, serial])
    if len(tables) > 6:
        fill_kv_table(
            tables[6],
            [device, purpose, length, category, pressure, diameter, thickness, medium, material, year, temp, location],
        )
    if len(tables) > 8:
        fill_documents_by_name(tables[8], data)
    if len(tables) > 9:
        _fill_results_summary(tables[9], data)

    # --- Прил. Б паспорт ---
    if len(tables) > 11:
        fill_documents_by_name(tables[11], data)
    if len(tables) > 12:
        fill_documents_by_name(tables[12], data)
    if len(tables) > 13:
        fill_kv_table(tables[13], [device, purpose, length, category, year, medium, material])
    if len(tables) > 14:
        for i, v in enumerate([None, diameter, pressure, temp, thickness, material, year]):
            if v is not None and i < len(tables[14].rows) and len(tables[14].rows[i].cells) > 1:
                _set(tables[14], i, 1, v)
    if len(tables) > 15:
        # результаты анализа документации
        conclusion_doc = str(g("documentation_conclusion", default="документация в полном объёме") or "")
        if len(tables[15].rows) > 1 and len(tables[15].rows[1].cells) > 2:
            _set(tables[15], 1, 2, conclusion_doc)

    # --- Прил. В ВИК ---
    for ti in (18, 24, 29, 36, 43, 48, 53, 61):
        if len(tables) > ti:
            fill_protocol_header_ids(
                tables[ti],
                contractor=contractor_name,
                customer=org_name,
                device=device,
                serial=serial,
                reg_no=reg_no,
                inv_no=inv_no,
                location=location if location != MISSING else "",
            )
    if len(tables) > 19:
        fill_instruments_table(tables[19], ve, data, ndt_methods)
    if len(tables) > 20:
        fill_vik_element_table(tables[20], data)
    if len(tables) > 21:
        fill_vik_weld_table(tables[21], data)

    # --- Прил. Г УЗТ ---
    if len(tables) > 25:
        fill_instruments_table(tables[25], ve, data, ndt_methods)
    if len(tables) > 26:
        fill_uzt_wide_table(tables[26], data)

    # --- Прил. Д МПК ---
    if len(tables) > 30:
        fill_instruments_table(tables[30], ve, data, ndt_methods)
    if len(tables) > 33:
        fill_mpk_table(tables[33], data)

    # --- Прил. Е ВТК ---
    if len(tables) > 37:
        fill_instruments_table(tables[37], ve, data, ndt_methods)
    if len(tables) > 40:
        fill_vtk_table(tables[40], data)

    # --- Прил. Ж твёрдость ---
    if len(tables) > 44:
        fill_instruments_table(tables[44], ve, data, ndt_methods)
    if len(tables) > 45:
        fill_hardness_pipeline_table(tables[45], data)

    # --- Прил. З геометрия ---
    if len(tables) > 50:
        fill_geometry_table(tables[50], data)

    # --- Прил. И ЭХЗ ---
    if len(tables) > 54:
        fill_instruments_table(tables[54], ve, data, ndt_methods)
    if len(tables) > 55:
        fill_ehz_station(tables[55], data)
    if len(tables) > 56:
        fill_ehz_points(tables[56], data)

    # --- Прил. К расчёт ---
    if len(tables) > 62:
        fill_pipeline_calc(tables[62], data)
    if len(tables) > 65:
        fill_pipeline_life(tables[65], data)

    _fill_paragraphs(doc, date_ru, device, serial, reg_no, inv_no, org_name, g)

    attachments = build_attachments_map(document_files)
    schemes = collect_scheme_paths(data, attachments)
    photos = collect_photo_paths(data, attachments)
    for title in (
        "Схемы с указанием номеров сварных соединений",
        "Схема контроля",
        "Схема визуального",
        "Схема магнитопорошкового",
        "Схема вихретокового",
    ):
        insert_media_block(doc, title, schemes, find_image=find_image, max_items=6)
    insert_media_block(doc, "Результаты контроля", photos, find_image=find_image, max_items=12)
    # Прил. М — акт шурфовки
    shurf = []
    ad = data.get("additional_data") if isinstance(data.get("additional_data"), dict) else {}
    for key in ("shurf_act_path", "shurf_act_scan", "pit_act_path"):
        p = ad.get(key) or data.get(key)
        if p:
            shurf.append(str(p))
    if shurf:
        insert_media_block(doc, "Акт", shurf, find_image=find_image, max_items=4)
        insert_media_block(doc, "шурф", shurf, find_image=find_image, max_items=4)

    specs = _extract_specialists(data, specialist_docs or [])
    for t in tables:
        if t.rows and "Контроль провел" in (t.rows[0].cells[0].text or ""):
            _fill_signatures(t, {"specialists": specs})

    finalize_official_form(doc, "to-33")

    doc.save(str(out))
    logger.info("Форма to-33 заполнена (full): %s", out)
    return str(out)


def _fill_results_summary(table, data: Dict[str, Any]) -> None:
    has_uzt = bool(data.get("thickness_measurements") or data.get("uzt_schemes"))
    has_vik = bool(data.get("visual_defects") or data.get("defects"))
    welds = data.get("weld_inspections") or []
    has_mpk = any(
        isinstance(w, dict)
        and ("MPK" in str(w.get("control_method") or "").upper() or "МПК" in str(w.get("control_method") or "").upper())
        for w in welds
    ) if isinstance(welds, list) else False
    has_uzk = any(
        isinstance(w, dict)
        and ("UZK" in str(w.get("control_method") or "").upper() or "УЗК" in str(w.get("control_method") or "").upper())
        for w in welds
    ) if isinstance(welds, list) else False
    has_hard = bool(data.get("hardness_tests"))
    has_ehz = bool((data.get("additional_data") or {}).get("ehz_points") if isinstance(data.get("additional_data"), dict) else data.get("ehz_points"))
    hints = {
        1: "выполнено",
        2: "см. приложение" if has_vik else "дефектов не выявлено",
        3: "см. приложение" if has_uzt else "не выполнялось",
        4: "см. приложение" if has_mpk else "не выполнялось",
        5: "см. приложение" if has_uzk else "не выполнялось",
        6: "см. приложение" if has_hard else "не выполнялось",
        7: "см. приложение" if has_ehz else "не выполнялось",
        8: "см. приложение",
    }
    for r, text in hints.items():
        if r < len(table.rows) and len(table.rows[r].cells) > 2 and text:
            if not (table.rows[r].cells[2].text or "").strip():
                _set(table, r, 2, text)


def _fill_paragraphs(doc, date_ru, device, serial, reg_no, inv_no, org_name, g) -> None:
    contract = str(g("contract_number", default="") or "")
    contract_date = _fmt_date_ru(g("contract_date", default="")) or ""
    period_from = _fmt_date_ru(g("work_period_from", "date_from", default="")) or ""
    period_to = _fmt_date_ru(g("work_period_to", "date_to", default="")) or date_ru
    state = str(g("technical_state", default="работоспособном") or "работоспособном")

    for p in doc.paragraphs:
        text = p.text or ""
        if not text.strip():
            continue
        new_text = text
        if "договором между" in text.lower():
            new_text = _replace_underscores(
                text, [org_name, contract_date or "____", contract or "____"]
            )
        elif "проведены в период" in text.lower():
            pf = period_from or "__.__.____"
            pt = period_to or "__.__.____"
            new_text = f"Работы по техническому диагностированию проведены в период с {pf} по {pt}."
        elif "зав." in text.lower() and ("рег." in text.lower() or "инв." in text.lower()):
            new_text = _replace_underscores(text, [device, serial, reg_no, inv_no])
        elif "находится в" in text.lower() and "состоян" in text.lower():
            if "____" in text or "______" in text:
                new_text = _replace_underscores(text, [state])
        else:
            continue
        if new_text != text:
            _set_paragraph_text(p, new_text)
