"""
Заполнение формы ТО to-3 «Обследование грузоподъемных механизмов» (полная глубина).
"""
from __future__ import annotations

import logging
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document

from form_media_helpers import (
    build_attachments_map,
    collect_photo_paths,
    collect_scheme_paths,
    insert_media_block,
)
from form_template_filler import (
    MISSING,
    _extract_specialists,
    _fill_signatures,
    _fmt_date_ru,
    _replace_underscores,
    _set,
    _set_paragraph_text,
    finalize_official_form,
)
from form_template_filler_common import (
    fill_contractor_table,
    fill_customer_table,
    fill_instruments_table,
    fill_kv_table,
    fill_specialists_table,
    g_data,
)
from form_template_filler_ndt_shared import (
    fill_crane_safety_checklists,
    fill_crane_safety_devices,
    fill_crane_uzk_table,
    fill_crane_vik_zones,
    fill_documents_by_name,
    fill_protocol_header_ids,
    fill_uzt_crane_table,
)
from report_forms_registry import resolve_form_path
from report_org_settings import load_report_org_settings

logger = logging.getLogger(__name__)


def fill_crane_form_to3(
    inspection_data: Dict[str, Any],
    equipment_data: Dict[str, Any],
    output_path: str,
    verification_equipment: Optional[List[Dict[str, Any]]] = None,
    org_settings: Optional[Dict[str, Any]] = None,
    specialist_docs: Optional[List[Dict[str, Any]]] = None,
    document_files: Optional[List[Dict[str, Any]]] = None,
    find_image: Optional[Any] = None,
    ndt_methods: Optional[List[Dict[str, Any]]] = None,
) -> str:
    template = resolve_form_path("to-3")
    if template is None or not template.exists():
        raise FileNotFoundError("Шаблон to-3 не найден в backend/report_forms/")

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, out)
    doc = Document(str(out))

    if org_settings is None:
        org_settings = load_report_org_settings()

    data = inspection_data.get("data") or {}
    if not isinstance(data, dict):
        data = {}
    attrs = equipment_data.get("attributes") or {}
    if not isinstance(attrs, dict):
        attrs = {}
    ad = data.get("additional_data") if isinstance(data.get("additional_data"), dict) else {}

    def g(*keys: str, default: Any = MISSING) -> Any:
        return g_data(data, attrs, *keys, default=default)

    date_ru = _fmt_date_ru(inspection_data.get("date_performed")) or datetime.now().strftime(
        "%d.%m.%Y"
    )
    device = str(
        g(
            "vessel_name",
            "crane_name",
            "equipment_device_name",
            default=equipment_data.get("name") or MISSING,
        )
    )
    serial = str(g("serial_number", default=equipment_data.get("serial_number") or MISSING))
    reg_no = str(g("reg_number", "account_number", default=MISSING))
    inv_no = str(g("inventory_number", "inv_number", default=MISSING))
    location = str(
        g("location", "equipment_location", default=equipment_data.get("location") or MISSING)
    )
    crane_type = str(g("crane_type", "lifting_type", default=MISSING))
    purpose = str(g("purpose", default=MISSING))
    capacity = str(g("crane_capacity", "lifting_capacity", "capacity", default=MISSING))
    mode = str(g("crane_mode", "duty_mode", "mode", default=MISSING))
    manufacturer = str(g("manufacturer", default=MISSING))
    year = str(g("manufacture_year", "commissioning_year", default=MISSING))

    cust = data.get("customer_info") if isinstance(data.get("customer_info"), dict) else {}
    org_name = str(
        cust.get("legal_name")
        or g(
            "organization",
            "customer_name",
            default=(org_settings.get("customer") or {}).get("legal_name") or MISSING,
        )
    )
    contractor = (org_settings.get("contractor") or {})
    contractor_name = contractor.get("legal_name") or contractor.get("name") or ""

    tables = doc.tables
    ve = verification_equipment or []

    if len(tables) > 1:
        fill_customer_table(
            tables[1],
            {"customer": {**(org_settings.get("customer") or {}), **cust}},
            location if location != MISSING else "",
        )
    if len(tables) > 2:
        fill_contractor_table(tables[2], org_settings)
    if len(tables) > 3:
        fill_specialists_table(tables[3], data, specialist_docs or [])
    if len(tables) > 5:
        fill_instruments_table(tables[5], ve, data, ndt_methods)

    if len(tables) > 6:
        fill_kv_table(
            tables[6],
            [device, serial, reg_no, inv_no, location, manufacturer, year, capacity, mode, crane_type, purpose],
        )
    if len(tables) > 7:
        fill_kv_table(
            tables[7],
            [
                crane_type,
                purpose,
                g("crane_execution", "execution", default=MISSING),
                capacity,
                mode,
                g("span", "crane_span", default=MISSING),
                g("lift_height", default=MISSING),
                g("hook_speed", default=MISSING),
                year,
                manufacturer,
                location,
                g("environment", "operating_environment", default=MISSING),
                g("voltage", default=MISSING),
                g("drive_type", default=MISSING),
                serial,
                reg_no,
                inv_no,
            ],
        )

    if len(tables) > 10:
        fill_documents_by_name(tables[10], data)
    if len(tables) > 14:
        fill_documents_by_name(tables[14], data, num_col=2, pages_col=4)

    # Прил.1 анализ
    if len(tables) > 15:
        remarks = str(g("supervisory_remarks", default="Нет") or "Нет")
        accidents = str(g("accidents_info", default="Нет") or "Нет")
        repair = str(g("repair_info", default="Нет") or "Нет")
        for r, val in ((1, remarks), (2, accidents), (3, repair)):
            if r < len(tables[15].rows) and len(tables[15].rows[r].cells) > 2:
                _set(tables[15], r, 2, val)

    # Прил.2 акт ПС
    for ti in (13, 18, 28, 31, 36, 41, 47):
        if len(tables) > ti:
            fill_protocol_header_ids(
                tables[ti],
                contractor=contractor_name,
                customer=org_name,
                device=device,
                serial=serial,
                reg_no=reg_no,
                inv_no=inv_no,
                location=location if location != MISSING else "",
            )

    if len(tables) > 19:
        # Тип ПС | значение уже в шаблоне — перезапишем
        pairs = [
            (0, crane_type),
            (1, manufacturer),
            (2, device),
            (3, serial),
            (4, reg_no),
            (5, location),
            (6, inv_no),
            (7, year),
        ]
        for r, v in pairs:
            if r < len(tables[19].rows) and len(tables[19].rows[r].cells) > 1 and v != MISSING:
                _set(tables[19], r, 1, v)

    if len(tables) > 21:
        pairs21 = [
            (0, capacity),
            (1, year),
            (2, mode),
            (3, g("wind_region", default="")),
            (4, g("temp_limits", default="")),
            (5, g("fire_hazard_ok", default="")),
            (6, g("explosion_hazard_ok", default="")),
            (7, g("aggressive_env", default="")),
        ]
        for r, v in pairs21:
            if r < len(tables[21].rows) and v not in (None, "", MISSING):
                col = 1 if len(tables[21].rows[r].cells) > 1 else 0
                if len(tables[21].rows[r].cells) > 1:
                    _set(tables[21], r, 1, v)

    act = ad.get("crane_act") if isinstance(ad.get("crane_act"), dict) else {}
    if not act:
        act = data.get("crane_act") if isinstance(data.get("crane_act"), dict) else {}
    overall = act.get("overall_state") or g("technical_state", default="работоспособное")
    if len(tables) > 23:
        _set(tables[23], 0, 2, overall if overall != MISSING else "работоспособное")
        if act.get("classification_limit"):
            _set(tables[23], 1, 2, act.get("classification_limit"))
        if act.get("residual_score"):
            _set(tables[23], 2, 2, act.get("residual_score"))
        defects_n = len(data.get("visual_defects") or [])
        _set(tables[23], 3, 2, str(act.get("defects_total") or defects_n))
        if act.get("defects_fixed"):
            _set(tables[23], 4, 2, act.get("defects_fixed"))
        if act.get("defects_before_use"):
            _set(tables[23], 5, 2, act.get("defects_before_use"))
    if len(tables) > 24:
        if act.get("allowed_until"):
            _set(tables[24], 0, 1, act.get("allowed_until"))
        if act.get("repair_required"):
            _set(tables[24], 1, 1, act.get("repair_required"))
    if len(tables) > 25 and act.get("recommendations"):
        _set(tables[25], 0, 1, act.get("recommendations"))

    # Прил.3 ВИК
    if len(tables) > 32:
        fill_instruments_table(tables[32], ve, data, ndt_methods)
    if len(tables) > 34:
        fill_crane_vik_zones(tables[34], data)

    # Прил.4 УЗТ
    if len(tables) > 37:
        fill_instruments_table(tables[37], ve, data, ndt_methods)
    if len(tables) > 38:
        fill_uzt_crane_table(tables[38], data)

    # Прил.5 УЗК
    if len(tables) > 42:
        fill_instruments_table(tables[42], ve, data, ndt_methods)
    if len(tables) > 44:
        fill_crane_uzk_table(tables[44], data)

    # Прил.6 безопасность
    if len(tables) > 52:
        fill_crane_safety_checklists(tables[48:53], data)
    if len(tables) > 53:
        fill_crane_safety_devices(tables[53], data)

    # Прил.7 баллы — если переданы
    scores = ad.get("crane_defect_scores") or data.get("crane_defect_scores")
    if isinstance(scores, dict) and len(tables) > 56:
        for r in range(3, len(tables[56].rows)):
            label = (tables[56].rows[r].cells[0].text or "").strip()
            if label and label in scores and len(tables[56].rows[r].cells) > 4:
                _set(tables[56], r, 4, scores[label])

    _fill_paragraphs(doc, date_ru, device, serial, reg_no, inv_no, org_name, g, overall)

    attachments = build_attachments_map(document_files)
    schemes = collect_scheme_paths(data, attachments)
    photos = collect_photo_paths(data, attachments)
    insert_media_block(doc, "Схема контроля", schemes, find_image=find_image, max_items=8)
    insert_media_block(doc, "Результаты контроля", photos, find_image=find_image, max_items=12)
    for key in ("static_dynamic_test_act", "work_character_certificate"):
        p = ad.get(key) or data.get(key)
        if p:
            insert_media_block(doc, "Акт", [str(p)], find_image=find_image, max_items=2)

    specs = _extract_specialists(data, specialist_docs or [])
    for t in tables:
        if t.rows and "Контроль провел" in (t.rows[0].cells[0].text or ""):
            _fill_signatures(t, {"specialists": specs})

    finalize_official_form(doc, "to-3")

    doc.save(str(out))
    logger.info("Форма to-3 заполнена (full): %s", out)
    return str(out)


def _fill_paragraphs(doc, date_ru, device, serial, reg_no, inv_no, org_name, g, overall) -> None:
    contract = str(g("contract_number", default="") or "")
    contract_date = _fmt_date_ru(g("contract_date", default="")) or ""
    period_from = _fmt_date_ru(g("work_period_from", default="")) or ""
    period_to = _fmt_date_ru(g("work_period_to", default="")) or date_ru

    for p in doc.paragraphs:
        text = p.text or ""
        if not text.strip():
            continue
        new_text = text
        if "договором между" in text.lower():
            new_text = _replace_underscores(
                text, [org_name, contract_date or "____", contract or "____"]
            )
        elif "проведены в период" in text.lower():
            pf = period_from or "__.__.____"
            pt = period_to or "__.__.____"
            new_text = f"Работы по обследованию грузоподъемных механизмов проведены в период с {pf} по {pt}."
        elif "зав." in text.lower() and ("инв" in text.lower() or "учет" in text.lower()):
            new_text = _replace_underscores(text, [device, serial, inv_no or reg_no])
        elif "находится в" in text.lower() and "состоянии" in text.lower():
            if "____" in text or "______" in text:
                new_text = _replace_underscores(text, [overall or "работоспособном"])
        else:
            continue
        if new_text != text:
            _set_paragraph_text(p, new_text)
