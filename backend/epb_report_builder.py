"""
Сборка основного текста заключения ЭПБ (структура по образцу 25-3173).

Разделы 1–9 + оглавление. Приложения А–З формируются в word_generator.py.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from typing import Any, Callable, Dict, List, Optional, Set

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from equipment_presets import preset_from_equipment_data

EPB_APPENDIX_LETTERS = "АБВГДЕЖЗИКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ"

EPB_TOC_ITEMS: List[str] = [
    "1. Вводная часть",
    (
        "2. Наименование объекта экспертизы, на которое распространяется "
        "действие заключения экспертизы промышленной безопасности"
    ),
    "3. Сведения о заказчике",
    "4. Цель экспертизы",
    "5. Сведения о документах, рассмотренных в процессе экспертизы",
    "6. Назначение и краткая характеристика объекта экспертизы",
    "7. Результаты проведённой экспертизы",
    "8. Выводы заключения экспертизы",
    "9. Регламент контроля технического состояния",
    "Приложение А Акт о проведении работ по техническому диагностированию",
    "Приложение Б Отчет по анализу технической документации",
    "Приложение В Результаты технического диагностирования",
    "Приложение Г Схема проведения неразрушающего контроля",
    "Приложение Д Акт проведения гидравлического испытания сосуда",
    (
        "Приложение Е Расчетные и аналитические процедуры оценки "
        "и прогнозирования технического состояния сосуда"
    ),
    "Приложение Ж Копия приказа экспертной организации о назначении эксперта",
    (
        "Приложение З Перечень использованной при экспертизе промышленной "
        "безопасности нормативной, технической и методической документации"
    ),
]

EPB_NORMATIVE_BULLETS = [
    (
        "Федеральный закон «О промышленной безопасности опасных производственных "
        "объектов», ст.9;"
    ),
    (
        "Федеральные нормы и правила в области промышленной безопасности "
        "«Правила промышленной безопасности при использовании оборудования, "
        "работающего под избыточным давлением», утвержденные приказом "
        "Ростехнадзора от 15.12.2020 №536, п.65, п.68, п.69, п.154, п.165, "
        "п.188, п.212, п.215–п.217, п.228, п.236, п.333, п.335, п.338–п.340, "
        "п.346, п.348, п.350, п.360, п.394, п.423, п.424, п.461, п.462, п.463 "
        "(далее по тексту приказ Ростехнадзора от 15.12.2020 №536);"
    ),
    (
        "Федеральные нормы и правила «Правила безопасности в нефтяной и газовой "
        "промышленности», утвержденные приказом Федеральной службы по "
        "экологическому, технологическому и атомному надзору от 15.12.2020 №534, "
        "п.6, п.32, п.125, п.126, п.128, п.129, п.131, п.228п, п.556, п.562, "
        "п.565, п.731, п.733, п.734, п.1148 (далее по тексту приказ "
        "Ростехнадзора от 15.12.2020 №534)."
    ),
]

EPB_EXPERT_ORG_ROWS = [
    ("Наименование", "Общество с ограниченной ответственностью «ЮТАР»"),
    (
        "Юридический (почтовый) адрес",
        "628285, Ханты-Мансийский автономный округ – Югра, г. Урай, "
        "улица Ивана Шестакова, строение 46Б",
    ),
    (
        "Лицензия",
        "Регистрационный номер Л043-00109-72/00514886 на осуществление "
        "деятельности по проведению экспертизы промышленной безопасности",
    ),
    (
        "Свидетельство",
        "№ АЦЛНК-5-00083 об аттестации лаборатории неразрушающих методов "
        "контроля ООО «ЮТАР»",
    ),
]

EPB_PURPOSE_BULLETS = [
    (
        "определение соответствия объекта экспертизы промышленной безопасности "
        "предъявляемым к нему требованиям промышленной безопасности;"
    ),
    "оценка фактического состояния;",
    "определение остаточного ресурса (срока службы);",
    (
        "установление срока и условий дальнейшей безопасной эксплуатации "
        "объекта экспертизы промышленной безопасности."
    ),
]

EPB_DAMAGE_MECHANISMS = [
    "общая поверхностная коррозия и локальные коррозионные повреждения (язвенная коррозия);",
    "механические повреждения, температурная деформация;",
    (
        "развитие трещин от непроектных нагрузок (подвижки фундаментов, "
        "деформации опор, повышенный уровень вибрации трубопроводов)."
    ),
]

EPB_DOCUMENT_NAMES: Dict[str, str] = {
    "1": "Лицензия на эксплуатацию взрывопожароопасных и химически опасных производственных объектов I, II и III классов опасности",
    "2": "Свидетельство о регистрации в государственном реестре опасных производственных объектов",
    "3": "Технологический регламент объектов опасных производственных объектов",
    "4": "План мероприятий по локализации и ликвидации последствий аварий на опасном производственном объекте",
    "5": "Положение о производственном контроле за соблюдением требований промышленной безопасности на ОПО",
    "6": "Журнал учета аварий и инцидентов на ОПО",
    "7": "Страховой полис обязательного страхования гражданской ответственности владельца опасного объекта",
    "8": "Приказ об организации безопасной эксплуатации оборудования, работающего под избыточным давлением",
    "9": "Приказ о назначении ответственного лица за осуществление производственного контроля",
    "10": "Паспорт сосуда (сборочный чертёж, расчёт на прочность, схема включения, инструкция по эксплуатации)",
    "11": "Инструкция по монтажу и эксплуатации",
    "12": "Паспорта на предохранительные клапаны",
    "13": "Паспорта на запорную арматуру",
    "14": "Документация на контрольно-измерительные приборы",
    "15": "Ремонтная (исполнительная) документация",
    "16": "Заключение экспертизы промышленной безопасности",
    "17": "Акты проведения УЗТ",
}


def epb_appendix_letter(app_no: int) -> str:
    idx = app_no - 1
    if 0 <= idx < len(EPB_APPENDIX_LETTERS):
        return EPB_APPENDIX_LETTERS[idx]
    return str(app_no)


def device_type_label(preset: str) -> str:
    return {
        "oil_settler": "Отстойник нефти",
        "gas_separator": "Газосепаратор",
        "underground_tank": "Ёмкость подземная",
        "receiver": "Ресивер",
    }.get(preset, "Сосуд")


@dataclass
class EpbReportContext:
    g: Callable[..., Any]
    opo_get: Callable[..., Any]
    device_name: str
    serial: str
    reg_no: str
    org: str
    location: str
    opo_name: str
    opo_reg: str
    opo_class: str
    date_perf_ru: str
    contractor: str
    director_title: str
    director_name: str
    purpose_default: Optional[str]
    equipment_data: Dict[str, Any]
    inspection_data: Dict[str, Any]
    ndt_methods: List[Dict[str, Any]]
    performed_codes: Set[str] = field(default_factory=set)
    inspectors: List[str] = field(default_factory=list)
    inspection_engineers: List[Dict[str, Any]] = field(default_factory=list)
    docs_dict: Dict[str, Any] = field(default_factory=dict)
    docs_info: Dict[str, Any] = field(default_factory=dict)
    doc_meta_fn: Optional[Callable[[str], tuple]] = None
    scheme_index: str = ""
    construction_type: str = "горизонтальный с эллиптическими днищами"
    residual_life_years: str = "10"
    residual_life_until: str = ""
    allowed_pressure: str = "1,0"
    allowed_temperature: str = "плюс 80"


def _add_grid_table(doc: Document, rows: List[tuple[str, str]]) -> None:
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v


def _has_method(codes: Set[str], *aliases: str) -> bool:
    up = {c.upper() for c in codes}
    return any(a.upper() in up for a in aliases)


def append_epb_toc(doc: Document) -> None:
    doc.add_heading("Содержание", level=1)
    for item in EPB_TOC_ITEMS:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = None
    doc.add_page_break()


def build_epb_main_body(doc: Document, ctx: EpbReportContext) -> None:
    """Разделы 1–9 заключения ЭПБ (образец 25-3173)."""
    g = ctx.g
    preset = preset_from_equipment_data(ctx.equipment_data)
    type_label = device_type_label(preset)
    obj_phrase = (
        f"{type_label} зав.№ {ctx.serial}, рег.№ {ctx.reg_no}, "
        f"эксплуатируемый {ctx.org}, установленный на {ctx.location}"
    )
    purpose = g("purpose", "vessel_purpose", default=ctx.purpose_default or "—")
    construction = g("construction_type", default=ctx.construction_type)
    contract = g(
        "contract_number",
        "basis",
        default="договор на выполнение работ по экспертизе промышленной безопасности",
    )

    # --- 1. Вводная часть ---
    doc.add_heading("Вводная часть", level=1)
    doc.add_paragraph(
        "1.1. Положение нормативных правовых актов в области промышленной "
        "безопасности, устанавливающих требования к объекту экспертизы, и на "
        "соответствие, которым проведена оценка соответствия объекта экспертизы:"
    )
    for bullet in EPB_NORMATIVE_BULLETS:
        doc.add_paragraph(f"- {bullet}")
    doc.add_paragraph(
        f"Экспертиза промышленной безопасности проводилась в рамках {contract}, "
        f"заключённого между {ctx.org} и {ctx.contractor}."
    )
    doc.add_paragraph()

    doc.add_paragraph("1.2. Сведения об экспертной организации")
    _add_grid_table(
        doc,
        [
            (k, g(f"contractor_{k.lower().replace(' ', '_')}", default=v))
            for k, v in EPB_EXPERT_ORG_ROWS
        ],
    )
    director = g("director_name", default=ctx.director_name)
    org_rows = doc.tables[-1].rows
    if len(org_rows) >= 3:
        org_rows[2].cells[1].text = (
            f"Генеральный директор {director}" if director and director != "__________________" else director
        )
    doc.add_paragraph()

    doc.add_paragraph("1.3. Сведения об экспертах и специалистах по неразрушающему контролю")
    order_no = g("expert_order_number", default="250")
    order_date = g("expert_order_date", default="08.09.2025")
    doc.add_paragraph(
        f"Для проведения экспертизы промышленной безопасности сосуда, работающего "
        f"под давлением, приказом генерального директора № {order_no} от {order_date}г., "
        "была назначена экспертная группа:"
    )
    if ctx.inspection_engineers:
        for ie in ctx.inspection_engineers:
            if not isinstance(ie, dict):
                continue
            name = ie.get("full_name") or ""
            cert = ie.get("certificate_number") or ""
            valid = ie.get("valid_until") or ""
            role = ie.get("role") or ie.get("expert_area") or "эксперт"
            if name:
                line = f"{name} – {role}"
                if cert:
                    line += f", удостоверение № {cert}"
                if valid:
                    line += f" срок аттестации до {valid}"
                doc.add_paragraph(line)
    elif ctx.inspectors:
        for name in ctx.inspectors:
            doc.add_paragraph(f"{name} – эксперт в области промышленной безопасности.")
    else:
        doc.add_paragraph("Состав экспертной группы – по приказу экспертной организации (приложение Ж).")
    doc.add_paragraph(
        "К проведению технического диагностирования допускались специалисты, "
        "имеющие квалификационные удостоверения по неразрушающим методам контроля."
    )
    doc.add_paragraph()

    # --- 2. Объект экспертизы ---
    doc.add_heading(
        "Наименование объекта экспертизы, на которое распространяется действие "
        "заключения экспертизы промышленной безопасности",
        level=1,
    )
    doc.add_paragraph(
        f"Объектом экспертизы промышленной безопасности является сосуд, "
        f"работающий под давлением – {obj_phrase}."
    )
    doc.add_paragraph()

    # --- 3. Заказчик ---
    doc.add_heading("Данные о заказчике", level=1)
    customer_rows = [
        ("Наименование организации", g("customer_legal_name", default=ctx.org)),
        ("Организационно-правовая форма", g("customer_legal_form", default="—")),
        ("Место нахождения", g("customer_address", default=ctx.location)),
        ("Телефон / факс", g("customer_phone", default="—")),
        ("Руководитель", g("customer_director", default="—")),
        ("Структурное подразделение", g("customer_department", default="—")),
        ("Место нахождения подразделения", g("customer_department_address", default=ctx.location)),
        ("Телефон / факс подразделения", g("customer_department_phone", default="—")),
        ("Руководитель подразделения", g("customer_department_head", default="—")),
    ]
    _add_grid_table(doc, customer_rows)
    doc.add_paragraph()

    # --- 4. Цель ---
    doc.add_heading("Цель экспертизы", level=1)
    doc.add_paragraph("Целью проведения экспертизы промышленной безопасности является:")
    for bullet in EPB_PURPOSE_BULLETS:
        doc.add_paragraph(f"- {bullet}")
    doc.add_paragraph()

    # --- 5. Документы ---
    doc.add_heading("Сведения о документах, рассмотренных в процессе экспертизы", level=1)
    doc.add_paragraph(
        "Перечень документации, рассмотренной в процессе экспертизы промышленной "
        "безопасности с указанием объёма материалов, представлен в таблице:"
    )
    doc_keys: List[str] = []
    if isinstance(ctx.docs_dict, dict):
        doc_keys.extend(str(k) for k in ctx.docs_dict.keys())
    if isinstance(ctx.docs_info, dict):
        doc_keys.extend(str(k) for k in ctx.docs_info.keys())
    doc_keys = sorted(set(doc_keys), key=lambda x: int(x) if x.isdigit() else 999)
    if doc_keys and ctx.doc_meta_fn:
        tbl = doc.add_table(rows=len(doc_keys) + 1, cols=2)
        tbl.style = "Table Grid"
        tbl.rows[0].cells[0].text = "Наименование документа"
        tbl.rows[0].cells[1].text = "Шифр, индикация для идентификации"
        for i, num in enumerate(doc_keys, 1):
            present, doc_number, doc_date = ctx.doc_meta_fn(num)
            if present is False:
                continue
            name = EPB_DOCUMENT_NAMES.get(num, f"Документ №{num}")
            ident = doc_number or "—"
            if doc_date:
                ident = f"{ident} от {doc_date}" if ident != "—" else f"от {doc_date}"
            tbl.rows[i].cells[0].text = name
            tbl.rows[i].cells[1].text = ident
    else:
        doc.add_paragraph("Перечень документов – по данным чек-листа обследования.")
    doc.add_paragraph(
        "Сведения об информации автоматизированных систем мониторинга технического "
        "состояния объекта экспертизы на опасном производственном объекте."
    )
    doc.add_paragraph(
        g(
            "asm_monitoring_info",
            default="Автоматизированные системы мониторинга технического состояния "
            "объекта экспертизы на опасном производственном объекте не применяются.",
        )
    )
    doc.add_paragraph()

    # --- 6. Характеристика ---
    doc.add_heading("6. Назначение и краткая характеристика объекта экспертизы", level=1)
    doc.add_paragraph(
        f"{type_label} {ctx.device_name} выполнен как {construction}, "
        "подводящими и отводящими патрубками, арматурой."
    )
    doc.add_paragraph(f"{type_label} предназначен для {purpose}.")
    doc.add_paragraph(
        "Регистрационные данные, эксплуатационные параметры и сведения об основных "
        "частях сосуда представлены в Таблицах 6.1 и 6.2."
    )

    doc.add_paragraph("Таблица 6.1. Регистрационные данные сосуда")
    reg_tbl = doc.add_table(rows=2, cols=7)
    reg_tbl.style = "Table Grid"
    reg_headers = [
        "Наименование сосуда",
        "Индекс по схеме",
        "Заводской номер",
        "Регистрационный номер",
        "Завод-изготовитель",
        "Год изготовления",
        "Год ввода в эксплуатацию",
    ]
    for i, h in enumerate(reg_headers):
        reg_tbl.rows[0].cells[i].text = h
    reg_tbl.rows[1].cells[0].text = type_label
    reg_tbl.rows[1].cells[1].text = g("scheme_index", default=ctx.scheme_index or "—")
    reg_tbl.rows[1].cells[2].text = str(ctx.serial)
    reg_tbl.rows[1].cells[3].text = str(ctx.reg_no)
    reg_tbl.rows[1].cells[4].text = g("manufacturer", default="—")
    reg_tbl.rows[1].cells[5].text = g("manufacturing_year", "manufacture_year", default="—")
    reg_tbl.rows[1].cells[6].text = g("commissioning_year", default="—")
    doc.add_paragraph()

    doc.add_paragraph("Таблица 6.2. Паспортные характеристики сосуда")
    passport_rows = [
        ("Наименование частей сосуда", "Корпус"),
        ("Рабочее давление по паспорту, МПа (кгс/см²)", g("working_pressure", default="—")),
        ("Расчетное давление, МПа (кгс/см²)", g("design_pressure", default="—")),
        ("Пробное гидравлическое давление, МПа (кгс/см²)", g("test_pressure", default="—")),
        ("Максимально допустимая рабочая температура стенки, ℃", g("working_temperature", default="—")),
        ("Минимально допустимая рабочая температура стенки, ℃", g("design_temperature", default="—")),
        ("Наименование рабочей среды", g("working_medium", default="—")),
        ("Характеристика рабочей среды — ядовитость", g("medium_toxicity", default="—")),
        ("Характеристика рабочей среды — пожароопасность", g("medium_fire_hazard", default="да")),
        ("Характеристика рабочей среды — взрывоопасность", g("medium_explosion_hazard", default="да")),
        ("Характеристика рабочей среды — максимальная температура, ℃", g("medium_max_temp", default="—")),
        ("Прибавка для компенсации коррозии (эрозии), мм", g("corrosion_allowance", default="—")),
        ("Внутренний объём, м³", g("volume", "internal_volume", default="—")),
        ("Расчетный срок службы сосуда, лет", g("design_service_life", default="20")),
    ]
    pass_tbl = doc.add_table(rows=len(passport_rows) + 1, cols=2)
    pass_tbl.style = "Table Grid"
    pass_tbl.rows[0].cells[0].text = "Наименование параметра, характеристики"
    pass_tbl.rows[0].cells[1].text = "Значение, величина"
    for i, (k, v) in enumerate(passport_rows, 1):
        pass_tbl.rows[i].cells[0].text = k
        pass_tbl.rows[i].cells[1].text = str(v)
    doc.add_paragraph(
        "Сведения об основных элементах сосуда представлены в Приложении Б отчёта "
        "технической документации."
    )
    doc.add_paragraph()

    # --- 7. Результаты экспертизы ---
    codes = ctx.performed_codes
    doc.add_heading("7. Результаты проведённой экспертизы", level=1)

    doc.add_paragraph(
        "7.1. Анализ технической документации проводился в соответствии с "
        "Федеральным законом №116-ФЗ, приказом Ростехнадзора от 15.12.2020 №536 "
        "и приказом Ростехнадзора от 15.12.2020 №534."
    )
    doc_analysis = g(
        "doc_analysis_conclusion",
        default=f"{type_label} эксплуатируется по назначению, в пределах паспортных "
        "(разрешённых) характеристик и соответствует технической и эксплуатационной "
        "документации. Замечания по наличию и ведению эксплуатационной документации отсутствуют.",
    )
    doc.add_paragraph(str(doc_analysis))
    doc.add_paragraph(
        "Предоставленная документация соответствует требованиям нормативных правовых актов."
    )
    doc.add_paragraph("Анализ технической документации представлен в Приложении Б.")
    doc.add_paragraph()

    doc.add_paragraph(
        "7.2. Оперативное (функциональное) диагностирование проводилось в соответствии "
        "с требованиями Р ИСО 17637-2024, приказа Ростехнадзора от 15.12.2020 №536."
    )
    func_concl = g(
        "functional_diagnosis_conclusion",
        default="В ходе оперативной (функциональной) диагностики сосуда и его элементов "
        "повреждений и других неисправностей не обнаружено.",
    )
    doc.add_paragraph(str(func_concl))
    doc.add_paragraph(
        "Сосуд соответствует требованиям п.69, п.338–п.340, п.346, п.348, п.350 "
        "приказа Ростехнадзора от 15.12.2020 №536."
    )
    doc.add_paragraph(
        "Результаты оперативной (функциональной) диагностики оформлены "
        "Протоколом № 1, Приложение В."
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "7.3 Определение действующих повреждающих факторов, механизмов повреждения "
        "и восприимчивости материала технического устройства к механизмам повреждения."
    )
    doc.add_paragraph(
        "По результатам проведённого анализа документации и оперативного "
        "(функционального) диагностирования определены действующие повреждающие факторы, "
        "механизмы повреждения, параметры эксплуатации сосуда."
    )
    doc.add_paragraph("Для обследуемого сосуда характерны следующие механизмы повреждения:")
    for m in EPB_DAMAGE_MECHANISMS:
        doc.add_paragraph(f"- {m}")
    doc.add_paragraph()

    if _has_method(codes, "ВИК", "VIK"):
        doc.add_paragraph(
            "7.4 Визуальный и измерительный контроль выполнен в соответствии с "
            "требованиями приказа Ростехнадзора от 15.12.2020 №536, ГОСТ 34347-2017."
        )
        vik_concl = g(
            "vik_conclusion",
            default="Сосуд соответствует требованиям п.65, п.68, п.154 приказа "
            "Ростехнадзора от 15.12.2020 №536.",
        )
        doc.add_paragraph(str(vik_concl))
        doc.add_paragraph(
            "Результаты визуального и измерительного контроля оформлены "
            "Протоколом № 2, Приложение В."
        )
        doc.add_paragraph()

    hardness = g("hardness_tests", default=[])
    if _has_method(codes, "ТВ", "TVI", "HARDNESS") or (isinstance(hardness, list) and hardness):
        doc.add_paragraph(
            "7.5. Оценка механических свойств элементов сосуда выполнена в соответствии "
            "с требованиями приказа Ростехнадзора от 15.12.2020 №536, СО 153-34.17.439-2003."
        )
        doc.add_paragraph(
            g(
                "hardness_conclusion",
                default="В результате контроля твердости металла элементов сосуда отклонений "
                "твердости металла от указанного в нормативной документации не выявлено.",
            )
        )
        doc.add_paragraph(
            "Результаты контроля твердости оформлены Протоколом № 3, Приложение В."
        )
        doc.add_paragraph()

    thickness = g("thickness_measurements", default=[])
    if _has_method(codes, "УЗТ", "UZT") or (isinstance(thickness, list) and thickness):
        doc.add_paragraph(
            "7.6. Ультразвуковой контроль толщины стенок элементов сосуда выполнен "
            "в соответствии с требованиями приказа Ростехнадзора от 15.12.2020 №536."
        )
        doc.add_paragraph(
            g(
                "uzt_conclusion",
                default="В результате ультразвукового контроля толщины стенок элементов "
                "сосуда утонений, вследствие коррозионного и эрозионного износа сверх "
                "минимального значения, установленного расчётом на прочность, не выявлено.",
            )
        )
        doc.add_paragraph(
            "Результаты ультразвукового контроля толщины оформлены "
            "Протоколом № 4, Приложение В."
        )
        doc.add_paragraph()

    if _has_method(codes, "МПК", "МК", "MK", "MPD", "PVK"):
        doc.add_paragraph(
            "7.7. Магнитопорошковый контроль сварных соединений выполнен в соответствии "
            "с требованиями приказа Ростехнадзора от 15.12.2020 №536, РД 13-05-2006."
        )
        doc.add_paragraph(
            g(
                "mpk_conclusion",
                default="В результате магнитопорошкового контроля индикационных следов "
                "удалённой формы не выявлено.",
            )
        )
        doc.add_paragraph(
            "Результаты магнитопорошкового контроля оформлены "
            "Протоколом № 5, Приложение В."
        )
        doc.add_paragraph()

    weld = g("weld_inspections", default=[])
    if _has_method(codes, "УЗК", "UZK") or (isinstance(weld, list) and weld):
        doc.add_paragraph(
            "7.8. Ультразвуковой контроль сварных соединений выполнен в соответствии "
            "с требованиями приказа Ростехнадзора от 15.12.2020 №536, ГОСТ Р ИСО 17640-2016."
        )
        doc.add_paragraph(
            g(
                "uzk_conclusion",
                default="При контроле сварных соединений ультразвуковым методом дефектов, "
                "превышающих браковочный уровень, не обнаружено.",
            )
        )
        doc.add_paragraph(
            "Результаты ультразвукового контроля сварных соединений оформлены "
            "Протоколом № 6, Приложение В."
        )
        doc.add_paragraph()

    doc.add_paragraph(
        "7.9. Гидравлическое испытание выполнено эксплуатирующей организацией "
        "в соответствии с требованиями приказа Ростехнадзора от 15.12.2020 №536."
    )
    doc.add_paragraph(
        f"{type_label} зав.№ {ctx.serial} рег.№ {ctx.reg_no} прошёл гидравлические испытания."
    )
    doc.add_paragraph("Сосуд соответствует требованиям п.188 приказа Ростехнадзора от 15.12.2020 №536.")
    doc.add_paragraph()

    doc.add_paragraph(
        "7.10. Расчётные и аналитические процедуры оценки и прогнозирования "
        "технического состояния сосуда выполнены в соответствии с требованиями "
        "ГОСТ 34233.1-2017, ГОСТ 34233.2-2017, приказа Ростехнадзора от 15.12.2020 №536."
    )
    doc.add_paragraph(
        g(
            "calculation_result",
            default="Сосуд обладает необходимым запасом прочности для работы согласно "
            "разрешённым характеристикам, так как допустимое давление больше рабочего.",
        )
    )
    residual = g("residual_life_text", default=f"более {ctx.residual_life_years} лет")
    doc.add_paragraph(f"Расчёт остаточного ресурса сосуда показал, что он может эксплуатироваться «{residual}».")
    doc.add_paragraph(
        "Результаты расчётных и аналитических процедур представлены в Приложении Е."
    )
    doc.add_paragraph()

    # --- 8. Выводы ---
    doc.add_heading("8. Выводы заключения экспертизы", level=1)
    conclusion = (
        ctx.inspection_data.get("conclusion")
        or g(
            "epb_conclusion",
            "conclusion",
            default=f"На основании результатов выполненного комплекса работ по экспертизе "
            f"промышленной безопасности объект экспертизы {type_label} зав.№ {ctx.serial} "
            f"рег.№ {ctx.reg_no} соответствует требованиям промышленной безопасности.",
        )
    )
    doc.add_paragraph(f"8.1 {conclusion}")
    until = ctx.residual_life_until or g("residual_life_until", default="—")
    doc.add_paragraph(
        f"По результатам расчётных и аналитических процедур установлен срок дальнейшей "
        f"безопасной эксплуатации {ctx.residual_life_years} лет до {until}."
    )
    doc.add_paragraph("8.2 Дальнейшая безопасная эксплуатация объекта экспертизы должна осуществляться при обеспечении следующих условий:")
    conditions = [
        f"соблюдение рабочих параметров: разрешённое давление не более {ctx.allowed_pressure} МПа "
        f"и разрешённая температура среды {ctx.allowed_temperature}°С;",
        "соблюдение требований Федерального закона №116, приказа Ростехнадзора от 15.12.2020 №536 и №534;",
        "внесение в реестр Ростехнадзора данного заключения экспертизы промышленной безопасности;",
        "выполнение регламента контроля технического состояния объекта экспертизы, указанного в разделе 9.",
    ]
    for c in conditions:
        doc.add_paragraph(f"- {c}")
    doc.add_paragraph()

    for ie in ctx.inspection_engineers[:2]:
        if isinstance(ie, dict) and ie.get("full_name"):
            area = ie.get("expert_area") or ie.get("role") or "Эксперт"
            cert = ie.get("certificate_number") or ""
            valid = ie.get("valid_until") or ""
            doc.add_paragraph(f"{area}     {ie['full_name']}")
            if cert:
                doc.add_paragraph(f"Удостоверение № {cert}" + (f" до {valid}" if valid else ""))
    doc.add_paragraph()

    # --- 9. Регламент ---
    doc.add_heading("9. Регламент контроля технического состояния", level=1)
    doc.add_paragraph(
        f"Регламент контроля технического состояния {type_label} зав.№ {ctx.serial}, "
        f"рег.№ {ctx.reg_no} в течение назначенного срока эксплуатации приведён в таблице 9.1."
    )
    doc.add_paragraph("Таблица 9.1. Регламент контроля технического состояния")
    reglament = g("inspection_reglament", default=None)
    if isinstance(reglament, list) and reglament:
        rt = doc.add_table(rows=len(reglament) + 1, cols=4)
        rt.style = "Table Grid"
        for i, h in enumerate(
            ["Наименование вида контроля", "Периодичность контроля", "Исполнитель", "Дата"]
        ):
            rt.rows[0].cells[i].text = h
        for i, row in enumerate(reglament, 1):
            if isinstance(row, dict):
                rt.rows[i].cells[0].text = str(row.get("control_type") or row.get("name") or "—")
                rt.rows[i].cells[1].text = str(row.get("periodicity") or "—")
                rt.rows[i].cells[2].text = str(row.get("executor") or "—")
                rt.rows[i].cells[3].text = str(row.get("date") or "—")
    else:
        default_reg = [
            (
                "Техническое освидетельствование (наружный и внутренний осмотр, УЗТ)",
                "1 раз в 2 года",
                "Ответственное лицо за исправное состояние и безопасную эксплуатацию",
                g("next_inspection_date", default="—"),
            ),
            (
                "Гидравлическое испытание пробным давлением",
                "через 8 лет",
                "Ответственное лицо за исправное состояние и безопасную эксплуатацию",
                g("next_hydrotest_date", default="—"),
            ),
            (
                "Экспертиза промышленной безопасности",
                f"через {ctx.residual_life_years} лет",
                "Специализированная организация, имеющая лицензию на ЭПБ",
                until,
            ),
        ]
        rt = doc.add_table(rows=len(default_reg) + 1, cols=4)
        rt.style = "Table Grid"
        for i, h in enumerate(
            ["Наименование вида контроля", "Периодичность контроля", "Исполнитель", "Дата"]
        ):
            rt.rows[0].cells[i].text = h
        for i, row in enumerate(default_reg, 1):
            for j, val in enumerate(row):
                rt.rows[i].cells[j].text = val
    doc.add_paragraph()


def append_epb_appendix_act(doc: Document, ctx: EpbReportContext) -> None:
    """Приложение А — акт о проведении работ (образец 25-3173)."""
    g = ctx.g
    preset = preset_from_equipment_data(ctx.equipment_data)
    type_label = device_type_label(preset)
    doc.add_paragraph(
        f"{ctx.date_perf_ru} года аттестованными специалистами лаборатории "
        f"неразрушающего контроля {ctx.contractor} было проведено техническое "
        f"диагностирование и оценка технического состояния {type_label} "
        f"зав.№ {ctx.serial}, рег.№ {ctx.reg_no}, эксплуатирующегося на опасном "
        f"производственном объекте – {ctx.opo_name}, {ctx.opo_class} класс опасности, "
        "проводимого в рамках экспертизы промышленной безопасности."
    )
    doc.add_paragraph(
        "Работы выполнены в соответствии с «Программой проведения работ по экспертизе "
        "промышленной безопасности сосудов, работающих под давлением»."
    )
    doc.add_paragraph(
        "Перечень работ, выполненных в ходе проведения технического диагностирования:"
    )
    for w in [
        "Анализ технической документации;",
        "Техническое диагностирование:",
        "Визуальный и измерительный контроль сосуда;",
        "Оперативное (функциональное) диагностирование сосуда;",
        "Ультразвуковой контроль толщины стенок элементов сосуда;",
        "Магнитопорошковый контроль качества сварных соединений сосуда;",
        "Ультразвуковой контроль качества сварных соединений сосуда;",
        "Оценка механических свойств элементов сосуда;",
        "Расчётные и аналитические процедуры оценки и прогнозирования технического состояния;",
        "Проведение гидравлического испытания сосуда.",
    ]:
        doc.add_paragraph(w)
    doc.add_paragraph(
        "По результатам работ составлен отчёт по анализу технической документации, "
        "оформлены протоколы неразрушающего контроля."
    )
    for ie in ctx.inspection_engineers[:3]:
        if isinstance(ie, dict) and ie.get("full_name"):
            role = ie.get("expert_area") or ie.get("role") or "Специалист"
            doc.add_paragraph(f"{role}   {ie['full_name']}")
    doc.add_paragraph()


def append_epb_appendix_doc_analysis(doc: Document, ctx: EpbReportContext) -> None:
    """Приложение Б — отчёт по анализу технической документации."""
    g = ctx.g
    preset = preset_from_equipment_data(ctx.equipment_data)
    type_label = device_type_label(preset)
    year = g("commissioning_year", default="—")
    doc.add_paragraph(
        f"{type_label} зав.№ {ctx.serial}, рег.№ {ctx.reg_no} введён в эксплуатацию "
        f"{year} году."
    )
    doc.add_paragraph(
        "Режимы эксплуатации сосуда ведутся в соответствии с требованиями приказа "
        "Ростехнадзора от 15.12.2020 №536 и требованиями к монтажу и эксплуатации."
    )
    from epb_protocol_tables import append_appendix_b_tables

    append_appendix_b_tables(doc, ctx)
    doc.add_paragraph(
        str(
            g(
                "doc_analysis_appendix_conclusion",
                default=f"{type_label} эксплуатируется по назначению, в пределах паспортных "
                "(разрешённых) характеристик и соответствует технической и эксплуатационной "
                "документации. Замечания по наличию и ведению эксплуатационной документации отсутствуют.",
            )
        )
    )
    doc.add_paragraph("Заключение по результатам анализа технической документации")
    doc.add_paragraph(
        "Предоставленная документация соответствует требованиям нормативных правовых актов."
    )
    doc.add_paragraph()
